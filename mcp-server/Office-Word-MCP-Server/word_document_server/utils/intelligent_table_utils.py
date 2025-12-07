"""
Intelligent table replacement utilities for Word Document Server.
Provides intelligent text distribution across table columns.
"""
import re
from typing import Dict, List, Tuple, Optional
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P


def find_table_containing_text(doc: Document, text: str, match_case: bool = True, table_index: Optional[int] = None) -> Optional[Tuple[int, int, int]]:
    """
    Find table, row, and column index containing the specified text.
    Prioritizes exact matches in first column (for abbreviation tables).
    
    Args:
        doc: Document object
        text: Text to search for
        match_case: Whether to match case
        table_index: Optional specific table index to search in (if specified, only this table is searched)
    
    Returns:
        Tuple of (table_index, row_index, col_index) or None if not found
    """
    search_text = text.strip() if match_case else text.strip().lower()
    
    # Determine which tables to search
    if table_index is not None:
        # Search only in specified table
        if table_index >= len(doc.tables):
            return None
        tables_to_search = [(table_index, doc.tables[table_index])]
    else:
        # Search in all tables
        tables_to_search = list(enumerate(doc.tables))
    
    # First pass: look for exact match in first column (abbreviation tables)
    for table_idx, table in tables_to_search:
        for row_idx, row in enumerate(table.rows):
            if len(row.cells) > 0:
                first_cell = row.cells[0]
                cell_text = first_cell.text.strip()
                if not match_case:
                    cell_text = cell_text.lower()
                
                # Exact match in first cell (abbreviation column)
                if cell_text == search_text:
                    return (table_idx, row_idx, 0)
    
    # Second pass: look for text as a separate word in first column
    for table_idx, table in tables_to_search:
        for row_idx, row in enumerate(table.rows):
            if len(row.cells) > 0:
                first_cell = row.cells[0]
                cell_text = first_cell.text.strip()
                if not match_case:
                    cell_text = cell_text.lower()
                
                # Check if search_text is a word boundary match
                pattern = r'(?:^|[\s,;])' + re.escape(search_text) + r'(?:[\s,;]|$)'
                if re.search(pattern, cell_text):
                    return (table_idx, row_idx, 0)
    
    # Third pass: look for substring match in any cell (fallback)
    for table_idx, table in tables_to_search:
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text
                if not match_case:
                    cell_text = cell_text.lower()
                
                if search_text in cell_text:
                    return (table_idx, row_idx, col_idx)
    
    return None


def analyze_table_structure(doc: Document, table_idx: int) -> Dict[str, any]:
    """
    Analyze table structure to determine column count and types.
    
    Args:
        doc: Document object
        table_idx: Index of the table
    
    Returns:
        Dictionary with table structure information
    """
    if table_idx >= len(doc.tables):
        return {"columns_count": 2, "column_types": ["key", "value"]}
    
    table = doc.tables[table_idx]
    columns_count = len(table.columns) if table.rows else 2
    
    # Analyze first few rows to understand column types
    column_types = []
    if table.rows:
        # Check header row if exists
        header_row = table.rows[0]
        for cell in header_row.cells:
            cell_text = cell.text.strip().lower()
            if any(keyword in cell_text for keyword in ['аббревиатура', 'сокращение', 'abbreviation']):
                column_types.append("key")
            elif any(keyword in cell_text for keyword in ['описание', 'наименование', 'description', 'name']):
                column_types.append("value")
            else:
                column_types.append("unknown")
        
        # If no headers detected, use default pattern
        if not column_types or all(t == "unknown" for t in column_types):
            column_types = ["key"] + ["value"] * (columns_count - 1) if columns_count > 1 else ["key", "value"]
    
    # Ensure we have types for all columns
    while len(column_types) < columns_count:
        if len(column_types) == 0:
            column_types.append("key")
        else:
            column_types.append("value")
    
    return {
        "columns_count": columns_count,
        "column_types": column_types[:columns_count],
        "rows_count": len(table.rows)
    }


def distribute_text_to_columns(new_text: str, columns_count: int, column_types: List[str]) -> List[str]:
    """
    Intelligently distribute text across table columns.
    For example: "ДКР Департамент кредитных рисков" -> ["ДКР", "Департамент кредитных рисков"]
    
    Args:
        new_text: Full new text to distribute
        columns_count: Number of columns in the table
        column_types: List of column types (e.g., ["key", "value"])
    
    Returns:
        List of text values for each column
    """
    if columns_count == 0:
        return []
    
    new_text = new_text.strip()
    if not new_text:
        return [""] * columns_count
    
    # For 2-column tables (most common case for abbreviation tables)
    if columns_count == 2:
        words = new_text.split()
        if len(words) == 1:
            # Single word - put in first column
            return [words[0], ""]
        elif len(words) >= 2:
            # Multiple words - first word is abbreviation, rest is description
            # Check if first word looks like abbreviation (short, uppercase, or all caps)
            first_word = words[0]
            if len(first_word) <= 5 or first_word.isupper() or first_word[0].isupper() and len(first_word) <= 8:
                # First word is likely abbreviation
                return [first_word, " ".join(words[1:])]
            else:
                # First word is too long - split differently
                # Try to find natural split point
                mid_point = len(words) // 2
                return [" ".join(words[:mid_point]), " ".join(words[mid_point:])]
    
    # For tables with more columns, try splitting by separators first
    # Common separators: space, dash, comma, semicolon
    parts = re.split(r'[^\w\s-]+', new_text)
    parts = [p.strip() for p in parts if p.strip()]
    
    if not parts:
        parts = [new_text]
    
    result = [""] * columns_count
    
    # If we have exactly the right number of parts
    if len(parts) == columns_count:
        return parts
    
    # If we have fewer parts than columns
    if len(parts) < columns_count:
        # First part goes to first column
        if parts:
            result[0] = parts[0]
        # Remaining parts go to second column
        if len(parts) > 1:
            result[1] = " ".join(parts[1:]) if columns_count > 1 else ""
        elif len(parts) == 1:
            # Only one part - try to split it intelligently
            words = parts[0].split()
            if len(words) >= 2:
                # Split words: first few for first column, rest for second
                split_point = min(2, len(words) // 2 + 1)
                result[0] = " ".join(words[:split_point])
                if columns_count > 1:
                    result[1] = " ".join(words[split_point:])
            else:
                result[0] = parts[0]
        
        return result
    
    # If we have more parts than columns, combine them intelligently
    # First column gets first part (or first few words)
    result[0] = parts[0]
    
    # Remaining columns get combined parts
    remaining_text = " ".join(parts[1:])
    
    if columns_count > 1:
        # Distribute remaining text across remaining columns
        if columns_count == 2:
            result[1] = remaining_text
        else:
            # For more columns, distribute evenly
            remaining_words = remaining_text.split()
            words_per_col = len(remaining_words) // (columns_count - 1)
            start_idx = 0
            for col_idx in range(1, columns_count):
                end_idx = start_idx + words_per_col if col_idx < columns_count - 1 else len(remaining_words)
                result[col_idx] = " ".join(remaining_words[start_idx:end_idx])
                start_idx = end_idx
    else:
        result[0] = new_text
    
    return result


def insert_paragraph_after_table(doc: Document, table_idx: int, text: str) -> bool:
    """
    Insert a paragraph immediately after a table.
    
    Args:
        doc: Document object
        table_idx: Index of the table
        text: Text for the new paragraph
    
    Returns:
        True if paragraph was inserted successfully
    """
    if table_idx >= len(doc.tables):
        return False
    
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        # Get table element
        table = doc.tables[table_idx]
        table_element = table._element
        
        # Get parent element (document body)
        parent = table_element.getparent()
        if parent is None:
            return False
        
        # Create a new paragraph element directly using OxmlElement
        # This avoids the problem of add_paragraph() adding to the end
        new_p = OxmlElement('w:p')
        
        # Create paragraph properties
        pPr = OxmlElement('w:pPr')
        
        # Add some styling for the comment (italic, colored)
        rPr = OxmlElement('w:rPr')
        italic = OxmlElement('w:i')
        rPr.append(italic)
        
        # Add color (blue for comments)
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)
        
        # Create run element
        run = OxmlElement('w:r')
        run.append(rPr.copy() if hasattr(rPr, 'copy') else rPr)
        
        # Create text element
        t = OxmlElement('w:t')
        t.text = f"[Комментарий: {text}]"
        run.append(t)
        
        new_p.append(run)
        
        # Insert the paragraph element directly after the table
        table_element.addnext(new_p)
        
        return True
    except Exception as e:
        import logging
        logging.getLogger(__name__).warning(f"Failed to insert paragraph after table: {e}")
        return False


def intelligent_replace_in_table(
    doc: Document,
    table_idx: int,
    row_idx: int,
    old_text: str,
    new_text: str,
    match_case: bool = True
) -> bool:
    """
    Intelligently replace text in a table cell, distributing across columns if needed.
    
    Args:
        doc: Document object
        table_idx: Index of the table
        row_idx: Index of the row
        old_text: Text to find
        new_text: New text to replace with
        match_case: Whether to match case
    
    Returns:
        True if replacement was successful
    """
    if table_idx >= len(doc.tables):
        return False
    
    table = doc.tables[table_idx]
    if row_idx >= len(table.rows):
        return False
    
    # Analyze table structure
    structure = analyze_table_structure(doc, table_idx)
    columns_count = structure["columns_count"]
    column_types = structure["column_types"]
    
    # Find the column containing old_text
    found_col = None
    row = table.rows[row_idx]
    
    search_text = old_text if match_case else old_text.lower()
    
    for col_idx, cell in enumerate(row.cells):
        cell_text = cell.text
        check_text = cell_text if match_case else cell_text.lower()
        
        if search_text in check_text:
            found_col = col_idx
            break
    
    if found_col is None:
        return False
    
    # Distribute new text across columns
    distributed_text = distribute_text_to_columns(new_text, columns_count, column_types)
    
    # Replace text in each column of the row
    for col_idx, cell in enumerate(row.cells):
        if col_idx < len(distributed_text) and distributed_text[col_idx]:
            # Clear existing paragraphs in cell
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.clear()
            # Clear all paragraphs except first
            while len(cell.paragraphs) > 1:
                cell._element.remove(cell.paragraphs[-1]._element)
            # Set new text in first paragraph
            if cell.paragraphs:
                cell.paragraphs[0].text = distributed_text[col_idx]
            else:
                cell.text = distributed_text[col_idx]
    
    return True

