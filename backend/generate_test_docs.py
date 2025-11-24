"""
Генерация 4 тестовых Word документов для демонстрации функционала
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def generate_basic_source(filepath: str) -> None:
    """
    Базовый исходный документ - простая структура
    """
    doc = Document()

    # Заголовок
    title = doc.add_heading('Руководство пользователя - Версия 1.0', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Дата создания: 2024-01-15')
    doc.add_paragraph('Статус: Черновик')
    doc.add_paragraph()

    # Раздел 1
    doc.add_heading('1. Введение', level=1)
    doc.add_paragraph(
        'Данное руководство предназначено для пользователей системы управления проектами. '
        'Система позволяет эффективно планировать задачи, отслеживать прогресс и '
        'управлять командой. Текущая версия системы 2.5 включает базовый функционал.'
    )

    # Раздел 2
    doc.add_heading('2. Установка', level=1)

    doc.add_heading('2.1 Системные требования', level=2)
    doc.add_paragraph(
        'Операционная система: Windows 10 или выше\n'
        'Оперативная память: минимум 4 GB\n'
        'Свободное место на диске: 500 MB\n'
        'Браузер: Chrome, Firefox, Safari'
    )

    doc.add_heading('2.2 Процесс установки', level=2)
    doc.add_paragraph(
        '1. Скачайте установочный файл\n'
        '2. Запустите установщик от имени администратора\n'
        '3. Следуйте инструкциям мастера установки\n'
        '4. Перезагрузите компьютер после завершения'
    )

    # Раздел 3
    doc.add_heading('3. Основные функции', level=1)

    doc.add_heading('3.1 Создание проекта', level=2)
    doc.add_paragraph(
        'Для создания нового проекта нажмите кнопку "Новый проект" в главном меню. '
        'Заполните обязательные поля: название, описание, срок выполнения.'
    )

    doc.add_heading('3.2 Управление задачами', level=2)
    doc.add_paragraph(
        'Задачи можно создавать, редактировать и удалять. Каждая задача имеет статус, '
        'приоритет и исполнителя. Используйте фильтры для быстрого поиска задач.'
    )

    doc.add_heading('3.3 Отчетность', level=2)
    doc.add_paragraph(
        'Система генерирует различные отчеты: по проектам, по исполнителям, '
        'сводные отчеты. Отчеты можно экспортировать в формате PDF или Excel.'
    )

    # Раздел 4
    doc.add_heading('4. Устаревшие функции', level=1)
    doc.add_paragraph(
        'Следующие функции помечены как устаревшие и будут удалены в версии 3.0:\n'
        '- Экспорт в формат XML\n'
        '- Старый интерфейс отчетов\n'
        '- Интеграция с устаревшим API v1.0'
    )

    # Раздел 5
    doc.add_heading('5. Техническая поддержка', level=1)
    doc.add_paragraph(
        'Email: support@example.com\n'
        'Телефон: +7 (495) 123-45-67\n'
        'Время работы: Пн-Пт, 9:00-18:00 МСК'
    )

    doc.save(filepath)
    print(f"✓ Создан: {filepath}")


def generate_basic_changes(filepath: str) -> None:
    """
    Базовые инструкции изменений
    """
    doc = Document()

    # Заголовок
    title = doc.add_heading('Инструкции по обновлению руководства', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Дата: 2024-11-12')
    doc.add_paragraph('Версия: 2.0')
    doc.add_paragraph()

    # Изменение 1
    doc.add_heading('Изменение 1: Обновление версии системы', level=2)
    doc.add_paragraph(
        'Измени текст "Текущая версия системы 2.5" на "Текущая версия системы 3.0"'
    )

    # Изменение 2
    doc.add_heading('Изменение 2: Увеличение требований к памяти', level=2)
    doc.add_paragraph(
        'Измени в разделе 2.1 текст "минимум 4 GB" на "минимум 8 GB"'
    )

    # Изменение 3
    doc.add_heading('Изменение 3: Удаление устаревших функций', level=2)
    doc.add_paragraph(
        'Удали весь раздел "4. Устаревшие функции"'
    )

    # Изменение 4
    doc.add_heading('Изменение 4: Обновление статуса документа', level=2)
    doc.add_paragraph(
        'Измени текст "Статус: Черновик" на "Статус: Утверждено"'
    )

    # Изменение 5
    doc.add_heading('Изменение 5: Добавление нового раздела', level=2)
    doc.add_paragraph(
        'Добавь новый раздел "2.3 Активация лицензии" после раздела 2.2 '
        'со следующим текстом: "После установки необходимо активировать лицензию. '
        'Для этого введите лицензионный ключ в меню Помощь → Активация."'
    )

    doc.save(filepath)
    print(f"✓ Создан: {filepath}")


def generate_complex_source(filepath: str) -> None:
    """
    Сложный исходный документ - API документация
    """
    doc = Document()

    # Заголовок
    title = doc.add_heading('API Documentation - E-Commerce Platform', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Version: 2.1.0')
    doc.add_paragraph('Last Updated: 2024-11-01')
    doc.add_paragraph('Status: Production')
    doc.add_paragraph()

    # Раздел 1: Overview
    doc.add_heading('1. Overview', level=1)
    doc.add_paragraph(
        'This document describes the REST API for the E-Commerce Platform. '
        'The API provides endpoints for managing products, orders, customers, '
        'and payments. All endpoints return data in JSON format. '
        'The current API version is v2 and is backward compatible with v1.'
    )

    # Раздел 2: Authentication
    doc.add_heading('2. Authentication', level=1)

    doc.add_heading('2.1 API Key Authentication', level=2)
    doc.add_paragraph(
        'Include your API key in the Authorization header:\n'
        'Authorization: Bearer YOUR_API_KEY\n\n'
        'API keys can be generated in the Dashboard under Settings → API Keys.'
    )

    doc.add_heading('2.2 OAuth 2.0', level=2)
    doc.add_paragraph(
        'For user-specific operations, use OAuth 2.0. The platform supports '
        'authorization code flow and refresh tokens. Token lifetime is 3600 seconds.'
    )

    # Раздел 3: Endpoints
    doc.add_heading('3. API Endpoints', level=1)

    doc.add_heading('3.1 Products', level=2)
    doc.add_paragraph(
        'GET /api/v2/products - Get all products\n'
        'GET /api/v2/products/{id} - Get product by ID\n'
        'POST /api/v2/products - Create new product\n'
        'PUT /api/v2/products/{id} - Update product\n'
        'DELETE /api/v2/products/{id} - Delete product'
    )

    doc.add_heading('3.2 Orders', level=2)
    doc.add_paragraph(
        'GET /api/v2/orders - Get all orders\n'
        'POST /api/v2/orders - Create new order\n'
        'GET /api/v2/orders/{id} - Get order details\n'
        'PATCH /api/v2/orders/{id}/status - Update order status'
    )

    doc.add_heading('3.3 Rate Limits', level=2)
    doc.add_paragraph(
        'Standard tier: 1000 requests per hour\n'
        'Premium tier: 5000 requests per hour\n'
        'Enterprise tier: unlimited requests\n\n'
        'Rate limit headers are included in every response.'
    )

    # Раздел 4: Response Codes
    doc.add_heading('4. HTTP Response Codes', level=1)
    doc.add_paragraph(
        '200 OK - Request succeeded\n'
        '201 Created - Resource created successfully\n'
        '400 Bad Request - Invalid request parameters\n'
        '401 Unauthorized - Invalid or missing API key\n'
        '403 Forbidden - Insufficient permissions\n'
        '404 Not Found - Resource not found\n'
        '429 Too Many Requests - Rate limit exceeded\n'
        '500 Internal Server Error - Server error occurred'
    )

    # Раздел 5: Deprecated
    doc.add_heading('5. Deprecated Endpoints', level=1)
    doc.add_paragraph(
        'The following endpoints are deprecated and will be removed in v3:\n'
        '- GET /api/v1/products (use /api/v2/products)\n'
        '- POST /api/v1/orders (use /api/v2/orders)\n'
        '- GET /api/legacy/customers (no replacement)'
    )

    # Раздел 6: Webhooks
    doc.add_heading('6. Webhooks', level=1)
    doc.add_paragraph(
        'Configure webhooks to receive real-time notifications about events. '
        'Supported events: order.created, order.updated, payment.completed, '
        'product.updated. Webhook timeout is set to 5 seconds.'
    )

    # Раздел 7: Support
    doc.add_heading('7. Support', level=1)
    doc.add_paragraph(
        'API Support: api-support@example.com\n'
        'Documentation: https://docs.example.com/api\n'
        'Status Page: https://status.example.com\n'
        'Response time: within 24 hours for standard tier'
    )

    doc.save(filepath)
    print(f"✓ Создан: {filepath}")


def generate_complex_changes(filepath: str) -> None:
    """
    Сложные инструкции изменений для API документации
    """
    doc = Document()

    # Заголовок
    title = doc.add_heading('API Documentation Updates - v3.0 Migration', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Date: 2024-11-12')
    doc.add_paragraph('Migration Version: 3.0')
    doc.add_paragraph()

    # Изменение 1
    doc.add_heading('Change 1: Update API Version', level=2)
    doc.add_paragraph(
        'Измени текст "The current API version is v2" на '
        '"The current API version is v3"'
    )

    # Изменение 2
    doc.add_heading('Change 2: Update Version Number in Header', level=2)
    doc.add_paragraph(
        'Измени в заголовке документа текст "Version: 2.1.0" на "Version: 3.0.0"'
    )

    # Изменение 3
    doc.add_heading('Change 3: Increase Standard Rate Limit', level=2)
    doc.add_paragraph(
        'Измени в разделе 3.3 текст "Standard tier: 1000 requests per hour" '
        'на "Standard tier: 2000 requests per hour"'
    )

    # Изменение 4
    doc.add_heading('Change 4: Update Premium Rate Limit', level=2)
    doc.add_paragraph(
        'Измени текст "Premium tier: 5000 requests per hour" '
        'на "Premium tier: 10000 requests per hour"'
    )

    # Изменение 5
    doc.add_heading('Change 5: Remove Deprecated Section', level=2)
    doc.add_paragraph(
        'Удали весь раздел "5. Deprecated Endpoints"'
    )

    # Изменение 6
    doc.add_heading('Change 6: Add GraphQL Section', level=2)
    doc.add_paragraph(
        'Добавь новый раздел "2.3 GraphQL Authentication" после раздела 2.2 '
        'со следующим текстом: "GraphQL endpoint supports the same authentication '
        'methods as REST API. Use the endpoint /graphql for all GraphQL queries. '
        'GraphQL introspection is enabled by default."'
    )

    # Изменение 7
    doc.add_heading('Change 7: Update Webhook Timeout', level=2)
    doc.add_paragraph(
        'Измени в разделе 6 текст "Webhook timeout is set to 5 seconds" '
        'на "Webhook timeout is set to 10 seconds"'
    )

    # Изменение 8
    doc.add_heading('Change 8: Add New Product Endpoint', level=2)
    doc.add_paragraph(
        'Добавь в раздел 3.1 после строки "DELETE /api/v2/products/{id}" '
        'новую строку: "PATCH /api/v2/products/{id}/inventory - Update product inventory"'
    )

    # Изменение 9
    doc.add_heading('Change 9: Update Support Response Time', level=2)
    doc.add_paragraph(
        'Измени в разделе 7 текст "Response time: within 24 hours for standard tier" '
        'на "Response time: within 12 hours for standard tier"'
    )

    # Изменение 10
    doc.add_heading('Change 10: Update Status', level=2)
    doc.add_paragraph(
        'Измени текст "Status: Production" на "Status: Stable"'
    )

    doc.save(filepath)
    print(f"✓ Создан: {filepath}")


def main():
    """
    Генерация всех 4 тестовых файлов
    """
    output_dir = "/mnt/user-data/outputs"

    print("=" * 60)
    print("🎯 Генерация тестовых Word файлов")
    print("=" * 60)
    print()

    # Набор 1: Базовый (простой для начала)
    print("📄 Набор 1: Базовые файлы (Руководство пользователя)")
    print("-" * 60)
    generate_basic_source(f"{output_dir}/1_source_basic.docx")
    generate_basic_changes(f"{output_dir}/1_changes_basic.docx")
    print()

    # Набор 2: Сложный (API документация)
    print("📄 Набор 2: Сложные файлы (API Documentation)")
    print("-" * 60)
    generate_complex_source(f"{output_dir}/2_source_complex.docx")
    generate_complex_changes(f"{output_dir}/2_changes_complex.docx")
    print()

    print("=" * 60)
    print("✅ Все файлы успешно сгенерированы!")
    print("=" * 60)
    print()
    print("📋 Сгенерированные файлы:")
    print()
    print("Набор 1 (Базовый) - Руководство пользователя:")
    print("  • 1_source_basic.docx    - Исходный документ (5 разделов)")
    print("  • 1_changes_basic.docx   - Инструкции (5 изменений)")
    print()
    print("Набор 2 (Сложный) - API Документация:")
    print("  • 2_source_complex.docx  - Исходный документ (7 разделов)")
    print("  • 2_changes_complex.docx - Инструкции (10 изменений)")
    print()
    print("🎯 Рекомендация:")
    print("  1. Начните с Набора 1 (базовый) для первого теста")
    print("  2. Затем протестируйте Набор 2 (сложный)")
    print()
    print("💡 Ожидаемые результаты:")
    print()
    print("Набор 1:")
    print("  • 5 изменений")
    print("  • ~4-5 успешных")
    print("  • Время: ~25 секунд")
    print()
    print("Набор 2:")
    print("  • 10 изменений")
    print("  • ~9-10 успешных")
    print("  • Время: ~50 секунд")
    print()


if __name__ == "__main__":
    main()