"""
Сервис для перевода документов между русским и казахским языками.
"""

import os
import uuid
import logging
from pathlib import Path
from typing import Optional, Dict, Any
from docx import Document
from docx.shared import Inches
import asyncio
from openai import AsyncOpenAI
import httpx
import certifi

logger = logging.getLogger(__name__)

class TranslationService:
    def __init__(self):
        """Инициализация сервиса перевода."""
        # Загружаем переменные окружения
        env_path = Path(__file__).parent.parent / '.env'
        if env_path.exists():
            from dotenv import load_dotenv
            load_dotenv(dotenv_path=env_path)
        
        # Настройка OpenAI клиента
        self.openai_api_key = os.getenv("OPENAI_API_KEY")
        self.openai_model = os.getenv("OPENAI_MODEL", "gpt-4o")
        
        # Настройка SSL
        verify_ssl = os.environ.get("OPENAI_VERIFY_SSL", "false").lower() == "true"
        
        if verify_ssl:
            try:
                cert_path = certifi.where()
                logger.info(f"Использование SSL сертификатов из certifi: {cert_path}")
                verify_param = cert_path
            except Exception as e:
                logger.warning(f"Не удалось получить путь к certifi: {e}. Отключаем проверку SSL.")
                verify_param = False
        else:
            logger.warning("Проверка SSL отключена для OpenAI API")
            verify_param = False

        self._openai_http_client = httpx.AsyncClient(
            timeout=300.0,
            verify=verify_param,
        )
        
        if self.openai_api_key:
            self.openai_client = AsyncOpenAI(
                api_key=self.openai_api_key,
                http_client=self._openai_http_client,
            )
        else:
            logger.warning("OpenAI API key не найден")
            self.openai_client = None

    async def translate_text(self, text: str, source_lang: str, target_lang: str) -> str:
        """
        Переводит текст с одного языка на другой.
        
        Args:
            text: Текст для перевода
            source_lang: Исходный язык ('ru' или 'kz')
            target_lang: Целевой язык ('ru' или 'kz')
            
        Returns:
            Переведенный текст
        """
        if not self.openai_client:
            raise RuntimeError("OpenAI клиент не инициализирован")
        
        if not text.strip():
            return text
            
        # Определяем названия языков
        lang_names = {
            'ru': 'русский',
            'kz': 'казахский'
        }
        
        source_name = lang_names.get(source_lang, source_lang)
        target_name = lang_names.get(target_lang, target_lang)
        
        # Создаем промпт для перевода
        system_prompt = f"""Ты профессиональный переводчик с {source_name} языка на {target_name} язык.

Правила перевода:
1. Переводи точно и сохраняй смысл оригинала
2. Сохраняй стиль и тон текста
3. Для официальных документов используй соответствующую терминологию
4. Если встречаются специальные термины, переводи их корректно
5. Сохраняй форматирование (заголовки, списки, нумерацию)
6. Не добавляй никаких комментариев или пояснений
7. Возвращай только переведенный текст

Переведи следующий текст с {source_name} языка на {target_name} язык:"""

        try:
            response = await self.openai_client.chat.completions.create(
                model=self.openai_model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": text}
                ],
                max_tokens=4000,
                temperature=0.3,
            )
            
            translated_text = response.choices[0].message.content
            if not translated_text:
                raise RuntimeError("Пустой ответ от OpenAI")
                
            return translated_text.strip()
            
        except Exception as e:
            logger.error(f"Ошибка при переводе текста: {e}")
            raise RuntimeError(f"Ошибка перевода: {str(e)}")

    async def translate_document(
        self, 
        input_file: str, 
        output_file: str, 
        source_lang: str, 
        target_lang: str
    ) -> Dict[str, Any]:
        """
        Переводит Word документ.
        
        Args:
            input_file: Путь к исходному файлу
            output_file: Путь к выходному файлу
            source_lang: Исходный язык ('ru' или 'kz')
            target_lang: Целевой язык ('ru' или 'kz')
            
        Returns:
            Информация о результате перевода
        """
        try:
            # Открываем исходный документ
            doc = Document(input_file)
            
            # Счетчики для статистики
            translated_paragraphs = 0
            translated_tables = 0
            total_characters = 0
            
            # Переводим параграфы
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    original_text = paragraph.text
                    total_characters += len(original_text)
                    
                    try:
                        translated_text = await self.translate_text(
                            original_text, source_lang, target_lang
                        )
                        
                        # Заменяем текст в параграфе, сохраняя форматирование
                        if paragraph.runs:
                            # Очищаем все runs кроме первого
                            for run in paragraph.runs[1:]:
                                run.clear()
                            # Заменяем текст в первом run
                            paragraph.runs[0].text = translated_text
                        else:
                            paragraph.text = translated_text
                            
                        translated_paragraphs += 1
                        
                    except Exception as e:
                        logger.warning(f"Ошибка перевода параграфа: {e}")
                        continue
            
            # Переводим таблицы
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                original_text = paragraph.text
                                total_characters += len(original_text)
                                
                                try:
                                    translated_text = await self.translate_text(
                                        original_text, source_lang, target_lang
                                    )
                                    
                                    # Заменяем текст в ячейке
                                    if paragraph.runs:
                                        for run in paragraph.runs[1:]:
                                            run.clear()
                                        paragraph.runs[0].text = translated_text
                                    else:
                                        paragraph.text = translated_text
                                        
                                except Exception as e:
                                    logger.warning(f"Ошибка перевода ячейки таблицы: {e}")
                                    continue
                
                translated_tables += 1
            
            # Сохраняем переведенный документ
            doc.save(output_file)
            
            logger.info(f"Документ переведен: {translated_paragraphs} параграфов, {translated_tables} таблиц")
            
            return {
                "success": True,
                "translated_paragraphs": translated_paragraphs,
                "translated_tables": translated_tables,
                "total_characters": total_characters,
                "source_language": source_lang,
                "target_language": target_lang,
                "output_file": output_file
            }
            
        except Exception as e:
            logger.error(f"Ошибка при переводе документа: {e}")
            raise RuntimeError(f"Ошибка перевода документа: {str(e)}")

    def get_supported_languages(self) -> Dict[str, str]:
        """Возвращает список поддерживаемых языков."""
        return {
            'ru': 'Русский',
            'kz': 'Казахский'
        }

# Глобальный экземпляр сервиса
translation_service = TranslationService()
