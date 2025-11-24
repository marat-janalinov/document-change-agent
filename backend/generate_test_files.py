"""
Генератор тестовых Word документов для демонстрации функционала
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def generate_source_document(filepath: str) -> None:
    """
    Генерация исходного документа с пронумерованными разделами
    """
    doc = Document()
    
    # Заголовок документа
    title = doc.add_heading('Техническая документация API v1.0', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Раздел 1
    doc.add_heading('1. Введение', level=1)
    p1 = doc.add_paragraph(
        'Данный документ описывает API версии 1.0 для системы управления заказами. '
        'API предоставляет REST интерфейс для работы с заказами, клиентами и продуктами.'
    )
    
    # Раздел 2
    doc.add_heading('2. Аутентификация', level=1)
    
    doc.add_heading('2.1 Базовая аутентификация', level=2)
    p2_1 = doc.add_paragraph(
        'Система поддерживает базовую HTTP аутентификацию. '
        'Необходимо передавать заголовок Authorization с каждым запросом.'
    )
    
    doc.add_heading('2.2 Token аутентификация', level=2)
    p2_2 = doc.add_paragraph(
        'Для получения токена необходимо отправить POST запрос на /api/auth/token '
        'с учетными данными пользователя.'
    )
    
    # Раздел 3
    doc.add_heading('3. Endpoints', level=1)
    
    doc.add_heading('3.1 Управление заказами', level=2)
    p3_1 = doc.add_paragraph(
        'GET /api/orders - получение списка заказов'
    )
    p3_1.add_run('\nPOST /api/orders - создание нового заказа')
    p3_1.add_run('\nPUT /api/orders/{id} - обновление заказа')
    
    doc.add_heading('3.2 Версия API', level=2)
    p3_2 = doc.add_paragraph(
        'Текущая версия API v1.2 является стабильной. '
        'Все endpoints возвращают данные в формате JSON.'
    )
    
    doc.add_heading('3.3 Коды ответов', level=2)
    p3_3 = doc.add_paragraph(
        '200 - успешный запрос\n'
        '400 - некорректный запрос\n'
        '401 - требуется аутентификация\n'
        '404 - ресурс не найден\n'
        '500 - внутренняя ошибка сервера'
    )
    
    # Раздел 4
    doc.add_heading('4. Примеры использования', level=1)
    
    doc.add_heading('4.1 Создание заказа', level=2)
    p4_1 = doc.add_paragraph(
        'Пример запроса для создания заказа:\n'
        'POST /api/orders\n'
        'Content-Type: application/json\n'
    )
    
    doc.add_heading('4.2 Получение списка заказов', level=2)
    p4_2 = doc.add_paragraph(
        'Для получения всех заказов отправьте GET запрос на /api/orders'
    )
    
    # Раздел 5
    doc.add_heading('5. Устаревшие методы', level=1)
    p5 = doc.add_paragraph(
        'Следующие методы помечены как устаревшие и будут удалены в версии 2.0:\n'
        '- GET /api/v1/legacy/orders\n'
        '- POST /api/v1/legacy/customers\n'
    )
    
    # Раздел 6
    doc.add_heading('6. Ограничения', level=1)
    p6 = doc.add_paragraph(
        'Максимальное количество запросов: 1000 в час.\n'
        'Максимальный размер запроса: 10 MB.\n'
        'Timeout запроса: 30 секунд.'
    )
    
    # Сохранение документа
    doc.save(filepath)
    print(f"✓ Сгенерирован исходный документ: {filepath}")


def generate_changes_document(filepath: str) -> None:
    """
    Генерация документа с инструкциями изменений
    """
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('Инструкции по изменению документации', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        'Ниже перечислены изменения, которые необходимо применить к технической документации API.'
    )
    
    # Изменение 1
    doc.add_heading('Изменение 1: Обновление версии API', level=2)
    p1 = doc.add_paragraph(
        'Измени в разделе 3.2 текст "версия API v1.2" на "версия API v2.0"'
    )
    
    # Изменение 2
    doc.add_heading('Изменение 2: Удаление устаревших методов', level=2)
    p2 = doc.add_paragraph(
        'Удали весь раздел "5. Устаревшие методы"'
    )
    
    # Изменение 3
    doc.add_heading('Изменение 3: Добавление нового раздела', level=2)
    p3 = doc.add_paragraph(
        'Добавь новый раздел "2.3 OAuth 2.0" после раздела 2.2 со следующим текстом:\n'
        '"Система поддерживает OAuth 2.0 аутентификацию. '
        'Для получения access token используйте authorization code flow."'
    )
    
    # Изменение 4
    doc.add_heading('Изменение 4: Обновление лимитов', level=2)
    p4 = doc.add_paragraph(
        'В разделе 6 измени "Максимальное количество запросов: 1000 в час" '
        'на "Максимальное количество запросов: 5000 в час"'
    )
    
    # Изменение 5
    doc.add_heading('Изменение 5: Добавление нового endpoint', level=2)
    p5 = doc.add_paragraph(
        'В разделе 3.1 после "PUT /api/orders/{id}" добавь строку:\n'
        '"DELETE /api/orders/{id} - удаление заказа"'
    )
    
    # Сохранение документа
    doc.save(filepath)
    print(f"✓ Сгенерирован файл с инструкциями: {filepath}")


def generate_test_files(data_dir: str = "/data") -> dict:
    """
    Генерация обоих тестовых файлов
    """
    uploads_dir = os.path.join(data_dir, "uploads")
    os.makedirs(uploads_dir, exist_ok=True)
    
    source_path = os.path.join(uploads_dir, "source.docx")
    changes_path = os.path.join(uploads_dir, "changes.docx")
    
    generate_source_document(source_path)
    generate_changes_document(changes_path)
    
    return {
        "source": source_path,
        "changes": changes_path
    }


if __name__ == "__main__":
    # Тест генератора
    files = generate_test_files("/tmp")
    print("\n✓ Все тестовые файлы сгенерированы успешно")
    print(f"  - Исходный документ: {files['source']}")
    print(f"  - Инструкции: {files['changes']}")
