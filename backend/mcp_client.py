"""
MCP Client для взаимодействия с Office-Word-MCP-Server через FastMCP transport.
"""
from __future__ import annotations

import asyncio
import json
import os
import shutil
import uuid
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional

import logging
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from fastmcp.client import Client
from fastmcp.exceptions import ToolError
from pydantic import BaseModel
from dotenv import load_dotenv

# Загрузка переменных окружения из .env файла
# Ищем .env файл в корне проекта (на уровень выше backend/)
env_path = Path(__file__).parent.parent / '.env'
if env_path.exists():
    load_dotenv(dotenv_path=env_path)
else:
    # Если .env не найден в корне, пробуем загрузить из текущей директории
    load_dotenv()

logger = logging.getLogger(__name__)


@dataclass
class ToolCallResult:
    """Вспомогательная структура для результатов MCP."""

    text: str
    structured: Optional[Dict[str, Any]] = None


class MCPTextMatch(BaseModel):
    """Результат поиска текста"""

    paragraph_index: int
    text: str


class MCPClient:
    """
    Клиент для взаимодействия с MCP Word Server.
    Использует FastMCP Streamable HTTP transport для вызова MCP инструментов.
    """

    def __init__(self, base_url: Optional[str] = None):
        host = os.environ.get("MCP_SERVER_HOST", "mcp-server")
        port = os.environ.get("MCP_SERVER_PORT", "8000")
        default_url = f"http://{host}:{port}/mcp"

        self.base_url = base_url or default_url
        self._client = Client(self.base_url, init_timeout=30)
        self._lock = asyncio.Lock()

    async def close(self):
        """Закрытие клиента (освобождение транспорта)."""
        await self._client.close()

    async def _call_tool(
        self,
        name: str,
        arguments: Optional[Dict[str, Any]] = None,
        *,
        expect_json: bool = False,
    ) -> ToolCallResult:
        """
        Унифицированный вызов MCP инструмента.

        Args:
            name: имя инструмента MCP
            arguments: аргументы вызова
            expect_json: необходимо ли распарсить ответ как JSON
        """

        async with self._lock:
            async with self._client:
                try:
                    result = await self._client.call_tool(
                        name, arguments or {}, raise_on_error=True
                    )
                except ToolError as exc:
                    raise RuntimeError(f"MCP tool '{name}' failed: {exc}") from exc

        text_output = self._extract_text(result.content)
        structured = result.data if isinstance(result.data, dict) else None

        if expect_json:
            try:
                parsed = json.loads(text_output)
                structured = parsed
            except json.JSONDecodeError as exc:  # noqa: PERF203
                raise RuntimeError(
                    f"MCP tool '{name}' returned invalid JSON: {text_output}"
                ) from exc

        return ToolCallResult(text=text_output, structured=structured)

    @staticmethod
    def _extract_text(content: Optional[List[Any]]) -> str:
        if not content:
            return ""
        texts: List[str] = []
        for block in content:
            text = getattr(block, "text", None)
            if text:
                texts.append(text)
        return "\n".join(texts).strip()

    # ------------------------------------------------------------------ #
    # Чтение и анализ документов через MCP
    # ------------------------------------------------------------------ #

    async def get_document_text(self, filename: str) -> str:
        result = await self._call_tool(
            "get_document_text", {"filename": filename}, expect_json=False
        )
        return result.text

    async def get_document_outline(self, filename: str) -> Dict[str, Any]:
        result = await self._call_tool(
            "get_document_outline", {"filename": filename}, expect_json=True
        )
        return result.structured or {}

    async def find_text_in_document(
        self,
        filename: str,
        text_to_find: str,
        match_case: bool = True,
    ) -> List[MCPTextMatch]:
        result = await self._call_tool(
            "find_text_in_document",
            {
                "filename": filename,
                "text_to_find": text_to_find,
                "match_case": match_case,
            },
            expect_json=True,
        )
        data = result.structured or {}

        if isinstance(data, dict):
            occurrences = data.get("occurrences", [])
        elif isinstance(data, list):
            occurrences = data
        else:
            occurrences = []

        matches: List[MCPTextMatch] = []
        for entry in occurrences:
            if not isinstance(entry, dict):
                continue

            paragraph_index = entry.get("paragraph_index")
            
            # Обрабатываем совпадения в таблицах
            if paragraph_index is None:
                # Проверяем, есть ли информация о таблице
                table_info = entry.get("table_info") or entry.get("location")
                if table_info:
                    # Для совпадений в таблицах используем специальный индекс
                    # Это позволит системе найти текст в таблицах
                    logger.info(f"Найден текст в таблице: {entry}")
                    # Используем отрицательный индекс для обозначения совпадений в таблицах
                    paragraph_index = -1
                else:
                    # Пропускаем только если совсем нет информации о местоположении
                    logger.warning(f"Пропущено совпадение без информации о местоположении: {entry}")
                    continue

            context = entry.get("context") or entry.get("text") or text_to_find
            matches.append(
                MCPTextMatch(
                    paragraph_index=int(paragraph_index),
                    text=context,
                )
            )
        return matches

    # ------------------------------------------------------------------ #
    # Модификация документов
    # ------------------------------------------------------------------ #

    async def replace_text(
        self,
        filename: str,
        old_text: str,
        new_text: str,
        paragraph_index: Optional[int] = None,
    ) -> bool:
        arguments: Dict[str, Any] = {
            "filename": filename,
            "find_text": old_text,
            "replace_text": new_text,
        }
        if paragraph_index is not None:
            arguments["paragraph_index"] = paragraph_index

        # Сначала пробуем стандартную замену через MCP
        result = await self._call_tool("replace_text", arguments)
        
        # Если стандартная замена не сработала, пробуем локальную замену с поддержкой таблиц
        if not self._message_is_successful(result.text):
            logger.info(f"Стандартная замена не сработала, пробуем локальную замену с поддержкой таблиц")
            return self._replace_text_locally_with_tables(filename, old_text, new_text, paragraph_index)
        
        return True

    async def delete_paragraph(self, filename: str, paragraph_index: int) -> bool:
        result = await self._call_tool(
            "delete_paragraph",
            {"filename": filename, "paragraph_index": paragraph_index},
        )
        return self._message_is_successful(result.text)

    async def add_paragraph(
        self,
        filename: str,
        text: str,
        position: Optional[int] = None,
        style: Optional[str] = None,
    ) -> bool:
        if position is None:
            result = await self._call_tool(
                "add_paragraph",
                {
                    "filename": filename,
                    "text": text,
                    "style": style,
                },
            )
            return self._message_is_successful(result.text)

        self._ensure_document_exists(filename)
        self._insert_paragraph_locally(filename, text, position, style)
        return True

    async def add_heading(
        self,
        filename: str,
        text: str,
        level: int = 1,
        position: Optional[int] = None,
    ) -> bool:
        if position is None:
            result = await self._call_tool(
                "add_heading",
                {"filename": filename, "text": text, "level": level},
            )
            return self._message_is_successful(result.text)

        self._ensure_document_exists(filename)
        style_name = f"Heading {level}"
        self._insert_paragraph_locally(filename, text, position, style_name)
        return True

    async def add_table(
        self,
        filename: str,
        rows: List[List[str]],
        position: Optional[int] = None,
    ) -> bool:
        """
        Добавление таблицы в документ.
        
        Args:
            filename: путь к файлу
            rows: список строк, каждая строка - список ячеек
            position: индекс параграфа после которого вставить (опционально)
        """
        if not rows:
            return False
        
        num_columns = max(len(row) for row in rows) if rows else 0
        if num_columns == 0:
            return False
        
        # Нормализуем строки - все должны иметь одинаковое количество колонок
        normalized_rows = []
        for row in rows:
            normalized_row = row[:num_columns]  # Берем только нужное количество колонок
            # Дополняем пустыми строками, если нужно
            while len(normalized_row) < num_columns:
                normalized_row.append("")
            normalized_rows.append(normalized_row)
        
        # Если указана позиция, вставляем таблицу локально через python-docx
        if position is not None:
            self._ensure_document_exists(filename)
            return self._insert_table_locally(filename, normalized_rows, num_columns, position)
        
        # Иначе используем MCP инструмент
        result = await self._call_tool(
            "add_table",
            {
                "filename": filename,
                "rows": normalized_rows,
                "columns": num_columns,
            },
        )
        return self._message_is_successful(result.text)
    
    def _insert_table_locally(
        self,
        filename: str,
        rows: List[List[str]],
        num_columns: int,
        position: int,
    ) -> bool:
        """
        Локальная вставка таблицы после указанного параграфа.
        """
        try:
            doc = Document(filename)
            paragraph_count = len(doc.paragraphs)
            
            # Создаем таблицу (временно в конце документа)
            table = doc.add_table(rows=len(rows), cols=num_columns)
            
            # Заполняем таблицу данными
            for row_idx, row_data in enumerate(rows):
                if row_idx < len(table.rows):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < len(table.rows[row_idx].cells):
                            table.rows[row_idx].cells[col_idx].text = str(cell_text) if cell_text else ""
            
            # Если указана позиция и она валидна, перемещаем таблицу
            if position < paragraph_count:
                reference_paragraph = doc.paragraphs[position]
                table_element = table._element
                reference_element = reference_paragraph._p  # noqa: SLF001
                parent = reference_element.getparent()
                if parent is not None:
                    # Удаляем таблицу из текущего места
                    table_element.getparent().remove(table_element)
                    # Вставляем после нужного параграфа
                    reference_element.addnext(table_element)
            
            doc.save(filename)
            return True
        except Exception as e:
            logger.error(f"Ошибка при вставке таблицы: {e}", exc_info=True)
            return False

    async def copy_document(
        self,
        source_filename: str,
        destination_filename: Optional[str] = None,
    ) -> str:
        if destination_filename:
            shutil.copy2(source_filename, destination_filename)
            return destination_filename

        result = await self._call_tool(
            "copy_document",
            {"source_filename": source_filename},
        )
        # Попытка извлечь путь из ответа
        for token in result.text.split():
            if token.endswith(".docx"):
                return token
        return result.text

    async def get_paragraph_text(self, filename: str, paragraph_index: int) -> str:
        result = await self._call_tool(
            "get_paragraph_text_from_document",
            {"filename": filename, "paragraph_index": paragraph_index},
        )
        return result.text

    async def add_comment(
        self,
        filename: str,
        paragraph_index: int,
        comment_text: str,
        author: str = "DocumentChangeAgent",
    ) -> str:
        """
        Добавление аннотации к параграфу посредством вставки отдельного параграфа.

        Поскольку текущая версия Office-Word-MCP-Server не предоставляет инструмент
        для создания комментариев, добавляем примечание непосредственно в документ.
        """

        self._ensure_document_exists(filename)
        document = Document(filename)

        if paragraph_index < 0 or paragraph_index >= len(document.paragraphs):
            raise RuntimeError(
                f"Paragraph index {paragraph_index} is out of bounds for document {filename}"
            )

        paragraph = document.paragraphs[paragraph_index]
        new_paragraph = self._insert_paragraph_xml(
            paragraph, f"[ANNOTATION by {author}] {comment_text}"
        )
        
        # Безопасная установка стиля с проверкой наличия
        style_set = False
        for style_name in ["Comment", "Normal", "Default Paragraph Font"]:
            try:
                if style_name in document.styles:
                    new_paragraph.style = document.styles[style_name]  # type: ignore[index]
                    style_set = True
                    break
            except (KeyError, ValueError):
                continue
        
        # Если ни один стиль не найден, оставляем без стиля (используется стиль по умолчанию)
        if not style_set:
            logger.warning(f"Не удалось установить стиль для комментария в документе {filename}, используется стиль по умолчанию")

        document.save(filename)
        return f"COMMENT-{uuid.uuid4()}"

    # ------------------------------------------------------------------ #
    # Вспомогательные методы
    # ------------------------------------------------------------------ #

    @staticmethod
    def _message_is_successful(message: str) -> bool:
        lowered = message.lower()
        return not any(
            error_token in lowered
            for error_token in ["failed", "cannot", "does not exist", "error"]
        )

    @staticmethod
    def _ensure_document_exists(filename: str) -> None:
        if not os.path.exists(filename):
            raise RuntimeError(f"Документ {filename} не найден")

    def _insert_paragraph_locally(
        self,
        filename: str,
        text: str,
        position: int,
        style: Optional[str] = None,
    ) -> None:
        document = Document(filename)
        paragraph_count = len(document.paragraphs)

        if position >= paragraph_count:
            paragraph = document.add_paragraph(text)
        else:
            anchor_paragraph = document.paragraphs[position - 1]
            paragraph = self._insert_paragraph_xml(anchor_paragraph, text)

        if style:
            # Безопасная установка стиля с проверкой наличия
            style_set = False
            # Пробуем сначала указанный стиль
            try:
                if style in document.styles:
                    paragraph.style = document.styles[style]  # type: ignore[index]
                    style_set = True
            except (KeyError, ValueError):
                pass
            
            # Если не удалось, пробуем альтернативные стили
            if not style_set:
                for fallback_style in ["Normal", "Default Paragraph Font"]:
                    try:
                        if fallback_style in document.styles:
                            paragraph.style = document.styles[fallback_style]  # type: ignore[index]
                            style_set = True
                            logger.debug(f"Использован альтернативный стиль '{fallback_style}' вместо '{style}'")
                            break
                    except (KeyError, ValueError):
                        continue
                
                if not style_set:
                    logger.warning(f"Не удалось установить стиль '{style}' и альтернативы для параграфа в документе {filename}")

        document.save(filename)

    def _replace_text_locally_with_tables(
        self,
        filename: str,
        old_text: str,
        new_text: str,
        paragraph_index: Optional[int] = None,
    ) -> bool:
        """
        Локальная замена текста с поддержкой таблиц через python-docx
        """
        try:
            from docx import Document
            
            doc = Document(filename)
            replacements_made = 0
            
            # Замена в обычных параграфах
            for paragraph in doc.paragraphs:
                if old_text in paragraph.text:
                    # Заменяем текст в параграфе
                    for run in paragraph.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)
                            replacements_made += 1
            
            # Замена в таблицах
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if old_text in cell.text:
                            # Заменяем текст в ячейке таблицы
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if old_text in run.text:
                                        run.text = run.text.replace(old_text, new_text)
                                        replacements_made += 1
            
            if replacements_made > 0:
                doc.save(filename)
                logger.info(f"Локальная замена выполнена: {replacements_made} замен текста '{old_text}' на '{new_text}'")
                return True
            else:
                logger.warning(f"Текст '{old_text}' не найден в документе для замены")
                return False
                
        except Exception as e:
            logger.error(f"Ошибка при локальной замене текста: {e}")
            return False

    def _replace_text_locally_with_tables(
        self,
        filename: str,
        old_text: str,
        new_text: str,
        paragraph_index: Optional[int] = None,
    ) -> bool:
        """
        Локальная замена текста с поддержкой таблиц через python-docx
        """
        try:
            from docx import Document
            
            doc = Document(filename)
            replacements_made = 0
            
            # Замена в обычных параграфах
            for paragraph in doc.paragraphs:
                if old_text in paragraph.text:
                    # Заменяем текст в параграфе
                    for run in paragraph.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)
                            replacements_made += 1
            
            # Замена в таблицах
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if old_text in cell.text:
                            # Заменяем текст в ячейке таблицы
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if old_text in run.text:
                                        run.text = run.text.replace(old_text, new_text)
                                        replacements_made += 1
            
            if replacements_made > 0:
                doc.save(filename)
                logger.info(f"Локальная замена выполнена: {replacements_made} замен текста '{old_text}' на '{new_text}'")
                return True
            else:
                logger.warning(f"Текст '{old_text}' не найден в документе для замены")
                return False
                
        except Exception as e:
            logger.error(f"Ошибка при локальной замене текста: {e}")
            return False

    @staticmethod
    def _insert_paragraph_xml(reference: Paragraph, text: str) -> Paragraph:
        reference_element = reference._p  # noqa: SLF001
        new_p = OxmlElement("w:p")
        reference_element.addnext(new_p)
        new_paragraph = Paragraph(new_p, reference._parent)  # noqa: SLF001
        run = new_paragraph.add_run(text)
        if reference.runs and getattr(reference.runs[0].font.color, "rgb", None):
            run.font.color.rgb = reference.runs[0].font.color.rgb  # type: ignore[attr-defined]
        return new_paragraph


# Глобальный экземпляр клиента
mcp_client = MCPClient()
