"""
LLM-агент для применения изменений к Word документам без зависимостей от Parlant runtime.
"""
import inspect
import json
import logging
import os
import re
from datetime import datetime
from functools import wraps
from pathlib import Path
from typing import Any, Awaitable, Callable, Dict, List, Optional

import httpx
import certifi
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from openai import AsyncOpenAI
from dotenv import load_dotenv

from mcp_client import MCPTextMatch, mcp_client

# Загрузка переменных окружения из .env файла
# Ищем .env файл в корне проекта (на уровень выше backend/)
env_path = Path(__file__).parent.parent / '.env'
if env_path.exists():
    load_dotenv(dotenv_path=env_path)
else:
    # Если .env не найден в корне, пробуем загрузить из текущей директории
    load_dotenv()

# Настройка логирования
logger = logging.getLogger(__name__)

OperationCallback = Optional[Callable[[Dict[str, Any]], Awaitable[None]]]


class DocumentChangeAgent:
    """
    LLM-агент, который парсит инструкции изменений и управляет операциями MCP Word Server.
    """

    def _load_prompt(self, filename: str) -> str:
        """
        Загрузка промпта из markdown файла.
        Файлы находятся в директории prompts/ относительно файла parlant_agent.py.
        """
        try:
            # Используем persistent volume для промптов
            data_dir = os.getenv("DATA_DIR", "/data")
            prompts_dir = os.path.join(data_dir, "prompts")
            # Если папка не существует, пробуем локальную
            if not os.path.exists(prompts_dir):
                current_dir = os.path.dirname(os.path.abspath(__file__))
                prompts_dir = os.path.join(current_dir, "prompts")
            prompt_path = os.path.join(prompts_dir, filename)
            
            if os.path.exists(prompt_path):
                with open(prompt_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    # Удаляем заголовок markdown (если есть)
                    lines = content.split('\n')
                    # Пропускаем строки, начинающиеся с # (заголовки markdown)
                    prompt_lines = [line for line in lines if not line.strip().startswith('#')]
                    return '\n'.join(prompt_lines).strip()
            else:
                logger.warning(f"Файл промпта не найден: {prompt_path}, используем дефолтный промпт")
                return self._get_default_prompt(filename)
        except Exception as e:
            logger.error(f"Ошибка загрузки промпта {filename}: {e}", exc_info=True)
            return self._get_default_prompt(filename)
    
    def _get_default_prompt(self, filename: str) -> str:
        """
        Возвращает дефолтный промпт, если файл не найден.
        """
        if "instruction_check_system" in filename:
            return (
                "Ты эксперт по анализу документов с инструкциями изменений. "
                "Твоя задача - проанализировать содержимое документа и распознать ВСЕ инструкции по изменению текста, "
                "независимо от формата их представления. "
                "Документ может содержать инструкции в любом формате: списки, параграфы, таблицы, свободный текст. "
                "Твоя задача - найти ВСЕ инструкции и преобразовать их в структурированный JSON. "
                "Допустимые операции: REPLACE_TEXT, DELETE_PARAGRAPH, INSERT_PARAGRAPH, INSERT_SECTION, ADD_COMMENT, REPLACE_POINT_TEXT. "
                "КРИТИЧЕСКИ ВАЖНО: Ответ должен быть валидным JSON без комментариев, trailing commas и других ошибок. "
                "Используй экранирование для специальных символов в строках (\\\", \\n, \\t). "
                "Будь внимательным и найди ВСЕ инструкции, даже если они написаны в нестандартном или неочевидном формате."
            )
        elif "instruction_check_user" in filename:
            return (
                "Проанализируй содержимое документа и найди ВСЕ инструкции по изменению текста. "
                "Инструкции могут быть представлены в любом формате: списки, параграфы, таблицы, свободный текст."
            )
        else:
            return ""

    SUPPORTED_OPERATIONS = {
        "REPLACE_TEXT",
        "DELETE_PARAGRAPH",
        "INSERT_PARAGRAPH",
        "INSERT_SECTION",
        "INSERT_TABLE",  # Вставка таблицы
        "ADD_COMMENT",
        "REPLACE_POINT_TEXT",  # Специальная операция для замены всего пункта
    }

    BASE_GUIDELINES: List[Dict[str, Any]] = [
        {
            "priority": "CRITICAL",
            "condition": "Получены текстовые инструкции изменений",
            "action": (
                "1. Считать файл инструкций через get_document_text.\n"
                "2. Вызвать LLM для структурирования изменений.\n"
                "3. Подготовить последовательность действий."
            ),
            "tools": ["get_document_text", "parse_changes_document"],
        },
        {
            "priority": "HIGH",
            "condition": "Необходим анализ структуры документа",
            "action": (
                "Используй get_document_outline и get_paragraph_text для понимания структуры "
                "и поиска точек привязки через find_text."
            ),
            "tools": ["get_document_outline", "get_paragraph_text", "find_text"],
        },
        {
            "priority": "HIGH",
            "condition": "Выполняются изменения документа",
            "action": (
                "Для замены текста используй replace_text, для вставки — add_paragraph/add_heading, "
                "для удаления — delete_paragraph. После успешного изменения добавь аннотацию через add_comment."
            ),
            "tools": ["replace_text", "add_paragraph", "add_heading", "delete_paragraph", "add_comment"],
        },
    ]

    AVAILABLE_TOOLS = [
        "parse_changes_document",
        "get_document_text",
        "get_document_outline",
        "find_text",
        "replace_text",
        "add_comment",
        "delete_paragraph",
        "add_paragraph",
        "add_heading",
        "copy_document",
        "get_paragraph_text",
    ]

    def __init__(self):
        self.openai_client: Optional[AsyncOpenAI] = None
        self._openai_http_client: Optional[httpx.AsyncClient] = None
        # Чтение модели из переменных окружения (загружены из .env)
        self.model_name: str = os.environ.get("OPENAI_MODEL", "gpt-4o")
        logger.info(f"Инициализация LLM агента с моделью: {self.model_name}")
        self._patch_openai_httpx()
        # Сохранение исходного текста инструкций для исправления target.text
        self._original_instructions_text: Optional[str] = None

    async def initialize(self) -> None:
        """
        Инициализация клиента OpenAI.
        """
        openai_key = os.environ.get("OPENAI_API_KEY")
        if not openai_key:
            raise RuntimeError("OPENAI_API_KEY не найден. Укажите ключ в .env.")

        # Увеличенный timeout для больших документов
        # Настройка SSL для работы с OpenAI API
        # Использование certifi для SSL сертификатов (решает проблему с сертификатами в Docker)
        verify_ssl = os.environ.get("OPENAI_VERIFY_SSL", "false").lower() == "true"
        
        if verify_ssl:
            # Использовать сертификаты из certifi
            try:
                cert_path = certifi.where()
                logger.info(f"Использование SSL сертификатов из certifi: {cert_path}")
                verify_param = cert_path
            except Exception as e:
                logger.warning(f"Не удалось получить путь к certifi: {e}. Отключаем проверку SSL.")
                verify_param = False
        else:
            logger.warning("Проверка SSL отключена (OPENAI_VERIFY_SSL=false или не установлен). Это небезопасно, но необходимо для работы на серверах с проблемами SSL.")
            verify_param = False
        
        self._openai_http_client = httpx.AsyncClient(
            timeout=300.0,  # 5 минут
            verify=verify_param,  # Использовать certifi сертификаты или отключить проверку
        )
        try:
            self.openai_client = AsyncOpenAI(
                api_key=openai_key,
                http_client=self._openai_http_client,
            )
        except Exception:
            await self._openai_http_client.aclose()
            self._openai_http_client = None
            raise

        if not self.model_name:
            self.model_name = "gpt-4o"
            logger.warning("OPENAI_MODEL не задан, используется модель по умолчанию: gpt-4o")

        logger.info("LLM агент инициализирован")

    def _get_mcp_tools_description(self) -> str:
        """
        Возвращает полное описание доступных MCP инструментов для LLM.
        """
        return """
## ПОЛНЫЙ СПИСОК MCP ИНСТРУМЕНТОВ ДЛЯ РАБОТЫ С ДОКУМЕНТАМИ:

### ОСНОВНЫЕ ИНСТРУМЕНТЫ ЧТЕНИЯ:

#### 1. get_document_text - ЧТЕНИЕ ПОЛНОГО ТЕКСТА
- **Назначение**: Получение всего содержимого документа
- **Параметры**: filename (string)
- **Возвращает**: Полный текст документа в виде строки
- **Применение**: Анализ структуры и содержимого, поиск информации
- **Пример**: get_document_text("document.docx")

#### 2. get_document_outline - ПОЛУЧЕНИЕ СТРУКТУРЫ
- **Назначение**: Анализ иерархической структуры документа
- **Параметры**: filename (string)
- **Возвращает**: JSON со структурой заголовков и разделов
- **Применение**: Навигация по разделам, анализ структуры
- **Пример**: get_document_outline("document.docx")

#### 3. get_paragraph_text - ЧТЕНИЕ КОНКРЕТНОГО ПАРАГРАФА
- **Назначение**: Получение текста определенного параграфа по индексу
- **Параметры**: filename (string), paragraph_index (integer)
- **Возвращает**: Текст указанного параграфа
- **Применение**: Точечный анализ содержимого пунктов
- **Пример**: get_paragraph_text("document.docx", 15)

### ИНСТРУМЕНТЫ ПОИСКА:

#### 4. find_text_in_document - ПОИСК ТЕКСТА С КОНТЕКСТОМ
- **Назначение**: Поиск всех вхождений текста с информацией о местоположении
- **Параметры**: filename (string), text_to_find (string), match_case (boolean, optional)
- **Возвращает**: Список объектов MCPTextMatch с полями:
  - location: местоположение (параграф, таблица, ячейка)
  - position: позиция в тексте
  - context: окружающий контекст
  - paragraph_index: индекс параграфа
- **Применение**: Локализация текста перед изменениями, анализ вхождений
- **Пример**: find_text_in_document("document.docx", "ДРМ", false)

### ИНСТРУМЕНТЫ ИЗМЕНЕНИЯ ТЕКСТА:

#### 5. replace_text - УНИВЕРСАЛЬНАЯ ЗАМЕНА ТЕКСТА
- **Назначение**: Поиск и замена текста во всем документе
- **Параметры**: filename (string), old_text (string), new_text (string), match_case (boolean, optional)
- **Возвращает**: Результат операции с количеством замен
- **Применение**: 
  - Замена аббревиатур в таблицах с интеллектуальным распределением по столбцам
  - Замена фраз в пунктах с сохранением номеров пунктов
  - Массовые замены по всему документу
- **Особенности**: 
  - Автоматически определяет контекст (таблица/параграф)
  - Интеллектуально распределяет изменения по столбцам таблиц
  - Сохраняет форматирование
- **Пример**: replace_text("document.docx", "ДРМ", "ДКР Департамент кредитных рисков", false)

### ИНСТРУМЕНТЫ УПРАВЛЕНИЯ ПАРАГРАФАМИ:

#### 6. delete_paragraph - УДАЛЕНИЕ ПАРАГРАФА
- **Назначение**: Полное удаление параграфа по индексу
- **Параметры**: filename (string), paragraph_index (integer)
- **Возвращает**: Результат операции
- **Применение**: Удаление целых пунктов, разделов, устаревшей информации
- **Пример**: delete_paragraph("document.docx", 25)

#### 7. add_paragraph - ДОБАВЛЕНИЕ НОВОГО ПАРАГРАФА
- **Назначение**: Вставка нового параграфа в указанную позицию
- **Параметры**: filename (string), text (string), position (integer, optional)
- **Возвращает**: Результат операции с индексом нового параграфа
- **Применение**: Добавление новых пунктов, дополнительной информации
- **Пример**: add_paragraph("document.docx", "Новый пункт документа", 10)

#### 8. add_heading - ДОБАВЛЕНИЕ ЗАГОЛОВКА
- **Назначение**: Вставка заголовка определенного уровня
- **Параметры**: filename (string), text (string), level (integer), position (integer, optional)
- **Возвращает**: Результат операции
- **Применение**: Создание новых разделов, подразделов, структурирование
- **Пример**: add_heading("document.docx", "Новый раздел", 2, 15)

### ИНСТРУМЕНТЫ РАБОТЫ С ТАБЛИЦАМИ:

#### 9. add_table - СОЗДАНИЕ ТАБЛИЦЫ
- **Назначение**: Создание новой таблицы с заголовками и данными
- **Параметры**: filename (string), headers (array), rows (array of arrays), position (integer, optional)
- **Возвращает**: Результат операции с информацией о созданной таблице
- **Применение**: Вставка структурированных данных, создание справочников
- **Пример**: add_table("document.docx", ["Аббревиатура", "Описание"], [["ДКР", "Департамент кредитных рисков"]], 20)

### ИНСТРУМЕНТЫ АННОТИРОВАНИЯ:

#### 10. add_comment - ДОБАВЛЕНИЕ КОММЕНТАРИЯ/АННОТАЦИИ
- **Назначение**: Вставка комментария или аннотации в документ
- **Параметры**: filename (string), text (string), position (integer, optional)
- **Возвращает**: Результат операции
- **Применение**: Пометки об изменениях, аннотации, примечания
- **Пример**: add_comment("document.docx", "Изменено согласно новым требованиям", 30)

### СЛУЖЕБНЫЕ ИНСТРУМЕНТЫ:

#### 11. copy_document - КОПИРОВАНИЕ ДОКУМЕНТА
- **Назначение**: Создание точной копии документа
- **Параметры**: source_filename (string), target_filename (string)
- **Возвращает**: Результат операции
- **Применение**: Создание резервных копий, версионирование
- **Пример**: copy_document("document.docx", "document_backup.docx")

## ПРАВИЛА ВЫБОРА ОПЕРАЦИЙ И ГЕНЕРАЦИИ JSON:

### ОБЯЗАТЕЛЬНЫЕ ПОЛЯ JSON:
```json
{
  "change_id": "CHG-001",
  "operation": "REPLACE_TEXT",
  "description": "Описание изменения",
  "target": {
    "text": "точный текст для поиска"
  },
  "payload": {
    "new_text": "новый текст для замены"
  }
}
```

### ПРАВИЛА ДЛЯ REPLACE_TEXT:

#### Для таблиц:
- **Инструкция**: "В таблице «Принятые сокращения» строку «ДРМ» изложить в редакции «ДКР Департамент кредитных рисков»"
- **target.text**: "ДРМ" (только аббревиатура для поиска)
- **payload.new_text**: "ДКР Департамент кредитных рисков" (полный новый текст)
- **Система автоматически**: распределит "ДКР" в первый столбец, "Департамент кредитных рисков" во второй

#### Для пунктов:
- **Инструкция**: "В пункте 32 слова «согласовывается с ДО и ДРМ» изложить в редакции «согласовывается с ДО»"
- **target.text**: "согласовывается с ДО и ДРМ" (точная фраза для замены)
- **payload.new_text**: "согласовывается с ДО" (новая фраза)
- **Система автоматически**: найдет пункт 32 и заменит только указанную фразу, не трогая номер

#### Для массовых замен:
- **Инструкция**: "По всему тексту заменить «ДРМ» на «ДКР»"
- **target.text**: "ДРМ"
- **payload.new_text**: "ДКР"
- **target.replace_all**: true

### КРИТИЧЕСКИ ВАЖНО:
1. **target.text** должен содержать ТОЧНЫЙ текст для поиска, БЕЗ кавычек
2. **payload.new_text** должен содержать ПОЛНЫЙ новый текст
3. НЕ используйте номера пунктов в target.text для замен внутри пунктов
4. Для таблиц указывайте только искомую аббревиатуру в target.text
5. Система автоматически определит контекст и применит интеллектуальную логику
"""

    async def _analyze_instruction_context(self, instruction_text: str, source_file: str) -> Dict[str, Any]:
        """
        Анализирует контекст инструкции: определяет тип элемента (параграф/таблица/ячейка)
        и рекомендует подходящий MCP инструмент.
        """
        logger.info(f"🔍 АНАЛИЗ КОНТЕКСТА: {instruction_text[:100]}...")
        
        context_analysis = {
            "instruction": instruction_text,
            "element_type": "unknown",  # paragraph, table, table_cell, document
            "recommended_tool": "replace_text",
            "reasoning": "",
            "target_location": None
        }
        
        # Анализируем ключевые слова для определения типа элемента
        instruction_lower = instruction_text.lower()
        
        # 1. Определяем тип элемента
        if "в таблице" in instruction_lower or "строку" in instruction_lower:
            context_analysis["element_type"] = "table_cell"
            context_analysis["reasoning"] = "Упоминается таблица или строка таблицы"
            
        elif any(word in instruction_lower for word in ["пункт", "пункте", "п.", "подпункт"]):
            context_analysis["element_type"] = "paragraph"
            context_analysis["reasoning"] = "Упоминается пункт или параграф"
            
        elif "по всему тексту" in instruction_lower:
            context_analysis["element_type"] = "document"
            context_analysis["reasoning"] = "Массовая замена по всему документу"
            
        # 2. Определяем рекомендуемый инструмент
        if any(word in instruction_lower for word in ["заменить", "изложить", "изменить"]):
            context_analysis["recommended_tool"] = "replace_text"
            
        elif any(word in instruction_lower for word in ["исключить", "удалить"]):
            context_analysis["recommended_tool"] = "delete_paragraph"
            
        elif any(word in instruction_lower for word in ["добавить", "вставить", "дополнить"]):
            if "таблиц" in instruction_lower:
                context_analysis["recommended_tool"] = "add_table"
            elif "заголов" in instruction_lower:
                context_analysis["recommended_tool"] = "add_heading"
            else:
                context_analysis["recommended_tool"] = "add_paragraph"
        
        # 3. Пытаемся найти целевой элемент в документе
        try:
            if context_analysis["element_type"] in ["table_cell", "paragraph"]:
                # Ищем упоминаемый текст в документе
                search_terms = []
                
                # Извлекаем потенциальные поисковые термины из инструкции
                quoted_text = re.findall(r'[«"](.*?)[»"]', instruction_text)
                if quoted_text:
                    search_terms.extend(quoted_text)
                
                # Ищем номера пунктов
                point_numbers = re.findall(r'пункт[е]?\s+(\d+)', instruction_lower)
                if point_numbers:
                    search_terms.extend([f"{num}." for num in point_numbers])
                
                if search_terms:
                    # Используем MCP для поиска
                    for term in search_terms[:2]:  # Ограничиваем поиск первыми двумя терминами
                        try:
                            matches = await mcp_client.find_text_in_document(source_file, term)
                            if matches:
                                context_analysis["target_location"] = {
                                    "search_term": term,
                                    "matches": len(matches),
                                    "first_match": matches[0] if matches else None
                                }
                                break
                        except Exception as e:
                            logger.debug(f"Ошибка поиска '{term}': {e}")
                            
        except Exception as e:
            logger.debug(f"Ошибка анализа целевого элемента: {e}")
        
        logger.info(f"📋 РЕЗУЛЬТАТ АНАЛИЗА: {context_analysis['element_type']} → {context_analysis['recommended_tool']}")
        return context_analysis

    async def _analyze_table_structure(self, source_file: str, target_text: str) -> Dict[str, Any]:
        """
        Анализирует структуру таблицы для правильного определения содержимого ячеек.
        """
        logger.info(f"🔍 АНАЛИЗ СТРУКТУРЫ ТАБЛИЦЫ для текста: {target_text}")
        
        table_analysis = {
            "found": False,
            "table_index": -1,
            "row_index": -1,
            "cell_index": -1,
            "full_cell_content": "",
            "recommended_target_text": target_text,
            "table_context": ""
        }
        
        try:
            # Используем MCP для поиска текста в документе
            matches = await mcp_client.find_text_in_document(source_file, target_text)
            
            if matches:
                for match in matches:
                    # Правильный доступ к атрибутам MCPTextMatch
                    if hasattr(match, 'location'):
                        location = match.location
                        context = match.context if hasattr(match, 'context') else ''
                    else:
                        # Для словарей используем .get()
                        location = match.get('location', '') if isinstance(match, dict) else ''
                        context = match.get('context', '') if isinstance(match, dict) else ''
                    
                    # Проверяем, находится ли текст в таблице
                    if 'Table' in location:
                        # Парсим информацию о местоположении
                        # Формат: 'Table 0, Row 3, Column 0'
                        parts = location.split(', ')
                        if len(parts) >= 3:
                            table_idx = int(parts[0].split(' ')[1])
                            row_idx = int(parts[1].split(' ')[1])
                            col_idx = int(parts[2].split(' ')[1])
                            
                            table_analysis.update({
                                "found": True,
                                "table_index": table_idx,
                                "row_index": row_idx,
                                "cell_index": col_idx,
                                "full_cell_content": context,
                                "table_context": f"Таблица {table_idx}, строка {row_idx}, ячейка {col_idx}"
                            })
                            
                            # Если контекст содержит больше информации, чем искомый текст,
                            # используем полный контекст как target_text
                            if len(context.strip()) > len(target_text.strip()) and target_text in context:
                                table_analysis["recommended_target_text"] = context.strip()
                                logger.info(f"📋 НАЙДЕНО ПОЛНОЕ СОДЕРЖИМОЕ ЯЧЕЙКИ: '{context.strip()}'")
                            
                            break
                            
            if not table_analysis["found"]:
                logger.warning(f"⚠️ Текст '{target_text}' не найден в таблицах")
            else:
                logger.info(f"✅ АНАЛИЗ ТАБЛИЦЫ ЗАВЕРШЕН: {table_analysis['table_context']}")
                
        except Exception as e:
            logger.error(f"Ошибка анализа структуры таблицы: {e}")
        
        return table_analysis

    async def _intelligent_table_analysis(self, source_file: str, instruction_text: str) -> Dict[str, Any]:
        """
        Динамический интеллектуальный анализ структуры таблицы.
        Определяет количество столбцов, их назначение и необходимые операции.
        Читает несколько строк таблицы для понимания структуры и содержания.
        """
        logger.info(f"🧠 ДИНАМИЧЕСКИЙ АНАЛИЗ ТАБЛИЦЫ для: {instruction_text[:50]}...")
        
        analysis = {
            "is_table_change": False,
            "table_structure": {
                "columns_count": 0,
                "column_types": [],
                "column_content": [],
                "sample_rows": []  # Добавляем образцы строк для анализа
            },
            "instruction_mapping": {
                "target_key": "",
                "new_values": [],
                "affected_columns": []
            },
            "recommended_operations": []
        }
        
        # Проверяем, касается ли инструкция таблицы
        if not ("таблице" in instruction_text.lower() and "строку" in instruction_text.lower()):
            return analysis
        
        analysis["is_table_change"] = True
        
        try:
            
            # Правильно извлекаем целевой текст из инструкции
            instruction_data = self._extract_target_and_new_text(instruction_text)
            
            if not instruction_data["target_text"]:
                logger.warning("Не удалось извлечь целевой текст из инструкции")
                return analysis
                
            target_key = instruction_data["target_text"]
            analysis["instruction_mapping"]["target_key"] = target_key
            analysis["instruction_mapping"]["instruction_type"] = instruction_data["instruction_type"]
            logger.info(f"🎯 Целевой текст для поиска: '{target_key}'")
            
            # Ищем строку в таблице
            matches = await mcp_client.find_text_in_document(source_file, target_key)
            
            if not matches:
                logger.warning(f"Строка с ключом '{target_key}' не найдена")
                return analysis
            
            # Анализируем первое совпадение в таблице
            for match in matches:
                # Исправляем ошибку: match может быть объектом, а не словарем
                if hasattr(match, 'location'):
                    location = match.location
                    context = match.context if hasattr(match, 'context') else ''
                else:
                    # Правильный доступ к атрибутам MCPTextMatch
                    if hasattr(match, 'location'):
                        location = match.location
                        context = match.context if hasattr(match, 'context') else ''
                    else:
                        # Для словарей используем .get()
                        location = match.get('location', '') if isinstance(match, dict) else ''
                        context = match.get('context', '') if isinstance(match, dict) else ''
                    
                if 'Table' in location:
                    logger.info(f"📍 Найдена строка в: {location}")
                    
                    # ДИНАМИЧЕСКИЙ АНАЛИЗ СТРУКТУРЫ СТРОКИ
                    # Здесь нужно получить всю строку и проанализировать ее столбцы
                    context = context
                    
                    # Парсим местоположение для получения координат
                    parts = location.split(', ')
                    if len(parts) >= 3:
                        table_idx = int(parts[0].split(' ')[1])
                        row_idx = int(parts[1].split(' ')[1])
                        
                        # Получаем структуру всей строки (это требует дополнительного MCP запроса)
                        row_structure = await self._analyze_table_row_structure(source_file, table_idx, row_idx)
                        
                        analysis["table_structure"] = row_structure
                        
                        # Извлекаем новые значения из инструкции
                        instruction_data = self._extract_target_and_new_text(instruction_text)
                        new_values = [instruction_data["new_text"]] if instruction_data["new_text"] else []
                        analysis["instruction_mapping"]["new_values"] = new_values
                        
                        # Сопоставляем новые значения со столбцами
                        affected_columns = self._map_values_to_columns(row_structure, new_values, target_key)
                        analysis["instruction_mapping"]["affected_columns"] = affected_columns
                        
                        # Создаем рекомендуемые операции
                        operations = self._create_adaptive_operations(row_structure, affected_columns, new_values)
                        analysis["recommended_operations"] = operations
                        
                        logger.info(f"📊 СТРУКТУРА: {row_structure['columns_count']} столбцов")
                        logger.info(f"🎯 ЗАТРОНУТЫЕ СТОЛБЦЫ: {affected_columns}")
                        logger.info(f"🔧 ОПЕРАЦИЙ: {len(operations)}")
                        
                        break
                        
        except Exception as e:
            logger.error(f"Ошибка динамического анализа таблицы: {e}")
        
        return analysis

    async def _analyze_table_row_structure(self, source_file: str, table_idx: int, row_idx: int) -> Dict[str, Any]:
        """
        Интеллектуальный анализ структуры конкретной строки таблицы.
        Читает несколько строк таблицы для понимания структуры и содержания.
        """
        logger.info(f"📊 ИНТЕЛЛЕКТУАЛЬНЫЙ АНАЛИЗ строки {row_idx} в таблице {table_idx}")
        
        structure = {
            "columns_count": 2,  # По умолчанию 2 столбца
            "column_types": ["key", "value"],
            "column_content": [],
            "sample_rows": [],
            "analysis_method": "default"
        }
        
        try:
            # Получаем текст документа для анализа
            doc_text = await mcp_client.get_document_text(source_file)
            
            # Читаем несколько строк таблицы для анализа структуры
            sample_rows = await self._read_table_sample_rows(source_file, doc_text, table_idx)
            structure["sample_rows"] = sample_rows
            
            if sample_rows:
                # Анализируем структуру на основе реальных данных
                analyzed_structure = self._analyze_real_table_structure(sample_rows)
                structure.update(analyzed_structure)
                structure["analysis_method"] = "real_data_analysis"
                logger.info(f"✅ Анализ реальных данных: {structure['columns_count']} столбцов, типы: {structure['column_types']}")
            else:
                # Используем эвристический анализ
                structure = self._heuristic_table_analysis()
                structure["analysis_method"] = "heuristic"
                logger.info(f"⚠️ Используется эвристический анализ: {structure['columns_count']} столбцов")
            
        except Exception as e:
            logger.error(f"Ошибка анализа структуры строки: {e}")
            # Возвращаем безопасные значения по умолчанию
            structure = {
                "columns_count": 2,
                "column_types": ["key", "value"],
                "column_content": [],
                "sample_rows": [],
                "analysis_method": "error_fallback"
            }
        
        return structure
    
    async def _read_table_sample_rows(self, source_file: str, doc_text: str, table_idx: int, max_rows: int = 3) -> List[Dict[str, Any]]:
        """
        Читает несколько строк таблицы для анализа структуры.
        
        Args:
            source_file: Путь к файлу
            doc_text: Текст документа
            table_idx: Индекс таблицы
            max_rows: Максимальное количество строк для чтения
            
        Returns:
            Список строк с их содержимым
        """
        sample_rows = []
        
        try:
            # Ищем таблицы в тексте по характерным паттернам
            
            # Паттерны для поиска таблиц
            table_patterns = [
                r'(\w+)\s+([^\n\r]+)',  # Простой паттерн: слово + описание
                r'([А-ЯЁ]{2,5})\s+([^\n\r]+)',  # Аббревиатура + описание
                r'(\d+\.?\d*)\s+([^\n\r]+)',  # Номер + описание
            ]
            
            lines = doc_text.split('\n')
            table_found = False
            rows_collected = 0
            
            for i, line in enumerate(lines):
                line = line.strip()
                if not line:
                    continue
                
                # Ищем начало таблицы по ключевым словам
                if any(keyword in line.lower() for keyword in ['таблица', 'сокращения', 'пояснения', 'обозначения']):
                    table_found = True
                    continue
                
                if table_found and rows_collected < max_rows:
                    # Пытаемся разобрать строку таблицы
                    for pattern in table_patterns:
                        match = re.match(pattern, line)
                        if match:
                            row_data = {
                                "row_index": rows_collected,
                                "raw_text": line,
                                "columns": list(match.groups()),
                                "column_count": len(match.groups())
                            }
                            sample_rows.append(row_data)
                            rows_collected += 1
                            logger.info(f"📋 Найдена строка таблицы {rows_collected}: {match.groups()}")
                            break
                
                # Прекращаем поиск если собрали достаточно строк
                if rows_collected >= max_rows:
                    break
                    
                # Прекращаем поиск если встретили конец таблицы
                if table_found and any(keyword in line.lower() for keyword in ['пункт', 'раздел', 'глава']):
                    break
            
            logger.info(f"📊 Собрано {len(sample_rows)} образцов строк таблицы")
            
        except Exception as e:
            logger.error(f"Ошибка чтения образцов строк таблицы: {e}")
        
        return sample_rows
    
    def _analyze_real_table_structure(self, sample_rows: List[Dict[str, Any]]) -> Dict[str, Any]:
        """
        Анализирует структуру таблицы на основе реальных данных.
        
        Args:
            sample_rows: Образцы строк таблицы
            
        Returns:
            Структура таблицы
        """
        if not sample_rows:
            return {
                "columns_count": 2,
                "column_types": ["key", "value"],
                "column_content": []
            }
        
        # Определяем количество столбцов
        columns_count = max(row["column_count"] for row in sample_rows)
        
        # Анализируем типы столбцов на основе содержимого
        column_types = []
        column_content = []
        
        for col_idx in range(columns_count):
            col_samples = []
            for row in sample_rows:
                if col_idx < len(row["columns"]):
                    col_samples.append(row["columns"][col_idx])
            
            # Определяем тип столбца
            col_type = self._determine_column_type(col_samples)
            column_types.append(col_type)
            column_content.append(col_samples[:3])  # Первые 3 образца
        
        logger.info(f"🧠 Анализ структуры: {columns_count} столбцов, типы: {column_types}")
        
        return {
            "columns_count": columns_count,
            "column_types": column_types,
            "column_content": column_content
        }
    
    def _determine_column_type(self, samples: List[str]) -> str:
        """
        Определяет тип столбца на основе образцов данных.
        
        Args:
            samples: Образцы данных столбца
            
        Returns:
            Тип столбца
        """
        if not samples:
            return "unknown"
        
        # Анализируем паттерны в образцах
        
        # Проверяем на аббревиатуры (короткие заглавные буквы)
        abbrev_count = sum(1 for s in samples if re.match(r'^[А-ЯЁ]{2,5}$', s.strip()))
        if abbrev_count > len(samples) * 0.5:
            return "abbreviation"
        
        # Проверяем на номера
        number_count = sum(1 for s in samples if re.match(r'^\d+\.?\d*$', s.strip()))
        if number_count > len(samples) * 0.5:
            return "number"
        
        # Проверяем на длинные описания
        desc_count = sum(1 for s in samples if len(s.strip()) > 10)
        if desc_count > len(samples) * 0.5:
            return "description"
        
        # По умолчанию - ключ
        return "key"
    
    def _analyze_table_patterns_in_text(self, doc_text: str, table_idx: int) -> Dict[str, Any]:
        """
        Анализирует паттерны таблицы в тексте документа.
        
        Args:
            doc_text: Полный текст документа
            table_idx: Индекс таблицы
            
        Returns:
            Результат анализа структуры таблицы
        """
        
        # Ищем признаки таблицы сокращений
        abbreviation_patterns = [
            r'сокращения?\s+и\s+пояснения',
            r'принятые\s+сокращения',
            r'список\s+сокращений',
            r'аббревиатур[ыа]'
        ]
        
        is_abbreviation_table = any(
            re.search(pattern, doc_text, re.IGNORECASE) 
            for pattern in abbreviation_patterns
        )
        
        if is_abbreviation_table:
            return {
                "success": True,
                "columns_count": 2,
                "column_types": ["abbreviation", "description"],
                "table_type": "abbreviations",
                "column_content": []
            }
        
        # Ищем другие паттерны таблиц
        # Если находим много коротких слов заглавными буквами - вероятно таблица сокращений
        uppercase_words = re.findall(r'\b[А-ЯA-Z]{2,6}\b', doc_text)
        if len(uppercase_words) > 5:
            return {
                "success": True,
                "columns_count": 2,
                "column_types": ["abbreviation", "description"],
                "table_type": "abbreviations_detected",
                "column_content": []
            }
        
        return {"success": False}
    
    def _heuristic_table_analysis(self) -> Dict[str, Any]:
        """
        Эвристический анализ структуры таблицы.
        Используется когда другие методы не сработали.
        
        Returns:
            Базовая структура таблицы
        """
        return {
            "columns_count": 2,
            "column_types": ["key", "value"],
            "table_type": "general",
            "column_content": []
        }

    def _extract_target_and_new_text(self, instruction_text: str) -> Dict[str, str]:
        """
        Правильно извлекает целевой текст и новый текст из инструкции.
        """
        result = {
            "target_text": "",
            "new_text": "",
            "instruction_type": "unknown"
        }
        
        logger.info(f"🔍 АНАЛИЗ ИНСТРУКЦИИ: {instruction_text}")
        
        # Тип 1: "В пункте X слова Y изложить в редакции Z"
        paragraph_match = re.search(r'пункте\s+(\d+)\s+слова\s*[«"\'](.*?)[»"\']\s+изложить.*?редакции:\s*[«"\'](.*?)[»"\']', instruction_text, re.IGNORECASE)
        if paragraph_match:
            paragraph_num = paragraph_match.group(1)
            target_phrase = paragraph_match.group(2).strip()
            new_phrase = paragraph_match.group(3).strip()
            
            result.update({
                "target_text": target_phrase,  # Ищем фразу, а не номер пункта!
                "new_text": new_phrase,
                "instruction_type": "paragraph_phrase_replacement",
                "paragraph_number": paragraph_num
            })
            
            logger.info(f"📋 ТИП: Замена фразы в пункте {paragraph_num}")
            logger.info(f"🎯 ЦЕЛЕВАЯ ФРАЗА: '{target_phrase}'")
            logger.info(f"📝 НОВАЯ ФРАЗА: '{new_phrase}'")
            return result
        
        # Тип 2: "В таблице строку X изложить в редакции Y"
        table_match = re.search(r'таблице.*?строку\s*[«"\'](.*?)[»"\']\s+изложить.*?редакции:\s*[«"\'](.*?)[»"\']', instruction_text, re.IGNORECASE)
        if table_match:
            target_key = table_match.group(1).strip()
            new_description = table_match.group(2).strip()
            
            result.update({
                "target_text": target_key,
                "new_text": new_description,
                "instruction_type": "table_row_replacement"
            })
            
            logger.info(f"📋 ТИП: Замена строки в таблице")
            logger.info(f"🎯 КЛЮЧ СТРОКИ: '{target_key}'")
            logger.info(f"📝 НОВОЕ ОПИСАНИЕ: '{new_description}'")
            return result
        
        # Тип 3: "По всему тексту X заменить на Y"
        mass_replace_match = re.search(r'всему тексту.*?[«"\'](.*?)[»"\'].*?заменить.*?[«"\'](.*?)[»"\']', instruction_text, re.IGNORECASE)
        if mass_replace_match:
            old_text = mass_replace_match.group(1).strip()
            new_text = mass_replace_match.group(2).strip()
            
            result.update({
                "target_text": old_text,
                "new_text": new_text,
                "instruction_type": "mass_replacement"
            })
            
            logger.info(f"📋 ТИП: Массовая замена")
            logger.info(f"🎯 СТАРЫЙ ТЕКСТ: '{old_text}'")
            logger.info(f"📝 НОВЫЙ ТЕКСТ: '{new_text}'")
            return result
        
        logger.warning(f"⚠️ НЕ УДАЛОСЬ РАСПОЗНАТЬ ТИП ИНСТРУКЦИИ: {instruction_text}")
        return result

    def _map_values_to_columns(self, row_structure: Dict[str, Any], new_values: List[str], target_key: str) -> List[int]:
        """
        Интеллектуально сопоставляет новые значения со столбцами таблицы.
        Анализирует содержимое инструкции, структуру таблицы и образцы строк.
        """
        affected_columns = []
        
        if not new_values or not row_structure.get("columns_count", 0):
            return affected_columns
        
        column_types = row_structure.get("column_types", [])
        sample_rows = row_structure.get("sample_rows", [])
        columns_count = row_structure.get("columns_count", 2)
        
        logger.info(f"🧠 ИНТЕЛЛЕКТУАЛЬНОЕ СОПОСТАВЛЕНИЕ:")
        logger.info(f"   Столбцов: {columns_count}")
        logger.info(f"   Типы столбцов: {column_types}")
        logger.info(f"   Новые значения: {new_values}")
        logger.info(f"   Целевой ключ: {target_key}")
        
        # Анализируем новое содержимое на основе реальной структуры
        new_content = new_values[0] if new_values else ""
        content_analysis = self._analyze_instruction_content(new_content, target_key, sample_rows)
        
        # Определяем затронутые столбцы на основе анализа содержимого
        affected_columns = self._determine_affected_columns(content_analysis, column_types, target_key)
        
        logger.info(f"🎯 РЕЗУЛЬТАТ СОПОСТАВЛЕНИЯ: столбцы {affected_columns}")
        return affected_columns
    
    def _analyze_instruction_content(self, new_content: str, target_key: str, sample_rows: List[Dict[str, Any]]) -> Dict[str, Any]:
        """
        Анализирует содержимое инструкции в контексте реальной структуры таблицы.
        
        Args:
            new_content: Новое содержимое из инструкции
            target_key: Ключ для поиска (например, аббревиатура)
            sample_rows: Образцы строк таблицы
            
        Returns:
            Анализ содержимого инструкции
        """
        analysis = {
            "has_key_change": False,
            "has_description_change": False,
            "key_part": "",
            "description_part": "",
            "change_type": "unknown"
        }
        
        try:
            # Анализируем содержимое инструкции
            
            # Ищем паттерны изменения ключа и описания
            # Например: "ПД Проектные дирекции 1,2,3,4,5,6."
            
            # Проверяем, содержит ли инструкция целевой ключ
            if target_key and target_key.upper() in new_content.upper():
                analysis["has_key_change"] = True
                analysis["key_part"] = target_key
                
                # Извлекаем описание после ключа
                pattern = rf'{re.escape(target_key)}\s+(.+)'
                match = re.search(pattern, new_content, re.IGNORECASE)
                if match:
                    analysis["description_part"] = match.group(1).strip()
                    analysis["has_description_change"] = True
                    analysis["change_type"] = "key_and_description"
            else:
                # Если ключ не найден, возможно это только изменение описания
                analysis["description_part"] = new_content.strip()
                analysis["has_description_change"] = True
                analysis["change_type"] = "description_only"
            
            # Дополнительный анализ на основе образцов строк
            if sample_rows:
                # Проверяем, похоже ли новое содержимое на существующие строки
                for row in sample_rows:
                    if len(row["columns"]) >= 2:
                        existing_key = row["columns"][0].strip()
                        existing_desc = row["columns"][1].strip()
                        
                        # Если новое содержимое содержит существующий ключ
                        if existing_key.upper() in new_content.upper():
                            analysis["key_part"] = existing_key
                            analysis["has_key_change"] = True
                            break
            
            logger.info(f"📝 АНАЛИЗ СОДЕРЖИМОГО: {analysis}")
            
        except Exception as e:
            logger.error(f"Ошибка анализа содержимого инструкции: {e}")
        
        return analysis
    
    def _determine_affected_columns(self, content_analysis: Dict[str, Any], column_types: List[str], target_key: str) -> List[int]:
        """
        Определяет какие столбцы должны быть изменены на основе анализа содержимого.
        
        Args:
            content_analysis: Анализ содержимого инструкции
            column_types: Типы столбцов
            target_key: Целевой ключ
            
        Returns:
            Список индексов затронутых столбцов
        """
        affected_columns = []
        
        try:
            change_type = content_analysis.get("change_type", "unknown")
            has_key_change = content_analysis.get("has_key_change", False)
            has_description_change = content_analysis.get("has_description_change", False)
            
            # Если есть изменение ключа, затрагиваем первый столбец (обычно ключ/аббревиатура)
            if has_key_change and len(column_types) > 0:
                if column_types[0] in ["abbreviation", "key", "number"]:
                    affected_columns.append(0)
            
            # Если есть изменение описания, затрагиваем столбец описания
            if has_description_change:
                # Ищем столбец с типом "description"
                desc_col_idx = -1
                for i, col_type in enumerate(column_types):
                    if col_type == "description":
                        desc_col_idx = i
                        break
                
                # Если не найден столбец описания, используем последний столбец
                if desc_col_idx == -1 and len(column_types) > 1:
                    desc_col_idx = len(column_types) - 1
                
                if desc_col_idx >= 0:
                    affected_columns.append(desc_col_idx)
            
            # Если ничего не определено, по умолчанию затрагиваем все столбцы
            if not affected_columns:
                affected_columns = list(range(len(column_types)))
            
            # Убираем дубликаты и сортируем
            affected_columns = sorted(list(set(affected_columns)))
            
            logger.info(f"🎯 ОПРЕДЕЛЕНЫ ЗАТРОНУТЫЕ СТОЛБЦЫ: {affected_columns} для типа изменения '{change_type}'")
            
        except Exception as e:
            logger.error(f"Ошибка определения затронутых столбцов: {e}")
            # По умолчанию затрагиваем все столбцы
            affected_columns = list(range(len(column_types))) if column_types else [0, 1]
        
        return affected_columns
    
    def _analyze_new_content(self, new_content: str, target_key: str) -> Dict[str, Any]:
        """
        Анализирует новое содержимое для понимания структуры.
        
        Args:
            new_content: Новое содержимое из инструкции
            target_key: Ключ для поиска
            
        Returns:
            Анализ содержимого
        """
        
        # Разбиваем содержимое на части
        parts = new_content.split()
        
        analysis = {
            "has_key": False,
            "has_description": False,
            "key_part": "",
            "description_part": "",
            "is_key_change": False,
            "is_description_change": False
        }
        
        # Проверяем, начинается ли с ключа
        if parts and parts[0].strip() == target_key.strip():
            analysis["has_key"] = True
            analysis["key_part"] = parts[0]
            if len(parts) > 1:
                analysis["has_description"] = True
                analysis["description_part"] = " ".join(parts[1:])
                analysis["is_description_change"] = True
        elif parts and len(parts[0]) <= 5 and parts[0].isupper():
            # Новый ключ (короткий и заглавными буквами)
            analysis["has_key"] = True
            analysis["key_part"] = parts[0]
            analysis["is_key_change"] = True
            if len(parts) > 1:
                analysis["has_description"] = True
                analysis["description_part"] = " ".join(parts[1:])
                analysis["is_description_change"] = True
        else:
            # Только описание
            analysis["has_description"] = True
            analysis["description_part"] = new_content
            analysis["is_description_change"] = True
        
        logger.info(f"📝 АНАЛИЗ СОДЕРЖИМОГО: {analysis}")
        return analysis
    
    def _map_abbreviation_table_columns(self, content_analysis: Dict[str, Any], target_key: str) -> List[int]:
        """
        Сопоставление для таблицы сокращений.
        
        Args:
            content_analysis: Анализ содержимого
            target_key: Целевой ключ
            
        Returns:
            Список столбцов для изменения
        """
        affected_columns = []
        
        # Если есть новый ключ - изменяем первый столбец
        if content_analysis["is_key_change"] or content_analysis["has_key"]:
            affected_columns.append(0)
        
        # Если есть описание - изменяем второй столбец
        if content_analysis["is_description_change"] or content_analysis["has_description"]:
            affected_columns.append(1)
        
        # Если ничего не определено, изменяем оба столбца (безопасный вариант)
        if not affected_columns:
            affected_columns = [0, 1]
        
        return affected_columns
    
    def _map_general_table_columns(self, content_analysis: Dict[str, Any], row_structure: Dict[str, Any]) -> List[int]:
        """
        Сопоставление для общих таблиц.
        
        Args:
            content_analysis: Анализ содержимого
            row_structure: Структура строки
            
        Returns:
            Список столбцов для изменения
        """
        # Для общих таблиц изменяем все столбцы
        return list(range(row_structure.get("columns_count", 2)))

    def _should_update_key_column(self, new_value: str, current_key: str) -> bool:
        """
        Универсально определяет, нужно ли обновлять ключевой столбец.
        """
        # Если в новом значении есть новый ключ (первое слово отличается)
        words = new_value.split()
        if words and words[0] != current_key:
            return True
        return False

    def _create_adaptive_operations(self, row_structure: Dict[str, Any], affected_columns: List[int], new_values: List[str]) -> List[Dict[str, Any]]:
        """
        Создает интеллектуальные операции для изменения таблицы с правильным распределением по столбцам.
        """
        operations = []
        
        if not new_values or not affected_columns:
            return operations
        
        # Анализируем новое содержимое
        new_content = new_values[0] if new_values else ""
        content_parts = self._split_content_for_columns(new_content, len(affected_columns))
        
        logger.info(f"🔧 СОЗДАНИЕ ОПЕРАЦИЙ:")
        logger.info(f"   Затронутые столбцы: {affected_columns}")
        logger.info(f"   Исходное содержимое: '{new_content}'")
        logger.info(f"   Разделенное содержимое: {content_parts}")
        
        for i, col_idx in enumerate(affected_columns):
            # Определяем значение для этого столбца
            if i < len(content_parts):
                column_value = content_parts[i]
            elif i == 0 and content_parts:
                # Для первого столбца используем первую часть
                column_value = content_parts[0].split()[0] if content_parts[0].split() else content_parts[0]
            else:
                # Для остальных столбцов используем оставшуюся часть
                column_value = " ".join(content_parts[0].split()[1:]) if content_parts and content_parts[0].split() else ""
            
            if column_value.strip():  # Только если есть значение
                operation = {
                    "column_index": col_idx,
                    "action": "replace",
                    "new_value": column_value.strip(),
                    "column_type": row_structure.get("column_types", [])[col_idx] if col_idx < len(row_structure.get("column_types", [])) else f"column_{col_idx}"
                }
                operations.append(operation)
                logger.info(f"   ✅ Операция для столбца {col_idx}: '{column_value.strip()}'")
        
        return operations
    
    def _split_content_for_columns(self, content: str, num_columns: int) -> List[str]:
        """
        Интеллектуально разделяет содержимое для распределения по столбцам.
        
        Args:
            content: Содержимое для разделения
            num_columns: Количество столбцов
            
        Returns:
            Список значений для каждого столбца
        """
        if not content.strip():
            return []
        
        parts = content.strip().split()
        
        if num_columns == 1:
            return [content.strip()]
        elif num_columns == 2:
            if len(parts) == 1:
                # Только одно слово - вероятно ключ
                return [parts[0], ""]
            elif len(parts) >= 2:
                # Первое слово - ключ, остальное - описание
                return [parts[0], " ".join(parts[1:])]
        else:
            # Для большего количества столбцов равномерно распределяем
            result = []
            words_per_column = max(1, len(parts) // num_columns)
            
            for i in range(num_columns):
                start_idx = i * words_per_column
                if i == num_columns - 1:
                    # Последний столбец получает все оставшиеся слова
                    column_words = parts[start_idx:]
                else:
                    column_words = parts[start_idx:start_idx + words_per_column]
                
                result.append(" ".join(column_words))
            
            return result
        
        return [content.strip()]

    async def _intelligent_text_search(self, source_file: str, instruction_text: str) -> Dict[str, Any]:
        """
        Интеллектуальный поиск текста с учетом контекста и вариаций.
        Улучшенное распознавание для пунктов документа.
        """
        logger.info(f"🔍 ИНТЕЛЛЕКТУАЛЬНЫЙ ПОИСК для: {instruction_text[:50]}...")
        
        search_result = {
            "found": False,
            "target_text": "",
            "context": "",
            "location": "",
            "search_variants": []
        }
        
        try:
            # Правильно извлекаем целевую фразу из инструкции
            instruction_data = self._extract_target_and_new_text(instruction_text)
            
            if instruction_data["instruction_type"] == "paragraph_phrase_replacement":
                target_phrase = instruction_data["target_text"]
                logger.info(f"🎯 Целевая фраза для поиска: '{target_phrase}'")
                
                # Создаем универсальные варианты поиска
                search_variants = [
                    target_phrase,  # Точная фраза
                    " ".join(target_phrase.split()),  # Нормализованная (убираем лишние пробелы)
                ]
                
                # Добавляем варианты с разными пробелами (универсально для любых фраз)
                # Нормализуем пробелы вокруг союзов и предлогов
                normalized = re.sub(r'\s+и\s+', ' и ', target_phrase)
                if normalized != target_phrase:
                    search_variants.append(normalized)
                
                search_result["search_variants"] = search_variants
                
                # Ищем каждый вариант
                for variant in search_variants:
                    if len(variant.strip()) < 3:  # Пропускаем слишком короткие
                        continue
                        
                    logger.info(f"🔎 Поиск варианта: '{variant}'")
                    
                    try:
                        matches = await mcp_client.find_text_in_document(source_file, variant)
                        if matches:
                            match = matches[0]  # Берем первое совпадение
                            # Правильный доступ к атрибутам MCPTextMatch
                            if hasattr(match, 'location'):
                                context = match.context if hasattr(match, 'context') else ""
                                location = match.location
                            else:
                                # Для словарей используем .get()
                                context = match.get("context", "") if isinstance(match, dict) else ""
                                location = match.get("location", "") if isinstance(match, dict) else ""
                            
                            search_result.update({
                                "found": True,
                                "target_text": variant,
                                "context": context,
                                "location": location
                            })
                            logger.info(f"✅ НАЙДЕНО: '{variant}' в {location}")
                            break
                    except Exception as e:
                        logger.debug(f"Ошибка поиска варианта '{variant}': {e}")
                        continue
                
                if not search_result["found"]:
                    logger.warning(f"⚠️ Ни один вариант не найден для фразы: '{target_phrase}'")
            
        except Exception as e:
            logger.error(f"Ошибка интеллектуального поиска: {e}")
        
        return search_result

    async def _add_change_annotations(self, source_file: str, results: List[Dict], session_id: str) -> Dict[str, Any]:
        """
        Добавляет автоматические аннотации в места изменений для отслеживания.
        """
        logger.info("📝 СОЗДАНИЕ АВТОМАТИЧЕСКИХ АННОТАЦИЙ")
        
        annotation_results = {
            "annotations_added": 0,
            "annotations_failed": 0,
            "details": []
        }
        
        try:
            from datetime import datetime
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
            
            for result in results:
                if result.get("status") == "SUCCESS":
                    change_id = result.get("change_id", "N/A")
                    operation = result.get("operation", "")
                    description = result.get("description", "")
                    
                    # Создаем текст аннотации
                    annotation_text = f"[ИЗМЕНЕНИЕ {change_id}] {description} ({timestamp})"
                    
                    # Определяем, где добавить аннотацию
                    target_text = ""
                    
                    # Пробуем разные способы извлечения target_text
                    if "target_text" in result and result["target_text"]:
                        target_text = result["target_text"]
                    elif "target" in result and isinstance(result["target"], dict):
                        target_text = result["target"].get("text", "")
                    elif "details" in result and isinstance(result["details"], dict):
                        # Пробуем извлечь из details
                        details = result["details"]
                        if "target_text" in details:
                            target_text = details["target_text"]
                        elif "target" in details and isinstance(details["target"], dict):
                            target_text = details["target"].get("text", "")
                    
                    # Если все еще не нашли, пробуем извлечь из description
                    if not target_text and description:
                        # Ищем текст в кавычках в description
                        quote_match = re.search(r'[«"]([^»"]+)[»"]', description)
                        if quote_match:
                            target_text = quote_match.group(1).strip()
                            logger.info(f"   📍 Извлечен target_text из description: '{target_text[:30]}...'")
                    
                    if target_text:
                        logger.info(f"📌 Добавление аннотации для {change_id}: '{target_text[:30]}...'")
                        
                        # Проверяем, было ли изменение в таблице и есть ли информация о местоположении
                        is_table_change = False
                        table_paragraph_index = None
                        
                        # Проверяем в details (куда попадает результат от _handle_replace_text)
                        details = result.get("details", {})
                        if isinstance(details, dict):
                            if details.get("is_table_change", False):
                                is_table_change = True
                                # Если есть информация о местоположении таблицы, используем её
                                if "table_location" in details and details["table_location"]:
                                    table_paragraph_index = details["table_location"].get("paragraph_index")
                                    logger.info(f"   📍 Используем paragraph_index из details.table_location: {table_paragraph_index}")
                                elif "paragraph_index" in details and details["paragraph_index"] >= 0:
                                    table_paragraph_index = details["paragraph_index"]
                                    logger.info(f"   📍 Используем paragraph_index из details: {table_paragraph_index}")
                            
                            # Также проверяем на верхнем уровне (на случай прямого возврата)
                            if not is_table_change and result.get("is_table_change", False):
                                is_table_change = True
                                if "table_location" in result and result["table_location"]:
                                    table_paragraph_index = result["table_location"].get("paragraph_index")
                                    logger.info(f"   📍 Используем paragraph_index из result.table_location: {table_paragraph_index}")
                                elif "paragraph_index" in result and result["paragraph_index"] >= 0:
                                    table_paragraph_index = result["paragraph_index"]
                                    logger.info(f"   📍 Используем paragraph_index из result: {table_paragraph_index}")
                        
                        # Создаем операцию ADD_COMMENT
                        comment_change = {
                            "change_id": f"ANN-{change_id}",
                            "operation": "ADD_COMMENT",
                            "target": {
                                "text": target_text
                            },
                            "payload": {
                                "comment_text": annotation_text,
                                "paragraph_hint": target_text[:50],  # Первые 50 символов как подсказка
                                "is_table_change": is_table_change,  # Флаг, что изменение было в таблице
                            },
                            "description": f"Аннотация для изменения {change_id}"
                        }
                        
                        # Если есть точный paragraph_index для таблицы, передаем его
                        if table_paragraph_index is not None and table_paragraph_index >= 0:
                            comment_change["payload"]["paragraph_index"] = table_paragraph_index
                        
                        # Выполняем добавление комментария
                        try:
                            logger.info(f"🔍 Попытка добавить аннотацию для {change_id}: target_text='{target_text[:50]}...'")
                            comment_result = await self._handle_add_comment(source_file, comment_change)
                            
                            if comment_result.get("success"):
                                annotation_results["annotations_added"] += 1
                                annotation_results["details"].append({
                                    "change_id": change_id,
                                    "annotation_id": f"ANN-{change_id}",
                                    "status": "SUCCESS",
                                    "text": annotation_text
                                })
                                logger.info(f"✅ Аннотация {change_id} добавлена успешно")
                            else:
                                error_msg = comment_result.get("message", comment_result.get("error", "Неизвестная ошибка"))
                                
                                # Если target_text не найден, пробуем использовать new_text
                                if "не найден" in error_msg.lower() or "anchor_not_found" in error_msg.lower():
                                    logger.info(f"   🔄 Пробуем использовать new_text для аннотации {change_id}")
                                    new_text = ""
                                    if "payload" in result and isinstance(result["payload"], dict):
                                        new_text = result["payload"].get("new_text", "")
                                    
                                    if new_text:
                                        # Используем new_text для поиска места добавления аннотации
                                        logger.info(f"   📍 Используем new_text для аннотации {change_id}: '{new_text[:50]}...'")
                                        comment_change_new = {
                                            "change_id": f"ANN-{change_id}",
                                            "operation": "ADD_COMMENT",
                                            "target": {
                                                "text": new_text
                                            },
                                            "payload": {
                                                "comment_text": annotation_text,
                                                "paragraph_hint": new_text[:50]
                                            },
                                            "description": f"Аннотация для изменения {change_id}"
                                        }
                                        
                                        try:
                                            comment_result_new = await self._handle_add_comment(source_file, comment_change_new)
                                            if comment_result_new.get("success"):
                                                annotation_results["annotations_added"] += 1
                                                annotation_results["details"].append({
                                                    "change_id": change_id,
                                                    "annotation_id": f"ANN-{change_id}",
                                                    "status": "SUCCESS",
                                                    "text": annotation_text
                                                })
                                                logger.info(f"✅ Аннотация {change_id} добавлена успешно (через new_text)")
                                            else:
                                                annotation_results["annotations_failed"] += 1
                                                annotation_results["details"].append({
                                                    "change_id": change_id,
                                                    "annotation_id": f"ANN-{change_id}",
                                                    "status": "FAILED",
                                                    "error": comment_result_new.get("message", "Неизвестная ошибка")
                                                })
                                                logger.warning(f"⚠️ Не удалось добавить аннотацию для {change_id} (через new_text)")
                                        except Exception as e:
                                            annotation_results["annotations_failed"] += 1
                                            logger.error(f"Ошибка добавления аннотации для {change_id} (через new_text): {e}")
                                    else:
                                        annotation_results["annotations_failed"] += 1
                                        annotation_results["details"].append({
                                            "change_id": change_id,
                                            "annotation_id": f"ANN-{change_id}",
                                            "status": "FAILED",
                                            "error": error_msg
                                        })
                                        logger.warning(f"⚠️ Не удалось добавить аннотацию для {change_id}: {error_msg}")
                                else:
                                    annotation_results["annotations_failed"] += 1
                                    annotation_results["details"].append({
                                        "change_id": change_id,
                                        "annotation_id": f"ANN-{change_id}",
                                        "status": "FAILED",
                                        "error": error_msg
                                    })
                                    logger.warning(f"⚠️ Не удалось добавить аннотацию для {change_id}: {error_msg}")
                                
                        except Exception as e:
                            annotation_results["annotations_failed"] += 1
                            logger.error(f"Ошибка добавления аннотации для {change_id}: {e}")
                    else:
                        # Пробуем использовать new_text из payload, если target_text не найден
                        new_text = ""
                        if "payload" in result and isinstance(result["payload"], dict):
                            new_text = result["payload"].get("new_text", "")
                        
                        if new_text:
                            # Используем new_text для поиска места добавления аннотации
                            logger.info(f"   📍 Используем new_text для аннотации {change_id}: '{new_text[:30]}...'")
                            target_text = new_text
                            
                            # Создаем операцию ADD_COMMENT с new_text
                            comment_change = {
                                "change_id": f"ANN-{change_id}",
                                "operation": "ADD_COMMENT",
                                "target": {
                                    "text": target_text
                                },
                                "payload": {
                                    "comment_text": annotation_text,
                                    "paragraph_hint": target_text[:50]
                                },
                                "description": f"Аннотация для изменения {change_id}"
                            }
                            
                            try:
                                comment_result = await self._handle_add_comment(source_file, comment_change)
                                
                                if comment_result.get("success"):
                                    annotation_results["annotations_added"] += 1
                                    annotation_results["details"].append({
                                        "change_id": change_id,
                                        "annotation_id": f"ANN-{change_id}",
                                        "status": "SUCCESS",
                                        "text": annotation_text
                                    })
                                    logger.info(f"✅ Аннотация {change_id} добавлена успешно (через new_text)")
                                else:
                                    annotation_results["annotations_failed"] += 1
                                    logger.warning(f"⚠️ Не удалось добавить аннотацию для {change_id} (через new_text)")
                            except Exception as e:
                                annotation_results["annotations_failed"] += 1
                                logger.error(f"Ошибка добавления аннотации для {change_id}: {e}")
                        else:
                            logger.warning(f"⚠️ Не удалось определить target_text для аннотации {change_id} (нет target_text и new_text)")
                            annotation_results["annotations_failed"] += 1
                        
        except Exception as e:
            logger.error(f"Ошибка создания аннотаций: {e}")
        
        logger.info(f"📊 ИТОГ АННОТАЦИЙ: добавлено={annotation_results['annotations_added']}, ошибок={annotation_results['annotations_failed']}")
        return annotation_results

    async def _intelligent_table_update(self, source_file: str, table_analysis: Dict[str, Any], target_text: str) -> bool:
        """
        Адаптивное обновление таблицы на основе динамического анализа структуры.
        """
        logger.info(f"🔧 АДАПТИВНОЕ ОБНОВЛЕНИЕ ТАБЛИЦЫ")
        
        if not table_analysis.get("is_table_change") or not table_analysis.get("recommended_operations"):
            logger.warning("Нет данных для интеллектуального обновления")
            return False
        
        try:
            target_key = table_analysis["instruction_mapping"]["target_key"]
            operations = table_analysis["recommended_operations"]
            
            logger.info(f"📍 Поиск строки с ключом: '{target_key}'")
            logger.info(f"🔧 Операций к выполнению: {len(operations)}")
            
            # Находим точное местоположение записи
            matches = await mcp_client.find_text_in_document(source_file, target_key)
            
            if not matches:
                logger.error(f"Строка с ключом '{target_key}' не найдена")
                return False
            
            success_count = 0
            
            for match in matches:
                # Правильный доступ к атрибутам MCPTextMatch
                if hasattr(match, 'location'):
                    location = match.location
                else:
                    # Для словарей используем .get()
                    location = match.get('location', '') if isinstance(match, dict) else ''
                    
                if 'Table' in location:
                    logger.info(f"📍 Обработка записи в: {location}")
                    
                    # Выполняем все операции для этой строки
                    for operation in operations:
                        column_idx = operation["column_index"]
                        action = operation["action"]
                        new_value = operation["new_value"]
                        column_type = operation["column_type"]
                        
                        logger.info(f"🔄 Столбец {column_idx} ({column_type}): {action} → '{new_value}'")
                        
                        try:
                            # ИНТЕЛЛЕКТУАЛЬНОЕ применение операции с учетом структуры таблицы
                            if action == "replace":
                                result = await self._smart_column_replace(
                                    source_file, match, column_idx, new_value, target_key, location
                                )
                                
                                if result:
                                    success_count += 1
                                    logger.info(f"✅ Столбец {column_idx} обновлен успешно: '{new_value}'")
                                else:
                                    logger.warning(f"⚠️ Не удалось обновить столбец {column_idx}")
                                    
                        except Exception as e:
                            logger.error(f"Ошибка обновления столбца {column_idx}: {e}")
                    
                    # Обрабатываем только первое совпадение
                    break
            
            logger.info(f"✅ РЕЗУЛЬТАТ: {success_count} из {len(operations)} операций выполнено успешно")
            return success_count > 0
                        
        except Exception as e:
            logger.error(f"Ошибка адаптивного обновления таблицы: {e}")
        
        return False
    
    async def _smart_column_replace(self, source_file: str, match: Any, column_idx: int, new_value: str, target_key: str, location: str) -> bool:
        """
        УПРОЩЕННАЯ замена содержимого таблицы.
        Заменяет всю строку таблицы новым значением.
        """
        logger.info(f"🎯 ПРОСТАЯ ЗАМЕНА строки таблицы: '{target_key}' → '{new_value}'")
        
        try:
            # Простая стратегия: заменяем ключ на полное новое значение
            # Это работает для инструкций типа: строку «ДРМ» изложить в следующей редакции: «ДКР Департамент кредитных рисков»
            result = await mcp_client.replace_text(source_file, target_key, new_value)
            if result:
                logger.info(f"✅ Успешная замена: '{target_key}' → '{new_value}'")
                return True
            else:
                logger.warning(f"❌ Не удалось заменить: '{target_key}' → '{new_value}'")
                return False
            
        except Exception as e:
            logger.error(f"Ошибка замены в таблице: {e}")
            return False
                    
        except Exception as e:
            logger.error(f"Ошибка умной замены столбца {column_idx}: {e}")
            return False
    
    async def _find_description_in_same_row(self, source_file: str, target_key: str, context: str) -> Optional[str]:
        """
        Находит описание во втором столбце той же строки таблицы.
        
        Args:
            source_file: Путь к файлу
            target_key: Ключ первого столбца
            context: Контекст найденной ячейки
            
        Returns:
            Текст описания или None
        """
        try:
            # Получаем весь текст документа
            doc_text = await mcp_client.get_document_text(source_file)
            
            # Ищем строку с ключом
            
            # Паттерн для поиска строки таблицы с ключом
            # Предполагаем, что строка содержит ключ и описание, разделенные табуляцией или пробелами
            patterns = [
                rf'{re.escape(target_key)}\s+([^\n\r\t]+)',  # Ключ + пробелы + описание
                rf'{re.escape(target_key)}\t+([^\n\r\t]+)',  # Ключ + табуляция + описание
                rf'{re.escape(target_key)}\s*\|\s*([^\n\r\|]+)',  # Ключ | описание
            ]
            
            for pattern in patterns:
                match = re.search(pattern, doc_text, re.IGNORECASE)
                if match:
                    description = match.group(1).strip()
                    if description and description != target_key:
                        logger.info(f"🔍 Найдено описание: '{description[:50]}...'")
                        return description
            
            # Если не нашли по паттернам, пробуем использовать контекст
            if context and len(context) > len(target_key) * 2:
                # Контекст длиннее ключа в 2 раза - вероятно содержит описание
                # Убираем ключ из контекста
                description = context.replace(target_key, '').strip()
                if description:
                    logger.info(f"🔍 Извлечено описание из контекста: '{description[:50]}...'")
                    return description
            
            return None
            
        except Exception as e:
            logger.error(f"Ошибка поиска описания: {e}")
            return None

    def _analyze_operation_order(self, changes: List[Dict[str, Any]], original_text: str) -> List[Dict[str, Any]]:
        """
        Анализирует порядок операций и выявляет потенциальные конфликты.
        Переставляет операции для предотвращения конфликтов.
        """
        logger.info("🔄 АНАЛИЗ ПОРЯДКА ОПЕРАЦИЙ")
        
        # Разделяем операции по типам
        mass_replacements = []
        specific_changes = []
        other_operations = []
        
        for i, change in enumerate(changes):
            description = change.get("description", "").lower()
            operation = change.get("operation", "")
            
            # Определяем тип операции
            if "по всему тексту" in description and operation == "REPLACE_TEXT":
                mass_replacements.append((i, change))
                logger.info(f"📋 МАССОВАЯ ЗАМЕНА: {change.get('change_id')} - {description[:50]}...")
            elif ("пункт" in description or "строку" in description) and "replace" in operation.lower():
                specific_changes.append((i, change))
                logger.info(f"📋 СПЕЦИФИЧЕСКОЕ ИЗМЕНЕНИЕ: {change.get('change_id')} - {description[:50]}...")
            else:
                other_operations.append((i, change))
        
        # Проверяем конфликты
        conflicts_detected = False
        if mass_replacements and specific_changes:
            for mass_idx, mass_change in mass_replacements:
                for spec_idx, spec_change in specific_changes:
                    if mass_idx < spec_idx:
                        conflicts_detected = True
                        logger.warning(f"⚠️ КОНФЛИКТ: Массовая замена {mass_change.get('change_id')} выполняется ДО специфического изменения {spec_change.get('change_id')}")
        
        # ПРИНУДИТЕЛЬНОЕ переупорядочивание если есть массовые замены
        if mass_replacements:
            logger.warning("🔄 ПРИНУДИТЕЛЬНОЕ ПЕРЕУПОРЯДОЧИВАНИЕ: Массовые замены перемещаются в конец")
            
            # Новый порядок: специфические изменения → другие операции → массовые замены
            reordered_changes = []
            
            # 1. Добавляем специфические изменения
            for _, change in specific_changes:
                reordered_changes.append(change)
                logger.info(f"✅ Перемещено в начало: {change.get('change_id')}")
            
            # 2. Добавляем другие операции
            for _, change in other_operations:
                reordered_changes.append(change)
            
            # 3. Добавляем массовые замены в конец
            for _, change in mass_replacements:
                reordered_changes.append(change)
                logger.info(f"✅ Перемещено в конец: {change.get('change_id')}")
            
            # Обновляем change_id для сохранения порядка
            for i, change in enumerate(reordered_changes, 1):
                change["change_id"] = f"CHG-{i:03d}"
                change["reordered"] = True
            
            logger.warning(f"🔄 ОПЕРАЦИИ ПЕРЕУПОРЯДОЧЕНЫ: {len(reordered_changes)} изменений")
            return reordered_changes
        else:
            logger.info("✅ ПОРЯДОК ОПЕРАЦИЙ корректен, массовых замен не обнаружено")
            return changes

    def _validate_and_fix_json(self, parsed_json: Dict[str, Any]) -> Dict[str, Any]:
        """
        Валидация и исправление JSON от LLM.
        
        Args:
            parsed_json: Распарсенный JSON от LLM
            
        Returns:
            Исправленный JSON
        """
        logger.info("🔍 ВАЛИДАЦИЯ JSON от LLM")
        
        # Исправляем операцию REPLACE_POINT_TEXT -> REPLACE_TEXT
        if parsed_json.get('operation') == 'REPLACE_POINT_TEXT':
            parsed_json['operation'] = 'REPLACE_TEXT'
            logger.info(f"🔧 Исправлена операция: REPLACE_POINT_TEXT -> REPLACE_TEXT для {parsed_json.get('change_id', 'неизвестно')}")
        
        # Проверяем основную структуру
        if not isinstance(parsed_json, dict):
            raise ValueError("JSON должен быть объектом")
        
        # Проверяем наличие массива changes
        if "changes" not in parsed_json:
            # НОВЫЙ ФУНКЦИОНАЛ: Попытка восстановления структуры
            logger.warning("⚠️ JSON не содержит массив 'changes', пытаемся восстановить структуру...")
            
            # КРИТИЧЕСКОЕ ИСПРАВЛЕНИЕ: Если сам словарь является объектом изменения
            # (имеет change_id, operation, description), оборачиваем в массив
            if "change_id" in parsed_json or ("operation" in parsed_json and "description" in parsed_json):
                logger.info("✅ Восстановлено: найден один объект изменения, оборачиваем в массив 'changes'")
                parsed_json = {"changes": [parsed_json]}
            else:
                # Ищем изменения в других ключах
                possible_keys = ["change", "modifications", "instructions", "updates", "edits", "items"]
                found = False
                for key in possible_keys:
                    if key in parsed_json and isinstance(parsed_json[key], list):
                        parsed_json["changes"] = parsed_json[key]
                        logger.info(f"✅ Восстановлено: найден массив изменений в ключе '{key}'")
                        found = True
                        break
                
                # Если не нашли в других ключах, проверяем верхний уровень
                if not found:
                    # Проверяем, может быть все ключи верхнего уровня - это изменения
                    all_items = []
                    for key, value in parsed_json.items():
                        if isinstance(value, dict) and ("operation" in value or "description" in value):
                            all_items.append(value)
                        elif isinstance(value, list):
                            # Может быть вложенный массив
                            for item in value:
                                if isinstance(item, dict) and ("operation" in item or "description" in item):
                                    all_items.append(item)
                    
                    if all_items:
                        parsed_json["changes"] = all_items
                        logger.info(f"✅ Восстановлено: найдено {len(all_items)} объектов изменений на верхнем уровне")
                    else:
                        # Последняя попытка - создаем пустой массив
                        logger.warning("⚠️ Не удалось найти изменения, создаем пустой массив")
                        parsed_json["changes"] = []
        
        # После восстановления структуры проверяем, что changes существует
        if "changes" not in parsed_json:
            logger.error("❌ КРИТИЧЕСКАЯ ОШИБКА: После восстановления структуры массив 'changes' все еще отсутствует")
            parsed_json["changes"] = []
        
        changes = parsed_json["changes"]
        if not isinstance(changes, list):
            raise ValueError("'changes' должен быть массивом")
        
        fixed_changes = []
        
        for i, change in enumerate(changes):
            if not isinstance(change, dict):
                logger.warning(f"⚠️ Изменение {i+1} не является объектом, пропускаем")
                continue
            
            # Исправляем обязательные поля
            fixed_change = self._fix_change_object(change, i+1)
            if fixed_change:
                fixed_changes.append(fixed_change)
        
        parsed_json["changes"] = fixed_changes
        logger.info(f"✅ JSON валидирован: {len(fixed_changes)} изменений")
        
        return parsed_json
    
    async def _recover_json_structure(
        self, 
        parsed: Any, 
        original_content: str, 
        changes_text: str
    ) -> Dict[str, Any]:
        """
        НОВЫЙ ФУНКЦИОНАЛ: Восстановление структуры JSON когда отсутствует массив 'changes'.
        
        Пытается восстановить правильную структуру JSON:
        1. Если parsed - список, оборачивает в {"changes": [...]}
        2. Если изменения в других ключах, переименовывает
        3. Если структура неправильная, пытается извлечь из текста
        
        Args:
            parsed: Распарсенный JSON
            original_content: Оригинальный текст ответа LLM
            changes_text: Исходный текст инструкций
            
        Returns:
            Восстановленный JSON с правильной структурой
        """
        logger.info("🔧 Восстановление структуры JSON...")
        logger.debug(f"🔍 Входной parsed (тип: {type(parsed).__name__}): {json.dumps(parsed, ensure_ascii=False, indent=2)[:500] if isinstance(parsed, (dict, list)) else str(parsed)[:500]}...")
        
        # Если parsed - список, оборачиваем в структуру
        if isinstance(parsed, list):
            logger.info(f"   ✅ Найден список из {len(parsed)} элементов, оборачиваем в структуру")
            return {"changes": parsed}
        
        # Если parsed - словарь, проверяем структуру
        if isinstance(parsed, dict):
            logger.debug(f"🔍 Проверяем словарь. Ключи: {list(parsed.keys())}")
            # Проверяем наличие ключа changes
            if "changes" in parsed:
                return parsed
            
            # КРИТИЧЕСКОЕ ИСПРАВЛЕНИЕ: Если сам словарь является объектом изменения
            # (имеет change_id, operation, description), оборачиваем в массив
            if "change_id" in parsed or ("operation" in parsed and "description" in parsed):
                logger.info("   ✅ Найден один объект изменения, оборачиваем в массив 'changes'")
                return {"changes": [parsed]}
            
            # Ищем изменения в других ключах
            possible_keys = ["change", "modifications", "instructions", "updates", "edits", "items"]
            for key in possible_keys:
                if key in parsed and isinstance(parsed[key], list):
                    logger.info(f"   ✅ Найден массив изменений в ключе '{key}', переименовываем в 'changes'")
                    parsed["changes"] = parsed[key]
                    return parsed
            
            # Проверяем, может быть все ключи верхнего уровня - это изменения
            all_items = []
            for key, value in parsed.items():
                if isinstance(value, dict) and ("operation" in value or "description" in value):
                    all_items.append(value)
            
            if all_items:
                logger.info(f"   ✅ Найдено {len(all_items)} объектов изменений на верхнем уровне")
                return {"changes": all_items}
        
        # Если ничего не помогло, возвращаем как есть
        logger.warning("   ⚠️ Не удалось автоматически восстановить структуру")
        return parsed if isinstance(parsed, dict) else {"changes": []}
    
    async def _extract_changes_from_text_directly(
        self, 
        llm_response_text: str, 
        changes_text: str
    ) -> Optional[Dict[str, Any]]:
        """
        НОВЫЙ ФУНКЦИОНАЛ: Прямое извлечение изменений из текста ответа LLM.
        
        Когда JSON не парсится или имеет неправильную структуру,
        пытается найти объекты изменений напрямую в тексте.
        
        Args:
            llm_response_text: Текст ответа от LLM
            changes_text: Исходный текст инструкций
            
        Returns:
            Словарь с массивом changes или None
        """
        logger.info("🔍 Прямое извлечение изменений из текста...")
        
        
        changes = []
        
        # Стратегия 1: Ищем JSON объекты с полями change_id, operation, description
        json_object_pattern = r'\{\s*["\']?change_id["\']?\s*:\s*["\']([^"\']+)["\']'
        matches = re.finditer(json_object_pattern, llm_response_text, re.IGNORECASE)
        
        for match in matches:
            start_pos = match.start()
            # Пытаемся найти полный JSON объект начиная с этой позиции
            brace_count = 0
            in_string = False
            escape_next = False
            obj_start = start_pos
            
            for i in range(start_pos, len(llm_response_text)):
                char = llm_response_text[i]
                
                if escape_next:
                    escape_next = False
                    continue
                
                if char == '\\':
                    escape_next = True
                    continue
                
                if char == '"' and not escape_next:
                    in_string = not in_string
                    continue
                
                if not in_string:
                    if char == '{':
                        brace_count += 1
                    elif char == '}':
                        brace_count -= 1
                        if brace_count == 0:
                            # Найден полный объект
                            obj_text = llm_response_text[obj_start:i+1]
                            try:
                                change_obj = json.loads(obj_text)
                                if isinstance(change_obj, dict) and ("operation" in change_obj or "description" in change_obj):
                                    changes.append(change_obj)
                                    logger.info(f"   ✅ Извлечено изменение: {change_obj.get('change_id', 'N/A')}")
                            except json.JSONDecodeError:
                                pass
                            break
        
        # Стратегия 2: Ищем структурированные блоки текста
        if not changes:
            # Пытаемся найти изменения по описанию
            description_pattern = r'["\']?description["\']?\s*:\s*["\']([^"\']+)["\']'
            desc_matches = re.finditer(description_pattern, llm_response_text, re.IGNORECASE)
            
            for desc_match in desc_matches:
                # Ищем объект, содержащий это описание
                start = max(0, desc_match.start() - 200)
                end = min(len(llm_response_text), desc_match.end() + 500)
                context = llm_response_text[start:end]
                
                # Пытаемся извлечь JSON объект из контекста
                json_obj_match = re.search(r'\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}', context, re.DOTALL)
                if json_obj_match:
                    try:
                        change_obj = json.loads(json_obj_match.group(0))
                        if isinstance(change_obj, dict):
                            changes.append(change_obj)
                            logger.info(f"   ✅ Извлечено изменение по описанию")
                    except json.JSONDecodeError:
                        pass
        
        if changes:
            logger.info(f"✅ Найдено {len(changes)} изменений при прямом извлечении")
            return {"changes": changes}
        
        logger.warning("   ❌ Не удалось извлечь изменения напрямую из текста")
        return None
    
    def _fix_change_object(self, change: Dict[str, Any], index: int) -> Optional[Dict[str, Any]]:
        """
        Исправление объекта изменения.
        
        Args:
            change: Объект изменения
            index: Индекс изменения (для логирования)
            
        Returns:
            Исправленный объект или None если не удалось исправить
        """
        try:
            # Обязательные поля (payload не обязателен для DELETE_PARAGRAPH)
            operation = change.get("operation", "")
            is_delete_paragraph = operation == "DELETE_PARAGRAPH"
            
            required_fields = ["change_id", "operation", "target", "description"]
            if not is_delete_paragraph:
                required_fields.append("payload")
            
            for field in required_fields:
                if field not in change:
                    logger.warning(f"⚠️ CHG-{index:03d}: отсутствует поле '{field}'")
                    if field == "change_id":
                        change[field] = f"CHG-{index:03d}"
                    elif field == "description":
                        change[field] = f"Изменение {index}"
                    elif field == "target":
                        change[field] = {}
                    elif field == "payload":
                        change[field] = {}
                    else:
                        logger.error(f"❌ CHG-{index:03d}: критическое поле '{field}' отсутствует")
                        return None
            
            # Для DELETE_PARAGRAPH создаем пустой payload, если его нет
            if is_delete_paragraph and "payload" not in change:
                change["payload"] = {}
            
            # Проверяем target.text
            target = change.get("target", {})
            if not isinstance(target, dict):
                logger.warning(f"⚠️ CHG-{index:03d}: target не является объектом")
                target = {}
                change["target"] = target
            
            # Если target.text отсутствует или пустой, пытаемся извлечь из описания
            operation = change.get("operation", "")
            if "text" not in target or not target.get("text"):
                logger.warning(f"⚠️ CHG-{index:03d}: target.text отсутствует, пытаемся извлечь из описания")
                original_description = change.get("description", "")
                
                # Для INSERT операций проверяем after_text или after_heading
                is_insert_operation = operation in ["INSERT_PARAGRAPH", "INSERT_SECTION", "INSERT_TABLE"]
                if is_insert_operation:
                    # Для INSERT_PARAGRAPH проверяем target.after_text
                    if operation == "INSERT_PARAGRAPH" and "after_text" in target and target.get("after_text"):
                        target["text"] = target["after_text"]
                        logger.info(f"🔧 CHG-{index:03d}: для INSERT_PARAGRAPH использован target.after_text: '{target['text']}'")
                    # Для INSERT_SECTION проверяем target.after_heading
                    elif operation == "INSERT_SECTION" and "after_heading" in target and target.get("after_heading"):
                        target["text"] = target["after_heading"]
                        logger.info(f"🔧 CHG-{index:03d}: для INSERT_SECTION использован target.after_heading: '{target['text']}'")
                    else:
                        # Пытаемся извлечь из описания для INSERT операций
                        extracted_text = self._extract_target_for_insert(original_description)
                        if extracted_text:
                            target["text"] = extracted_text
                            logger.info(f"🔧 CHG-{index:03d}: извлечен target.text для {operation} из описания: '{extracted_text}'")
                        else:
                            # Если не удалось, создаем пустой target.text (для INSERT это допустимо)
                            target["text"] = ""
                            logger.warning(f"⚠️ CHG-{index:03d}: для {operation} target.text будет пустым, используется позиция из payload/target")
                else:
                    # Для других операций пытаемся извлечь target.text из описания
                    extracted_text = self._extract_target_from_description(original_description)
                    if extracted_text:
                        target["text"] = extracted_text
                        logger.info(f"🔧 CHG-{index:03d}: извлечен target.text из описания: '{extracted_text}'")
                    else:
                        # Альтернативные методы
                        alternative_text = self._extract_alternative_target(original_description, "")
                        if alternative_text:
                            target["text"] = alternative_text
                            logger.info(f"🔧 CHG-{index:03d}: найден альтернативный target.text: '{alternative_text}'")
                        else:
                            # Для DELETE_PARAGRAPH и "Изложить пункт" разрешаем номер пункта как target.text
                            description_lower = original_description.lower()
                            is_delete_paragraph = operation == "DELETE_PARAGRAPH"
                            is_full_paragraph_replacement = (
                                operation == "REPLACE_TEXT" and
                                "изложить" in description_lower and 
                                "пункт" in description_lower and 
                                ("редакции" in description_lower or "редакция" in description_lower)
                            )
                            
                            # Извлекаем номер пункта для DELETE_PARAGRAPH или "Изложить пункт X"
                            if is_delete_paragraph or is_full_paragraph_replacement:
                                paragraph_num_match = re.search(r'пункт[е]?\s+(\d+)', description_lower)
                                if paragraph_num_match:
                                    paragraph_num = paragraph_num_match.group(1)
                                    # Используем формат с точкой для совместимости
                                    target["text"] = f"{paragraph_num}."
                                    logger.info(f"🔧 CHG-{index:03d}: извлечен номер пункта для {operation}: '{target['text']}'")
                                else:
                                    logger.error(f"❌ CHG-{index:03d}: target.text отсутствует и не удалось извлечь из описания")
                                    return None
                            else:
                                logger.error(f"❌ CHG-{index:03d}: target.text отсутствует или пустой")
                                return None
            
            # СТРОГАЯ ВАЛИДАЦИЯ target.text
            target_text = target["text"]
            description = change.get("description", "").lower()
            original_description = change.get("description", "")
            operation = change.get("operation", "")
            
            # Проверяем что target.text не является номером пункта (кроме DELETE_PARAGRAPH и "Изложить пункт")
            is_delete_paragraph = operation == "DELETE_PARAGRAPH"
            is_full_paragraph_replacement = (
                operation == "REPLACE_TEXT" and
                "изложить" in description and 
                "пункт" in description and 
                ("редакции" in description or "редакция" in description)
            )
            
            # Для DELETE_PARAGRAPH и "Изложить пункт X в новой редакции" номер пункта допустим
            if self._is_paragraph_number(target_text) and not (is_delete_paragraph or is_full_paragraph_replacement):
                logger.warning(f"⚠️ CHG-{index:03d}: target.text '{target_text}' похож на номер пункта")
                # Пытаемся извлечь правильный target.text из description
                corrected_text = self._extract_target_from_description(original_description)
                if corrected_text:
                    target["text"] = corrected_text
                    logger.info(f"🔧 CHG-{index:03d}: исправлено target.text: '{target_text}' → '{corrected_text}'")
                else:
                    # Если не удалось извлечь, пробуем альтернативные методы
                    alternative_text = self._extract_alternative_target(original_description, target_text)
                    if alternative_text:
                        target["text"] = alternative_text
                        logger.info(f"🔧 CHG-{index:03d}: найден альтернативный target.text: '{target_text}' → '{alternative_text}'")
                    else:
                        # Последняя попытка: ищем в исходном тексте инструкций
                        if self._original_instructions_text:
                            logger.info(f"🔍 CHG-{index:03d}: последняя попытка - поиск в исходных инструкциях")
                            # Ищем паттерн для пунктов: "В пункте N слова «...»"
                            paragraph_num_match = re.search(r'\d+', target_text)
                            if paragraph_num_match:
                                paragraph_num = paragraph_num_match.group(0)
                                # Ищем в исходных инструкциях
                                patterns = [
                                    rf'пункте\s+{paragraph_num}\s+слова\s*[«"](.*?)[»"]',
                                    rf'пункте\s+{paragraph_num}\s+слова\s+([^изложить]+?)(?:\s+изложить|\s+в\s+следующей)',
                                ]
                                for pattern in patterns:
                                    match = re.search(pattern, self._original_instructions_text, re.IGNORECASE | re.DOTALL)
                                    if match:
                                        extracted = match.group(1).strip().rstrip('«»"')
                                        if extracted and not self._is_paragraph_number(extracted):
                                            target["text"] = extracted
                                            logger.info(f"🔧 CHG-{index:03d}: исправлено из исходных инструкций: '{target_text}' → '{extracted}'")
                                            break
                        
                        if target["text"] == target_text:  # Если не исправили
                            logger.error(f"❌ CHG-{index:03d}: не удалось исправить target.text")
                            # НЕ отклоняем изменение, а продолжаем с исходным target.text
                            logger.warning(f"⚠️ CHG-{index:03d}: продолжаем с исходным target.text: '{target_text}'")
            
            # Дополнительная проверка: если target.text пустой или слишком короткий
            elif len(target_text.strip()) < 2:
                logger.warning(f"⚠️ CHG-{index:03d}: target.text '{target_text}' слишком короткий")
                corrected_text = self._extract_target_from_description(original_description)
                if corrected_text:
                    target["text"] = corrected_text
                    logger.info(f"🔧 CHG-{index:03d}: исправлено короткий target.text: '{target_text}' → '{corrected_text}'")
                else:
                    logger.error(f"❌ CHG-{index:03d}: не удалось исправить короткий target.text")
                    return None
            
            # Проверка на кавычки в target.text (убираем их если есть)
            elif any(quote in target_text for quote in ['«', '»', '"', '"', "'", '„']):
                # Убираем все виды кавычек для поиска в документе
                cleaned_text = re.sub(r'[«»""\'„]', '', target_text).strip()
                if cleaned_text != target_text:
                    target["text"] = cleaned_text
                    logger.info(f"🔧 CHG-{index:03d}: убраны кавычки из target.text для поиска: '{target_text}' → '{cleaned_text}'")
            
            # Проверяем payload.new_text для REPLACE_TEXT
            payload = change.get("payload", {})
            if not isinstance(payload, dict):
                logger.warning(f"⚠️ CHG-{index:03d}: payload не является объектом")
                payload = {}
                change["payload"] = payload
            
            operation = change.get("operation", "")
            
            # Для INSERT операций payload проверяется отдельно
            is_insert_operation = operation in ["INSERT_PARAGRAPH", "INSERT_SECTION", "INSERT_TABLE"]
            
            if operation == "REPLACE_TEXT":
                # Исправляем неправильное поле "text" на "new_text"
                if "text" in payload and "new_text" not in payload:
                    payload["new_text"] = payload["text"]
                    del payload["text"]
                    logger.info(f"🔧 CHG-{index:03d}: исправлено payload.text → payload.new_text")
                
                # Для инструкций "Изложить пункт X в новой редакции" payload.new_text может быть пустым,
                # так как новое содержимое (включая таблицы) извлекается из документа инструкций
                description_lower = change.get("description", "").lower()
                is_full_paragraph_replacement = (
                    "изложить" in description_lower and 
                    "пункт" in description_lower and 
                    ("редакции" in description_lower or "редакция" in description_lower)
                )
                
                if "new_text" not in payload or not payload["new_text"]:
                    if is_full_paragraph_replacement:
                        logger.info(f"✅ CHG-{index:03d}: для 'Изложить пункт в новой редакции' новый текст будет извлечен из инструкции")
                        # Устанавливаем пустую строку, чтобы не было ошибок
                        payload["new_text"] = ""
                    else:
                        logger.error(f"❌ CHG-{index:03d}: payload.new_text отсутствует или пустой")
                        return None
            
            # Для INSERT_PARAGRAPH проверяем payload.text или payload.new_text
            elif operation == "INSERT_PARAGRAPH":
                if "text" not in payload and "new_text" not in payload:
                    logger.error(f"❌ CHG-{index:03d}: для INSERT_PARAGRAPH необходим payload.text или payload.new_text")
                    return None
                # Нормализуем: используем payload.text, если есть, иначе payload.new_text
                if "text" in payload and payload.get("text"):
                    if "new_text" not in payload:
                        payload["new_text"] = payload["text"]
                elif "new_text" not in payload or not payload.get("new_text"):
                    logger.error(f"❌ CHG-{index:03d}: для INSERT_PARAGRAPH payload.text или payload.new_text должны быть заполнены")
                    return None
            
            # Проверяем валидность операции
            valid_operations = ["REPLACE_TEXT", "DELETE_PARAGRAPH", "INSERT_PARAGRAPH", "INSERT_SECTION", "INSERT_TABLE", "ADD_COMMENT"]
            if operation not in valid_operations:
                logger.warning(f"⚠️ CHG-{index:03d}: неизвестная операция '{operation}', заменяем на REPLACE_TEXT")
                change["operation"] = "REPLACE_TEXT"
            
            # Для INSERT_SECTION проверяем наличие необходимых полей
            if operation == "INSERT_SECTION":
                if "payload" not in change or not isinstance(change["payload"], dict):
                    change["payload"] = {}
                payload = change["payload"]
                # Если нет payload.new_text, но есть payload.heading_text и payload.paragraphs
                if "new_text" not in payload or not payload.get("new_text"):
                    if "heading_text" in payload and payload.get("heading_text"):
                        # Создаем new_text из heading_text и paragraphs
                        heading = payload["heading_text"]
                        paragraphs = payload.get("paragraphs", [])
                        if isinstance(paragraphs, list):
                            new_text = heading + "\n" + "\n".join(paragraphs)
                        else:
                            new_text = heading + "\n" + str(paragraphs)
                        payload["new_text"] = new_text
                        logger.info(f"🔧 CHG-{index:03d}: создан payload.new_text для INSERT_SECTION из heading_text и paragraphs")
                    else:
                        logger.warning(f"⚠️ CHG-{index:03d}: для INSERT_SECTION отсутствуют payload.heading_text или payload.paragraphs")
                        payload["new_text"] = ""
            
            logger.info(f"✅ CHG-{index:03d}: валидирован ({operation})")
            return change
            
        except Exception as e:
            logger.error(f"❌ CHG-{index:03d}: ошибка валидации - {e}")
            return None

    def _is_paragraph_number(self, text: str) -> bool:
        """
        Проверяет, является ли текст номером пункта.
        
        Args:
            text: Текст для проверки
            
        Returns:
            True если текст похож на номер пункта
        """
        
        # Паттерны для номеров пунктов
        patterns = [
            r'^\d+\.$',  # 32.
            r'^\d+\)$',  # 32)
            r'^п\.\s*\d+$',  # п.32, п. 32
            r'^\d+\.\d+\.$',  # 32.1.
            r'^\d+$'  # просто число
        ]
        
        text_clean = text.strip()
        for pattern in patterns:
            if re.match(pattern, text_clean, re.IGNORECASE):
                return True
        
        return False
    
    def _extract_target_from_description(self, description: str) -> Optional[str]:
        """
        Извлекает правильный target.text из описания инструкции.
        Улучшенная версия с поддержкой различных паттернов.
        
        Args:
            description: Описание инструкции
            
        Returns:
            Извлеченный target.text или None
        """
        
        logger.info(f"🔍 ИЗВЛЕЧЕНИЕ TARGET из описания: '{description}'")
        
        # Сначала проверяем специальные случаи: "Удалить пункт X", "Изложить пункт X"
        description_lower = description.lower()
        
        # ВАЖНО: Сначала проверяем наличие конкретного текста в кавычках (слова, фразы)
        # Это нужно, чтобы отличить "Изложить слова «...» в пункте X" от "Изложить пункт X в новой редакции"
        has_specific_text = any(keyword in description_lower for keyword in ["слова", "фразу", "строку", "текст"]) and \
                           any(quote in description for quote in ['«', '"', "'", '"'])
        
        # Для "Удалить пункт X" - возвращаем номер пункта с точкой
        if "удалить" in description_lower and "пункт" in description_lower:
            paragraph_num_match = re.search(r'пункт[е]?\s+(\d+)', description_lower)
            if paragraph_num_match:
                paragraph_num = paragraph_num_match.group(1)
                result = f"{paragraph_num}."
                logger.info(f"🎯 Извлечен номер пункта для удаления: '{result}'")
                return result
        
        # Для "Изложить пункт X в новой редакции" - возвращаем номер пункта с точкой
        # НО: ТОЛЬКО если НЕТ конкретного текста в кавычках (слова, фразы)
        # Если есть "слова «...»" или "фразу «...»", это замена конкретного текста, а не полная замена пункта
        if not has_specific_text and "изложить" in description_lower and "пункт" in description_lower and "редакции" in description_lower:
            paragraph_num_match = re.search(r'пункт[е]?\s+(\d+)', description_lower)
            if paragraph_num_match:
                paragraph_num = paragraph_num_match.group(1)
                result = f"{paragraph_num}."
                logger.info(f"🎯 Извлечен номер пункта для полной замены пункта: '{result}' (без конкретного текста)")
                return result
        
        # Расширенные паттерны для извлечения текста
        patterns = [
            # Универсальные паттерны для пунктов - разные порядки слов
            r'пункте\s+\d+\s+слова\s*[«"](.*?)[»"]',  # В пункте N слова «текст»
            r'слова\s*[«"](.*?)[»"]\s+в\s+пункте\s+\d+',  # слова «текст» в пункте N
            r'слова\s*[«"](.*?)[»"]\s+пункте\s+\d+',  # слова «текст» пункте N
            r'изложить\s+слова\s*[«"](.*?)[»"]',  # изложить слова «текст»
            
            # Основные паттерны с контекстом
            r'строку\s*[«"](.*?)[»"]',  # строку «текст»
            r'слова\s*[«"](.*?)[»"]',   # слова «текст»
            r'фразу\s*[«"](.*?)[»"]',   # фразу «текст»
            r'текст\s*[«"](.*?)[»"]',   # текст «текст»
            r'аббревиатуру\s*[«"](.*?)[»"]',  # аббревиатуру «текст»
            
            # Паттерны для разных типов кавычек
            r'[«"](.*?)[»"]',  # основные кавычки
            r'"(.*?)"',  # обычные двойные кавычки
            r"'(.*?)'",  # одинарные кавычки
            
            # Паттерны без кавычек (как резерв)
            r'строку\s+([А-ЯЁа-яё\s]+?)(?:\s+изложить|\s+заменить|$)',  # строку ТЕКСТ изложить
            r'слова\s+([А-ЯЁа-яё\s]+?)(?:\s+изложить|\s+заменить|$)',   # слова ТЕКСТ изложить
            r'аббревиатуру\s+([А-ЯЁ]+)',  # аббревиатуру СЛОВО
        ]
        
        for pattern in patterns:
            match = re.search(pattern, description, re.IGNORECASE | re.DOTALL)
            if match:
                extracted = match.group(1).strip()
                # Убираем лишние пробелы и символы
                extracted = re.sub(r'\s+', ' ', extracted)
                if extracted and not self._is_paragraph_number(extracted):
                    logger.info(f"🎯 Извлечен target.text: '{extracted}' (паттерн: {pattern[:30]}...)")
                    return extracted
        
        # Дополнительная попытка: ищем ключевые слова для пунктов
        if 'пункте' in description.lower() and 'слова' in description.lower():
            # Для любого пункта ищем фразу после "слова"
            match = re.search(r'слова\s+([^изложить]+?)(?:\s+изложить|$)', description, re.IGNORECASE)
            if match:
                extracted_text = match.group(1).strip().rstrip('«»"')
                if extracted_text and not self._is_paragraph_number(extracted_text):
                    logger.info(f"🎯 Извлечен target.text для пункта: '{extracted_text}'")
                    return extracted_text
            
            # Если не нашли в description, ищем в исходном тексте инструкций
            if self._original_instructions_text:
                logger.info("🔍 Поиск target.text в исходном тексте инструкций для пункта")
                # Универсальный паттерн для любого пункта: "В пункте N слова «...»"
                patterns = [
                    r'пункте\s+\d+\s+слова\s*[«"](.*?)[»"]',  # В пункте N слова «текст»
                    r'пункте\s+\d+\s+слова\s+([^изложить]+?)(?:\s+изложить|\s+в\s+следующей)',  # В пункте N слова текст изложить
                ]
                for pattern in patterns:
                    match = re.search(pattern, self._original_instructions_text, re.IGNORECASE | re.DOTALL)
                    if match:
                        extracted_text = match.group(1).strip().rstrip('«»"')
                        if extracted_text and not self._is_paragraph_number(extracted_text):
                            logger.info(f"🎯 Извлечен target.text из исходных инструкций: '{extracted_text}'")
                            return extracted_text
        
        logger.warning(f"⚠️ Не удалось извлечь target.text из описания: '{description}'")
        return None
    
    def _extract_target_for_insert(self, description: str) -> Optional[str]:
        """
        Извлекает target.text для INSERT операций из описания.
        
        Args:
            description: Описание инструкции
            
        Returns:
            Извлеченный target.text или None
        """
        
        logger.info(f"🔍 ИЗВЛЕЧЕНИЕ TARGET для INSERT из описания: '{description}'")
        description_lower = description.lower()
        
        # Для "Добавь новый раздел X после раздела Y" - извлекаем Y
        if "после раздела" in description_lower or "после" in description_lower:
            patterns = [
                r'после раздела\s+([^«"]+)',
                r'после\s+([^«"]+)',
                r'после\s+раздела\s*[«"]([^»"]+)[»"]',
            ]
            for pattern in patterns:
                match = re.search(pattern, description, re.IGNORECASE)
                if match:
                    extracted = match.group(1).strip()
                    if extracted:
                        logger.info(f"🎯 Извлечен target.text для INSERT: '{extracted}'")
                        return extracted
        
        # Для "Добавь новый endpoint после X" - извлекаем X
        if "после" in description_lower:
            patterns = [
                r'после\s+([^«"]+?)(?:\s+со следующим|\s+с текстом|$)',
                r'после\s*[«"]([^»"]+)[»"]',
            ]
            for pattern in patterns:
                match = re.search(pattern, description, re.IGNORECASE)
                if match:
                    extracted = match.group(1).strip()
                    if extracted:
                        logger.info(f"🎯 Извлечен target.text для INSERT: '{extracted}'")
                        return extracted
        
        # Если не нашли, возвращаем None (для INSERT это допустимо)
        return None

    def _extract_alternative_target(self, description: str, current_target: str) -> Optional[str]:
        """
        Альтернативные методы извлечения target.text для сложных случаев.
        
        Args:
            description: Описание инструкции
            current_target: Текущий неправильный target.text
            
        Returns:
            Альтернативный target.text или None
        """
        
        logger.info(f"🔍 АЛЬТЕРНАТИВНОЕ ИЗВЛЕЧЕНИЕ для: '{description}'")
        
        # Универсальные паттерны для пунктов
        if "пункте" in description.lower() and "слова" in description.lower():
            # Ищем фразы после "слова" с кавычками и без
            patterns = [
                r'слова\s*[«"\'„]([^»"\']+)[»"\'"]',  # слова «текст»
                r'слова\s+([^изложить]+?)(?:\s+изложить|$)',  # слова ТЕКСТ изложить
                r'слова\s+(.*?)(?:\s+изложить|\s+в\s+следующей|\s+заменить|$)',  # более широкий поиск
                r'\d+\s+слова\s+(.*?)(?:\s+изложить|$)',  # пункте N слова ТЕКСТ
            ]
            
            for pattern in patterns:
                match = re.search(pattern, description, re.IGNORECASE)
                if match:
                    extracted = match.group(1).strip()
                    # Убираем кавычки и лишние символы
                    extracted = re.sub(r'[«»"\'„]', '', extracted).strip()
                    if extracted and len(extracted) > 3 and not self._is_paragraph_number(extracted):
                        logger.info(f"🎯 Альтернативное извлечение: '{extracted}'")
                        return extracted
        
        # Для таблиц - ищем аббревиатуры в кавычках
        if "таблице" in description.lower():
            patterns = [
                r'строку\s*[«"\'„]([А-ЯЁ]{2,6})[»"\'"]',  # строку «ДРМ»
                r'[«"\'„]([А-ЯЁ]{2,6})[»"\'"]',  # просто «ДРМ»
            ]
            
            for pattern in patterns:
                match = re.search(pattern, description)
                if match:
                    extracted = match.group(1).strip()
                    logger.info(f"🎯 Альтернативное извлечение для таблицы: '{extracted}'")
                    return extracted
        
        logger.warning(f"⚠️ Альтернативное извлечение не дало результатов")
        return None

    def _optimize_operation_order(self, changes: List[Dict]) -> List[Dict]:
        """
        Оптимизирует порядок операций для предотвращения конфликтов.
        
        Правила оптимизации:
        1. Локальные изменения в таблицах - ПЕРВЫМИ
        2. Локальные изменения в пунктах - ВТОРЫМИ
        3. Глобальные замены - ПОСЛЕДНИМИ, НО:
           - Если локальное изменение содержит текст, который изменяется глобальной заменой,
             то локальное изменение должно быть выполнено ДО глобальной замены
        
        Args:
            changes: Список изменений
            
        Returns:
            Оптимизированный список изменений
        """
        logger.info(f"🔄 ОПТИМИЗАЦИЯ ПОРЯДКА ОПЕРАЦИЙ для {len(changes)} изменений")
        
        global_changes = []
        table_changes = []
        paragraph_changes = []
        other_changes = []
        
        for change in changes:
            description = change.get("description", "").lower()
            target_text = change.get("target", {}).get("text", "")
            
            # Определяем тип изменения
            is_global = any(keyword in description for keyword in ["по всему тексту", "по всему документу", "везде в документе"]) or change.get("target", {}).get("replace_all", False)
            
            if is_global:
                global_changes.append(change)
                logger.info(f"   🌍 Глобальное изменение: {change.get('change_id', 'N/A')} (заменяет '{target_text}')")
            elif "таблице" in description:
                table_changes.append(change)
                logger.info(f"   📊 Изменение в таблице: {change.get('change_id', 'N/A')}")
            elif any(keyword in description for keyword in ["пункте", "разделе", "параграфе"]):
                paragraph_changes.append(change)
                logger.info(f"   📄 Изменение в пункте: {change.get('change_id', 'N/A')}")
            else:
                other_changes.append(change)
                logger.info(f"   ❓ Другое изменение: {change.get('change_id', 'N/A')}")
        
        # НОВЫЙ ФУНКЦИОНАЛ: Проверяем зависимости между изменениями
        # Если локальное изменение использует текст, который изменяется глобальной заменой,
        # то локальное изменение должно быть выполнено ДО глобальной замены
        global_target_texts = {}  # Словарь: текст для замены -> глобальное изменение
        for global_change in global_changes:
            target_text = global_change.get("target", {}).get("text", "")
            if target_text:
                global_target_texts[target_text.lower()] = global_change
        
        # Разделяем локальные изменения на зависимые и независимые
        dependent_local_changes = []  # Зависят от глобальных замен (должны быть ДО них)
        independent_local_changes = []  # Не зависят (могут быть до или после)
        
        all_local_changes = table_changes + paragraph_changes + other_changes
        for local_change in all_local_changes:
            target_text = local_change.get("target", {}).get("text", "")
            description = local_change.get("description", "").lower()
            payload_new_text = local_change.get("payload", {}).get("new_text", "").lower()
            # Также проверяем исходный target_text из описания (может быть указан в кавычках)
            quoted_texts = re.findall(r'[«"](.*?)[»"]', description)
            
            # Проверяем, содержит ли локальное изменение текст, который изменяется глобальной заменой
            is_dependent = False
            if target_text:
                target_lower = target_text.lower()
                # Проверяем, используется ли этот текст в глобальной замене
                for global_target, global_change in global_target_texts.items():
                    if global_target in target_lower or target_lower in global_target:
                        is_dependent = True
                        logger.info(f"   ⚠️ Найдена зависимость: {local_change.get('change_id', 'N/A')} зависит от глобальной замены {global_change.get('change_id', 'N/A')}")
                        logger.info(f"      Локальное: '{target_text}' может быть изменено глобальной заменой '{global_target}'")
                        break
            
            # Проверяем описание на наличие текста, который может быть изменен глобальной заменой
            if not is_dependent:
                for global_target, global_change in global_target_texts.items():
                    if global_target in description:
                        is_dependent = True
                        logger.info(f"   ⚠️ Найдена зависимость (в описании): {local_change.get('change_id', 'N/A')} зависит от глобальной замены {global_change.get('change_id', 'N/A')}")
                        logger.info(f"      Описание содержит '{global_target}', который изменяется глобальной заменой")
                        break
            
            # Проверяем тексты в кавычках из описания
            if not is_dependent:
                for quoted_text in quoted_texts:
                    quoted_lower = quoted_text.lower()
                    for global_target, global_change in global_target_texts.items():
                        if global_target in quoted_lower or quoted_lower in global_target:
                            is_dependent = True
                            logger.info(f"   ⚠️ Найдена зависимость (в кавычках описания): {local_change.get('change_id', 'N/A')} зависит от глобальной замены {global_change.get('change_id', 'N/A')}")
                            logger.info(f"      Текст в кавычках '{quoted_text}' содержит '{global_target}', который изменяется глобальной заменой")
                            break
                    if is_dependent:
                        break
            
            # Проверяем payload.new_text (может содержать исходный текст для замены)
            if not is_dependent and payload_new_text:
                for global_target, global_change in global_target_texts.items():
                    if global_target in payload_new_text:
                        is_dependent = True
                        logger.info(f"   ⚠️ Найдена зависимость (в payload): {local_change.get('change_id', 'N/A')} зависит от глобальной замены {global_change.get('change_id', 'N/A')}")
                        logger.info(f"      Payload содержит '{global_target}', который изменяется глобальной заменой")
                        break
            
            if is_dependent:
                dependent_local_changes.append(local_change)
            else:
                independent_local_changes.append(local_change)
        
        # Оптимальный порядок: зависимые локальные → независимые локальные → глобальные
        optimized = dependent_local_changes + independent_local_changes + global_changes
        
        logger.info(f"📋 ОПТИМИЗИРОВАННЫЙ ПОРЯДОК:")
        logger.info(f"   Зависимые локальные изменения: {len(dependent_local_changes)}")
        logger.info(f"   Независимые локальные изменения: {len(independent_local_changes)}")
        logger.info(f"   Глобальные изменения: {len(global_changes)}")
        for i, change in enumerate(optimized, 1):
            logger.info(f"   {i}. {change.get('change_id', 'N/A')}: {change.get('description', 'N/A')[:50]}...")
        
        return optimized

    async def _simple_parse_changes_with_llm(
        self, 
        changes_text: str, 
        initial_changes: Optional[List[Dict[str, Any]]] = None
    ) -> tuple[List[Dict[str, Any]], Dict[str, int]]:
        """
        Простое преобразование текстовых инструкций в структурированный JSON через LLM.
        БЕЗ двухэтапного анализа - только базовый парсинг и валидация.
        
        Returns:
            Tuple[список изменений, словарь с информацией о токенах]
        """
        if not self.openai_client:
            raise RuntimeError("OpenAI клиент не инициализирован")

        logger.info(f"📝 Простой парсинг инструкций: {len(changes_text)} символов")
        logger.info(f"🔍 СОДЕРЖИМОЕ ИНСТРУКЦИЙ (первые 500 символов): {changes_text[:500]}...")

        # Загружаем промпты
        system_prompt = self._load_prompt("instruction_check_system.md")
        user_prompt = self._load_prompt("instruction_check_user.md")
        
        # Формируем запрос к LLM
        user_message = f"{user_prompt}\n\nТекст инструкций:\n{changes_text}"
        
        logger.info(f"Отправка запроса к LLM: модель=gpt-4o, длина промпта={len(user_message)} символов")
        
        try:
            response = await self.openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_message}
                ],
                temperature=0.1,
                max_tokens=4000
            )
            
            logger.info("Ответ от LLM получен успешно")
            
            # Извлекаем JSON из ответа
            response_text = response.choices[0].message.content.strip()
            logger.info(f"📥 Сырой ответ от LLM (первые 1000 символов): {response_text[:1000]}...")
            logger.info(f"📥 Полный сырой ответ от LLM (длина: {len(response_text)} символов): {response_text}")
            # КРИТИЧЕСКОЕ: Всегда логируем полный ответ для диагностики проблем с множественными инструкциями
            
            # КРИТИЧЕСКОЕ: Подсчитываем примерное количество инструкций для сравнения с ответом LLM
            instruction_count_indicators = [
                "CHG-", "1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.",
                "изложить", "заменить", "исключить", "добавить", "удалить", "изменить"
            ]
            estimated_instructions = sum(1 for indicator in instruction_count_indicators if indicator in changes_text.lower())
            logger.info(f"📊 ОЦЕНКА КОЛИЧЕСТВА ИНСТРУКЦИЙ В ИСХОДНОМ ТЕКСТЕ: {estimated_instructions} (по индикаторам)")
            
            # Парсим JSON
            changes_json = self._extract_json_from_response(response_text)
            if not changes_json:
                raise ValueError("Не удалось извлечь JSON из ответа LLM")
            
            logger.info(f"🔍 ИЗВЛЕЧЕННЫЙ JSON: {changes_json}")
            
            # КРИТИЧЕСКОЕ: Проверяем количество найденных инструкций
            if isinstance(changes_json, list):
                logger.info(f"📊 LLM вернул {len(changes_json)} изменений")
                if estimated_instructions > len(changes_json):
                    logger.warning(
                        f"⚠️ КРИТИЧЕСКОЕ ВНИМАНИЕ: В тексте, вероятно, содержится {estimated_instructions} инструкций "
                        f"(найдено индикаторов: {estimated_instructions}), но LLM вернул только {len(changes_json)} изменений. "
                        f"LLM мог остановиться на первой инструкции или не прочитать весь документ!"
                    )
                    logger.warning(
                        f"⚠️ ПРОВЕРКА: Первые 500 символов текста инструкций: {changes_text[:500]}"
                    )
            else:
                logger.warning(f"⚠️ LLM вернул не список: {type(changes_json)}")
            
            # Исправляем операции REPLACE_POINT_TEXT -> REPLACE_TEXT
            for change in changes_json:
                if isinstance(change, dict) and change.get('operation') == 'REPLACE_POINT_TEXT':
                    change['operation'] = 'REPLACE_TEXT'
                    logger.info(f"🔧 Исправлена операция: REPLACE_POINT_TEXT -> REPLACE_TEXT для {change.get('change_id', 'неизвестно')}")
            
            # Простая валидация JSON (список изменений)
            logger.info("🔍 ПРОСТАЯ ВАЛИДАЦИЯ JSON от LLM")
            if not isinstance(changes_json, list):
                logger.error("JSON должен быть списком изменений")
                raise ValueError("JSON должен быть списком изменений")
            
            # Валидируем и исправляем каждое изменение
            validated_changes = []
            for idx, change in enumerate(changes_json, start=1):
                if isinstance(change, dict):
                    fixed_change = self._fix_change_object(change, idx)
                    if fixed_change:
                        validated_changes.append(fixed_change)
                    else:
                        logger.warning(f"⚠️ Изменение {idx} не прошло валидацию и будет пропущено")
                else:
                    logger.warning(f"⚠️ Изменение {idx} не является объектом и будет пропущено")
            
            # Подсчитываем токены
            tokens_info = {
                "prompt_tokens": response.usage.prompt_tokens,
                "completion_tokens": response.usage.completion_tokens,
                "total_tokens": response.usage.total_tokens
            }
            
            logger.info(f"✅ Простой парсинг завершен: {len(validated_changes)} изменений")
            logger.info(f"Использовано токенов: {tokens_info['total_tokens']} (prompt: {tokens_info['prompt_tokens']}, completion: {tokens_info['completion_tokens']})")
            
            return validated_changes, tokens_info
            
        except Exception as e:
            logger.error(f"Ошибка простого парсинга с LLM: {e}")
            raise

    def _extract_json_from_response(self, response_text: str) -> Optional[List[Dict]]:
        """
        Извлекает JSON из ответа LLM.
        
        Args:
            response_text: Текст ответа от LLM
            
        Returns:
            Список изменений в формате JSON или None
        """
        import json
        
        try:
            # Ищем JSON блок в ответе
            json_match = re.search(r'```json\s*(.*?)\s*```', response_text, re.DOTALL)
            if json_match:
                json_text = json_match.group(1).strip()
            else:
                # Если нет блока ```json```, ищем массив JSON
                json_match = re.search(r'\[\s*\{.*?\}\s*\]', response_text, re.DOTALL)
                if json_match:
                    json_text = json_match.group(0)
                else:
                    # Пытаемся парсить весь ответ как JSON
                    json_text = response_text.strip()
            
            # Парсим JSON
            changes = json.loads(json_text)
            logger.info(f"🔍 ПАРСИНГ JSON: тип={type(changes)}, длина={len(changes) if isinstance(changes, (list, dict)) else 'N/A'}")
            
            # Проверяем различные форматы JSON от LLM
            if isinstance(changes, list):
                logger.info(f"🔍 JSON является списком, длина={len(changes)}")
                # Если список содержит один элемент-словарь с полем 'changes'
                if len(changes) == 1 and isinstance(changes[0], dict) and 'changes' in changes[0]:
                    logger.info("✅ JSON содержит список с полем 'changes', извлекаем массив изменений")
                    changes = changes[0]['changes']
                    logger.info(f"✅ Извлечено {len(changes)} изменений из поля 'changes'")
                # Если список уже содержит массив изменений - оставляем как есть
                else:
                    logger.info(f"✅ JSON является списком изменений, длина={len(changes)}")
            elif isinstance(changes, dict):
                if 'changes' in changes:
                    logger.info("JSON содержит поле 'changes', извлекаем массив изменений")
                    changes = changes['changes']
                else:
                    logger.warning("JSON не является списком и не содержит 'changes', пытаемся обернуть в список")
                    changes = [changes]
            else:
                logger.warning("JSON имеет неожиданный формат, пытаемся обернуть в список")
                changes = [changes] if changes else []
            
            logger.info(f"Успешно извлечен JSON с {len(changes)} изменениями")
            return changes
            
        except json.JSONDecodeError as e:
            logger.error(f"Ошибка парсинга JSON: {e}")
            logger.debug(f"Проблемный JSON: {response_text[:500]}...")
            return None
        except Exception as e:
            logger.error(f"Ошибка извлечения JSON: {e}")
            return None

    def _validate_and_correct_operations(self, changes: List[Dict[str, Any]], original_text: str) -> List[Dict[str, Any]]:
        """
        Валидация и автокоррекция неправильно выбранных операций LLM.
        
        Args:
            changes: Список изменений от LLM
            original_text: Исходный текст инструкций
            
        Returns:
            Скорректированный список изменений
        """
        logger.info(f"🔍 НАЧАЛО ВАЛИДАЦИИ: получено {len(changes)} изменений")
        corrected_changes = []
        corrections_made = 0
        
        for change in changes:
            operation = change.get("operation", "")
            description = change.get("description", "").lower()
            change_id = change.get("change_id", "")
            
            logger.info(f"🔍 ВАЛИДАЦИЯ {change_id}: operation={operation}, description='{description[:50]}...'")
            
            # Создаем копию изменения для модификации
            corrected_change = change.copy()
            original_operation = operation
            
            # ПРАВИЛА АВТОКОРРЕКЦИИ:
            
            # 1. Если в описании есть "слова", "строку", "фразу" - это REPLACE_TEXT
            if operation == "REPLACE_POINT_TEXT":
                keywords_for_replace_text = ["слова", "строку", "фразу", "текст", "аббревиатуру"]
                if any(keyword in description for keyword in keywords_for_replace_text):
                    corrected_change["operation"] = "REPLACE_TEXT"
                    logger.warning(f"🔧 АВТОКОРРЕКЦИЯ {change_id}: {operation} → REPLACE_TEXT (найдено ключевое слово)")
                    corrections_made += 1
            
            # 2. Если упоминается "в таблице" - это REPLACE_TEXT
            if operation == "REPLACE_POINT_TEXT" and "в таблице" in description:
                corrected_change["operation"] = "REPLACE_TEXT"
                logger.warning(f"🔧 АВТОКОРРЕКЦИЯ {change_id}: {operation} → REPLACE_TEXT (таблица)")
                corrections_made += 1
            
            # 3. Если упоминается "по всему тексту" - это REPLACE_TEXT с replace_all=true
            if "по всему тексту" in description:
                corrected_change["operation"] = "REPLACE_TEXT"
                if "target" in corrected_change and isinstance(corrected_change["target"], dict):
                    corrected_change["target"]["replace_all"] = True
                logger.warning(f"🔧 АВТОКОРРЕКЦИЯ {change_id}: установлен replace_all=true (массовая замена)")
                corrections_made += 1
            
            # 4. Проверяем исходный текст инструкций для дополнительной валидации
            if change_id and original_text:
                # Ищем соответствующую инструкцию в исходном тексте
                lines = original_text.split('\n')
                instruction_text = ""
                
                # Пробуем разные способы найти инструкцию
                for line in lines:
                    line_clean = line.strip().lower()
                    # Ищем строку с номером изменения (1., 2., 3., 4.)
                    if any(marker in line_clean for marker in [f"{change_id[-1]}.", f"{change_id[-1]} "]):
                        instruction_text = line_clean
                        break
                
                if instruction_text:
                    # Дополнительная проверка по исходной инструкции
                    if operation == "REPLACE_POINT_TEXT":
                        correction_triggers = [
                            "слова", "строку", "фразу", "в таблице", 
                            "аббревиатуру", "по всему тексту"
                        ]
                        if any(trigger in instruction_text for trigger in correction_triggers):
                            corrected_change["operation"] = "REPLACE_TEXT"
                            logger.warning(f"🔧 АВТОКОРРЕКЦИЯ {change_id}: {operation} → REPLACE_TEXT (анализ исходной инструкции: '{instruction_text[:50]}...')")
                            corrections_made += 1
            
            # 5. Специальные правила для конкретных паттернов
            if operation == "REPLACE_POINT_TEXT":
                # Если в описании упоминается конкретная замена текста
                replace_patterns = ["изложить в следующей редакции", "заменить на", "изменить на"]
                table_patterns = ["в таблице", "строку"]
                
                has_replace_pattern = any(pattern in description for pattern in replace_patterns)
                has_table_pattern = any(pattern in description for pattern in table_patterns)
                
                if has_replace_pattern and has_table_pattern:
                    corrected_change["operation"] = "REPLACE_TEXT"
                    logger.warning(f"🔧 АВТОКОРРЕКЦИЯ {change_id}: {operation} → REPLACE_TEXT (паттерн замены в таблице)")
                    corrections_made += 1
            
            # Логируем если операция была изменена
            if corrected_change["operation"] != original_operation:
                logger.info(f"✅ Операция скорректирована: {change_id} {original_operation} → {corrected_change['operation']}")
            
            corrected_changes.append(corrected_change)
        
        # Детальное логирование результатов валидации
        if corrections_made > 0:
            logger.warning(f"🔧 ВАЛИДАЦИЯ: Выполнено {corrections_made} автокоррекций операций")
            logger.info("📋 ИТОГОВЫЕ ОПЕРАЦИИ ПОСЛЕ ВАЛИДАЦИИ:")
            for change in corrected_changes:
                logger.info(f"  {change.get('change_id')}: {change.get('operation')} - {change.get('description', '')[:60]}...")
        else:
            logger.info("✅ ВАЛИДАЦИЯ: Все операции корректны")
        
        return corrected_changes

    async def _enhanced_parse_changes_with_llm(
        self, 
        changes_text: str, 
        source_file: str,
        initial_changes: Optional[List[Dict[str, Any]]] = None
    ) -> tuple[List[Dict[str, Any]], Dict[str, int]]:
        """
        УЛУЧШЕННЫЙ двухэтапный парсинг инструкций с анализом контекста и MCP инструментов.
        
        ЭТАП 1: Анализ контекста каждой инструкции
        ЭТАП 2: Выбор правильной MCP операции на основе контекста
        """
        logger.info("🚀 ЗАПУСК УЛУЧШЕННОГО ДВУХЭТАПНОГО АНАЛИЗА")
        
        # Сначала выполняем стандартный парсинг
        changes, tokens_info = await self._parse_changes_with_llm(changes_text, initial_changes)
        
        # КРИТИЧЕСКИ ВАЖНО: Сохраняем исходный порядок инструкций из файла
        # Изменения должны выполняться строго в том порядке, в котором они указаны в файле изменений
        logger.info(f"📋 ПОРЯДОК ОПЕРАЦИЙ: Сохранен исходный порядок из файла ({len(changes)} изменений)")
        for i, change in enumerate(changes, 1):
            change_id = change.get("change_id", f"CHG-{i:03d}")
            description = change.get("description", "N/A")
            logger.info(f"   {i}. {change_id}: {description[:60]}...")
        
        # ОТКЛЮЧЕНО: Оптимизация порядка операций - изменения выполняются в исходном порядке
        # Если в будущем понадобится умная оптимизация, можно включить с учетом зависимостей:
        # optimized_changes = self._optimize_operation_order(changes)
        # Но только если пользователь явно запросит оптимизацию
        
        # Затем анализируем контекст каждой инструкции и корректируем операции
        enhanced_changes = []
        corrections_made = 0
        
        logger.info(f"🔍 АНАЛИЗ КОНТЕКСТА ДЛЯ {len(changes)} ИНСТРУКЦИЙ")
        
        for change in changes:
            change_id = change.get("change_id", "")
            description = change.get("description", "")
            operation = change.get("operation", "")
            
            # Создаем копию изменения для обработки
            enhanced_change = change.copy()
            
            logger.info(f"📋 АНАЛИЗ {change_id}: {description[:60]}...")
            
            # Анализируем контекст инструкции
            try:
                # Ищем исходную инструкцию в тексте
                instruction_lines = changes_text.split('\n')
                instruction_text = ""
                
                for line in instruction_lines:
                    if any(marker in line for marker in [f"{change_id[-1]}.", f"{change_id[-1]} "]):
                        instruction_text = line.strip()
                        break
                
                if not instruction_text:
                    instruction_text = description
                
                # Выполняем анализ контекста
                context = await self._analyze_instruction_context(instruction_text, source_file)
                
                # Дополнительный анализ для таблиц
                table_analysis = None
                intelligent_table_analysis = None
                
                if context["element_type"] == "table_cell" and change.get("target", {}).get("text"):
                    target_text = change["target"]["text"]
                    table_analysis = await self._analyze_table_structure(source_file, target_text)
                    
                    # НОВЫЙ: Интеллектуальный анализ структуры таблицы
                    intelligent_table_analysis = await self._intelligent_table_analysis(source_file, instruction_text)
                    
                    # Если найдено более полное содержимое ячейки, обновляем target.text
                    if table_analysis["found"] and table_analysis["recommended_target_text"] != target_text:
                        enhanced_change["target"]["text"] = table_analysis["recommended_target_text"]
                        logger.warning(f"🔧 КОРРЕКЦИЯ TARGET.TEXT {change_id}: '{target_text}' → '{table_analysis['recommended_target_text']}'")
                        corrections_made += 1
                    
                    # Добавляем информацию об интеллектуальном анализе
                    if intelligent_table_analysis["is_table_change"]:
                        enhanced_change["intelligent_table_analysis"] = intelligent_table_analysis
                        logger.info(f"🧠 ИНТЕЛЛЕКТУАЛЬНЫЙ АНАЛИЗ {change_id}: тип таблицы = {intelligent_table_analysis['table_type']}")
                        corrections_made += 1
                
                # НОВЫЙ: Интеллектуальный поиск для улучшения распознавания текста
                elif context["element_type"] == "paragraph" and "слова" in description:
                    intelligent_search = await self._intelligent_text_search(source_file, instruction_text)
                    
                    if intelligent_search["found"] and intelligent_search["target_text"] != enhanced_change.get("target", {}).get("text", ""):
                        old_target = enhanced_change.get("target", {}).get("text", "")
                        enhanced_change.setdefault("target", {})["text"] = intelligent_search["target_text"]
                        enhanced_change["intelligent_search"] = intelligent_search
                        
                        logger.warning(f"🔍 ИНТЕЛЛЕКТУАЛЬНЫЙ ПОИСК {change_id}: '{old_target}' → '{intelligent_search['target_text']}'")
                        logger.info(f"   Найдено в: {intelligent_search['location']}")
                        corrections_made += 1
                
                # Используем уже созданную копию изменения
                original_operation = operation
                
                # Корректируем операцию на основе анализа контекста
                if context["recommended_tool"] != operation.lower().replace("_", ""):
                    # Маппинг MCP инструментов к операциям
                    tool_to_operation = {
                        "replace_text": "REPLACE_TEXT",
                        "delete_paragraph": "DELETE_PARAGRAPH", 
                        "add_paragraph": "ADD_PARAGRAPH",
                        "add_heading": "ADD_HEADING",
                        "add_table": "ADD_TABLE",
                        "add_comment": "ADD_COMMENT"
                    }
                    
                    recommended_operation = tool_to_operation.get(context["recommended_tool"], operation)
                    
                    if recommended_operation != operation:
                        enhanced_change["operation"] = recommended_operation
                        corrections_made += 1
                        
                        logger.warning(f"🔧 КОРРЕКЦИЯ {change_id}: {operation} → {recommended_operation}")
                        logger.info(f"   Причина: {context['reasoning']}")
                        logger.info(f"   Тип элемента: {context['element_type']}")
                
                # ДОПОЛНИТЕЛЬНАЯ КОРРЕКЦИЯ для частых ошибок
                if operation == "REPLACE_POINT_TEXT":
                    # Проверяем ключевые слова в описании
                    correction_keywords = ["слова", "строку", "фразу", "в таблице", "аббревиатуру"]
                    if any(keyword in description for keyword in correction_keywords):
                        enhanced_change["operation"] = "REPLACE_TEXT"
                        corrections_made += 1
                        logger.warning(f"🔧 ДОПОЛНИТЕЛЬНАЯ КОРРЕКЦИЯ {change_id}: REPLACE_POINT_TEXT → REPLACE_TEXT (ключевые слова)")
                    
                    # Проверяем исходную инструкцию
                    elif instruction_text and any(keyword in instruction_text.lower() for keyword in correction_keywords):
                        enhanced_change["operation"] = "REPLACE_TEXT"
                        corrections_made += 1
                        logger.warning(f"🔧 ДОПОЛНИТЕЛЬНАЯ КОРРЕКЦИЯ {change_id}: REPLACE_POINT_TEXT → REPLACE_TEXT (анализ инструкции)")
                
                # Добавляем информацию о контексте в изменение
                enhanced_change["context_analysis"] = {
                    "element_type": context["element_type"],
                    "recommended_tool": context["recommended_tool"],
                    "reasoning": context["reasoning"]
                }
                
                # Добавляем информацию о таблице, если есть
                if table_analysis and table_analysis["found"]:
                    enhanced_change["table_analysis"] = {
                        "table_index": table_analysis["table_index"],
                        "row_index": table_analysis["row_index"],
                        "cell_index": table_analysis["cell_index"],
                        "full_cell_content": table_analysis["full_cell_content"],
                        "table_context": table_analysis["table_context"]
                    }
                
                enhanced_changes.append(enhanced_change)
                
            except Exception as e:
                logger.error(f"Ошибка анализа контекста для {change_id}: {e}")
                enhanced_changes.append(change)  # Используем исходное изменение
        
        if corrections_made > 0:
            logger.warning(f"🎯 ДВУХЭТАПНЫЙ АНАЛИЗ: Выполнено {corrections_made} коррекций операций")
            logger.info("📋 ИТОГОВЫЕ ОПЕРАЦИИ ПОСЛЕ АНАЛИЗА КОНТЕКСТА:")
            for change in enhanced_changes:
                logger.info(f"  {change.get('change_id')}: {change.get('operation')} ({change.get('context_analysis', {}).get('element_type', 'unknown')})")
        else:
            logger.info("✅ ДВУХЭТАПНЫЙ АНАЛИЗ: Все операции корректны")
        
        return enhanced_changes, tokens_info

    async def _parse_changes_with_llm(
        self, 
        changes_text: str, 
        initial_changes: Optional[List[Dict[str, Any]]] = None
    ) -> tuple[List[Dict[str, Any]], Dict[str, int]]:
        """
        Преобразование текстовых инструкций в структурированный JSON через LLM.
        
        Returns:
            Tuple[список изменений, словарь с информацией о токенах]
        """
        """
        Преобразование текстовых инструкций в структурированный JSON через LLM.
        """
        if not self.openai_client:
            raise RuntimeError("OpenAI клиент не инициализирован")

        # Обрабатываем весь документ без обрезания
        logger.info(f"Обработка полного текста инструкций: {len(changes_text)} символов")

        # Формируем контекст о уже найденных изменениях (для будущего использования)
        initial_context = ""
        if initial_changes:
            initial_context = (
                f"\n\nУЖЕ РАСПОЗНАННЫЕ ИЗМЕНЕНИЯ (для справки, не дублируй их):\n"
                f"Найдено {len(initial_changes)} изменений. "
                f"Твоя задача - найти ВСЕ остальные изменения, которые могли быть пропущены.\n"
            )
        
        # Загрузка system prompt из файла с добавлением описания MCP инструментов
        # Файл: /data/prompts/instruction_check_system.md (или backend/prompts/instruction_check_system.md)
        system_prompt = self._load_prompt("instruction_check_system.md")
        
        # Добавляем описание доступных MCP инструментов для лучшего выбора операций
        mcp_tools_description = self._get_mcp_tools_description()
        system_prompt += "\n\n" + mcp_tools_description

        # Загрузка user prompt из файла
        # Файл: /data/prompts/instruction_check_user.md (или backend/prompts/instruction_check_user.md)
        user_prompt_template = self._load_prompt("instruction_check_user.md")

        # Подготовка промпта с учетом уже найденных изменений
        # Форматируем user_prompt_template, подставляя changes_text
        if "{changes_list}" in user_prompt_template:
            user_prompt = user_prompt_template.format(changes_list=changes_text)
        else:
            user_prompt = user_prompt_template
        
        full_prompt = f"{user_prompt}{initial_context}\n\nИНСТРУКЦИИ ДЛЯ АНАЛИЗА:\n'''{changes_text}'''"
        
        logger.info(f"Отправка запроса к LLM: модель={self.model_name}, длина промпта={len(full_prompt)} символов")
        logger.info(f"📤 System prompt (первые 500 символов): {system_prompt[:500]}...")
        logger.info(f"📤 User prompt (первые 500 символов): {full_prompt[:500]}...")
        logger.debug(f"System prompt длина: {len(system_prompt)} символов")
        logger.debug(f"User prompt длина: {len(user_prompt)} символов")
        logger.debug(f"Changes text длина: {len(changes_text)} символов")
        logger.debug(f"📋 ПОЛНЫЙ SYSTEM PROMPT:\n{system_prompt}")
        logger.debug(f"📋 ПОЛНЫЙ USER PROMPT:\n{full_prompt}")
        
        try:
            # OpenAI SDK использует timeout из http_client, который уже установлен в 300 секунд
            # КРИТИЧЕСКОЕ ИСПРАВЛЕНИЕ: response_format={"type": "json_object"} заставляет LLM вернуть один объект
            # Это может быть причиной, почему возвращается только одна инструкция!
            # Убираем response_format или явно указываем, что нужен массив
            response = await self.openai_client.chat.completions.create(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": full_prompt},
                ],
                temperature=0,
                max_tokens=16384,  # Максимальное значение для completion tokens (gpt-4o поддерживает до 16384)
                # УБРАНО: response_format={"type": "json_object"} - это может заставлять LLM возвращать один объект вместо массива!
            )
            logger.info("Ответ от LLM получен успешно")
        except Exception as e:
            logger.error(f"Ошибка при запросе к LLM: {e}", exc_info=True)
            raise RuntimeError(
                f"Не удалось получить ответ от LLM: {str(e)}. "
                f"Возможно, документ слишком большой или произошел таймаут."
            ) from e

        content = response.choices[0].message.content if response.choices else None
        if isinstance(content, list):
            # new SDKs may return content as list of dicts
            content = "".join(
                segment.get("text", "")
                for segment in content
                if isinstance(segment, dict)
            )
        if not isinstance(content, str) or not content.strip():
            raise RuntimeError("LLM не вернул корректный JSON для парсинга инструкций")

        # ЛОГИРОВАНИЕ сырого ответа от LLM для диагностики
        logger.info(f"📥 Сырой ответ от LLM (первые 1000 символов): {content[:1000]}")
        logger.info(f"📥 Полный сырой ответ от LLM (длина: {len(content)} символов): {content}")
        # КРИТИЧЕСКОЕ: Всегда логируем полный ответ для диагностики проблем с множественными инструкциями
        
        # КРИТИЧЕСКОЕ: Проверяем, сколько инструкций должно быть в ответе
        # Подсчитываем количество инструкций в исходном тексте
        instruction_indicators = [
            "CHG-", "1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.",
            "изложить", "заменить", "исключить", "добавить", "удалить", "изменить"
        ]
        estimated_count = sum(1 for indicator in instruction_indicators if indicator in changes_text)
        logger.info(f"📊 ОЦЕНКА КОЛИЧЕСТВА ИНСТРУКЦИЙ в исходном тексте: {estimated_count} (по индикаторам)")

        # Попытка очистки JSON от возможных проблем
        content_cleaned = content.strip()
        
        # Удаление markdown code blocks, если есть
        if content_cleaned.startswith("```"):
            lines = content_cleaned.split("\n")
            # Удаляем первую строку (```json или ```)
            if len(lines) > 1:
                lines = lines[1:]
            # Удаляем последнюю строку (```)
            if lines and lines[-1].strip() == "```":
                lines = lines[:-1]
            content_cleaned = "\n".join(lines).strip()
        
        # Попытка парсинга JSON
        try:
            parsed = json.loads(content_cleaned)
            logger.debug(f"📋 Распарсенный JSON (после json.loads): {json.dumps(parsed, ensure_ascii=False, indent=2)[:1000]}...")
            
            # НОВЫЙ ФУНКЦИОНАЛ: Попытка восстановления структуры JSON перед валидацией
            parsed = await self._recover_json_structure(parsed, content_cleaned, changes_text)
            logger.debug(f"📋 JSON после восстановления структуры: {json.dumps(parsed, ensure_ascii=False, indent=2)[:1000]}...")
            
            # НОВАЯ ВАЛИДАЦИЯ JSON
            parsed = self._validate_and_fix_json(parsed)
            logger.info(f"📋 Финальный JSON после валидации: количество изменений = {len(parsed.get('changes', []))}")
        except (json.JSONDecodeError, ValueError) as e:
            # Логируем проблемный JSON для отладки
            error_pos = e.pos if hasattr(e, 'pos') else None
            if error_pos:
                start = max(0, error_pos - 100)
                end = min(len(content_cleaned), error_pos + 100)
                context = content_cleaned[start:end]
                logger.error(f"Ошибка парсинга JSON на позиции {error_pos}")
                logger.error(f"Контекст: ...{context}...")
                logger.debug(f"Полный ответ LLM (первые 500 символов): {content_cleaned[:500]}")
            
            # Попытка исправить распространенные проблемы
            try:
                # Удаление trailing commas
                content_fixed = re.sub(r',\s*}', '}', content_cleaned)
                content_fixed = re.sub(r',\s*]', ']', content_fixed)
                parsed = json.loads(content_fixed)
                # НОВЫЙ ФУНКЦИОНАЛ: Попытка восстановления структуры JSON перед валидацией
                parsed = await self._recover_json_structure(parsed, content_fixed, changes_text)
                # НОВАЯ ВАЛИДАЦИЯ JSON
                parsed = self._validate_and_fix_json(parsed)
                logger.info("JSON исправлен автоматически (удалены trailing commas)")
            except (json.JSONDecodeError, ValueError) as e2:
                # Если не удалось исправить, пробуем извлечь JSON из текста
                try:
                    # Ищем JSON объект в тексте
                    json_match = re.search(r'\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}', content_cleaned, re.DOTALL)
                    if json_match:
                        parsed = json.loads(json_match.group(0))
                        # НОВЫЙ ФУНКЦИОНАЛ: Попытка восстановления структуры JSON перед валидацией
                        parsed = await self._recover_json_structure(parsed, json_match.group(0), changes_text)
                        # НОВАЯ ВАЛИДАЦИЯ JSON
                        parsed = self._validate_and_fix_json(parsed)
                        logger.info("JSON извлечен из текста")
                    else:
                        # НОВЫЙ ФУНКЦИОНАЛ: Последняя попытка - извлечь изменения напрямую из текста
                        logger.info("🔧 Попытка прямого извлечения изменений из текста...")
                        try:
                            recovered_parsed = await self._extract_changes_from_text_directly(content_cleaned, changes_text)
                            if recovered_parsed and recovered_parsed.get("changes"):
                                parsed = recovered_parsed
                                logger.info("✅ Изменения успешно извлечены из текста напрямую")
                            else:
                                raise RuntimeError("Не удалось извлечь изменения")
                        except Exception as recover_error:
                            logger.warning(f"Не удалось извлечь изменения напрямую: {recover_error}")
                            raise RuntimeError(
                                f"Не удалось распарсить JSON от LLM. Ошибка: {str(e)}. "
                                f"Позиция: {error_pos}. "
                                f"Попробуйте упростить инструкции или разбить их на части."
                            ) from e
                        raise RuntimeError(
                            f"Не удалось распарсить JSON от LLM. Ошибка: {str(e)}. "
                            f"Позиция: {error_pos}. "
                            f"Попробуйте упростить инструкции или разбить их на части."
                        ) from e
                except (json.JSONDecodeError, AttributeError, ValueError) as e3:
                    # НОВЫЙ ФУНКЦИОНАЛ: Последняя попытка восстановления
                    logger.info("🔧 Последняя попытка восстановления структуры JSON...")
                    try:
                        recovered_parsed = await self._extract_changes_from_text_directly(content_cleaned, changes_text)
                        if recovered_parsed and recovered_parsed.get("changes"):
                            parsed = recovered_parsed
                            logger.info("✅ Изменения успешно восстановлены из текста")
                        else:
                            raise RuntimeError("Не удалось восстановить изменения")
                    except Exception as recover_error:
                        logger.error(f"Все попытки восстановления JSON провалились: {recover_error}")
                        raise RuntimeError(
                            f"Не удалось распарсить JSON от LLM. Ошибка: {str(e)}. "
                            f"Позиция: {error_pos}. "
                            f"Ответ LLM (первые 1000 символов): {content_cleaned[:1000]}"
                        ) from e
        
        changes = parsed.get("changes", [])

        # Валидация результата
        if not isinstance(changes, list):
            raise RuntimeError(
                f"LLM вернул некорректный формат: 'changes' должен быть массивом, "
                f"получен: {type(changes).__name__}"
            )
        
        # ДОПОЛНИТЕЛЬНАЯ ВАЛИДАЦИЯ ОПЕРАЦИЙ
        changes = self._validate_and_correct_operations(changes, changes_text)
        
        # Извлечение информации о токенах
        tokens_info = {
            "prompt_tokens": 0,
            "completion_tokens": 0,
            "total_tokens": 0
        }
        if hasattr(response, 'usage') and response.usage:
            tokens_info = {
                "prompt_tokens": response.usage.prompt_tokens or 0,
                "completion_tokens": response.usage.completion_tokens or 0,
                "total_tokens": response.usage.total_tokens or 0
            }
        
        if not changes:
            logger.warning("LLM не вернул ни одного изменения. Проверьте инструкции.")
            logger.warning(f"Длина исходного текста инструкций: {len(changes_text)} символов")
            logger.warning(f"Первые 500 символов текста: {changes_text[:500]}")
            return [], tokens_info
        
        # Проверка на возможные пропущенные инструкции
        # Улучшенный подсчет инструкций: ищем паттерны, указывающие на отдельные инструкции
        
        # Паттерны, которые могут указывать на отдельные инструкции
        instruction_patterns = [
            r'CHG-\d+',  # Явные номера инструкций
            r'\d+\.\s+[А-Я]',  # Номер с точкой и заглавной буквой (начало новой инструкции)
            r'(?:^|\n)\s*\d+[\.\)]\s+[А-Я]',  # Номер с точкой/скобкой и заглавной буквой в начале строки
            r'Инструкция\s+\d+',  # Явное упоминание "Инструкция N"
        ]
        
        # Подсчитываем явные паттерны инструкций
        pattern_count = 0
        for pattern in instruction_patterns:
            matches = re.findall(pattern, changes_text, re.MULTILINE | re.IGNORECASE)
            if matches:
                logger.debug(f"Паттерн '{pattern}' найден: {len(matches)} совпадений - {matches[:5]}")
            pattern_count = max(pattern_count, len(matches))
        
        # Подсчитываем ключевые слова действий (каждое может быть отдельной инструкцией)
        action_keywords = [
            "изложить", "заменить", "исключить", "добавить", "удалить", 
            "изменить", "в редакции", "в новой редакции", "в следующей редакции"
        ]
        action_count = 0
        action_found = []
        for keyword in action_keywords:
            count = changes_text.lower().count(keyword.lower())
            if count > 0:
                action_count += count
                action_found.append(f"{keyword}:{count}")
        
        # Оценка количества инструкций (берем максимальное из паттернов или ключевых слов)
        estimated_instructions = max(pattern_count, action_count // 2) if pattern_count == 0 else pattern_count
        
        logger.info(f"📊 АНАЛИЗ ИНСТРУКЦИЙ: найдено паттернов: {pattern_count}, ключевых слов действий: {action_count} ({', '.join(action_found[:10])}), оценка количества инструкций: {estimated_instructions}")
        logger.info(f"📊 LLM распознал изменений: {len(changes)}")
        
        # КРИТИЧЕСКОЕ ПРЕДУПРЕЖДЕНИЕ если оценка значительно больше найденных
        if estimated_instructions > len(changes):
            logger.warning(
                f"⚠️ КРИТИЧЕСКОЕ ВНИМАНИЕ: В тексте, вероятно, содержится {estimated_instructions} инструкций "
                f"(найдено паттернов: {pattern_count}, ключевых слов: {action_count}), "
                f"но LLM распознал только {len(changes)} изменений. "
                f"Это серьезное несоответствие! Возможно, LLM пропустил инструкции или вернул неправильный формат."
            )
            logger.warning(
                f"⚠️ ПРОВЕРКА: Первые 500 символов текста инструкций: {changes_text[:500]}"
            )
            logger.warning(
                f"⚠️ ПРОВЕРКА: Полный сырой ответ LLM был залогирован выше в DEBUG режиме"
            )
        
        # Валидация и нормализация каждого изменения
        validated_changes = []
        for idx, change in enumerate(changes, start=1):
            if not isinstance(change, dict):
                logger.warning(f"Пропущено изменение {idx}: не является объектом")
                continue
            
            # Установка обязательных полей
            change.setdefault("change_id", f"CHG-{idx:03d}")
            change.setdefault("annotation", True)
            change.setdefault("operation", "UNKNOWN")
            change.setdefault("description", f"Изменение {idx}")
            
            # Проверка обязательных полей
            operation = change.get("operation", "").upper()
            if operation not in self.SUPPORTED_OPERATIONS:
                logger.warning(f"Пропущено изменение {idx}: неподдерживаемая операция '{operation}'")
                continue
            
            # Автоматическое определение replace_all для массовых замен
            if operation == "REPLACE_TEXT":
                target = change.get("target", {})
                description = change.get("description", "").lower()
                # Если в описании есть "по всему тексту" или match_case=false, устанавливаем replace_all
                if "по всему тексту" in description or target.get("match_case") is False:
                    target.setdefault("replace_all", True)
                    change["target"] = target
                    logger.info(f"Автоматически установлен replace_all=true для {change.get('change_id')}")
            
            validated_changes.append(change)
        
        logger.info(f"Успешно распарсено {len(validated_changes)} изменений из {len(changes)} полученных")
        logger.info(f"Использовано токенов: {tokens_info['total_tokens']} (prompt: {tokens_info['prompt_tokens']}, completion: {tokens_info['completion_tokens']})")
        return validated_changes, tokens_info

    def _validate_and_correct_operations(self, changes: List[Dict[str, Any]], original_text: str) -> List[Dict[str, Any]]:
        """
        Валидация и автокоррекция неправильно выбранных операций LLM.
        
        Args:
            changes: Список изменений от LLM
            original_text: Исходный текст инструкций
            
        Returns:
            Скорректированный список изменений
        """
        logger.info(f"🔍 НАЧАЛО ВАЛИДАЦИИ: получено {len(changes)} изменений")
        corrected_changes = []
        corrections_made = 0
        
        for change in changes:
            operation = change.get("operation", "")
            description = change.get("description", "").lower()
            change_id = change.get("change_id", "")
            
            logger.info(f"🔍 ВАЛИДАЦИЯ {change_id}: operation={operation}, description='{description[:50]}...'")
            
            # Создаем копию изменения для модификации
            corrected_change = change.copy()
            original_operation = operation
            
            # ПРАВИЛА АВТОКОРРЕКЦИИ:
            
            # 1. Если в описании есть "слова", "строку", "фразу" - это REPLACE_TEXT
            if operation == "REPLACE_POINT_TEXT":
                keywords_for_replace_text = ["слова", "строку", "фразу", "текст", "аббревиатуру"]
                if any(keyword in description for keyword in keywords_for_replace_text):
                    corrected_change["operation"] = "REPLACE_TEXT"
                    logger.warning(f"🔧 АВТОКОРРЕКЦИЯ {change_id}: {operation} → REPLACE_TEXT (найдено ключевое слово)")
                    corrections_made += 1
            
            # 2. Если упоминается "в таблице" - это REPLACE_TEXT
            if operation == "REPLACE_POINT_TEXT" and "в таблице" in description:
                corrected_change["operation"] = "REPLACE_TEXT"
                logger.warning(f"🔧 АВТОКОРРЕКЦИЯ {change_id}: {operation} → REPLACE_TEXT (таблица)")
                corrections_made += 1
            
            # 3. Если упоминается "по всему тексту" - это REPLACE_TEXT с replace_all=true
            if "по всему тексту" in description:
                corrected_change["operation"] = "REPLACE_TEXT"
                if "target" in corrected_change and isinstance(corrected_change["target"], dict):
                    corrected_change["target"]["replace_all"] = True
                logger.warning(f"🔧 АВТОКОРРЕКЦИЯ {change_id}: установлен replace_all=true (массовая замена)")
                corrections_made += 1
            
            # 4. Проверяем исходный текст инструкций для дополнительной валидации
            if change_id and original_text:
                # Ищем соответствующую инструкцию в исходном тексте
                lines = original_text.split('\n')
                instruction_text = ""
                
                # Пробуем разные способы найти инструкцию
                for line in lines:
                    line_clean = line.strip().lower()
                    # Ищем строку с номером изменения (1., 2., 3., 4.)
                    if any(marker in line_clean for marker in [f"{change_id[-1]}.", f"{change_id[-1]} "]):
                        instruction_text = line_clean
                        break
                
                if instruction_text:
                    # Дополнительная проверка по исходной инструкции
                    if operation == "REPLACE_POINT_TEXT":
                        correction_triggers = [
                            "слова", "строку", "фразу", "в таблице", 
                            "аббревиатуру", "по всему тексту"
                        ]
                        if any(trigger in instruction_text for trigger in correction_triggers):
                            corrected_change["operation"] = "REPLACE_TEXT"
                            logger.warning(f"🔧 АВТОКОРРЕКЦИЯ {change_id}: {operation} → REPLACE_TEXT (анализ исходной инструкции: '{instruction_text[:50]}...')")
                            corrections_made += 1
            
            # 5. Специальные правила для конкретных паттернов
            if operation == "REPLACE_POINT_TEXT":
                # Если в описании упоминается конкретная замена текста
                replace_patterns = ["изложить в следующей редакции", "заменить на", "изменить на"]
                table_patterns = ["в таблице", "строку"]
                
                has_replace_pattern = any(pattern in description for pattern in replace_patterns)
                has_table_pattern = any(pattern in description for pattern in table_patterns)
                
                if has_replace_pattern and has_table_pattern:
                    corrected_change["operation"] = "REPLACE_TEXT"
                    logger.warning(f"🔧 АВТОКОРРЕКЦИЯ {change_id}: {operation} → REPLACE_TEXT (паттерн замены в таблице)")
                    corrections_made += 1
            
            # Логируем если операция была изменена
            if corrected_change["operation"] != original_operation:
                logger.info(f"✅ Операция скорректирована: {change_id} {original_operation} → {corrected_change['operation']}")
            
            corrected_changes.append(corrected_change)
        
        # Детальное логирование результатов валидации
        if corrections_made > 0:
            logger.warning(f"🔧 ВАЛИДАЦИЯ: Выполнено {corrections_made} автокоррекций операций")
            logger.info("📋 ИТОГОВЫЕ ОПЕРАЦИИ ПОСЛЕ ВАЛИДАЦИИ:")
            for change in corrected_changes:
                logger.info(f"  {change.get('change_id')}: {change.get('operation')} - {change.get('description', '')[:60]}...")
        else:
            logger.info("✅ ВАЛИДАЦИЯ: Все операции корректны")
        
        return corrected_changes

    async def process_documents(
        self,
        source_file: str,
        changes_file: str,
        session_id: str,
        progress_callback: OperationCallback = None,
        operation_id: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        Главный сценарий: создание бэкапа, парсинг инструкций, применение изменений.
        """
        logger.info(f"Начало обработки документов: session_id={session_id}, source={source_file}, changes={changes_file}")
        
        try:
            # Получаем базовое имя файла (без пути)
            source_basename = os.path.basename(source_file)
            logger.info(f"Исходный файл: {source_basename}, полный путь: {source_file}")
            
            # Определяем директории
            # source_file может быть в /data/uploads/{username}/source/{filename}
            # или в другой структуре
            if os.path.dirname(source_file).endswith('source'):
                # Файл в подпапке source
                uploads_dir = os.path.dirname(os.path.dirname(source_file))
            else:
                uploads_dir = os.path.dirname(source_file)
            
            root_dir = os.path.dirname(uploads_dir) if not uploads_dir.endswith('uploads') else uploads_dir
            backup_dir = os.path.join(root_dir, "backups")
            os.makedirs(backup_dir, exist_ok=True)
            logger.info(f"Директория для бэкапов: {backup_dir}")

            backup_filename = os.path.splitext(source_basename)[0] + "_backup.docx"
            backup_path = os.path.join(backup_dir, backup_filename)

            logger.info(f"Создание резервной копии: {backup_path}")
            try:
                await mcp_client.copy_document(source_file, backup_path)
                logger.info("Резервная копия создана успешно")
            except RuntimeError as e:
                logger.warning(f"MCP сервер недоступен, используем локальное копирование: {e}")
                import shutil
                shutil.copy2(source_file, backup_path)
                logger.info("Резервная копия создана локально")

            logger.info("Извлечение текста из файла с инструкциями")
            changes_text = await self._safe_get_document_text(changes_file)
            logger.debug(f"Извлечено {len(changes_text)} символов инструкций")
            
            # Простое распознавание изменений без двухэтапного анализа
            logger.info("🚀 Запуск простого распознавания инструкций")
            # Сохраняем исходный текст инструкций для исправления target.text
            self._original_instructions_text = changes_text
            changes, tokens_info_parse = await self._simple_parse_changes_with_llm(changes_text, initial_changes=[])
            
            # Сохраняем исходный порядок операций из файла инструкций
            logger.info("📋 ПОРЯДОК ОПЕРАЦИЙ: Сохраняется исходный порядок из файла инструкций")
            logger.info(f"🎯 Анализ завершен: {len(changes)} изменений")
            logger.info(f"Использовано токенов при парсинге: {tokens_info_parse.get('total_tokens', 0)}")
            
            # Нумерация изменений с сохранением исходного порядка
            logger.info("📋 ИСХОДНЫЙ ПОРЯДОК ИЗМЕНЕНИЙ (как указано в файле):")
            for idx, change in enumerate(changes, start=1):
                change["change_id"] = f"CHG-{idx:03d}"
                description = change.get("description", "N/A")
                logger.info(f"   {idx}. {change['change_id']}: {description[:60]}...")

            if not changes:
                logger.warning("Не найдено изменений для применения")
                return {
                    "session_id": session_id,
                    "total_changes": 0,
                    "successful": 0,
                    "failed": 0,
                    "changes": [],
                    "processed_filename": source_basename,
                    "backup_filename": backup_filename,
                    "warning": "Не найдено изменений для применения",
                }

            results: List[Dict[str, Any]] = []
            total = len(changes)
            
            # КРИТИЧЕСКОЕ: Создаем ОДИН объект Document() для всех изменений
            # Все изменения будут применяться к этому объекту, и файл сохранится только один раз в конце
            logger.info(f"📄 Создание единого объекта Document() для всех {total} изменений")
            master_doc = Document(source_file)
            logger.info(f"✅ Документ загружен: {len(master_doc.paragraphs)} параграфов, {len(master_doc.tables)} таблиц")
            
            logger.info(f"🚀 Начало последовательного применения {total} изменений в ИСХОДНОМ порядке")
            logger.info("📋 ПОРЯДОК ВЫПОЛНЕНИЯ:")
            for idx, change in enumerate(changes, start=1):
                change_id = change.get("change_id", f"CHG-{idx:03d}")
                description = change.get("description", "N/A")
                logger.info(f"   {idx}. {change_id}: {description[:60]}...")
            
            # Последовательное выполнение изменений строго в исходном порядке
            for idx, change in enumerate(changes, start=1):
                change_id = change.get("change_id", f"CHG-{idx:03d}")
                operation = change.get("operation", "UNKNOWN")
                logger.info(f"Обработка {change_id}: {operation}")
                
                try:
                    execution_result = await self._execute_change(
                        source_file,
                        change,
                        progress_callback=progress_callback,
                        changes_file=changes_file,
                        master_doc=master_doc,  # Передаем единый объект документа
                    )
                    results.append(execution_result)
                    
                    if execution_result["status"] == "SUCCESS":
                        logger.info(f"{change_id}: успешно выполнено")
                    else:
                        error_msg = execution_result.get("details", {}).get("message", "Неизвестная ошибка")
                        error_type = execution_result.get("details", {}).get("error", "UNKNOWN")
                        logger.warning(f"{change_id}: ошибка - {error_msg} (тип: {error_type})")
                        
                        # НОВЫЙ ФУНКЦИОНАЛ: Дополнительная диагностика для неудачных изменений
                        if error_type in ["TEXT_NOT_FOUND", "TEXT_NOT_FOUND_IN_PARAGRAPH"]:
                            logger.info(f"🔍 ДИАГНОСТИКА для {change_id}:")
                            logger.info(f"   Описание: {change.get('description', 'N/A')}")
                            logger.info(f"   Target text: {change.get('target', {}).get('text', 'N/A')}")
                            logger.info(f"   New text: {change.get('payload', {}).get('new_text', 'N/A')}")
                            logger.info(f"   Operation: {change.get('operation', 'N/A')}")
                            
                            # Попытка найти альтернативные варианты текста
                            target_text = change.get("target", {}).get("text", "")
                            if target_text:
                                logger.info(f"   🔍 Поиск альтернативных вариантов для '{target_text}'...")
                                try:
                                    alt_matches = await self._safe_find_text(source_file, target_text[:20] if len(target_text) > 20 else target_text, match_case=False)
                                    if alt_matches:
                                        logger.info(f"   ✅ Найдено {len(alt_matches)} альтернативных совпадений")
                                    else:
                                        logger.info(f"   ❌ Альтернативные совпадения не найдены")
                                except Exception as e:
                                    logger.debug(f"   Ошибка при поиске альтернатив: {e}")

                except Exception as exc:  # noqa: BLE001
                    logger.error(f"{change_id}: исключение при выполнении - {exc}", exc_info=True)
                    results.append({
                        "change_id": change_id,
                        "operation": operation,
                        "description": change.get("description", ""),
                        "status": "FAILED",
                        "details": {
                            "success": False,
                            "error": "EXCEPTION",
                            "message": str(exc),
                        },
                    })

                if progress_callback:
                    await progress_callback(
                        {
                            "type": "progress",
                            "data": {
                                "status": f"Выполнено {idx} из {total} изменений",
                                "progress": int(idx / max(total, 1) * 100),
                            },
                        }
                    )

            successful = sum(1 for r in results if r["status"] == "SUCCESS")
            failed = sum(1 for r in results if r["status"] == "FAILED")

            logger.info(f"Обработка завершена: успешно={successful}, ошибок={failed}")
            
            # КРИТИЧЕСКОЕ: Сохраняем файл ОДИН РАЗ после всех изменений
            # Все изменения применены к master_doc, теперь сохраняем его в файл
            logger.info(f"💾 ФИНАЛЬНОЕ СОХРАНЕНИЕ документа после всех {total} изменений...")
            try:
                master_doc.save(source_file)
                logger.info(f"✅ Документ успешно сохранен: {source_file}")
                
                # Проверяем, что файл существует и доступен
                if os.path.exists(source_file):
                    file_size = os.path.getsize(source_file)
                    logger.info(f"✅ Файл подтвержден: размер {file_size} байт")
                else:
                    logger.error(f"❌ Файл не найден после сохранения: {source_file}")
            except Exception as final_save_e:
                logger.error(f"❌ Ошибка при финальном сохранении документа: {final_save_e}", exc_info=True)
            
            # ФИНАЛЬНАЯ ПРОВЕРКА: Проверяем, что изменения действительно присутствуют в файле
            # Используем master_doc напрямую (он уже сохранен), а не читаем с диска
            try:
                logger.info(f"💾 ПРОВЕРКА финального состояния документа после всех изменений...")
                
                    # Используем master_doc напрямую для проверки, а не читаем с диска
                    # master_doc уже сохранен, поэтому можем проверять его напрямую
                    verify_doc = master_doc
                    
                    # Ищем хотя бы один успешный результат и проверяем, что новое содержимое присутствует
                    for result in results:
                        if result.get("status") == "SUCCESS":
                            payload = result.get("payload", {})
                            new_text = payload.get("new_text", "")
                            if new_text:
                                # Проверяем, есть ли новый текст в документе
                                all_text = "\n".join([p.text for p in verify_doc.paragraphs])
                                if new_text[:50] in all_text:
                                    logger.info(f"✅ Изменение подтверждено в финальном файле: '{new_text[:50]}...' найден")
                                else:
                                    logger.warning(f"⚠️ Изменение НЕ найдено в финальном файле: '{new_text[:50]}...' отсутствует")
                                    logger.warning(f"⚠️ Это критическая проблема - изменения потеряны!")
                    logger.info(f"✅ Финальная проверка документа завершена")
                    except Exception as verify_e:
                        logger.warning(f"⚠️ Не удалось проверить содержимое файла: {verify_e}")
                else:
                    logger.error(f"❌ Файл не найден: {source_file}")
            except Exception as final_check_e:
                logger.error(f"❌ Ошибка при проверке финального состояния документа: {final_check_e}", exc_info=True)

            # НОВАЯ ФУНКЦИЯ: Добавляем автоматические аннотации для отслеживания изменений
            logger.info(f"🔍 ПРОВЕРКА АННОТАЦИЙ: successful={successful}, total_results={len(results)}")
            
            if successful > 0:
                # Проверяем настройку для аннотаций (по умолчанию включено)
                add_annotations = os.getenv("ADD_CHANGE_ANNOTATIONS", "true").lower() == "true"
                logger.info(f"🔍 НАСТРОЙКА АННОТАЦИЙ: ADD_CHANGE_ANNOTATIONS={add_annotations}")
                
                if add_annotations:
                    logger.info("📝 ДОБАВЛЕНИЕ АВТОМАТИЧЕСКИХ АННОТАЦИЙ")
                    try:
                        annotation_results = await self._add_change_annotations(source_file, results, session_id)
                        logger.info(f"✅ Добавлено аннотаций: {annotation_results.get('annotations_added', 0)}")
                    except Exception as e:
                        logger.error(f"❌ Ошибка добавления аннотаций: {e}")
                        import traceback
                        logger.error(f"Трассировка: {traceback.format_exc()}")
                else:
                    logger.info("📝 Автоматические аннотации отключены (ADD_CHANGE_ANNOTATIONS=false)")
            else:
                logger.info("📝 Аннотации не добавляются: нет успешных изменений")

            # Собираем информацию о токенах
            tokens_total = tokens_info_parse.get("total_tokens", 0)
            tokens_prompt = tokens_info_parse.get("prompt_tokens", 0)
            tokens_completion = tokens_info_parse.get("completion_tokens", 0)

            return {
                "session_id": session_id,
                "total_changes": total,
                "successful": successful,
                "failed": failed,
                "changes": results,
                "processed_filename": source_basename,
                "backup_filename": backup_filename,
                "tokens_used": tokens_total,
                "tokens_prompt": tokens_prompt,
                "tokens_completion": tokens_completion,
            }

        except Exception as exc:  # noqa: BLE001
            error_msg = str(exc)
            logger.error(f"Критическая ошибка при обработке документов: {error_msg}", exc_info=True)
            
            # Специальная обработка таймаутов
            if "timeout" in error_msg.lower() or "timed out" in error_msg.lower():
                error_msg = (
                    "Превышено время ожидания ответа от LLM. "
                    "Возможно, документ с инструкциями слишком большой. "
                    "Попробуйте разбить инструкции на несколько файлов или упростить их."
                )
            
            return {
                "session_id": session_id,
                "status": "FAILED",
                "error": error_msg,
                "total_changes": 0,
                "successful": 0,
                "failed": 0,
                "changes": [],
                "tokens_used": 0,
                "tokens_prompt": 0,
                "tokens_completion": 0,
            }

    async def _execute_change(
        self,
        filename: str,
        change: Dict[str, Any],
        progress_callback: OperationCallback = None,
        changes_file: Optional[str] = None,
        master_doc: Optional[Document] = None,  # Единый объект документа для всех изменений
    ) -> Dict[str, Any]:
        """
        Выполнение одного изменения в документе.
        """
        change_id = change.get("change_id", "UNKNOWN")
        operation = change.get("operation", "").upper()
        description = change.get("description", "")
        
        # Извлекаем target_text для аннотаций
        target_text = ""
        if "target" in change and isinstance(change["target"], dict):
            target_text = change["target"].get("text", "")
        elif "target_text" in change:
            target_text = change["target_text"]
        
        result: Dict[str, Any] = {
            "change_id": change_id,
            "operation": operation,
            "description": description,
            "status": "FAILED",
            "details": {},
            "target_text": target_text,  # Сохраняем для аннотаций
            "target": change.get("target", {}),  # Сохраняем весь объект target
            "payload": change.get("payload", {}),  # Сохраняем payload для аннотаций
        }

        if operation not in self.SUPPORTED_OPERATIONS:
            error_msg = f"Операция {operation} не поддерживается"
            logger.warning(f"{change_id}: {error_msg}")
            result["details"] = {
                "success": False,
                "error": "UNSUPPORTED_OPERATION",
                "message": error_msg,
            }
            return result

        try:
            logger.debug(f"{change_id}: выполнение операции {operation}")
            
            if operation == "REPLACE_TEXT":
                details = await self._handle_replace_text(filename, change, changes_file=changes_file, master_doc=master_doc)
            elif operation == "REPLACE_POINT_TEXT":
                details = await self._handle_replace_point_text(filename, change, master_doc=master_doc)
            elif operation == "DELETE_PARAGRAPH":
                details = await self._handle_delete_paragraph(filename, change, master_doc=master_doc)
            elif operation == "INSERT_PARAGRAPH":
                details = await self._handle_insert_paragraph(filename, change, master_doc=master_doc)
            elif operation == "INSERT_SECTION":
                details = await self._handle_insert_section(filename, change, master_doc=master_doc)
            elif operation == "INSERT_TABLE":
                details = await self._handle_insert_table(filename, change, master_doc=master_doc)
            elif operation == "ADD_COMMENT":
                # Комментарии используют MCP, который сохраняет файл сам - передаем master_doc для совместимости
                details = await self._handle_add_comment(filename, change, master_doc=master_doc)
            else:
                error_msg = f"Операция {operation} не реализована"
                logger.warning(f"{change_id}: {error_msg}")
                details = {
                    "success": False,
                    "error": "UNSUPPORTED_OPERATION",
                    "message": error_msg,
                }
        except Exception as exc:  # noqa: BLE001
            logger.error(f"{change_id}: исключение при выполнении операции {operation}: {exc}", exc_info=True)
            details = {"success": False, "error": "EXCEPTION", "message": str(exc)}

        # НОВЫЙ ФУНКЦИОНАЛ: Повторная попытка применения изменения при неудаче
        if not details.get("success") and operation in ["REPLACE_TEXT", "DELETE_PARAGRAPH"]:
            logger.info(f"🔄 {change_id}: Первая попытка не удалась, пробуем альтернативные стратегии...")
            retry_details = await self._retry_change_application(filename, change, operation, details)
            if retry_details.get("success"):
                logger.info(f"✅ {change_id}: Успешно применено после повторной попытки")
                details = retry_details
            else:
                logger.warning(f"❌ {change_id}: Повторная попытка также не удалась")

        result["details"] = details
        result["status"] = "SUCCESS" if details.get("success") else "FAILED"

        if progress_callback:
            await progress_callback(
                {
                    "type": "operation_completed",
                    "data": {
                        "change_id": change_id,
                        "operation": operation,
                        "description": description,
                        "status": result["status"],
                        "details": details,
                    },
                }
            )

        return result
    
    async def _retry_change_application(
        self,
        filename: str,
        change: Dict[str, Any],
        operation: str,
        original_details: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        НОВЫЙ ФУНКЦИОНАЛ: Повторная попытка применения изменения с альтернативными стратегиями.
        
        Когда первая попытка применения изменения не удалась, пробует альтернативные подходы:
        1. Для REPLACE_TEXT: пробует различные варианты поиска и замены
        2. Для DELETE_PARAGRAPH: пробует альтернативные способы поиска параграфа
        3. Детальное логирование для диагностики
        
        Args:
            filename: Путь к файлу
            change: Объект изменения
            operation: Тип операции
            original_details: Детали первой неудачной попытки
            
        Returns:
            Детали результата повторной попытки
        """
        change_id = change.get("change_id", "UNKNOWN")
        logger.info(f"🔄 ПОВТОРНАЯ ПОПЫТКА для {change_id}: операция {operation}")
        logger.info(f"   Оригинальная ошибка: {original_details.get('error', 'UNKNOWN')}")
        logger.info(f"   Сообщение: {original_details.get('message', 'N/A')}")
        
        if operation == "REPLACE_TEXT":
            return await self._retry_replace_text(filename, change, original_details)
        elif operation == "DELETE_PARAGRAPH":
            return await self._retry_delete_paragraph(filename, change, original_details)
        
        # Для других операций возвращаем оригинальный результат
        return original_details
    
    async def _select_best_match_for_local_change(
        self,
        filename: str,
        matches: List[MCPTextMatch],
        target_text: str,
        description: str
    ) -> Optional[MCPTextMatch]:
        """
        НОВЫЙ ФУНКЦИОНАЛ: Универсальный выбор наиболее подходящего совпадения для локального изменения.
        
        Когда найдено несколько совпадений для локального изменения, анализирует контекст каждого
        совпадения и выбирает наиболее подходящее на основе:
        1. Ключевых слов из описания в контексте совпадений
        2. Структурных признаков (заголовки, пункты, разделы)
        3. Позиции в документе (первые вхождения чаще являются заголовками)
        4. Схожести контекста с описанием изменения
        
        Args:
            filename: Путь к файлу
            matches: Список найденных совпадений
            target_text: Искомый текст
            description: Описание изменения
            
        Returns:
            Наиболее подходящее совпадение или None, если выбрать не удалось
        """
        if len(matches) <= 1:
            return matches[0] if matches else None
        
        logger.info(f"🔍 Анализ {len(matches)} совпадений для выбора наиболее подходящего...")
        logger.info(f"   Описание: '{description}'")
        logger.info(f"   Искомый текст: '{target_text}'")
        
        try:
            doc = Document(filename)
            
            # Извлекаем ключевые слова из описания (исключая стоп-слова)
            description_lower = description.lower()
            stop_words = {'в', 'на', 'из', 'к', 'с', 'для', 'по', 'от', 'за', 'под', 'над', 
                         'при', 'о', 'об', 'изложить', 'изменить', 'заменить', 'строку', 
                         'текст', 'новой', 'редакции', 'редакцию', 'пункт', 'пункте', 'глава', 'главе'}
            
            description_words = [
                word for word in description_lower.split() 
                if len(word) > 2 and word not in stop_words
            ]
            
            logger.info(f"   Ключевые слова из описания: {description_words}")
            
            # Оценка каждого совпадения
            match_scores = []
            
            for match_idx, match in enumerate(matches):
                score = 0
                para_idx = match.paragraph_index
                
                if para_idx >= len(doc.paragraphs):
                    continue
                
                para = doc.paragraphs[para_idx]
                para_text = para.text
                
                # Собираем расширенный контекст: текущий параграф + предыдущий + следующий
                context_text = para_text.lower()
                
                if para_idx > 0:
                    prev_para_text = doc.paragraphs[para_idx - 1].text.lower()
                    context_text = prev_para_text + " " + context_text
                
                if para_idx < len(doc.paragraphs) - 1:
                    next_para_text = doc.paragraphs[para_idx + 1].text.lower()
                    context_text = context_text + " " + next_para_text
                
                # Критерий 1: Наличие ключевых слов из описания в контексте
                words_found = sum(1 for word in description_words if word in context_text)
                score += words_found * 10
                
                # Критерий 2: Проверка на заголовок/главу (обычно в начале строки, короткие параграфы)
                is_heading_like = (
                    para_text.strip().startswith(target_text) or
                    para_text.strip() == target_text or
                    len(para_text.split()) <= 10
                )
                if is_heading_like:
                    score += 5
                
                # Критерий 3: Позиция в документе (первые вхождения чаще являются заголовками)
                # Чем раньше в документе, тем выше оценка
                position_score = max(0, (len(doc.paragraphs) - para_idx) / len(doc.paragraphs) * 3)
                score += position_score
                
                # Критерий 4: Точное совпадение с описанием (если есть упоминание номера пункта/главы)
                # Ищем номера в описании
                numbers_in_desc = re.findall(r'\d+', description)
                if numbers_in_desc:
                    # Ищем эти номера в контексте
                    for num in numbers_in_desc:
                        if num in para_text or num in (doc.paragraphs[para_idx - 1].text if para_idx > 0 else ""):
                            score += 15
                
                # Критерий 5: Проверка на структурные элементы (нумерация, заголовки)
                if re.match(r'^\s*\d+[\.\)]\s*', para_text) or re.match(r'^\s*[А-ЯЁ]+\s*\d+', para_text):
                    # Это может быть пункт или раздел
                    score += 3
                
                match_scores.append((score, match_idx, match, para_text[:50]))
                logger.info(f"   Совпадение {match_idx + 1} (параграф {para_idx}): оценка {score:.1f}, текст: '{para_text[:50]}...'")
            
            if not match_scores:
                logger.warning("   ⚠️ Не удалось оценить совпадения")
                # В крайнем случае возвращаем первое совпадение
                return matches[0] if matches else None
            
            # Сортируем по оценке (высшая оценка = лучший выбор)
            match_scores.sort(key=lambda x: x[0], reverse=True)
            best_score, best_idx, best_match, best_text = match_scores[0]
            
            # Если есть значительная разница в оценках (более 5 баллов), выбираем лучший
            if len(match_scores) > 1:
                second_score = match_scores[1][0]
                if best_score - second_score >= 5:
                    logger.info(f"   ✅ Выбрано совпадение {best_idx + 1} (разница в оценке: {best_score - second_score:.1f})")
                    return best_match
                else:
                    # Если оценки близки, используем дополнительные критерии
                    logger.info(f"   ⚠️ Оценки близки (лучшая: {best_score:.1f}, вторая: {second_score:.1f}), используем дополнительные критерии")
                    
                    # Предпочитаем первое вхождение, если разница небольшая
                    first_match = matches[0]
                    logger.info(f"   ✅ Выбрано первое совпадение как наиболее вероятное для локального изменения")
                    return first_match
            
            logger.info(f"   ✅ Выбрано совпадение {best_idx + 1} с оценкой {best_score:.1f}")
            return best_match
            
        except Exception as e:
            logger.warning(f"   ⚠️ Ошибка при выборе совпадения: {e}")
            # В случае ошибки возвращаем первое совпадение
            return matches[0] if matches else None
    
    async def _retry_replace_text(
        self,
        filename: str,
        change: Dict[str, Any],
        original_details: Dict[str, Any],
        master_doc: Optional[Document] = None  # Единый объект документа для всех изменений
    ) -> Dict[str, Any]:
        """
        НОВЫЙ ФУНКЦИОНАЛ: Повторная попытка замены текста с альтернативными стратегиями.
        """
        change_id = change.get("change_id", "UNKNOWN")
        target = change.get("target", {})
        payload = change.get("payload", {})
        target_text = target.get("text", "")
        new_text = payload.get("new_text", "")
        description = change.get("description", "")
        
        logger.info(f"🔄 Стратегия 1: Попытка замены через прямое обращение к документу")
        try:
            # Используем master_doc, если передан, иначе создаем новый
            if master_doc is not None:
                doc = master_doc
                logger.info(f"📄 Используем master_doc для повторной попытки замены")
            else:
                doc = Document(filename)
            
            # Стратегия 1: Прямой поиск и замена в параграфах
            replaced_count = 0
            for para in doc.paragraphs:
                para_text = para.text
                if target_text in para_text:
                    # Заменяем напрямую в параграфе
                    para.clear()
                    para.add_run(para_text.replace(target_text, new_text))
                    replaced_count += 1
                    logger.info(f"   ✅ Найдено и заменено в параграфе: '{target_text}' → '{new_text}'")
            
            # Стратегия 1.2: Поиск в таблицах
            if replaced_count == 0:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            cell_text = cell.text
                            if target_text in cell_text:
                                # Заменяем в ячейке
                                cell.text = cell_text.replace(target_text, new_text)
                                replaced_count += 1
                                logger.info(f"   ✅ Найдено и заменено в таблице: '{target_text}' → '{new_text}'")
            
            if replaced_count > 0:
                # КРИТИЧЕСКОЕ: НЕ сохраняем файл здесь, если используется master_doc
                # Файл будет сохранен один раз в конце после всех изменений
                if master_doc is None:
                    doc.save(filename)
                    logger.info(f"💾 Файл сохранен после повторной попытки (master_doc не использовался)")
                logger.info(f"✅ {change_id}: Успешно применено после повторной попытки ({replaced_count} замен)")
                return {
                    "success": True,
                    "replacements_count": replaced_count,
                    "retry_method": "direct_document_access"
                }
        except Exception as e:
            logger.warning(f"   ⚠️ Стратегия 1 не удалась: {e}")
        
        # Стратегия 2: Попытка через расширенный поиск
        logger.info(f"🔄 Стратегия 2: Расширенный поиск с различными вариантами")
        try:
            # Пробуем найти текст различными способами
            matches = await self._enhanced_text_search(filename, target_text, description, match_case=False)
            if matches:
                logger.info(f"   ✅ Найдено {len(matches)} совпадений через расширенный поиск")
                doc = Document(filename)
                replaced_count = 0
                
                for match in matches:
                    para_idx = match.paragraph_index
                    if para_idx < len(doc.paragraphs):
                        para = doc.paragraphs[para_idx]
                        para_text = para.text
                        if target_text in para_text:
                            para.clear()
                            para.add_run(para_text.replace(target_text, new_text))
                            replaced_count += 1
                
                if replaced_count > 0:
                    doc.save(filename)
                    logger.info(f"✅ {change_id}: Успешно применено через расширенный поиск ({replaced_count} замен)")
                    return {
                        "success": True,
                        "replacements_count": replaced_count,
                        "retry_method": "enhanced_search"
                    }
        except Exception as e:
            logger.warning(f"   ⚠️ Стратегия 2 не удалась: {e}")
        
        # Стратегия 3: Частичная замена (если текст длинный)
        if len(target_text) > 10:
            logger.info(f"🔄 Стратегия 3: Попытка частичной замены")
            try:
                words = target_text.split()
                if len(words) > 2:
                    # Берем ключевые слова
                    key_words = " ".join(words[:3])
                    logger.info(f"   Поиск по ключевым словам: '{key_words}'")
                    
                    # Используем master_doc, если передан, иначе создаем новый
                    if master_doc is not None:
                        doc = master_doc
                        logger.info(f"📄 Используем master_doc для частичной замены")
                    else:
                        doc = Document(filename)
                    replaced_count = 0
                    
                    for para in doc.paragraphs:
                        para_text = para.text
                        if key_words in para_text and target_text in para_text:
                            para.clear()
                            para.add_run(para_text.replace(target_text, new_text))
                            replaced_count += 1
                            logger.info(f"   ✅ Частичная замена успешна")
                    
                    if replaced_count > 0:
                        # КРИТИЧЕСКОЕ: НЕ сохраняем файл здесь, если используется master_doc
                        if master_doc is None:
                            doc.save(filename)
                            logger.info(f"💾 Файл сохранен после частичной замены (master_doc не использовался)")
                        logger.info(f"✅ {change_id}: Успешно применено через частичную замену ({replaced_count} замен)")
                        return {
                            "success": True,
                            "replacements_count": replaced_count,
                            "retry_method": "partial_replacement"
                        }
            except Exception as e:
                logger.warning(f"   ⚠️ Стратегия 3 не удалась: {e}")
        
        logger.warning(f"❌ {change_id}: Все стратегии повторной попытки не удались")
        return original_details
    
    async def _retry_delete_paragraph(
        self,
        filename: str,
        change: Dict[str, Any],
        original_details: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        НОВЫЙ ФУНКЦИОНАЛ: Повторная попытка удаления параграфа с альтернативными стратегиями.
        """
        change_id = change.get("change_id", "UNKNOWN")
        target = change.get("target", {})
        target_text = target.get("text", "")
        description = change.get("description", "")
        
        logger.info(f"🔄 Повторная попытка удаления параграфа для {change_id}")
        
        try:
            doc = Document(filename)
            
            # Стратегия 1: Прямой поиск по тексту
            for idx, para in enumerate(doc.paragraphs):
                para_text = para.text
                if target_text in para_text or (target_text.isdigit() and target_text in para_text):
                    # Удаляем параграф
                    para_element = para._element
                    para_element.getparent().remove(para_element)
                    # КРИТИЧЕСКОЕ: НЕ сохраняем файл здесь, если используется master_doc
                    if master_doc is None:
                        doc.save(filename)
                        logger.info(f"💾 Файл сохранен после удаления параграфа (master_doc не использовался)")
                    logger.info(f"✅ {change_id}: Параграф успешно удален (индекс {idx})")
                    return {
                        "success": True,
                        "paragraph_index": idx,
                        "retry_method": "direct_search"
                    }
            
            # Стратегия 2: Поиск по номеру пункта из описания
            punkt_match = re.search(r'пункт[еа]?\s+(\d+)', description, re.IGNORECASE)
            if punkt_match:
                punkt_num = punkt_match.group(1)
                logger.info(f"   Поиск пункта {punkt_num} из описания")
                
                for idx, para in enumerate(doc.paragraphs):
                    para_text = para.text
                    # Ищем различные форматы номера пункта
                    if (f"{punkt_num}." in para_text or 
                        f"{punkt_num})" in para_text or 
                        f" {punkt_num} " in para_text):
                        para_element = para._element
                        para_element.getparent().remove(para_element)
                        doc.save(filename)
                        logger.info(f"✅ {change_id}: Параграф удален по номеру пункта (индекс {idx})")
                        return {
                            "success": True,
                            "paragraph_index": idx,
                            "retry_method": "punkt_number_search"
                        }
        except Exception as e:
            logger.warning(f"   ⚠️ Повторная попытка удаления не удалась: {e}")
        
        return original_details

    async def _handle_replace_text(self, filename: str, change: Dict[str, Any], changes_file: Optional[str] = None, master_doc: Optional[Document] = None) -> Dict[str, Any]:
        """
        Обработка замены текста с поддержкой массовых замен и интеллектуальным анализом таблиц.
        """
        target = change.get("target", {})
        payload = change.get("payload", {})
        target_text = target.get("text")
        new_text = payload.get("new_text", "")
        match_case = target.get("match_case", False)
        replace_all = target.get("replace_all", False)  # Флаг для массовых замен
        description = change.get("description", "")

        # Для инструкций "Изложить пункт X в новой редакции" new_text может быть пустым,
        # так как новое содержимое (включая таблицы) извлекается из документа инструкций
        is_full_paragraph_replacement = (
            "изложить" in description.lower() and 
            "пункт" in description.lower() and 
            ("редакции" in description.lower() or "редакция" in description.lower())
        )
        
        if not target_text:
            return {
                "success": False,
                "error": "INVALID_PAYLOAD",
                "message": "Для REPLACE_TEXT необходим target.text",
            }
        
        # Разрешаем пустой new_text только для полной замены пункта
        if not new_text and not is_full_paragraph_replacement:
            return {
                "success": False,
                "error": "INVALID_PAYLOAD",
                "message": "Для REPLACE_TEXT необходим payload.new_text (кроме случаев 'Изложить пункт в новой редакции')",
            }
        
        # Для полной замены пункта логируем особый случай
        if is_full_paragraph_replacement and not new_text:
            logger.info(f"🔍 ИНСТРУКЦИЯ 'Изложить пункт в новой редакции': новый текст будет извлечен из документа инструкций")

        # ИНТЕЛЛЕКТУАЛЬНАЯ ЛОГИКА: Сначала определяем, где находится пункт - в таблице или в параграфе
        description_lower = description.lower()
        
        # Проверяем, есть ли в инструкции номер пункта
        punkt_in_instruction = False
        punkt_number = None
        punkt_match = re.search(r'пункт[еа]?\s+(\d+)', description, re.IGNORECASE)
        if punkt_match:
            punkt_number = punkt_match.group(1)
            punkt_in_instruction = True
            logger.info(f"📋 ОБНАРУЖЕН НОМЕР ПУНКТА в инструкции: {punkt_number}")
        
        # Если есть номер пункта, определяем, где он находится - в таблице или в параграфе
        is_table_change = False
        if punkt_in_instruction:
            logger.info(f"🔍 ОПРЕДЕЛЕНИЕ МЕСТОПОЛОЖЕНИЯ ПУНКТА {punkt_number}: проверяем таблицы и параграфы...")
            
            # Ищем номер пункта в документе
            punkt_patterns = [f"{punkt_number}.", f"{punkt_number})", f"{punkt_number}."]
            punkt_location = None  # "table" или "paragraph" или None
            
            try:
                doc = Document(filename)
                
                # Сначала проверяем таблицы - ищем номер пункта в первой ячейке строк
                for table_idx, table in enumerate(doc.tables):
                    for row_idx, row in enumerate(table.rows):
                        if len(row.cells) > 0:
                            first_cell_text = row.cells[0].text.strip()
                            # Проверяем, начинается ли первая ячейка с номера пункта
                            for pattern in punkt_patterns:
                                if first_cell_text.startswith(pattern) or first_cell_text == punkt_number:
                                    punkt_location = "table"
                                    logger.info(f"   ✅ Пункт {punkt_number} найден в ТАБЛИЦЕ {table_idx}, строка {row_idx}")
                                    is_table_change = True
                                    break
                            if punkt_location == "table":
                                break
                    if punkt_location == "table":
                        break
                
                # Если не нашли в таблицах, проверяем параграфы
                if punkt_location != "table":
                    for para_idx, para in enumerate(doc.paragraphs):
                        para_text = para.text.strip()
                        # Проверяем, начинается ли параграф с номера пункта
                        for pattern in punkt_patterns:
                            if para_text.startswith(pattern) or para_text == punkt_number:
                                punkt_location = "paragraph"
                                logger.info(f"   ✅ Пункт {punkt_number} найден в ПАРАГРАФЕ {para_idx}")
                                break
                        if punkt_location == "paragraph":
                            break
                
                if punkt_location:
                    logger.info(f"📍 МЕСТОПОЛОЖЕНИЕ ПУНКТА {punkt_number}: {punkt_location.upper()}")
                else:
                    logger.warning(f"⚠️ Пункт {punkt_number} не найден ни в таблицах, ни в параграфах")
                    
            except Exception as e:
                logger.warning(f"⚠️ Ошибка при определении местоположения пункта: {e}")
                # Fallback: используем стандартную логику
                is_table_change = "таблице" in description_lower
        else:
            # Если номера пункта нет, используем стандартную проверку
            is_table_change = "таблице" in description_lower
        
        logger.info(f"🔍 ПРОВЕРКА ТАБЛИЦЫ: is_table_change={is_table_change}, description='{description[:50]}...'")
        
        # Для инструкций "Изложить пункт X в новой редакции" используем специальную обработку
        if is_full_paragraph_replacement:
            logger.info(f"🔍 ИНСТРУКЦИЯ 'Изложить пункт {target_text} в новой редакции': ищем текст пункта в документе")
            logger.info(f"    Используем интеллектуальную замену для пункта {target_text}")
            # Используем интеллектуальную замену для обработки полной замены пункта
            # с поддержкой извлечения таблиц из инструкций
            matches = await self._safe_find_text(filename, target_text, match_case)
            if matches:
                result = await self._intelligent_paragraph_replacement(
                    filename, target_text, new_text, description, matches, changes_file=changes_file
                )
                if result.get("success"):
                    return result
                else:
                    logger.warning(f"⚠️ Интеллектуальная замена не удалась, продолжаем со стандартной логикой: {result.get('message', 'Неизвестная ошибка')}")
        
        if is_table_change:
            logger.info("🧠 ОБНАРУЖЕНО ИЗМЕНЕНИЕ В ТАБЛИЦЕ - запуск интеллектуальной замены")
            logger.info(f"   Target: '{target_text}', New: '{new_text}'")
            
            try:
                # Интеллектуальная замена в таблице с анализом структуры
                result = await self._intelligent_table_replacement(filename, target_text, new_text, description)
                logger.info(f"🧠 РЕЗУЛЬТАТ интеллектуальной замены: {result.get('success', False)}")
                if result["success"]:
                    logger.info(f"✅ Интеллектуальная замена в таблице успешна")
                    return result
                else:
                    logger.warning("⚠️ Интеллектуальная замена не удалась, используем стандартную логику")
            except Exception as e:
                logger.error(f"Ошибка интеллектуальной замены в таблице: {e}")
                logger.info("Переключение на стандартную логику замены")

        # Нормализация текста для поиска (удаление лишних пробелов)
        normalized_target = " ".join(target_text.split())
        logger.debug(f"Поиск текста: '{normalized_target}' (оригинал: '{target_text}')")
        
        # Поиск всех вхождений
        matches = await self._safe_find_text(filename, normalized_target, match_case)
        
        # Если не найдено с нормализованным текстом, пробуем оригинальный
        if not matches and normalized_target != target_text:
            logger.debug(f"Повторный поиск с оригинальным текстом: '{target_text}'")
            matches = await self._safe_find_text(filename, target_text, match_case)

        if not matches:
            # Попытка найти похожий текст (для пунктов типа "36." или "36)")
            if target_text.isdigit() or (target_text.replace(".", "").replace(")", "").isdigit()):
                # Пробуем найти пункт с разными форматами
                for variant in [f"{target_text}.", f"{target_text})", f"{target_text}."]:
                    logger.debug(f"Попытка найти вариант: '{variant}'")
                    variant_matches = await self._safe_find_text(filename, variant, match_case=False)
                    if variant_matches:
                        matches = variant_matches
                        logger.info(f"Найдено совпадение для варианта '{variant}'")
                        break
            
            if not matches:
                # НОВЫЙ ФУНКЦИОНАЛ: Расширенный поиск текста с различными вариантами
                logger.info(f"🔍 Расширенный поиск текста '{target_text}' с различными вариантами...")
                matches = await self._enhanced_text_search(filename, target_text, description, match_case)
                
            if not matches:
                logger.warning(f"Текст '{target_text}' не найден в документе")
                return {
                    "success": False,
                    "error": "TEXT_NOT_FOUND",
                    "message": f"Текст '{target_text}' не найден в документе. Попробуйте использовать более точный текст для поиска.",
                }

        # КРИТИЧЕСКОЕ: Используем master_doc, если он передан, иначе создаем новый
        # Это позволяет работать с одним объектом документа для всех изменений
        if master_doc is not None:
            doc = master_doc
            logger.info(f"📄 Используем единый объект документа (master_doc) для изменения")
        else:
            # Fallback для совместимости (если master_doc не передан)
            doc = Document(filename)
            logger.info(f"⚠️ master_doc не передан, создан новый объект Document() - файл будет сохранен отдельно")
        
        # ПРОВЕРКА ОБЛАСТИ ПРИМЕНЕНИЯ: локальные vs глобальные изменения
        is_global_change = self._is_global_change(description)
        logger.info(f"📍 Область применения: {'ГЛОБАЛЬНАЯ' if is_global_change else 'ЛОКАЛЬНАЯ'}")
        
        # Для массовых замен или глобальных изменений
        if replace_all or is_global_change or (len(matches) > 1 and is_global_change):
            # Используем doc (master_doc или новый)
            replaced_count = 0
            affected_paragraphs = set()

            # Проходим по всем параграфам и заменяем текст
            for idx, para in enumerate(doc.paragraphs):
                # Пробуем сначала стандартную замену
                if self._replace_in_paragraph(para, target_text, new_text):
                    replaced_count += 1
                    affected_paragraphs.add(idx)
                # НОВЫЙ ФУНКЦИОНАЛ: Если стандартная не сработала, пробуем надежную
                elif self._robust_replace_in_paragraph(para, target_text, new_text):
                    replaced_count += 1
                    affected_paragraphs.add(idx)

            # УЛУЧШЕНИЕ: Также проверяем таблицы для замен
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if self._replace_in_cell(cell, target_text, new_text):
                            replaced_count += 1

            if replaced_count == 0:
                return {
                    "success": False,
                    "error": "TEXT_NOT_FOUND_IN_PARAGRAPH",
                    "message": f"Не удалось заменить '{target_text}' в документе (проверены параграфы и таблицы)",
                }

            # КРИТИЧЕСКОЕ: НЕ сохраняем файл здесь, если используется master_doc
            # Файл будет сохранен один раз в конце после всех изменений
            if master_doc is None:
                doc.save(filename)
                logger.info(f"💾 Файл сохранен после массовой замены (master_doc не использовался)")

            # Добавляем аннотацию к первому затронутому параграфу
            if change.get("annotation", True) and affected_paragraphs:
                # Добавляем аннотации к каждому измененному параграфу, а не только к первому
                for para_idx in affected_paragraphs:
                    await self._add_annotation(
                        filename,
                        para_idx,
                        change,
                        extra=f'"{target_text}" → "{new_text}"',
                    )

            return {
                "success": True,
                "replacements_count": replaced_count,
                "affected_paragraphs": sorted(affected_paragraphs),
            }

        # ИНТЕЛЛЕКТУАЛЬНАЯ ЗАМЕНА В ПАРАГРАФАХ: Проверяем, не является ли это заменой в пункте
        if "пункте" in description and len(matches) >= 1:
            logger.info("📋 ОБНАРУЖЕНО ИЗМЕНЕНИЕ В ПУНКТЕ - используем интеллектуальный поиск")
            
            # Ищем правильный параграф и заменяем только нужную часть
            result = await self._intelligent_paragraph_replacement(filename, target_text, new_text, description, matches, changes_file=changes_file)
            if result["success"]:
                return result
            else:
                logger.warning("⚠️ Интеллектуальная замена в пункте не удалась, используем стандартную логику")

        # Для единичной замены (точное совпадение)
        if len(matches) != 1:
            # НОВЫЙ ФУНКЦИОНАЛ: Универсальная обработка множественных совпадений для локальных изменений
            logger.info(f"🔍 Найдено {len(matches)} совпадений для локального изменения, пытаемся выбрать наиболее подходящее...")
            selected_match = await self._select_best_match_for_local_change(filename, matches, target_text, description)
            
            if selected_match is not None:
                logger.info(f"✅ Выбрано наиболее подходящее совпадение (индекс параграфа: {selected_match.paragraph_index})")
                # Создаем новый список с одним выбранным совпадением
                matches = [selected_match]
            else:
                # Если не удалось выбрать, возвращаем ошибку
                return {
                    "success": False,
                    "error": "TEXT_NOT_UNIQUE",
                    "message": f"Ожидалось ровно одно совпадение, найдено: {len(matches)}. "
                               f"Используйте replace_all=true для массовых замен.",
                }

        paragraph_index = matches[0].paragraph_index
        
        # КРИТИЧЕСКОЕ ИСПРАВЛЕНИЕ: НЕ создаем Document() здесь!
        # Это прочитает файл с диска и может перезаписать изменения от предыдущих операций
        # Вместо этого сначала пробуем локальную замену (которая создаст свой Document() и сохранит)
        
        # Проверяем валидность paragraph_index через быстрое чтение файла
        try:
            check_doc = Document(filename)
            if paragraph_index >= len(check_doc.paragraphs):
                return {
                    "success": False,
                    "error": "PARAGRAPH_INDEX_OUT_OF_RANGE",
                    "message": f"Неверный индекс параграфа: {paragraph_index}",
                }
            # Получаем текст параграфа для проверки
            para_text_check = check_doc.paragraphs[paragraph_index].text
            del check_doc  # Освобождаем память
        except Exception as check_e:
            logger.warning(f"⚠️ Не удалось проверить paragraph_index: {check_e}")
        
        # Для работы с параграфом создадим doc только если локальная замена не сработает
        doc = None
        para = None
        
        # Инициализируем переменную replaced
        replaced = False
        
        # КРИТИЧЕСКОЕ ИЗМЕНЕНИЕ: Если используется master_doc, работаем напрямую с ним
        # НЕ используем локальную замену, которая сохраняет файл - файл будет сохранен в конце
        if master_doc is not None:
            logger.info(f"📄 Работа с единым объектом документа (master_doc) - замена напрямую в памяти")
            # Не вызываем локальную замену, которая сохраняет файл
            local_replaced_first = False
        else:
            # Fallback: если master_doc не передан, используем старую логику с локальной заменой
            logger.info(f"🔄 Fallback: пробуем локальную замену (master_doc не передан)")
            local_replaced_first = mcp_client._replace_text_locally_with_tables(
                filename, target_text, new_text, paragraph_index
            )
        
        if local_replaced_first:
            # Проверяем результат локальной замены
            # КРИТИЧЕСКОЕ: Если используется master_doc, локальная замена не вызывается
            # Если же вызвана (fallback режим), проверяем файл с диска
            if master_doc is None:
                verify_doc_local = Document(filename)
                verify_success_local = False
                
                if paragraph_index is not None and paragraph_index >= 0 and paragraph_index < len(verify_doc_local.paragraphs):
                    verify_para_text_local = verify_doc_local.paragraphs[paragraph_index].text
                    if new_text in verify_para_text_local or target_text not in verify_para_text_local:
                        verify_success_local = True
                        replaced = True
                        logger.info(f"✅ Локальная замена выполнена успешно и подтверждена в параграфе {paragraph_index}")
                    else:
                        logger.warning(f"⚠️ Локальная замена вернула успех, но текст не найден в параграфе {paragraph_index}")
                else:
                    # Проверяем по всему документу
                    all_text_local = "\n".join([p.text for p in verify_doc_local.paragraphs])
                    if new_text in all_text_local or target_text not in all_text_local:
                        verify_success_local = True
                        replaced = True
                        logger.info(f"✅ Локальная замена выполнена успешно и подтверждена (по всему документу)")
                    else:
                        logger.warning(f"⚠️ Локальная замена вернула успех, но текст не найден в документе")
                
                del verify_doc_local  # Освобождаем память
        
        # Если master_doc используется, работаем с ним напрямую
        # Если нет - создаем Document() только если локальная замена не сработала
        if not replaced:
            if master_doc is not None:
                logger.info(f"📄 Используем master_doc для работы с параграфом")
                doc = master_doc
            else:
                logger.info(f"🔄 Локальная замена не сработала, создаем Document() для работы с параграфом")
                doc = Document(filename)
            para = doc.paragraphs[paragraph_index]
        
        # Инициализируем переменные для специальной обработки заголовков (до их использования)
        is_heading = False
        is_heading_by_description = False
        description_lower = description.lower() if not replaced else ""
        
        # НОВЫЙ ФУНКЦИОНАЛ: Специальная обработка для заголовков/разделов (только если локальная не сработала)
        if not replaced:
            is_heading = self._is_heading(para)
            # Проверяем описание на наличие явных указаний на замену заголовка/наименования раздела
            # Важно: применяем специальную обработку только если явно указано, что нужно изменить заголовок/наименование
            description_lower = description.lower()
        
        # Проверяем явные указания на замену заголовка/наименования
        explicit_heading_keywords = [
            "заголовок главы", "заголовок раздела",  # Явные указания на заголовок
            "наименование раздела", "наименование главы",  # Явные указания на наименование
            "название раздела", "название главы",  # Явные указания на название
            "изменить заголовок", "заменить заголовок",  # Изменение заголовка
            "изменить наименование", "заменить наименование",  # Изменение наименования
            "изменить название", "заменить название",  # Изменение названия
            "изложить заголовок", "изложить наименование",  # Изложение заголовка/наименования
        ]
        
        # Проверяем, есть ли явное указание на заголовок/наименование
        has_explicit_heading_indication = any(
            keyword in description_lower 
            for keyword in explicit_heading_keywords
        )
        
        # Проверяем, что описание НЕ говорит о замене внутри главы/раздела (исключаем такие случаи)
        has_internal_replacement_indication = any(
            exclusion in description_lower 
            for exclusion in [
                "в главе", "в разделе", "в пункте",  # Эти фразы указывают на замену внутри, а не заголовка
                "текст в главе", "текст в разделе", "текст в",  # Явно не заголовок
                "строку в главе", "строку в разделе", "строку в",  # Замена строки внутри
                "слова в главе", "слова в разделе", "слова в",  # Замена слов внутри
            ]
        )
        
        # Специальная обработка применяется только если:
        # 1. Есть явное указание на заголовок/наименование И
        # 2. НЕТ указания на замену внутри главы/раздела
        is_heading_by_description = has_explicit_heading_indication and not has_internal_replacement_indication
        
        if is_heading or is_heading_by_description:
            logger.info(f"📌 Обнаружен заголовок/раздел (стиль: {para.style.name if para.style else 'N/A'}, по описанию: {is_heading_by_description}), используем специальную обработку")
            # Для заголовков используем прямую замену через paragraph.text
            # Это сохранит стиль заголовка, но пересоздаст runs
            try:
                para_text = para.text
                # Проверяем точное совпадение или совпадение без учета регистра
                text_found = target_text in para_text
                if not text_found and match_case:
                    # Пробуем без учета регистра
                    text_found = target_text.lower() in para_text.lower()
                
                if text_found:
                    # Сохраняем стиль заголовка
                    heading_style = para.style
                    # Выполняем замену (с учетом регистра или без)
                    if target_text in para_text:
                        new_para_text = para_text.replace(target_text, new_text, 1)
                    else:
                        # Замена без учета регистра
                        pattern = re.escape(target_text)
                        new_para_text = re.sub(pattern, new_text, para_text, count=1, flags=re.IGNORECASE)
                    
                    para.text = new_para_text
                    # Восстанавливаем стиль заголовка (на случай, если он был потерян)
                    if heading_style:
                        para.style = heading_style
                    
                    # Проверяем результат
                    if new_text in para.text:
                        logger.info(f"✅ Замена в заголовке выполнена успешно")
                        replaced = True
                        # КРИТИЧЕСКОЕ: НЕ сохраняем файл здесь, если используется master_doc
                        # Файл будет сохранен один раз в конце после всех изменений
                        if master_doc is None:
                            doc.save(filename)
                            logger.info(f"💾 Файл сохранен после замены в заголовке (master_doc не использовался)")
                        
                        # НОВЫЙ ФУНКЦИОНАЛ: Синхронизация с содержанием (оглавлением)
                        # Передаем master_doc для синхронизации, чтобы не создавать новый Document()
                        await self._sync_heading_with_table_of_contents(
                            filename, target_text, new_text, is_heading_change=True, master_doc=master_doc
                        )
                        
                        # Проверяем результат в master_doc (если используется) или в файле (если нет)
                        if master_doc is not None:
                            # Проверяем напрямую в master_doc
                            if paragraph_index < len(master_doc.paragraphs):
                                verify_para = master_doc.paragraphs[paragraph_index]
                                if new_text in verify_para.text:
                                    logger.info(f"✅ Замена в заголовке подтверждена в master_doc")
                        elif master_doc is None:
                            # Fallback: проверяем в файле
                            try:
                                verify_doc = Document(filename)
                                if paragraph_index < len(verify_doc.paragraphs):
                                    verify_para = verify_doc.paragraphs[paragraph_index]
                                    if new_text in verify_para.text:
                                        logger.info(f"✅ Замена в заголовке подтверждена после сохранения")
                                    else:
                                        logger.warning(f"⚠️ Замена в заголовке не обнаружена после сохранения, пробуем еще раз")
                                        # Пробуем еще раз прямую замену
                                        if target_text in verify_para.text:
                                            verify_para.text = verify_para.text.replace(target_text, new_text, 1)
                                            if verify_para.style:
                                                verify_para.style = heading_style
                                            verify_doc.save(filename)
                                            logger.info(f"✅ Повторная замена в заголовке выполнена")
                            except Exception as verify_e:
                                logger.warning(f"⚠️ Не удалось проверить замену в заголовке после сохранения: {verify_e}")
                    else:
                        logger.warning(f"⚠️ Замена в заголовке не подтверждена")
                else:
                    logger.warning(f"⚠️ Текст '{target_text}' не найден в заголовке '{para_text}'")
            except Exception as e:
                logger.error(f"❌ Ошибка при замене в заголовке: {e}", exc_info=True)
        
        if not replaced:
            replaced = self._replace_in_paragraph(para, target_text, new_text)

        if not replaced:
            # НОВЫЙ ФУНКЦИОНАЛ: Попытка надежной замены для текста, разбитого на runs
            logger.info(f"🔍 Стандартная замена не удалась, пробуем надежную замену для текста, разбитого на runs")
            replaced = self._robust_replace_in_paragraph(para, target_text, new_text)
            
            if replaced:
                logger.info(f"✅ Надежная замена выполнена успешно")
        
        # НОВЫЙ ФУНКЦИОНАЛ: Попытка с нормализацией пробелов, если замена не удалась
        if not replaced:
            logger.info(f"🔍 Попытка замены с нормализацией пробелов")
            normalized_target = " ".join(target_text.split())
            para_text = para.text
            normalized_para_text = " ".join(para_text.split())
            
            if normalized_target in normalized_para_text and normalized_target != target_text:
                # Находим позицию в нормализованном тексте
                norm_pos = normalized_para_text.find(normalized_target)
                # Пробуем заменить в оригинальном тексте
                if normalized_target in para_text:
                    try:
                        para.text = para_text.replace(normalized_target, new_text, 1)
                        # Проверяем, что замена действительно произошла
                        if new_text in para.text:
                            logger.info(f"✅ Замена выполнена с нормализацией пробелов")
                            replaced = True
                    except Exception as e:
                        logger.warning(f"⚠️ Ошибка при замене с нормализацией: {e}")
        
        # НОВЫЙ ФУНКЦИОНАЛ: Последняя попытка - прямая замена через очистку runs
        if not replaced and doc is not None and para is not None and target_text in para.text:
            logger.info(f"🔍 Последняя попытка: прямая замена через очистку всех runs")
            try:
                para_text = para.text
                new_para_text = para_text.replace(target_text, new_text, 1)
                
                # Очищаем все runs
                for run in para.runs:
                    run.text = ""
                
                # Если есть хотя бы один run, записываем в него
                if para.runs:
                    para.runs[0].text = new_para_text
                else:
                    # Если нет runs, создаем новый параграф (но это не должно случиться)
                    para.add_run(new_para_text)
                
                # Проверяем результат
                if new_text in para.text:
                    logger.info(f"✅ Замена выполнена через очистку runs")
                    replaced = True
            except Exception as e:
                logger.error(f"❌ Ошибка при замене через очистку runs: {e}")
        
        if not replaced and doc is not None:
            # НОВЫЙ ФУНКЦИОНАЛ: Расширенная попытка замены в других параграфах
            logger.info(f"🔍 Расширенный поиск для замены текста '{target_text}'")
            replaced = await self._enhanced_replace_attempt(doc, target_text, new_text, paragraph_index)
            
            if not replaced:
                # Пробуем найти в других параграфах (старая логика)
                for para_other in doc.paragraphs:
                    if self._replace_in_paragraph(para_other, target_text, new_text):
                        replaced = True
                        break
                    
                    # НОВЫЙ ФУНКЦИОНАЛ: Также пробуем надежную замену в других параграфах
                    if self._robust_replace_in_paragraph(para_other, target_text, new_text):
                        replaced = True
                        break
        
        # НОВЫЙ ФУНКЦИОНАЛ: Попытка замены без учета регистра, если match_case=True и замена не удалась
        if not replaced and doc is not None and para is not None and match_case:
            logger.info(f"🔧 Попытка замены без учета регистра (match_case=True, но замена не удалась)")
            try:
                para_text = para.text
                # Ищем текст без учета регистра
                if target_text.lower() in para_text.lower():
                    # Находим точное вхождение (с учетом регистра как в документе)
                    pattern = re.escape(target_text)
                    match = re.search(pattern, para_text, re.IGNORECASE)
                    if match:
                        actual_text = match.group(0)  # Реальный текст из документа
                        new_para_text = para_text.replace(actual_text, new_text, 1)
                        para.text = new_para_text
                        if new_text in para.text:
                            logger.info(f"✅ Замена выполнена без учета регистра")
                            replaced = True
            except Exception as e:
                logger.warning(f"⚠️ Ошибка при замене без учета регистра: {e}")
        
        # КРИТИЧЕСКОЕ ИСПРАВЛЕНИЕ: Если локальная замена в начале не сработала и мы попробовали другие методы,
        # пробуем еще раз локальную замену и MCP как финальные попытки
        if not replaced:
            logger.info(f"🔧 Финальная попытка: сначала локальная замена (гарантированно сохраняет), затем MCP")
            try:
                # ШАГ 1: Пробуем локальную замену СНАЧАЛА (которая гарантированно сохраняет файл)
                logger.info(f"🔄 ШАГ 1: Локальная замена для '{target_text}' → '{new_text}' (параграф {paragraph_index})")
                local_replaced_first = mcp_client._replace_text_locally_with_tables(
                    filename, target_text, new_text, paragraph_index
                )
                
                if local_replaced_first:
                    # Проверяем результат локальной замены
                    verify_doc_local = Document(filename)
                    verify_success_local = False
                    
                    if paragraph_index is not None and paragraph_index >= 0 and paragraph_index < len(verify_doc_local.paragraphs):
                        verify_para_text_local = verify_doc_local.paragraphs[paragraph_index].text
                        if new_text in verify_para_text_local or target_text not in verify_para_text_local:
                            verify_success_local = True
                            replaced = True
                            logger.info(f"✅ Локальная замена выполнена успешно и подтверждена в параграфе {paragraph_index}")
                        else:
                            logger.warning(f"⚠️ Локальная замена вернула успех, но текст не найден в параграфе {paragraph_index}")
                    else:
                        # Проверяем по всему документу
                        all_text_local = "\n".join([p.text for p in verify_doc_local.paragraphs])
                        if new_text in all_text_local or target_text not in all_text_local:
                            verify_success_local = True
                            replaced = True
                            logger.info(f"✅ Локальная замена выполнена успешно и подтверждена (по всему документу)")
                        else:
                            logger.warning(f"⚠️ Локальная замена вернула успех, но текст не найден в документе")
                
                # ШАГ 2: Если локальная замена не сработала, пробуем MCP replace_text
                if not replaced:
                    logger.info(f"🔄 ШАГ 2: Локальная замена не сработала, пробуем MCP replace_text")
                    # КРИТИЧЕСКОЕ: НЕ сохраняем документ перед MCP - это сохранит старую версию!
                    
                    # Используем MCP replace_text для замены в указанном параграфе
                    mcp_replaced = await mcp_client.replace_text(
                        filename=filename,
                        old_text=target_text,
                        new_text=new_text,
                        paragraph_index=paragraph_index
                    )
                
                    logger.info(f"📊 Результат MCP replace_text: {mcp_replaced} (параграф {paragraph_index})")
                    
                    if mcp_replaced:
                        logger.info(f"✅ MCP replace_text вернул успех")
                        # КРИТИЧЕСКОЕ: После MCP replace_text ВСЕГДА проверяем результат
                        # MCP может вернуть успех, но не сохранить файл!
                        try:
                            verify_doc = Document(filename)
                            verify_success = False
                            
                            if paragraph_index is not None and paragraph_index >= 0 and paragraph_index < len(verify_doc.paragraphs):
                                verify_para_text = verify_doc.paragraphs[paragraph_index].text
                                # Проверяем: новый текст присутствует ИЛИ старый текст отсутствует
                                if new_text in verify_para_text or target_text not in verify_para_text:
                                    replaced = True
                                    verify_success = True
                                    logger.info(f"✅ Замена подтверждена после MCP replace_text в параграфе {paragraph_index}")
                                else:
                                    logger.warning(f"⚠️ MCP replace_text вернул успех, но замена не обнаружена в параграфе {paragraph_index}")
                                    logger.info(f"   Параграф {paragraph_index}: старый текст найден={target_text[:50] in verify_para_text}, новый текст найден={new_text[:50] in verify_para_text}")
                            else:
                                # Проверяем по всему документу
                                all_text = "\n".join([p.text for p in verify_doc.paragraphs])
                                old_found = target_text in all_text
                                new_found = new_text in all_text
                                
                                if new_found or not old_found:
                                    replaced = True
                                    verify_success = True
                                    logger.info(f"✅ Замена подтверждена после MCP replace_text (по всему документу)")
                                else:
                                    logger.warning(f"⚠️ MCP replace_text вернул успех, но замена не обнаружена (старый текст найден, новый отсутствует)")
                            
                            # КРИТИЧЕСКОЕ: Если верификация не прошла, MCP не сохранил файл - делаем локальную замену
                            if not verify_success:
                                logger.warning(f"🔄 MCP вернул успех, но верификация не прошла - MCP не сохранил файл. Делаем локальную замену для гарантии сохранения")
                                local_replaced_after = mcp_client._replace_text_locally_with_tables(
                                    filename, target_text, new_text, paragraph_index
                                )
                                
                                if local_replaced_after:
                                    # Повторно проверяем результат локальной замены
                                    verify_doc_after = Document(filename)
                                    verify_success_after = False
                                    
                                    if paragraph_index is not None and paragraph_index >= 0 and paragraph_index < len(verify_doc_after.paragraphs):
                                        verify_para_text_after = verify_doc_after.paragraphs[paragraph_index].text
                                        if new_text in verify_para_text_after or target_text not in verify_para_text_after:
                                            verify_success_after = True
                                    else:
                                        all_text_after = "\n".join([p.text for p in verify_doc_after.paragraphs])
                                        if new_text in all_text_after or target_text not in all_text_after:
                                            verify_success_after = True
                                    
                                    if verify_success_after:
                                        replaced = True
                                        logger.info(f"✅ Локальная замена выполнена успешно после неудачной верификации MCP")
                                    else:
                                        logger.warning(f"⚠️ Локальная замена вернула успех, но верификация не прошла")
                                else:
                                    logger.warning(f"⚠️ Локальная замена также не удалась после MCP")
                        except Exception as verify_e:
                            logger.warning(f"⚠️ Не удалось проверить результат MCP replace_text: {verify_e}")
                            # Если не удалось проверить, делаем локальную замену для гарантии сохранения
                            logger.info(f"🔄 Делаем локальную замену как fallback после ошибки верификации MCP")
                            local_replaced_fallback = mcp_client._replace_text_locally_with_tables(
                                filename, target_text, new_text, paragraph_index
                            )
                            if local_replaced_fallback:
                                replaced = True
                                logger.info(f"✅ Локальная замена выполнена как fallback после ошибки верификации MCP")
                    else:
                        logger.warning(f"⚠️ MCP replace_text не удалась с paragraph_index={paragraph_index}")
                        # Локальная замена уже была попробована в ШАГ 1, больше ничего не делаем
            except Exception as e:
                logger.error(f"❌ Ошибка при использовании MCP replace_text: {e}", exc_info=True)

        # НОВЫЙ ФУНКЦИОНАЛ: Проверка, находится ли текст в содержании (оглавлении)
        # Если текст найден в таблице и это похоже на содержание, синхронизируем с заголовком
        is_in_table_of_contents = False
        if not replaced and doc is not None:
            # Проверяем, может быть текст находится в содержании
            try:
                # Ищем текст в таблицах
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if target_text in cell_text or cell_text in target_text:
                                # Проверяем, похоже ли это на содержание (есть номер страницы или точки)
                                # Содержание обычно имеет формат: "1. Название ........ 5"
                                if re.search(r'[. ]+\d+$', cell_text) or re.match(r'^\d+\.', cell_text):
                                    is_in_table_of_contents = True
                                    logger.info(f"📋 Обнаружено, что текст находится в содержании (оглавлении)")
                                    # Выполняем замену в ячейке
                                    if target_text in cell_text:
                                        new_cell_text = cell_text.replace(target_text, new_text, 1)
                                    else:
                                        # Ячейка содержит только часть, заменяем с сохранением форматирования
                                        page_match = re.search(r'([. ]+)(\d+)$', cell_text)
                                        if page_match:
                                            separator = page_match.group(1)
                                            page_num = page_match.group(2)
                                            heading_num_match = re.match(r'^(\d+\.?\s*)', cell_text)
                                            if heading_num_match:
                                                heading_num = heading_num_match.group(1)
                                                new_cell_text = heading_num + new_text.replace(heading_num, '').strip() + separator + page_num
                                            else:
                                                new_cell_text = new_text + separator + page_num
                                        else:
                                            heading_num_match = re.match(r'^(\d+\.?\s*)', cell_text)
                                            if heading_num_match:
                                                heading_num = heading_num_match.group(1)
                                                new_cell_text = heading_num + new_text.replace(heading_num, '').strip() if heading_num in new_text else heading_num + new_text
                                            else:
                                                new_cell_text = new_text
                                    
                                    cell.text = new_cell_text
                                    replaced = True
                                    doc.save(filename)
                                    logger.info(f"✅ Замена в содержании выполнена: '{cell_text}' → '{new_cell_text}'")
                                    
                                    # Синхронизируем с заголовком раздела
                                    await self._sync_heading_with_table_of_contents(filename, target_text, new_text, is_heading_change=False)
                                    break
                        if is_in_table_of_contents:
                            break
                    if is_in_table_of_contents:
                        break
            except Exception as toc_e:
                logger.warning(f"⚠️ Ошибка при проверке содержания: {toc_e}")
        
        # НОВЫЙ ФУНКЦИОНАЛ: Финальная проверка замены перед сохранением
        if not replaced and doc is not None and para is not None:
            # Последняя попытка: проверяем, может быть текст уже заменен (например, через paragraph.text)
            final_para_text = para.text
            if new_text in final_para_text and target_text not in final_para_text:
                logger.info(f"✅ Текст уже заменен (возможно, через paragraph.text)")
                replaced = True
            else:
                # Еще одна попытка: прямое использование paragraph.text для полной замены
                logger.info(f"🔧 Последняя попытка: прямая замена через paragraph.text")
                try:
                    if target_text in final_para_text:
                        para.text = final_para_text.replace(target_text, new_text, 1)
                        # Проверяем сразу после замены
                        if new_text in para.text:
                            logger.info(f"✅ Замена выполнена через прямое присваивание paragraph.text")
                            replaced = True
                except Exception as e:
                    logger.error(f"❌ Ошибка при прямой замене через paragraph.text: {e}")
        
        # Если все еще не заменено, возвращаем ошибку
        if not replaced:
            return {
                "success": False,
                "error": "TEXT_NOT_FOUND_IN_PARAGRAPH",
                "message": f"Не удалось заменить '{target_text}' в найденном параграфе. Испробованы все методы замены.",
            }

        # КРИТИЧЕСКОЕ: НЕ сохраняем файл здесь, если используется master_doc
        # Файл будет сохранен один раз в конце после всех изменений в process_documents
        # Сохраняем только если master_doc не используется (fallback режим)
        if doc is not None and master_doc is None:
            doc.save(filename)
            logger.info(f"💾 Документ сохранен после замены через Document() (fallback режим, master_doc не использовался)")
        elif replaced and master_doc is None:
            # Если локальная замена сработала (replaced = True) в fallback режиме,
            # файл уже сохранен локальной заменой - ничего не делаем
            logger.info(f"💾 Файл уже сохранен локальной заменой (fallback режим)")
        
        # Финальная проверка результата (только если master_doc не используется)
        # Если используется master_doc, проверяем напрямую в нем
        if master_doc is not None:
            # Проверяем напрямую в master_doc
            try:
                if paragraph_index < len(master_doc.paragraphs):
                    verify_para_text = master_doc.paragraphs[paragraph_index].text
                    if new_text in verify_para_text or target_text not in verify_para_text:
                        logger.info(f"✅ Замена подтверждена в master_doc")
                    else:
                        logger.warning(f"⚠️ Замена не обнаружена в master_doc после всех попыток")
            except Exception as e:
                logger.warning(f"⚠️ Не удалось проверить замену в master_doc: {e}")
        elif master_doc is None:
            # Fallback: проверяем в файле
            try:
                verify_doc = Document(filename)
                if paragraph_index < len(verify_doc.paragraphs):
                    verify_para_text = verify_doc.paragraphs[paragraph_index].text
                    if new_text not in verify_para_text and target_text in verify_para_text:
                        logger.warning(f"⚠️ Замена не обнаружена после сохранения, пробуем еще раз...")
                        # Пробуем еще раз прямую замену
                        verify_para = verify_doc.paragraphs[paragraph_index]
                        if target_text in verify_para.text:
                            verify_para.text = verify_para.text.replace(target_text, new_text, 1)
                            verify_doc.save(filename)
                            logger.info(f"✅ Повторная замена выполнена успешно")
                    elif new_text in verify_para_text:
                        logger.info(f"✅ Замена подтверждена после сохранения")
            except Exception as e:
                logger.warning(f"⚠️ Не удалось проверить замену после сохранения: {e}")

        if change.get("annotation", True):
            await self._add_annotation(
                filename,
                paragraph_index,
                change,
                extra=f'"{target_text}" → "{new_text}"',
            )

        return {"success": True, "paragraph_index": paragraph_index}

    async def _intelligent_table_replacement(self, filename: str, target_text: str, new_text: str, description: str) -> Dict[str, Any]:
        """
        УЛУЧШЕННАЯ интеллектуальная замена в таблице:
        1. Находит строку с target_text
        2. Анализирует структуру строки (количество столбцов, содержимое)
        3. Сопоставляет новый текст с существующей структурой
        4. Правильно распределяет по столбцам
        
        Args:
            filename: Путь к файлу
            target_text: Искомый текст (например, "ДРМ")
            new_text: Новый текст (например, "ДКР Департамент кредитных рисков")
            description: Описание инструкции
            
        Returns:
            Результат операции
        """
        logger.info(f"🧠 УЛУЧШЕННАЯ ИНТЕЛЛЕКТУАЛЬНАЯ ЗАМЕНА В ТАБЛИЦЕ:")
        logger.info(f"   Ищем строку с: '{target_text}'")
        logger.info(f"   Новое содержимое: '{new_text}'")
        logger.info(f"   Описание: '{description}'")
        
        try:
            # Открываем документ для анализа
            doc = Document(filename)
            replacements_made = 0
            
            # Извлекаем название таблицы из описания (если указано)
            table_name = None
            if "таблице" in description.lower():
                # Ищем паттерн "в таблице «название»" или "таблице 'название'"
                patterns = [
                    r'таблице\s*[«"](.*?)[»"]',  # таблице «название»
                    r'таблице\s*[\']([^\']+)[\']',  # таблице 'название'
                    r'таблиц[еи]\s+[«"](.*?)[»"]',  # таблице/таблицы «название»
                    r'таблиц[еи]\s+[\']([^\']+)[\']',  # таблице/таблицы 'название'
                ]
                for pattern in patterns:
                    match = re.search(pattern, description, re.IGNORECASE)
                    if match:
                        table_name = match.group(1).strip()
                        logger.info(f"📋 Извлечено название таблицы из описания: '{table_name}' (паттерн: {pattern})")
                        break
                
                if not table_name:
                    logger.warning(f"⚠️ Не удалось извлечь название таблицы из описания: '{description[:100]}...'")
            
            # ИНТЕЛЛЕКТУАЛЬНОЕ ОПРЕДЕЛЕНИЕ ЦЕЛЕВОЙ ТАБЛИЦЫ ЧЕРЕЗ LLM
            llm_target_table_indices = None
            try:
                llm_target_table_indices = await self._identify_target_table_with_llm(
                    doc=doc,
                    description=description,
                    target_text=target_text,
                    table_name=table_name
                )
                if llm_target_table_indices:
                    logger.info(f"   🎯 LLM определил целевые таблицы: {llm_target_table_indices}")
            except Exception as e:
                logger.warning(f"   ⚠️ Ошибка при LLM определении таблицы, продолжаем с алгоритмическим подходом: {e}")
            
            # Если указано название таблицы, сначала ищем его в тексте документа
            table_name_found_in_text = False
            table_name_paragraph_index = -1  # Позиция названия таблицы в документе
            if table_name:
                logger.info(f"🔍 Поиск названия таблицы '{table_name}' в тексте документа...")
                try:
                    # Ищем название таблицы в тексте через MCP
                    matches = await mcp_client.find_text_in_document(filename, table_name, match_case=False)
                    if matches:
                        # Берем первое вхождение названия таблицы (обычно это заголовок таблицы)
                        first_match = matches[0]
                        if hasattr(first_match, 'paragraph_index'):
                            table_name_paragraph_index = first_match.paragraph_index
                        elif isinstance(first_match, dict):
                            table_name_paragraph_index = first_match.get('paragraph_index', -1)
                        
                        logger.info(f"   ✅ Найдено {len(matches)} упоминаний названия таблицы в тексте")
                        logger.info(f"   📍 Название найдено в параграфе {table_name_paragraph_index}")
                        table_name_found_in_text = True
                        logger.info(f"   📍 Название найдено в тексте, ищем таблицы после этого параграфа")
                    else:
                        logger.info(f"   ⚠️ Название таблицы не найдено в тексте, используем проверку заголовка")
                    
                except Exception as e:
                    logger.warning(f"   ⚠️ Ошибка при поиске названия таблицы в тексте: {e}")
                    # Продолжаем без ограничения
            
            # Ищем таблицы и анализируем их структуру
            # Если LLM определил целевые таблицы, используем их как приоритет
            # Если название найдено в тексте, обрабатываем только первую таблицу с target_text, которая идет ПОСЛЕ названия
            first_table_processed = False
            table_location = None  # Информация о местоположении замены для аннотаций
            
            # НОВЫЙ ФУНКЦИОНАЛ: Предварительный отбор таблиц по позиции относительно названия
            # Если название таблицы найдено в тексте, собираем все таблицы, которые идут после названия и содержат target_text
            candidate_tables = []  # Список кандидатов: (table_idx, table_paragraph_index, contains_target)
            if table_name_found_in_text and table_name_paragraph_index >= 0:
                logger.info(f"   🔍 Предварительный отбор таблиц по позиции (название в параграфе {table_name_paragraph_index})...")
                for table_idx, table in enumerate(doc.tables):
                    table_paragraph_index = self._find_paragraph_for_table(doc, table_idx, after_table=False)
                    
                    # Проверяем, что таблица идет ПОСЛЕ названия
                    if table_paragraph_index >= table_name_paragraph_index:
                        # Проверяем, содержит ли таблица target_text
                        contains_target = False
                        for row in table.rows:
                            for cell in row.cells:
                                if target_text in cell.text:
                                    contains_target = True
                                    break
                            if contains_target:
                                break
                        
                        if contains_target:
                            candidate_tables.append((table_idx, table_paragraph_index, True))
                            logger.info(f"   ✅ Таблица {table_idx} - кандидат (после названия, содержит target_text, параграф {table_paragraph_index})")
                        else:
                            logger.info(f"   ⏭️ Таблица {table_idx} - пропущена (после названия, но не содержит target_text)")
                    else:
                        logger.info(f"   ⏭️ Таблица {table_idx} - пропущена (перед названием, параграф {table_paragraph_index})")
                
                # Если нашли кандидатов, выбираем первую (ближайшую к названию)
                if candidate_tables:
                    # Сортируем по позиции (ближайшая к названию)
                    candidate_tables.sort(key=lambda x: x[1])
                    best_table_idx = candidate_tables[0][0]
                    logger.info(f"   🎯 Выбрана таблица {best_table_idx} как наиболее подходящая (ближайшая к названию)")
            
            for table_idx, table in enumerate(doc.tables):
                logger.info(f"📊 Анализ таблицы {table_idx}")
                
                should_process_this_table = True
                
                # ПРИОРИТЕТ 0: Если есть предварительно отобранные кандидаты, используем их
                if candidate_tables:
                    candidate_indices = [t[0] for t in candidate_tables]
                    if table_idx not in candidate_indices:
                        logger.info(f"   ⏭️ Пропускаем таблицу {table_idx} (не прошла предварительный отбор по позиции)")
                        continue
                    else:
                        logger.info(f"   ✅ Таблица {table_idx} прошла предварительный отбор по позиции")
                
                # ПРИОРИТЕТ 1: Если LLM определил целевые таблицы, обрабатываем только их (если нет предварительного отбора)
                if llm_target_table_indices is not None and not candidate_tables:
                    if table_idx not in llm_target_table_indices:
                        logger.info(f"   ⏭️ Пропускаем таблицу {table_idx} (не определена LLM как целевая)")
                        continue
                    else:
                        logger.info(f"   ✅ Таблица {table_idx} определена LLM как целевая")
                
                # ПРИОРИТЕТ 2: Проверяем название таблицы (если указано и LLM не определил таблицы)
                # Если LLM определил таблицы, используем их и пропускаем проверку по названию
                if table_name and llm_target_table_indices is None:
                    # Если название найдено в тексте, проверяем только таблицы, которые идут ПОСЛЕ названия
                    if table_name_found_in_text:
                        # Если уже обработали первую подходящую таблицу, пропускаем остальные
                        if first_table_processed:
                            logger.info(f"   ⏭️ Пропускаем таблицу {table_idx} (уже обработана первая таблица с названием)")
                            continue
                        
                        # НОВЫЙ ФУНКЦИОНАЛ: Проверяем, что таблица идет ПОСЛЕ названия в документе
                        if table_name_paragraph_index >= 0:
                            table_paragraph_index = self._find_paragraph_for_table(doc, table_idx, after_table=False)
                            
                            if table_paragraph_index < table_name_paragraph_index:
                                logger.info(f"   ⏭️ Пропускаем таблицу {table_idx} (находится ДО названия в параграфе {table_name_paragraph_index}, таблица в параграфе {table_paragraph_index})")
                                continue
                            else:
                                logger.info(f"   ✅ Таблица {table_idx} идет ПОСЛЕ названия (название в {table_name_paragraph_index}, таблица в {table_paragraph_index})")
                        
                        # Проверяем, содержит ли эта таблица target_text
                        table_contains_target = False
                        for row in table.rows:
                            for cell in row.cells:
                                if target_text in cell.text:
                                    table_contains_target = True
                                    break
                            if table_contains_target:
                                break
                        
                        if not table_contains_target:
                            logger.info(f"   ⏭️ Пропускаем таблицу {table_idx} (не содержит target_text '{target_text}')")
                            should_process_this_table = False
                        else:
                            logger.info(f"   ✅ Таблица {table_idx} содержит target_text и идет после названия в тексте")
                            # first_table_processed будет установлен после успешной замены
                    
                    # Если не нашли через поиск в тексте, проверяем заголовок таблицы
                    if not table_name_found_in_text and should_process_this_table:
                        # Проверяем заголовок таблицы (первые 3 строки для более надежного поиска)
                        table_header = ""
                        for i, row in enumerate(table.rows[:3]):
                            for cell in row.cells:
                                table_header += cell.text + " "
                        table_header = table_header.strip().lower()
                        table_name_lower = table_name.lower()
                        
                        # Гибкая проверка: ищем ключевые слова из названия таблицы
                        # Разбиваем название на слова и проверяем, что хотя бы 2-3 ключевых слова присутствуют
                        # Убираем служебные слова
                        stop_words = {'и', 'в', 'на', 'с', 'по', 'для', 'от', 'до', 'из', 'к', 'о', 'об', 'обо', 'со', 'во'}
                        table_name_words = [w for w in re.findall(r'\b\w+\b', table_name_lower) if w not in stop_words and len(w) > 2]
                        
                        # Проверяем совпадение ключевых слов
                        matched_words = [word for word in table_name_words if word in table_header]
                        match_ratio = len(matched_words) / len(table_name_words) if table_name_words else 0
                        
                        # Также проверяем точное совпадение (на случай коротких названий)
                        exact_match = table_name_lower in table_header
                        
                        # Принимаем таблицу, если:
                        # 1. Точное совпадение ИЛИ
                        # 2. Совпало больше половины ключевых слов (минимум 2 слова)
                        if not exact_match and (match_ratio < 0.5 or len(matched_words) < 2):
                            logger.info(f"   ⏭️ Пропускаем таблицу {table_idx} (не соответствует названию '{table_name}')")
                            logger.info(f"      Заголовок: '{table_header[:100]}...'")
                            logger.info(f"      Совпало слов: {len(matched_words)}/{len(table_name_words)} ({matched_words})")
                            continue
                        else:
                            logger.info(f"   ✅ Таблица {table_idx} соответствует названию '{table_name}'")
                            logger.info(f"      Совпало слов: {len(matched_words)}/{len(table_name_words)} ({matched_words})")
                    else:
                        logger.info(f"   ✅ Таблица {table_idx} - целевая таблица (найдена через поиск в тексте)")
                
                # Пропускаем таблицу, если она не должна обрабатываться
                if not should_process_this_table:
                    continue
                
                # Если LLM определил эту таблицу как целевую, проверяем наличие target_text
                if llm_target_table_indices and table_idx in llm_target_table_indices:
                    table_contains_target = False
                    for row in table.rows:
                        for cell in row.cells:
                            if target_text in cell.text:
                                table_contains_target = True
                                break
                        if table_contains_target:
                            break
                    
                    if not table_contains_target:
                        logger.warning(f"   ⚠️ Таблица {table_idx} определена LLM как целевая, но не содержит target_text '{target_text}'")
                        logger.info(f"   🔍 Продолжаем поиск в таблице {table_idx} (возможно, target_text в другой форме или требуется более глубокий поиск)")
                        # НЕ пропускаем таблицу, если LLM определила её как целевую - продолжаем поиск
                        # Это позволяет найти target_text даже если он в другой форме
                    else:
                        logger.info(f"   ✅ Таблица {table_idx} определена LLM как целевая и содержит target_text")
                
                for row_idx, row in enumerate(table.rows):
                    # 1. НАХОДИМ СТРОКУ с target_text
                    # Сначала проверяем, если есть номер пункта в инструкции - ищем строку с этим номером
                    target_found = False
                    target_cell_idx = -1
                    
                    # Если есть номер пункта, сначала проверяем, начинается ли первая ячейка строки с этого номера
                    punkt_in_row = False
                    if punkt_number and len(row.cells) > 0:
                        first_cell_text = row.cells[0].text.strip()
                        punkt_patterns = [f"{punkt_number}.", f"{punkt_number})", f"{punkt_number}."]
                        for pattern in punkt_patterns:
                            if first_cell_text.startswith(pattern) or first_cell_text == punkt_number:
                                punkt_in_row = True
                                logger.info(f"   📋 Найдена строка {row_idx} с номером пункта {punkt_number} в первой ячейке")
                                # Если пункт найден в этой строке, ищем target_text в ячейках этой строки
                                for cell_idx, cell in enumerate(row.cells):
                                    if target_text in cell.text:
                                        target_found = True
                                        target_cell_idx = cell_idx
                                        logger.info(f"   ✅ Найдена строка {row_idx} с пунктом {punkt_number} и '{target_text}' в ячейке {cell_idx}")
                                        break
                                break
                    
                    # Если пункт не найден в строке или пункт найден но target_text не найден, 
                    # или номер пункта не указан - ищем target_text в любой ячейке строки
                    if not target_found:
                        for cell_idx, cell in enumerate(row.cells):
                            if target_text in cell.text:
                                # Если номер пункта указан, но мы его не нашли в первой ячейке, 
                                # пропускаем эту строку (строго по инструкции)
                                if punkt_number and not punkt_in_row:
                                    logger.info(f"   ⏭️ Пропускаем строку {row_idx} (target_text найден, но номер пункта {punkt_number} не совпадает)")
                                    continue
                                target_found = True
                                target_cell_idx = cell_idx
                                logger.info(f"   ✅ Найдена строка {row_idx} с '{target_text}' в ячейке {cell_idx}")
                                break
                    
                    if target_found:
                        # 2. АНАЛИЗИРУЕМ СТРУКТУРУ СТРОКИ
                        row_structure = self._analyze_row_structure(row, row_idx)
                        logger.info(f"   📋 Структура строки: {row_structure}")
                        
                        # 2.1. ПОЛУЧАЕМ КОНТЕКСТ ТАБЛИЦЫ (для LLM)
                        table_context = self._get_table_context(table, row_idx)
                        
                        # 3. СОПОСТАВЛЯЕМ НОВЫЙ ТЕКСТ СО СТРУКТУРОЙ (алгоритм + LLM проверка)
                        distribution = await self._map_new_text_to_structure(
                            new_text=new_text,
                            target_text=target_text,
                            row_structure=row_structure,
                            description=description,
                            table_context=table_context
                        )
                        logger.info(f"   🎯 Распределение по столбцам: {distribution}")
                        
                        # 4. ПРИМЕНЯЕМ ИЗМЕНЕНИЯ ПО СТОЛБЦАМ
                        if self._apply_structured_replacement(row, target_text, distribution):
                            replacements_made += 1
                            logger.info(f"   ✅ Структурированная замена в строке {row_idx}")
                            
                            # Сохраняем информацию о местоположении замены для аннотаций
                            # Ищем параграф, соответствующий этой таблице
                            if table_location is None:  # Сохраняем только информацию о первой замене
                                table_paragraph_index = self._find_paragraph_for_table(doc, table_idx)
                                if table_paragraph_index >= 0:
                                    table_location = {
                                        "table_idx": table_idx,
                                        "row_idx": row_idx,
                                        "cell_idx": target_cell_idx,
                                        "paragraph_index": table_paragraph_index
                                    }
                                    logger.info(f"   📍 Сохранено местоположение для аннотации: Table {table_idx}, Row {row_idx}, Para {table_paragraph_index}")
                            
                            # Если LLM определил точечное изменение в конкретной таблице, завершаем после первой замены
                            if llm_target_table_indices and len(llm_target_table_indices) == 1:
                                logger.info(f"   ✅ Точечное изменение выполнено в целевой таблице {table_idx}, завершаем обработку")
                                first_table_processed = True
                                # Выходим из цикла по строкам
                                break
                            # Если название найдено в тексте, обрабатываем только первую найденную строку
                            elif table_name_found_in_text:
                                logger.info(f"   ✅ Найдена целевая таблица с названием в тексте, завершаем обработку")
                                # Устанавливаем флаг, что первая таблица обработана
                                first_table_processed = True
                                # Выходим из цикла по строкам
                                break
                
                # Выходим из цикла по таблицам, если уже сделали замену и:
                # - LLM определил точечное изменение (одна таблица), или
                # - название найдено в тексте
                if (llm_target_table_indices and len(llm_target_table_indices) == 1 and first_table_processed) or (table_name_found_in_text and first_table_processed):
                    if llm_target_table_indices and len(llm_target_table_indices) == 1:
                        logger.info(f"   ✅ Завершаем обработку таблиц (точечное изменение в таблице {llm_target_table_indices[0]} выполнено)")
                    elif table_name_found_in_text:
                        logger.info(f"   ✅ Завершаем обработку таблиц (найдена целевая таблица с названием и выполнена замена)")
                    break
            
            if replacements_made > 0:
                doc.save(filename)
                result = {
                    "success": True,
                    "message": f"Структурированная замена выполнена в {replacements_made} строках",
                    "replacements_made": replacements_made,
                    "method": "structured_table_replace",
                    "is_table_change": True,  # Флаг, что изменение было в таблице
                }
                # Добавляем информацию о местоположении для аннотаций
                if table_location:
                    result["table_location"] = table_location
                    result["paragraph_index"] = table_location.get("paragraph_index", -1)
                return result
            else:
                return {
                    "success": False,
                    "error": "NO_REPLACEMENTS",
                    "message": f"Строка с '{target_text}' не найдена в таблицах"
                }
                
        except Exception as e:
            logger.error(f"Ошибка структурированной замены в таблице: {e}")
            return {
                "success": False,
                "error": "STRUCTURED_REPLACE_ERROR",
                "message": f"Ошибка структурированной замены: {e}"
            }

    def _get_text_before_table(self, doc: Document, table_idx: int, max_paragraphs: int = 3) -> str:
        """
        Получает текст параграфов перед указанной таблицей (для поиска названия таблицы).
        
        Args:
            doc: Документ
            table_idx: Индекс таблицы в документе
            max_paragraphs: Максимальное количество параграфов перед таблицей для анализа
            
        Returns:
            Текст параграфов перед таблицей
        """
        try:
            table_count = 0
            para_count = 0
            paragraphs_before = []
            
            # Проходим по элементам документа
            for i, element in enumerate(doc.element.body):
                if element.tag.endswith('p'):  # Параграф
                    para_count += 1
                elif element.tag.endswith('tbl'):  # Таблица
                    if table_count == table_idx:
                        # Нашли нужную таблицу, собираем параграфы перед ней
                        # Ищем параграфы перед этим элементом (до max_paragraphs)
                        found_paragraphs = 0
                        for j in range(i - 1, -1, -1):
                            if doc.element.body[j].tag.endswith('p'):
                                # Получаем текст параграфа
                                para_element = doc.element.body[j]
                                para_text = ""
                                for t in para_element.iter():
                                    if t.text:
                                        para_text += t.text
                                
                                if para_text.strip():
                                    paragraphs_before.insert(0, para_text.strip())
                                    found_paragraphs += 1
                                    if found_paragraphs >= max_paragraphs:
                                        break
                        break
                    table_count += 1
            
            # Объединяем параграфы в один текст
            text_before = "\n".join(paragraphs_before)
            return text_before
            
        except Exception as e:
            logger.warning(f"Ошибка получения текста перед таблицей {table_idx}: {e}")
            return ""
    
    def _find_paragraph_for_table(self, doc: Document, table_idx: int, after_table: bool = True) -> int:
        """
        Находит индекс параграфа, соответствующего указанной таблице.
        
        Args:
            doc: Документ
            table_idx: Индекс таблицы в документе
            after_table: Если True, ищет параграф ПОСЛЕ таблицы (для аннотаций),
                        если False, ищет параграф ПЕРЕД таблицей
            
        Returns:
            Индекс параграфа после таблицы (или перед, если after_table=False),
            или -1 если не найден
        """
        try:
            table_count = 0
            
            # Проходим по элементам документа
            for i, element in enumerate(doc.element.body):
                if element.tag.endswith('tbl'):  # Таблица
                    if table_count == table_idx:
                        # Нашли нужную таблицу
                        if after_table:
                            # Ищем параграф ПОСЛЕ таблицы
                            for j in range(i + 1, len(doc.element.body)):
                                if doc.element.body[j].tag.endswith('p'):
                                    # Подсчитываем индекс параграфа (сколько параграфов до этого элемента включительно)
                                    para_idx = sum(1 for k in range(j + 1) if doc.element.body[k].tag.endswith('p')) - 1
                                    logger.info(f"   📍 Найден параграф {para_idx} ПОСЛЕ таблицы {table_idx}")
                                    return para_idx
                            # Если не нашли параграф после таблицы, возвращаем последний параграф в документе
                            last_para_idx = sum(1 for k in range(len(doc.element.body)) if doc.element.body[k].tag.endswith('p')) - 1
                            if last_para_idx >= 0:
                                logger.info(f"   📍 Параграф после таблицы {table_idx} не найден, используем последний параграф {last_para_idx}")
                                return last_para_idx
                        else:
                            # Ищем параграф ПЕРЕД таблицей (старая логика для обратной совместимости)
                            para_count = sum(1 for k in range(i) if doc.element.body[k].tag.endswith('p'))
                            for j in range(i - 1, -1, -1):
                                if doc.element.body[j].tag.endswith('p'):
                                    para_idx = sum(1 for k in range(j + 1) if doc.element.body[k].tag.endswith('p')) - 1
                                    logger.info(f"   📍 Найден параграф {para_idx} перед таблицей {table_idx}")
                                    return para_idx
                            logger.warning(f"   ⚠️ Параграф перед таблицей {table_idx} не найден, используем {para_count - 1}")
                            return max(0, para_count - 1)
                    table_count += 1
            
            # Если не нашли таблицу, возвращаем -1
            logger.warning(f"   ⚠️ Таблица {table_idx} не найдена")
            return -1
        except Exception as e:
            logger.error(f"Ошибка поиска параграфа для таблицы {table_idx}: {e}")
            return -1

    def _analyze_new_text_for_table(self, new_text: str, target_text: str) -> Dict[str, str]:
        """
        Анализирует новый текст и разделяет его на части для столбцов таблицы.
        
        Args:
            new_text: Новый текст (например, "ДКР Департамент кредитных рисков")
            target_text: Исходный текст (например, "ДРМ")
            
        Returns:
            Словарь с частями текста для разных столбцов
        """
        parts = {
            "abbreviation": "",
            "description": "",
            "full_text": new_text
        }
        
        # Разделяем текст на аббревиатуру и описание
        words = new_text.split()
        if len(words) >= 2:
            # Первое слово - аббревиатура
            first_word = words[0]
            if len(first_word) <= 5 and first_word.isupper():
                parts["abbreviation"] = first_word
                parts["description"] = " ".join(words[1:])
            else:
                # Если первое слово не похоже на аббревиатуру, используем исходный target_text
                parts["abbreviation"] = target_text
                parts["description"] = new_text
        else:
            # Если только одно слово, используем его как аббревиатуру
            parts["abbreviation"] = new_text
            parts["description"] = ""
        
        logger.info(f"🔍 Анализ нового текста:")
        logger.info(f"   Аббревиатура: '{parts['abbreviation']}'")
        logger.info(f"   Описание: '{parts['description']}'")
        
        return parts

    def _replace_intelligently_in_row(self, row, target_text: str, parts: Dict[str, str], target_cell_idx: int) -> bool:
        """
        Интеллектуально заменяет текст в строке таблицы, распределяя по столбцам.
        
        Args:
            row: Строка таблицы
            target_text: Искомый текст
            parts: Части нового текста для разных столбцов
            target_cell_idx: Индекс ячейки, где найден target_text
            
        Returns:
            True если замена была выполнена
        """
        replaced = False
        
        try:
            # Определяем количество столбцов
            num_columns = len(row.cells)
            logger.info(f"   Строка имеет {num_columns} столбцов")
            
            if num_columns >= 2:
                # Для таблиц с 2+ столбцами: первый - аббревиатура, второй - описание
                
                # Заменяем в первом столбце (аббревиатура)
                first_cell = row.cells[0]
                if target_text in first_cell.text and parts["abbreviation"]:
                    logger.info(f"   Замена в столбце 0: '{target_text}' → '{parts['abbreviation']}'")
                    if self._replace_in_cell(first_cell, target_text, parts["abbreviation"]):
                        replaced = True
                
                # Заменяем во втором столбце (описание), если есть
                if num_columns > 1 and parts["description"]:
                    second_cell = row.cells[1]
                    # Ищем старое описание или добавляем новое
                    if target_text in second_cell.text:
                        logger.info(f"   Замена в столбце 1: '{target_text}' → '{parts['description']}'")
                        if self._replace_in_cell(second_cell, target_text, parts["description"]):
                            replaced = True
                    else:
                        # Если во втором столбце нет target_text, но есть описание связанное с аббревиатурой
                        old_description = second_cell.text.strip()
                        if old_description and parts["abbreviation"] in first_cell.text:
                            # Заменяем все содержимое второго столбца
                            logger.info(f"   Полная замена описания в столбце 1: '{old_description}' → '{parts['description']}'")
                            second_cell.text = parts["description"]
                            replaced = True
            else:
                # Для таблиц с 1 столбцом: заменяем полным текстом
                if self._replace_in_cell(row.cells[0], target_text, parts["full_text"]):
                    replaced = True
                    
        except Exception as e:
            logger.error(f"Ошибка интеллектуальной замены в строке: {e}")
        
        return replaced

    def _is_global_change(self, description: str) -> bool:
        """
        Определяет, является ли изменение глобальным (по всему документу) или локальным.
        
        Args:
            description: Описание инструкции
            
        Returns:
            True если изменение глобальное
        """
        description_lower = description.lower()
        
        # Ключевые фразы для глобальных изменений
        global_keywords = [
            "по всему тексту",
            "по всему документу", 
            "во всем документе",
            "везде в документе",
            "по всему файлу",
            "глобально заменить",
            "массовая замена"
        ]
        
        # Ключевые фразы для локальных изменений
        local_keywords = [
            "в пункте",
            "в разделе",
            "в таблице",
            "в строке",
            "в параграфе",
            "в части",
            "в главе"
        ]
        
        # Проверяем на глобальные ключевые слова
        for keyword in global_keywords:
            if keyword in description_lower:
                logger.info(f"   🌍 Обнаружено глобальное изменение: '{keyword}'")
                return True
        
        # Проверяем на локальные ключевые слова
        for keyword in local_keywords:
            if keyword in description_lower:
                logger.info(f"   📍 Обнаружено локальное изменение: '{keyword}'")
                return False
        
        # По умолчанию считаем изменение локальным
        logger.info(f"   📍 По умолчанию: локальное изменение")
        return False

    def _analyze_row_structure(self, row, row_idx: int) -> Dict[str, Any]:
        """
        Детальный анализ структуры строки таблицы.
        
        Args:
            row: Строка таблицы
            row_idx: Индекс строки
            
        Returns:
            Структура строки с информацией о столбцах
        """
        structure = {
            "row_index": row_idx,
            "columns_count": len(row.cells),
            "columns": [],
            "table_type": "unknown"
        }
        
        # Анализируем каждый столбец
        for col_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            column_info = {
                "index": col_idx,
                "content": cell_text,
                "type": self._determine_column_type_enhanced(cell_text, col_idx),
                "length": len(cell_text)
            }
            structure["columns"].append(column_info)
        
        # Определяем тип таблицы на основе структуры
        structure["table_type"] = self._determine_table_type(structure["columns"])
        
        logger.info(f"   📋 Анализ строки {row_idx}:")
        for col in structure["columns"]:
            logger.info(f"      Столбец {col['index']}: '{col['content'][:20]}...' (тип: {col['type']})")
        
        return structure
    
    def _determine_column_type_enhanced(self, content: str, col_index: int) -> str:
        """
        Улучшенное определение типа столбца на основе содержимого.
        
        Args:
            content: Содержимое ячейки
            col_index: Индекс столбца
            
        Returns:
            Тип столбца
        """
        
        if not content:
            return "empty"
        
        # Аббревиатуры (короткие заглавные буквы)
        if re.match(r'^[А-ЯЁ]{2,6}$', content):
            return "abbreviation"
        
        # Номера или коды
        if re.match(r'^\d+\.?\d*$', content):
            return "number"
        
        # Длинные описания
        if len(content) > 15 and ' ' in content:
            return "description"
        
        # Короткие ключи
        if len(content) <= 10:
            return "key"
        
        # По позиции столбца
        if col_index == 0:
            return "primary_key"
        elif col_index == 1:
            return "secondary_info"
        else:
            return "additional_info"
    
    def _determine_table_type(self, columns: List[Dict]) -> str:
        """
        Определяет тип таблицы на основе анализа столбцов.
        
        Args:
            columns: Информация о столбцах
            
        Returns:
            Тип таблицы
        """
        if len(columns) >= 2:
            first_col_type = columns[0]["type"]
            second_col_type = columns[1]["type"]
            
            if first_col_type == "abbreviation" and second_col_type == "description":
                return "abbreviations_table"
            elif first_col_type == "number" and second_col_type == "description":
                return "numbered_list"
            elif first_col_type == "key" and second_col_type in ["description", "secondary_info"]:
                return "key_value_table"
        
        return "general_table"
    
    async def _map_new_text_to_structure(
        self, 
        new_text: str, 
        target_text: str, 
        row_structure: Dict,
        description: str = "",
        table_context: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Сопоставляет новый текст со структурой строки таблицы.
        Использует алгоритмический подход + LLM для проверки и корректировки.
        
        Args:
            new_text: Новый текст для распределения
            target_text: Исходный текст для замены
            row_structure: Структура строки
            description: Описание инструкции (для контекста LLM)
            table_context: Контекст таблицы (заголовки, соседние строки)
            
        Returns:
            Распределение текста по столбцам
        """
        # 1. АЛГОРИТМИЧЕСКОЕ РАСПРЕДЕЛЕНИЕ (основной подход)
        distribution = {
            "columns_mapping": {},
            "strategy": "auto"
        }
        
        table_type = row_structure.get("table_type", "general_table")
        columns = row_structure.get("columns", [])
        
        logger.info(f"   🎯 Сопоставление для типа таблицы: {table_type}")
        
        if table_type == "abbreviations_table" and len(columns) >= 2:
            # Для таблиц сокращений: разделяем на аббревиатуру и описание
            parts = self._split_abbreviation_text(new_text, target_text)
            
            # Первый столбец - аббревиатура
            if columns[0]["type"] in ["abbreviation", "primary_key"]:
                distribution["columns_mapping"][0] = parts["abbreviation"]
            
            # Второй столбец - описание
            if len(columns) > 1 and columns[1]["type"] in ["description", "secondary_info"]:
                distribution["columns_mapping"][1] = parts["description"]
                
            distribution["strategy"] = "abbreviation_split"
            
        elif table_type == "key_value_table" and len(columns) >= 2:
            # Для таблиц ключ-значение
            parts = self._split_key_value_text(new_text, target_text)
            distribution["columns_mapping"][0] = parts["key"]
            if len(columns) > 1:
                distribution["columns_mapping"][1] = parts["value"]
            distribution["strategy"] = "key_value_split"
            
        else:
            # Для общих таблиц - равномерное распределение
            parts = self._split_general_text(new_text, len(columns))
            for i, part in enumerate(parts):
                if i < len(columns):
                    distribution["columns_mapping"][i] = part
            distribution["strategy"] = "general_split"
        
        logger.info(f"   📝 Алгоритмическая стратегия: {distribution['strategy']}")
        for col_idx, content in distribution["columns_mapping"].items():
            logger.info(f"      Столбец {col_idx}: '{content}'")
        
        # 2. ПРОВЕРКА И КОРРЕКТИРОВКА ЧЕРЕЗ LLM
        try:
            llm_distribution = await self._map_text_with_llm(
                new_text=new_text,
                target_text=target_text,
                row_structure=row_structure,
                algorithmic_distribution=distribution,
                description=description,
                table_context=table_context
            )
            
            if llm_distribution and llm_distribution.get("confidence", 0) >= 0.7:
                logger.info(f"   ✅ LLM корректировка применена (уверенность: {llm_distribution.get('confidence', 0):.2f})")
                distribution = llm_distribution.get("distribution", distribution)
                distribution["strategy"] = distribution.get("strategy", "llm_corrected")
                distribution["llm_corrected"] = True
                distribution["llm_reasoning"] = llm_distribution.get("reasoning", "")
            else:
                logger.info(f"   ⚠️ LLM корректировка не применена (низкая уверенность или ошибка), используется алгоритмический результат")
                distribution["llm_corrected"] = False
                
        except Exception as e:
            logger.warning(f"   ⚠️ Ошибка при LLM проверке, используется алгоритмический результат: {e}")
            distribution["llm_corrected"] = False
            distribution["llm_error"] = str(e)
        
        return distribution
    
    def _split_abbreviation_text(self, text: str, target_text: str) -> Dict[str, str]:
        """Разделяет текст на аббревиатуру и описание."""
        words = text.split()
        if len(words) >= 2:
            first_word = words[0]
            if len(first_word) <= 6 and first_word.isupper():
                return {
                    "abbreviation": first_word,
                    "description": " ".join(words[1:])
                }
        
        # Если не удалось разделить, используем target_text как аббревиатуру
        return {
            "abbreviation": target_text,
            "description": text
        }
    
    def _split_key_value_text(self, text: str, target_text: str) -> Dict[str, str]:
        """Разделяет текст на ключ и значение."""
        # Простое разделение по первому пробелу или двоеточию
        if ':' in text:
            parts = text.split(':', 1)
            return {"key": parts[0].strip(), "value": parts[1].strip()}
        elif ' ' in text:
            parts = text.split(' ', 1)
            return {"key": parts[0].strip(), "value": parts[1].strip()}
        else:
            return {"key": text, "value": ""}
    
    def _split_general_text(self, text: str, num_columns: int) -> List[str]:
        """Разделяет текст на части для общих таблиц."""
        if num_columns <= 1:
            return [text]
        
        words = text.split()
        if len(words) <= num_columns:
            # Если слов меньше или равно количеству столбцов
            result = words + [""] * (num_columns - len(words))
            return result[:num_columns]
        else:
            # Распределяем слова по столбцам
            words_per_col = len(words) // num_columns
            result = []
            for i in range(num_columns):
                start_idx = i * words_per_col
                if i == num_columns - 1:  # Последний столбец получает оставшиеся слова
                    end_idx = len(words)
                else:
                    end_idx = (i + 1) * words_per_col
                result.append(" ".join(words[start_idx:end_idx]))
            return result
    
    async def _identify_target_table_with_llm(
        self,
        doc: Document,
        description: str,
        target_text: str,
        table_name: Optional[str] = None
    ) -> Optional[List[int]]:
        """
        Использует LLM для точного определения целевой таблицы на основе семантического анализа.
        
        Args:
            doc: Документ python-docx
            description: Описание инструкции
            target_text: Искомый текст
            table_name: Название таблицы (если указано в инструкции)
            
        Returns:
            Список индексов целевых таблиц, или None при ошибке
        """
        if not self.openai_client:
            logger.warning("LLM клиент не инициализирован, пропускаем интеллектуальное определение таблицы")
            return None
        
        if not doc.tables:
            logger.info("В документе нет таблиц")
            return None
        
        try:
            # Подготовка информации о всех таблицах
            tables_info = []
            for table_idx, table in enumerate(doc.tables):
                # Получаем заголовки таблицы (первые 2 строки)
                headers = []
                for i in range(min(2, len(table.rows))):
                    header_row = []
                    for cell in table.rows[i].cells:
                        header_row.append(cell.text.strip()[:100])  # Ограничиваем длину
                    headers.append(" | ".join(header_row))
                
                # Получаем текст перед таблицей (для поиска названия)
                text_before_table = self._get_text_before_table(doc, table_idx, max_paragraphs=3)
                
                # Получаем структуру таблицы
                num_rows = len(table.rows)
                num_cols = len(table.rows[0].cells) if table.rows else 0
                
                # Проверяем наличие target_text
                contains_target = False
                target_cells_info = []
                for row_idx, row in enumerate(table.rows[:5]):  # Проверяем первые 5 строк
                    for col_idx, cell in enumerate(row.cells):
                        if target_text in cell.text:
                            contains_target = True
                            target_cells_info.append({
                                "row": row_idx,
                                "col": col_idx,
                                "content": cell.text.strip()[:100]
                            })
                
                table_info = {
                    "index": table_idx,
                    "headers": headers,
                    "text_before": text_before_table[:300] if text_before_table else "",  # Ограничиваем длину
                    "num_rows": num_rows,
                    "num_cols": num_cols,
                    "contains_target": contains_target,
                    "target_cells": target_cells_info[:3]  # Первые 3 совпадения
                }
                tables_info.append(table_info)
            
            # Формируем промпт для LLM
            system_prompt = """Ты эксперт по анализу структуры документов Word. 
Твоя задача - определить, какая таблица из списка является целевой для выполнения изменения на основе:
1. Семантики описания инструкции
2. Названия таблицы (если указано) - ВАЖНО: ищи название не только в заголовках таблицы, но и в тексте ПЕРЕД таблицей
3. Наличия искомого текста в таблице
4. Структуры и содержимого таблиц

Верни JSON с индексами целевых таблиц (может быть одна или несколько, если изменение касается нескольких таблиц)."""

            tables_summary = "\n".join([
                f"Таблица {t['index']}:\n"
                + (f"  Текст перед таблицей: {t['text_before']}\n" if t['text_before'] else "  Текст перед таблицей: (нет)\n")
                + f"  Заголовки таблицы: {'; '.join(t['headers'][:2]) if t['headers'] else '(нет)'}\n"
                + f"  Размер: {t['num_rows']} строк × {t['num_cols']} столбцов\n"
                + f"  Содержит target_text: {'Да' if t['contains_target'] else 'Нет'}\n"
                + (f"  Совпадения: {', '.join(['Row ' + str(c['row']) + ', Col ' + str(c['col']) for c in t['target_cells']])}\n" if t['target_cells'] else "")
                for t in tables_info
            ])
            
            user_prompt = f"""ИНСТРУКЦИЯ: {description}

ИСКОМЫЙ ТЕКСТ: "{target_text}"
НАЗВАНИЕ ТАБЛИЦЫ (если указано): {table_name if table_name else 'Не указано'}

ДОСТУПНЫЕ ТАБЛИЦЫ В ДОКУМЕНТЕ:
{tables_summary}

ПРОАНАЛИЗИРУЙ инструкцию и определи, какая таблица (или таблицы) является целевой для этого изменения.

КРИТЕРИИ ВЫБОРА (в порядке приоритета):
1. НАЗВАНИЕ ТАБЛИЦЫ: Если указано название (полностью или частично), ищи его:
   - В тексте ПЕРЕД таблицей (в параграфах перед таблицей) - ВЫСОКИЙ ПРИОРИТЕТ
   - В заголовках самой таблицы
   - Название может быть сокращено или обрезано, ищи частичные совпадения
   
2. НАЛИЧИЕ ИСКОМОГО ТЕКСТА: Таблица должна содержать искомый текст "{target_text}"

3. СЕМАНТИКА ИНСТРУКЦИИ: Структура и тип таблицы должны соответствовать описанию инструкции

4. ТИП ТАБЛИЦЫ: Для инструкций типа "Изменение строки" в таблице сокращений - выбирай таблицы с аббревиатурами

ВАЖНО: 
- Если название таблицы указано (даже частично), приоритет отдавай таблицам, где это название найдено в тексте перед таблицей
- Если название не указано, но есть искомый текст - выбирай таблицу, где этот текст найден
- Если несколько критериев указывают на одну таблицу - confidence должен быть высоким (>= 0.8)
- Если есть неопределенность, но один вариант более вероятен - установи confidence >= 0.6 и верни этот вариант

Верни JSON:
{{
  "target_table_indices": [0, 1, ...],
  "confidence": 0.95,
  "reasoning": "детальное объяснение выбора с указанием всех критериев"
}}

Если изменение точечное и касается конкретной таблицы, верни только её индекс.
Если изменение должно применяться к нескольким таблицам, верни все соответствующие индексы.
Если название таблицы указано и найдено в тексте перед таблицей - устанавливай confidence >= 0.8."""

            logger.info(f"   🤖 Отправка запроса к LLM для определения целевой таблицы...")
            logger.info(f"   📝 Параметры запроса:")
            logger.info(f"      - Описание: '{description[:200]}...'")
            logger.info(f"      - Искомый текст: '{target_text}'")
            logger.info(f"      - Название таблицы: '{table_name if table_name else 'Не указано'}'")
            logger.info(f"      - Количество таблиц в документе: {len(doc.tables)}")
            
            response = await self.openai_client.chat.completions.create(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=0.2,  # Низкая температура для более точного определения
                max_tokens=1024,
                response_format={"type": "json_object"},
            )
            
            content = response.choices[0].message.content if response.choices else None
            if isinstance(content, list):
                content = "".join(
                    segment.get("text", "")
                    for segment in content
                    if isinstance(segment, dict)
                )
            
            if not isinstance(content, str) or not content.strip():
                logger.warning("LLM не вернул корректный ответ для определения таблицы")
                return None
            
            # Очистка JSON
            content_cleaned = content.strip()
            if content_cleaned.startswith("```"):
                lines = content_cleaned.split("\n")
                if len(lines) > 1:
                    lines = lines[1:]
                if lines and lines[-1].strip() == "```":
                    lines = lines[:-1]
                content_cleaned = "\n".join(lines).strip()
            
            # Парсинг JSON
            result = json.loads(content_cleaned)
            
            # Валидация результата
            if "target_table_indices" not in result:
                logger.warning("LLM вернул некорректную структуру для определения таблицы")
                return None
            
            target_indices = result["target_table_indices"]
            if not isinstance(target_indices, list):
                logger.warning("target_table_indices должен быть списком")
                return None
            
            confidence = result.get("confidence", 0)
            reasoning = result.get("reasoning", "")
            
            logger.info(f"   📊 Результат LLM: target_table_indices={target_indices}, confidence={confidence:.2f}")
            
            # Фильтруем индексы (проверяем, что они валидны)
            valid_indices = [idx for idx in target_indices if isinstance(idx, int) and 0 <= idx < len(doc.tables)]
            
            # Логируем результаты валидации
            if len(target_indices) != len(valid_indices):
                invalid_indices = [idx for idx in target_indices if idx not in valid_indices]
                logger.warning(f"   ⚠️ Некоторые индексы не прошли валидацию: {invalid_indices} (всего таблиц в документе: {len(doc.tables)})")
            
            if valid_indices:
                logger.info(f"   ✅ LLM определил целевые таблицы: {valid_indices} (confidence: {confidence:.2f})")
                if reasoning:
                    logger.info(f"   💭 LLM reasoning: {reasoning[:200]}...")
                return valid_indices
            else:
                logger.warning(f"   ⚠️ LLM не смог точно определить целевую таблицу:")
                logger.warning(f"      - target_table_indices от LLM: {target_indices}")
                logger.warning(f"      - confidence: {confidence:.2f}")
                logger.warning(f"      - valid_indices после фильтрации: {valid_indices}")
                logger.warning(f"      - всего таблиц в документе: {len(doc.tables)}")
                if reasoning:
                    logger.warning(f"      - reasoning: {reasoning[:200]}...")
                return None
            
        except json.JSONDecodeError as e:
            logger.error(f"Ошибка парсинга JSON от LLM при определении таблицы: {e}")
            return None
        except Exception as e:
            logger.error(f"Ошибка при запросе к LLM для определения таблицы: {e}")
            return None
    
    def _get_table_context(self, table, row_idx: int, max_header_rows: int = 2, max_sample_rows: int = 2) -> Dict[str, Any]:
        """
        Получает контекст таблицы: заголовки и соседние строки для анализа LLM.
        
        Args:
            table: Таблица из python-docx
            row_idx: Индекс текущей строки
            max_header_rows: Максимальное количество строк заголовка
            max_sample_rows: Максимальное количество соседних строк для примера
            
        Returns:
            Контекст таблицы с заголовками и образцами строк
        """
        context = {
            "headers": [],
            "sample_rows_before": [],
            "current_row": [],
            "sample_rows_after": [],
            "total_columns": 0
        }
        
        try:
            if not table.rows:
                return context
            
            # Получаем заголовки (первые max_header_rows строк)
            for i in range(min(max_header_rows, len(table.rows))):
                header_row = []
                for cell in table.rows[i].cells:
                    header_row.append(cell.text.strip())
                context["headers"].append(header_row)
            
            # Определяем общее количество столбцов
            if table.rows:
                context["total_columns"] = len(table.rows[0].cells)
            
            # Получаем текущую строку
            if row_idx < len(table.rows):
                for cell in table.rows[row_idx].cells:
                    context["current_row"].append(cell.text.strip())
            
            # Получаем строки перед текущей
            start_idx = max(0, row_idx - max_sample_rows)
            for i in range(start_idx, row_idx):
                sample_row = []
                for cell in table.rows[i].cells:
                    sample_row.append(cell.text.strip())
                context["sample_rows_before"].append(sample_row)
            
            # Получаем строки после текущей
            end_idx = min(len(table.rows), row_idx + 1 + max_sample_rows)
            for i in range(row_idx + 1, end_idx):
                sample_row = []
                for cell in table.rows[i].cells:
                    sample_row.append(cell.text.strip())
                context["sample_rows_after"].append(sample_row)
                
        except Exception as e:
            logger.warning(f"Ошибка получения контекста таблицы: {e}")
        
        return context
    
    async def _map_text_with_llm(
        self,
        new_text: str,
        target_text: str,
        row_structure: Dict,
        algorithmic_distribution: Dict,
        description: str = "",
        table_context: Optional[Dict[str, Any]] = None
    ) -> Optional[Dict[str, Any]]:
        """
        Использует LLM для проверки и корректировки алгоритмического распределения текста по столбцам.
        
        Args:
            new_text: Новый текст для распределения
            target_text: Исходный текст для замены
            row_structure: Структура строки
            algorithmic_distribution: Результат алгоритмического распределения
            description: Описание инструкции
            table_context: Контекст таблицы (заголовки, соседние строки)
            
        Returns:
            Словарь с корректированным распределением и уверенностью, или None при ошибке
        """
        if not self.openai_client:
            logger.warning("LLM клиент не инициализирован, пропускаем проверку")
            return None
        
        try:
            # Формируем контекст для LLM
            columns = row_structure.get("columns", [])
            table_type = row_structure.get("table_type", "unknown")
            
            # Подготовка информации о текущей строке
            current_row_info = "\n".join([
                f"  Столбец {col['index']}: '{col['content']}' (тип: {col['type']})"
                for col in columns
            ])
            
            # Подготовка алгоритмического результата
            algo_result = "\n".join([
                f"  Столбец {col_idx}: '{content}'"
                for col_idx, content in algorithmic_distribution.get("columns_mapping", {}).items()
            ])
            
            # Подготовка контекста таблицы
            table_info = ""
            if table_context:
                if table_context.get("headers"):
                    headers_text = "\n".join([
                        f"  {' | '.join(header)}"
                        for header in table_context["headers"]
                    ])
                    table_info += f"\nЗаголовки таблицы:\n{headers_text}\n"
                
                if table_context.get("sample_rows_before"):
                    sample_text = "\n".join([
                        f"  {' | '.join(row)}"
                        for row in table_context["sample_rows_before"]
                    ])
                    table_info += f"\nСтроки перед текущей:\n{sample_text}\n"
                
                if table_context.get("sample_rows_after"):
                    sample_text = "\n".join([
                        f"  {' | '.join(row)}"
                        for row in table_context["sample_rows_after"]
                    ])
                    table_info += f"\nСтроки после текущей:\n{sample_text}\n"
            
            # Формируем промпт для LLM
            system_prompt = """Ты эксперт по анализу структуры таблиц в Word документах. 
Твоя задача - проверить и скорректировать распределение текста по столбцам таблицы на основе:
1. Семантики текста инструкции
2. Структуры и содержимого текущей строки
3. Контекста таблицы (заголовки, соседние строки)

Верни JSON с корректированным распределением и оценкой уверенности (0.0-1.0).
Если алгоритмический результат корректен, подтверди его с высокой уверенностью.
Если нужна корректировка, предложи улучшенный вариант с объяснением."""

            user_prompt = f"""ИНСТРУКЦИЯ: {description}

ТЕКУЩАЯ СТРОКА ТАБЛИЦЫ:
{current_row_info}

НОВЫЙ ТЕКСТ ДЛЯ РАСПРЕДЕЛЕНИЯ: "{new_text}"
ИСХОДНЫЙ ТЕКСТ: "{target_text}"

АЛГОРИТМИЧЕСКОЕ РАСПРЕДЕЛЕНИЕ:
{algo_result}
Стратегия: {algorithmic_distribution.get('strategy', 'unknown')}

КОНТЕКСТ ТАБЛИЦЫ:
{table_info}

Тип таблицы: {table_type}
Количество столбцов: {len(columns)}

ПРОАНАЛИЗИРУЙ распределение и верни JSON:
{{
  "distribution": {{
    "columns_mapping": {{
      "0": "текст для столбца 0",
      "1": "текст для столбца 1",
      ...
    }},
    "strategy": "название стратегии",
    "llm_corrected": true
  }},
  "confidence": 0.95,
  "reasoning": "объяснение корректировки или подтверждения"
}}

Если распределение корректно, установи confidence >= 0.9 и подтверди результат.
Если нужна корректировка, установи confidence >= 0.7 и предложи улучшенный вариант."""

            logger.info(f"   🤖 Отправка запроса к LLM для проверки распределения...")
            
            response = await self.openai_client.chat.completions.create(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=0.3,  # Немного творчества для корректировки
                max_tokens=2048,
                response_format={"type": "json_object"},
            )
            
            content = response.choices[0].message.content if response.choices else None
            if isinstance(content, list):
                content = "".join(
                    segment.get("text", "")
                    for segment in content
                    if isinstance(segment, dict)
                )
            
            if not isinstance(content, str) or not content.strip():
                logger.warning("LLM не вернул корректный ответ")
                return None
            
            # Очистка JSON от markdown code blocks
            content_cleaned = content.strip()
            if content_cleaned.startswith("```"):
                lines = content_cleaned.split("\n")
                if len(lines) > 1:
                    lines = lines[1:]
                if lines and lines[-1].strip() == "```":
                    lines = lines[:-1]
                content_cleaned = "\n".join(lines).strip()
            
            # Парсинг JSON
            result = json.loads(content_cleaned)
            
            # Валидация результата
            if "distribution" not in result or "confidence" not in result:
                logger.warning("LLM вернул некорректную структуру")
                return None
            
            # Преобразование columns_mapping в правильный формат
            distribution = result["distribution"]
            if "columns_mapping" in distribution:
                # Преобразуем строковые ключи в int
                columns_mapping = {}
                for key, value in distribution["columns_mapping"].items():
                    try:
                        col_idx = int(key)
                        columns_mapping[col_idx] = value
                    except ValueError:
                        logger.warning(f"Некорректный индекс столбца: {key}")
                distribution["columns_mapping"] = columns_mapping
            
            logger.info(f"   ✅ LLM вернул результат (confidence: {result.get('confidence', 0):.2f})")
            if result.get("reasoning"):
                logger.info(f"   💭 LLM reasoning: {result['reasoning'][:200]}...")
            
            return {
                "distribution": distribution,
                "confidence": float(result.get("confidence", 0)),
                "reasoning": result.get("reasoning", "")
            }
            
        except json.JSONDecodeError as e:
            logger.error(f"Ошибка парсинга JSON от LLM: {e}")
            logger.debug(f"Содержимое ответа: {content_cleaned[:500] if 'content_cleaned' in locals() else 'N/A'}")
            return None
        except Exception as e:
            logger.error(f"Ошибка при запросе к LLM для проверки распределения: {e}")
            return None
    
    def _apply_structured_replacement(self, row, target_text: str, distribution: Dict) -> bool:
        """
        Применяет структурированную замену в строке таблицы.
        
        Args:
            row: Строка таблицы
            target_text: Исходный текст
            distribution: Распределение по столбцам
            
        Returns:
            True если замена была выполнена
        """
        replaced = False
        columns_mapping = distribution.get("columns_mapping", {})
        strategy = distribution.get("strategy", "")
        
        try:
            # Для структурированной замены заменяем все столбцы из mapping
            # независимо от того, содержится ли target_text в каждом столбце
            for col_idx, new_content in columns_mapping.items():
                if col_idx < len(row.cells):
                    cell = row.cells[col_idx]
                    old_content = cell.text.strip()
                    
                    # Для структурированной замены всегда заменяем все столбцы из mapping
                    # Это гарантирует, что вся строка будет обновлена корректно
                    logger.info(f"      Замена в столбце {col_idx}: '{old_content}' → '{new_content}'")
                    
                    # Сохраняем форматирование через runs
                    if cell.paragraphs:
                        para = cell.paragraphs[0]
                        if para.runs:
                            # Заменяем текст в первом run, сохраняя форматирование
                            para.runs[0].text = new_content
                            # Удаляем остальные runs, если они есть
                            for run in para.runs[1:]:
                                para._element.remove(run._element)
                        else:
                            # Если нет runs, создаем новый
                            para.text = new_content
                    else:
                        # Если нет параграфов, создаем новый
                        cell.text = new_content
                    
                    replaced = True
        
        except Exception as e:
            logger.error(f"Ошибка применения структурированной замены: {e}")
        
        return replaced

    def _should_use_structured_replacement(self, description: str) -> bool:
        """
        Определяет, нужно ли использовать структурированную замену (распределение по столбцам)
        на основе анализа описания инструкции.
        
        Args:
            description: Описание инструкции
            
        Returns:
            True если требуется структурированная замена (распределение по столбцам)
        """
        description_lower = description.lower()
        
        # Ключевые фразы, указывающие на распределение по столбцам (замена строки таблицы)
        structured_keywords = [
            "строку",
            "строки",
            "в таблице",
            "таблице строку",
            "таблице строки",
        ]
        
        # Ключевые фразы, указывающие на простую замену фразы
        simple_keywords = [
            "слова",
            "слово",
            "фразу",
            "фразы",
            "текст",
            "в пункте",
            "в разделе",
            "в параграфе",
        ]
        
        # Проверяем наличие ключевых слов для структурированной замены
        for keyword in structured_keywords:
            if keyword in description_lower:
                # Дополнительная проверка: должно быть слово "изложить" или "заменить"
                if "изложить" in description_lower or "заменить" in description_lower:
                    logger.info(f"   🔍 Обнаружена инструкция на распределение по столбцам: '{keyword}'")
                    return True
        
        # Если найдены ключевые слова для простой замены
        for keyword in simple_keywords:
            if keyword in description_lower:
                logger.info(f"   🔍 Обнаружена инструкция на простую замену фразы: '{keyword}'")
                return False
        
        # По умолчанию - простая замена
        logger.info(f"   🔍 По умолчанию: простая замена фразы")
        return False

    async def _find_paragraph_location_with_llm(
        self,
        doc: Document,
        description: str,
        target_text: str,
        punkt_number: Optional[str] = None
    ) -> Optional[Dict[str, Any]]:
        """
        Использует LLM для поиска пункта и определения его местоположения (в таблице или параграфе).
        
        Args:
            doc: Документ python-docx
            description: Описание инструкции
            target_text: Искомый текст для замены
            punkt_number: Номер пункта (если указан)
            
        Returns:
            Словарь с информацией о местоположении пункта, или None при ошибке
        """
        if not self.openai_client:
            logger.warning("LLM клиент не инициализирован, пропускаем поиск пункта через LLM")
            return None
        
        if not punkt_number:
            logger.info("Номер пункта не указан, пропускаем LLM поиск")
            return None
        
        try:
            # Подготовка информации о документе для LLM
            # Ищем все упоминания номера пункта
            punkt_locations = []
            
            # Ищем в параграфах с более точным поиском
            for para_idx, para in enumerate(doc.paragraphs):  # Ищем во всех параграфах
                para_text = para.text.strip()
                # Используем точные regex паттерны для начала строки
                punkt_patterns = [
                    rf"^{re.escape(punkt_number)}\.",
                    rf"^{re.escape(punkt_number)}\)",
                    rf"^{re.escape(punkt_number)}:",
                    rf"^{re.escape(punkt_number)}\s",
                    rf"\bпункт\s+{re.escape(punkt_number)}\b",
                    rf"\bп\.\s*{re.escape(punkt_number)}\b",
                ]
                punkt_found = False
                for pattern in punkt_patterns:
                    if re.search(pattern, para_text, re.IGNORECASE):
                        punkt_locations.append({
                            "type": "paragraph",
                            "index": para_idx,
                            "text": para_text[:200],
                            "contains_target": target_text in para_text
                        })
                        punkt_found = True
                        break
                # Если нашли достаточно совпадений, останавливаемся
                if len(punkt_locations) >= 10:
                    break
            
            # Ищем в таблицах с более точным поиском
            table_info_list = []
            for table_idx, table in enumerate(doc.tables):
                # Получаем текст перед таблицей (для контекста)
                text_before_table = self._get_text_before_table(doc, table_idx, max_paragraphs=3)
                
                table_rows_info = []
                for row_idx, row in enumerate(table.rows):  # Ищем во всех строках
                    row_text = ""
                    contains_punkt = False
                    contains_target = False
                    punkt_cell_idx = None
                    target_cell_idx = None
                    
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        row_text += f" | {cell_text}"
                        
                        # Используем точные regex паттерны для начала строки/ячейки
                        punkt_patterns = [
                            rf"^{re.escape(punkt_number)}\.",
                            rf"^{re.escape(punkt_number)}\)",
                            rf"^{re.escape(punkt_number)}:",
                            rf"^{re.escape(punkt_number)}\s",
                            rf"\bпункт\s+{re.escape(punkt_number)}\b",
                            rf"\bп\.\s*{re.escape(punkt_number)}\b",
                        ]
                        for pattern in punkt_patterns:
                            if re.search(pattern, cell_text, re.IGNORECASE):
                                contains_punkt = True
                                punkt_cell_idx = cell_idx
                                break
                        
                        # Проверяем наличие target_text
                        if target_text in cell_text:
                            contains_target = True
                            target_cell_idx = cell_idx
                    
                    if contains_punkt or contains_target:
                        table_rows_info.append({
                            "row_index": row_idx,
                            "text": row_text[:300],
                            "contains_punkt": contains_punkt,
                            "contains_target": contains_target,
                            "punkt_cell": punkt_cell_idx,
                            "target_cell": target_cell_idx
                        })
                
                if table_rows_info:
                    # Получаем заголовки таблицы
                    headers = []
                    for i in range(min(2, len(table.rows))):
                        header_row = []
                        for cell in table.rows[i].cells:
                            header_row.append(cell.text.strip()[:50])
                        headers.append(" | ".join(header_row))
                    
                    table_info_list.append({
                        "table_index": table_idx,
                        "headers": headers,
                        "text_before": text_before_table[:200] if text_before_table else "",
                        "rows": table_rows_info[:10]  # Первые 10 релевантных строк
                    })
            
            # Формируем промпт для LLM
            system_prompt = """Ты эксперт по анализу структуры документов Word. 
Твоя задача - найти указанный пункт в документе и определить:
1. Находится ли пункт в таблице или в обычном параграфе
2. Если в таблице - в какой таблице (индекс), в какой строке (индекс строки) и в какой ячейке (индекс столбца) находится номер пункта
3. В какой ячейке (индекс столбца) находится target_text для замены (ВАЖНО: target_text НЕ должен быть в той же ячейке, где номер пункта)
4. Если в параграфе - в каком параграфе (индекс)

ВАЖНО:
- Номер пункта может быть в первой ячейке строки таблицы
- target_text для замены должен находиться в других ячейках той же строки
- Если пункт найден в таблице, но target_text не найден в той же строке, проверь соседние строки
- Учитывай текст перед таблицей - он может содержать название таблицы или контекст

Верни JSON с точной информацией о местоположении пункта и текста для замены."""

            paragraphs_info = "\n".join([
                f"Параграф {loc['index']}: {loc['text']} (содержит target_text: {loc['contains_target']})"
                for loc in punkt_locations[:10]  # Первые 10 найденных
            ]) if punkt_locations else "Параграфы с номером пункта не найдены"
            
            tables_info = "\n".join([
                f"Таблица {t['table_index']}:\n"
                + (f"  Текст перед таблицей: {t.get('text_before', '')}\n" if t.get('text_before') else "")
                + f"  Заголовки: {'; '.join(t['headers'])}\n"
                + "\n".join([
                    f"  Строка {r['row_index']}: {r['text']} (пункт в ячейке {r.get('punkt_cell', 'N/A')}: {r['contains_punkt']}, target в ячейке {r.get('target_cell', 'N/A')}: {r['contains_target']})"
                    for r in t['rows']
                ])
                for t in table_info_list
            ]) if table_info_list else "Таблицы с номером пункта не найдены"
            
            user_prompt = f"""ИНСТРУКЦИЯ: {description}

НОМЕР ПУНКТА: {punkt_number}
ИСКОМЫЙ ТЕКСТ ДЛЯ ЗАМЕНЫ: "{target_text}"

ПАРАГРАФЫ С НОМЕРОМ ПУНКТА:
{paragraphs_info}

ТАБЛИЦЫ С НОМЕРОМ ПУНКТА:
{tables_info}

ПРОАНАЛИЗИРУЙ и определи:
1. Находится ли пункт {punkt_number} в таблице или в обычном параграфе?
2. Если в таблице:
   - В какой таблице (индекс)?
   - В какой строке (индекс строки) находится номер пункта {punkt_number}?
   - В какой ячейке (индекс столбца) находится номер пункта?
   - В какой ячейке (индекс столбца) находится target_text "{target_text}" для замены?
   - ВАЖНО: target_text должен быть в другой ячейке, не в той, где номер пункта
3. Если в параграфе - в каком параграфе (индекс)?

КРИТЕРИИ ВЫБОРА:
- Если пункт найден в таблице и target_text найден в той же строке - confidence >= 0.8
- Если пункт найден в таблице, но target_text не найден в той же строке - проверь соседние строки, confidence >= 0.6
- Если пункт найден в параграфе и target_text найден в том же параграфе - confidence >= 0.8
- Если пункт не найден, но есть похожие совпадения - confidence < 0.7, location_type: "unknown"

Верни JSON:
{{
  "location_type": "table" или "paragraph" или "unknown",
  "table_index": 0 (если location_type == "table"),
  "row_index": 5 (если location_type == "table"),
  "cell_index": 1 (индекс ячейки с target_text, если location_type == "table"),
  "paragraph_index": 10 (если location_type == "paragraph"),
  "confidence": 0.95,
  "reasoning": "детальное объяснение определения местоположения с указанием всех найденных совпадений"
}}

Если пункт не найден или не удалось точно определить, верни location_type: "unknown" с confidence < 0.7."""

            logger.info(f"   🤖 Отправка запроса к LLM для поиска пункта {punkt_number}...")
            
            response = await self.openai_client.chat.completions.create(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=0.2,
                max_tokens=1024,
                response_format={"type": "json_object"},
            )
            
            content = response.choices[0].message.content if response.choices else None
            if isinstance(content, list):
                content = "".join(
                    segment.get("text", "")
                    for segment in content
                    if isinstance(segment, dict)
                )
            
            if not isinstance(content, str) or not content.strip():
                logger.warning("LLM не вернул корректный ответ для поиска пункта")
                return None
            
            # Очистка JSON
            content_cleaned = content.strip()
            if content_cleaned.startswith("```"):
                lines = content_cleaned.split("\n")
                if len(lines) > 1:
                    lines = lines[1:]
                if lines and lines[-1].strip() == "```":
                    lines = lines[:-1]
                content_cleaned = "\n".join(lines).strip()
            
            # Парсинг JSON
            result = json.loads(content_cleaned)
            
            # Валидация результата
            if "location_type" not in result:
                logger.warning("LLM вернул некорректную структуру для поиска пункта")
                return None
            
            location_type = result.get("location_type", "unknown")
            confidence = result.get("confidence", 0)
            reasoning = result.get("reasoning", "")
            
            logger.info(f"   ✅ LLM определил местоположение: {location_type} (confidence: {confidence:.2f})")
            if reasoning:
                logger.info(f"   💭 LLM reasoning: {reasoning[:200]}...")
            
            if confidence >= 0.7 and location_type != "unknown":
                return {
                    "location_type": location_type,
                    "table_index": result.get("table_index"),
                    "row_index": result.get("row_index"),
                    "cell_index": result.get("cell_index"),
                    "paragraph_index": result.get("paragraph_index"),
                    "confidence": confidence,
                    "reasoning": reasoning
                }
            else:
                logger.info(f"   ⚠️ LLM не смог точно определить местоположение (confidence: {confidence:.2f})")
                return None
            
        except json.JSONDecodeError as e:
            logger.error(f"Ошибка парсинга JSON от LLM при поиске пункта: {e}")
            return None
        except Exception as e:
            logger.error(f"Ошибка при запросе к LLM для поиска пункта: {e}")
            return None
    
    def _extract_tables_from_instructions(self, changes_file: str) -> Dict[str, List]:
        """
        Извлекает таблицы из документа инструкций и связывает их с предшествующим текстом.
        
        Returns:
            Словарь, где ключ - текст перед таблицей (например, номер пункта), значение - список таблиц
        """
        if not changes_file or not os.path.exists(changes_file):
            return {}
        
        try:
            doc = Document(changes_file)
            tables_info = {}
            
            # Проходим по всем элементам документа
            prev_text = ""
            for element in doc.element.body:
                if element.tag.endswith('p'):  # Параграф
                    para_idx = doc.element.body.index(element)
                    if para_idx < len(doc.paragraphs):
                        para_text = doc.paragraphs[para_idx].text.strip()
                        if para_text:
                            # Сохраняем текст как контекст для следующей таблицы
                            prev_text = para_text
                
                elif element.tag.endswith('tbl'):  # Таблица
                    table_idx = sum(1 for i, e in enumerate(doc.element.body) 
                                   if e.tag.endswith('tbl') and i <= doc.element.body.index(element))
                    if table_idx < len(doc.tables):
                        table = doc.tables[table_idx]
                        
                        # Извлекаем данные таблицы
                        table_data = []
                        for row in table.rows:
                            row_data = [cell.text.strip() for cell in row.cells]
                            table_data.append(row_data)
                        
                        # Связываем таблицу с предшествующим текстом
                        key = prev_text[:100]  # Берем первые 100 символов как ключ
                        if key not in tables_info:
                            tables_info[key] = []
                        tables_info[key].append(table_data)
                        logger.info(f"📊 Извлечена таблица {table_idx} из инструкций (контекст: '{key[:50]}...')")
            
            return tables_info
        except Exception as e:
            logger.error(f"Ошибка при извлечении таблиц из инструкций: {e}")
            return {}
    
    def _extract_content_for_paragraph_replacement(
        self,
        changes_file: str,
        paragraph_num: str
    ) -> Dict[str, Any]:
        """
        Извлекает содержимое (таблицу или текст) для замены пункта из документа инструкций.
        Ищет текст после фразы "пункт X изложить в новой/следующей редакции".
        
        Returns:
            Словарь с ключами:
            - 'table_data': данные таблицы (если есть) или None
            - 'text_content': текстовое содержимое (если есть) или None
        """
        if not changes_file or not os.path.exists(changes_file):
            return {"table_data": None, "text_content": None}
        
        try:
            doc = Document(changes_file)
            result = {"table_data": None, "text_content": None}
            
            # Паттерн для поиска инструкции "пункт X изложить в новой/следующей редакции"
            # Учитываем возможные варианты: "редакции:", "редакции и далее", "редакции." и т.д.
            pattern = re.compile(
                rf'пункт[е]?\s+{re.escape(paragraph_num)}\s+изложить\s+в\s+(новой|следующей)\s+редакции[:\.,]?\s*(?:и\s+далее)?',
                re.IGNORECASE
            )
            
            # Ищем инструкцию в документе
            instruction_para_idx = None
            for idx, para in enumerate(doc.paragraphs):
                para_text = para.text.strip()
                if pattern.search(para_text):
                    instruction_para_idx = idx
                    logger.info(f"✅ Найдена инструкция для пункта {paragraph_num} в параграфе {idx}: '{para_text[:100]}...'")
                    break
            
            if instruction_para_idx is None:
                logger.warning(f"⚠️ Инструкция для пункта {paragraph_num} не найдена в документе инструкций")
                # Попробуем альтернативный поиск без строгого паттерна
                logger.info(f"🔍 Попытка альтернативного поиска инструкции для пункта {paragraph_num}...")
                for idx, para in enumerate(doc.paragraphs):
                    para_text = para.text.strip().lower()
                    if f"пункт {paragraph_num}" in para_text and "изложить" in para_text and "редакции" in para_text:
                        instruction_para_idx = idx
                        logger.info(f"✅ Найдена инструкция для пункта {paragraph_num} (альтернативный поиск) в параграфе {idx}")
                        break
                
                if instruction_para_idx is None:
                    logger.error(f"❌ Инструкция для пункта {paragraph_num} не найдена даже альтернативным поиском")
                    return result
            
            # Ищем содержимое после инструкции
            # Проходим по параграфам после инструкции
            text_parts = []
            
            for idx in range(instruction_para_idx + 1, len(doc.paragraphs)):
                para = doc.paragraphs[idx]
                para_text = para.text.strip()
                
                # Проверяем, не началась ли новая инструкция
                # Новая инструкция начинается с номера или содержит ключевые слова изменений
                if re.match(r'^\d+[\.\):]', para_text) or \
                   ("изложить" in para_text.lower() and "редакции" in para_text.lower()) or \
                   ("заменить" in para_text.lower() and idx != instruction_para_idx + 1) or \
                   ("удалить" in para_text.lower() and idx != instruction_para_idx + 1):
                    # Началась новая инструкция, останавливаемся
                    break
                
                # Добавляем текст (если не пустой)
                if para_text:
                    text_parts.append(para_text)
            
            # Проверяем, есть ли таблицы между параграфами
            # Ищем таблицы, которые находятся после инструкции
            # Для этого используем XML структуру документа
            instruction_para_element = doc.paragraphs[instruction_para_idx]._p
            parent = instruction_para_element.getparent()
            
            if parent is not None:
                # Находим позицию параграфа инструкции
                instruction_pos = None
                for i, elem in enumerate(parent):
                    if elem == instruction_para_element:
                        instruction_pos = i
                        break
                
                logger.info(f"🔍 Поиск таблиц после инструкции (позиция параграфа: {instruction_pos}, всего элементов: {len(parent)})")
                
                # Проверяем элементы после инструкции
                if instruction_pos is not None:
                    for i in range(instruction_pos + 1, len(parent)):
                        element = parent[i]
                        if element.tag.endswith('tbl'):
                            # Найдена таблица - извлекаем её
                            # Находим индекс таблицы в doc.tables
                            # Считаем все таблицы до текущей позиции
                            table_num = sum(1 for j, e in enumerate(parent[:i+1]) if e.tag.endswith('tbl')) - 1
                            logger.info(f"📋 Найдена таблица в XML на позиции {i}, индекс в doc.tables: {table_num}")
                            
                            if table_num < len(doc.tables):
                                table = doc.tables[table_num]
                                table_data = []
                                for row in table.rows:
                                    row_data = [cell.text.strip() for cell in row.cells]
                                    table_data.append(row_data)
                                result["table_data"] = table_data
                                logger.info(f"✅ Извлечена таблица для пункта {paragraph_num} ({len(table_data)} строк, {len(table_data[0]) if table_data else 0} столбцов)")
                            else:
                                logger.warning(f"⚠️ Индекс таблицы {table_num} выходит за пределы doc.tables (размер: {len(doc.tables)})")
                            # Таблица найдена, останавливаем поиск текста (таблица имеет приоритет)
                            break
                        elif element.tag.endswith('p'):
                            # Если встречаем параграф, проверяем, не началась ли новая инструкция
                            para_text = ""
                            for text_elem in element.iter():
                                if text_elem.tag.endswith('t'):
                                    para_text += text_elem.text or ""
                            para_text = para_text.strip()
                            
                            # Если это новая инструкция, останавливаем поиск
                            if re.match(r'^\d+[\.\):]', para_text) or \
                               ("изложить" in para_text.lower() and "редакции" in para_text.lower() and f"пункт {paragraph_num}" not in para_text.lower()):
                                logger.info(f"⏹️ Найдена новая инструкция в параграфе, останавливаем поиск таблицы: '{para_text[:50]}...'")
                                break
            
            # Альтернативный поиск: если таблица не найдена через XML, ищем первую таблицу после инструкции по индексу
            if result["table_data"] is None:
                logger.info(f"🔍 Альтернативный поиск таблицы: проверяем таблицы документа...")
                # Ищем первую таблицу, которая может быть связана с инструкцией
                # Для этого просто берём первую таблицу в документе после параграфа инструкции
                # (если таблиц несколько, берём ту, что ближе к инструкции)
                if doc.tables:
                    logger.info(f"📊 Найдено таблиц в документе: {len(doc.tables)}")
                    # Пока просто берём первую таблицу, если она есть
                    # В будущем можно улучшить логику поиска
                    if len(doc.tables) > 0:
                        table = doc.tables[0]
                        table_data = []
                        for row in table.rows:
                            row_data = [cell.text.strip() for cell in row.cells]
                            table_data.append(row_data)
                        result["table_data"] = table_data
                        logger.info(f"✅ Извлечена первая таблица документа для пункта {paragraph_num} (альтернативный метод)")
            
            # Если нет таблицы, используем текст
            if result["table_data"] is None and text_parts:
                result["text_content"] = "\n".join(text_parts)
                logger.info(f"✅ Извлечен текст для пункта {paragraph_num} ({len(result['text_content'])} символов)")
            elif result["table_data"]:
                # Если есть таблица, но есть и текст, сохраняем текст тоже (он может быть после таблицы)
                if text_parts:
                    result["text_content"] = "\n".join(text_parts)
                    logger.info(f"✅ Извлечен дополнительный текст для пункта {paragraph_num}")
            
            return result
        except Exception as e:
            logger.error(f"Ошибка при извлечении содержимого для замены пункта {paragraph_num}: {e}")
            return {"table_data": None, "text_content": None}
    
    async def _replace_entire_paragraph(
        self,
        filename: str,
        paragraph_num: str,
        new_content: Optional[str] = None,
        table_data: Optional[List[List[str]]] = None,
        changes_file: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Заменяет весь контент пункта, сохраняя номер пункта.
        
        Args:
            filename: Путь к файлу
            paragraph_num: Номер пункта (например, "7")
            new_content: Новый текстовый контент (если есть)
            table_data: Данные таблицы для вставки (если есть)
            changes_file: Путь к файлу инструкций для извлечения таблиц
            
        Returns:
            Результат операции
        """
        doc = Document(filename)
        
        # Ищем параграф с номером пункта
        target_para_idx = None
        for idx, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            # Проверяем различные форматы номера пункта
            patterns = [
                rf"^{re.escape(paragraph_num)}\.",
                rf"^{re.escape(paragraph_num)}\)",
                rf"^{re.escape(paragraph_num)}:",
            ]
            if any(re.match(p, para_text) for p in patterns):
                target_para_idx = idx
                logger.info(f"✅ Найден пункт {paragraph_num} в параграфе {idx}")
                break
        
        if target_para_idx is None:
            return {
                "success": False,
                "error": "PARAGRAPH_NOT_FOUND",
                "message": f"Пункт {paragraph_num} не найден в документе"
            }
        
        target_para = doc.paragraphs[target_para_idx]
        original_text = target_para.text
        
        # Определяем границы пункта
        start_idx = target_para_idx
        end_idx = start_idx + 1
        
        # Ищем конец пункта (следующий пункт или раздел)
        for idx in range(start_idx + 1, len(doc.paragraphs)):
            para = doc.paragraphs[idx]
            para_text = para.text.strip()
            # Проверяем, не начинается ли следующий параграф с номера
            if re.match(r'^\d+[\.\):]', para_text):
                end_idx = idx
                break
            # Если параграф пустой, это может быть конец
            if not para_text:
                end_idx = idx
                break
        
        # Удаляем содержимое пункта (кроме номера)
        # Сохраняем номер пункта
        num_pattern = None
        for pattern in [rf"^{re.escape(paragraph_num)}\.", rf"^{re.escape(paragraph_num)}\)", rf"^{re.escape(paragraph_num)}:"]:
            if re.match(pattern, original_text):
                num_pattern = re.match(pattern, original_text).group(0)
                break
        
        if not num_pattern:
            num_pattern = f"{paragraph_num}."
        
        # Удаляем параграфы пункта (кроме первого, где номер)
        for idx in range(end_idx - 1, start_idx, -1):
            DocumentChangeAgent._delete_paragraph(doc.paragraphs[idx])
        
        # Очищаем первый параграф, оставляя только номер
        for run in target_para.runs:
            run.text = ""
        target_para.add_run(num_pattern)
        
        # Если есть новая таблица, вставляем её после номера пункта
        if table_data:
            # Создаем таблицу
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]) if table_data else 0)
            
            # Заполняем таблицу
            for row_idx, row_data in enumerate(table_data):
                if row_idx < len(table.rows):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < len(table.rows[row_idx].cells):
                            table.rows[row_idx].cells[col_idx].text = str(cell_text) if cell_text else ""
            
            # Перемещаем таблицу после параграфа с номером
            table_element = table._element
            para_element = target_para._p
            parent = para_element.getparent()
            if parent is not None:
                # Удаляем таблицу из текущего места
                table_element.getparent().remove(table_element)
                # Вставляем после параграфа
                para_element.addnext(table_element)
                logger.info(f"✅ Таблица вставлена после пункта {paragraph_num}")
        
        # Если есть текстовый контент, добавляем его после номера (или после таблицы)
        if new_content:
            if table_data:
                # Текст добавляем после таблицы
                new_para = doc.add_paragraph(new_content)
                # Перемещаем параграф после таблицы
                para_element = target_para._p
                # Находим элемент таблицы
                table_element = None
                for sibling in para_element.itersiblings():
                    if sibling.tag.endswith('tbl'):
                        table_element = sibling
                        break
                if table_element is not None:
                    new_para_element = new_para._p
                    table_element.addnext(new_para_element)
            else:
                # Текст добавляем сразу после номера
                new_para = doc.add_paragraph(new_content)
                para_element = target_para._p
                new_para_element = new_para._p
                para_element.addnext(new_para_element)
            
            logger.info(f"✅ Текст вставлен после пункта {paragraph_num}")
        
        doc.save(filename)
        
        return {
            "success": True,
            "paragraph_index": start_idx,
            "message": f"Пункт {paragraph_num} заменен успешно"
        }
    
    async def _intelligent_paragraph_replacement(self, filename: str, target_text: str, new_text: str, description: str, matches: List, changes_file: Optional[str] = None) -> Dict[str, Any]:
        """
        Интеллектуальная замена в параграфах - заменяет содержимое, а не номер пункта.
        Использует LLM для поиска пункта и определения его местоположения.
        
        Args:
            filename: Путь к файлу
            target_text: Искомый текст (например, "согласовывается с ДО и ДРМ")
            new_text: Новый текст (например, "согласовывается с ДО")
            description: Описание инструкции
            matches: Найденные совпадения
            
        Returns:
            Результат операции
        """
        logger.info(f"🧠 ИНТЕЛЛЕКТУАЛЬНАЯ ЗАМЕНА В ПУНКТЕ:")
        logger.info(f"   Ищем: '{target_text}'")
        logger.info(f"   Заменяем на: '{new_text}'")
        logger.info(f"   Описание: '{description}'")
        
        try:
            # Открываем документ для анализа
            doc = Document(filename)
            replacements_made = 0
            
            # Извлекаем номер пункта из описания
            punkt_match = re.search(r'пункт[е]?\s+(\d+)', description, re.IGNORECASE)
            punkt_number = punkt_match.group(1) if punkt_match else None
            if not punkt_number and target_text:
                # Пробуем извлечь из target_text (может быть "7.")
                num_match = re.match(r'^(\d+)', target_text.replace(".", "").replace(")", "").replace(":", ""))
                if num_match:
                    punkt_number = num_match.group(1)
            
            logger.info(f"   Номер пункта из описания: {punkt_number}")
            
            # Проверяем, является ли это полной заменой пункта "Изложить пункт X в новой редакции"
            # Эту проверку делаем ПЕРВОЙ, так как для полной замены target_text может быть номером пункта
            is_full_replacement = (
                "изложить" in description.lower() and 
                "пункт" in description.lower() and 
                ("редакции" in description.lower() or "редакция" in description.lower())
            )
            
            # КРИТИЧЕСКАЯ ПРОВЕРКА: Если target_text является номером пункта, это ошибка
            # НО: для полной замены пункта (is_full_replacement) это допустимо - пропускаем проверку
            # Номер пункта НИКОГДА не должен быть target_text для замены, КРОМЕ случаев полной замены пункта
            if not is_full_replacement and punkt_number and target_text:
                # Проверяем, является ли target_text номером пункта (например, "32", "32.", "32)")
                target_clean = target_text.strip().replace(".", "").replace(")", "").replace(":", "").replace(" ", "")
                if target_clean == punkt_number:
                    logger.error(f"   ❌ ОШИБКА: target_text '{target_text}' является номером пункта {punkt_number}!")
                    logger.error(f"   ❌ Номер пункта НИКОГДА не должен быть target_text для замены!")
                    logger.error(f"   ❌ Нужно извлечь правильный target_text из описания: '{description}'")
                    return {
                        "success": False,
                        "error": "INVALID_TARGET_TEXT",
                        "message": f"target_text '{target_text}' является номером пункта, а не текстом для замены. Нужно извлечь правильный текст из описания."
                    }
                
                # Также проверяем, что target_text не начинается с номера пункта
                if target_text.strip().startswith(f"{punkt_number}.") or \
                   target_text.strip().startswith(f"{punkt_number})") or \
                   target_text.strip().startswith(f"{punkt_number}:"):
                    logger.error(f"   ❌ ОШИБКА: target_text '{target_text}' начинается с номера пункта!")
                    return {
                        "success": False,
                        "error": "INVALID_TARGET_TEXT",
                        "message": f"target_text '{target_text}' содержит номер пункта, который не должен заменяться."
                    }
            
            # Если это полная замена пункта, используем специальную функцию
            if is_full_replacement and punkt_number:
                logger.info(f"🔍 ПОЛНАЯ ЗАМЕНА ПУНКТА {punkt_number}: извлечение содержимого из инструкций")
                
                # Извлекаем содержимое (таблицу и/или текст) из инструкций
                table_data = None
                extracted_text = None
                if changes_file:
                    content = self._extract_content_for_paragraph_replacement(changes_file, punkt_number)
                    table_data = content.get("table_data")
                    extracted_text = content.get("text_content")
                    
                    if table_data:
                        logger.info(f"✅ Найдена таблица для пункта {punkt_number} в инструкциях")
                    if extracted_text:
                        logger.info(f"✅ Найден текст для пункта {punkt_number} в инструкциях ({len(extracted_text)} символов)")
                
                # Используем извлеченный текст, если он есть, иначе используем new_text из LLM
                final_text = extracted_text if extracted_text else (new_text if new_text else None)
                
                # Выполняем замену пункта
                result = await self._replace_entire_paragraph(
                    filename=filename,
                    paragraph_num=punkt_number,
                    new_content=final_text,
                    table_data=table_data,
                    changes_file=changes_file
                )
                
                if result.get("success"):
                    return result
                else:
                    logger.warning(f"⚠️ Замена всего пункта не удалась, продолжаем с обычной логикой: {result.get('message')}")
            
            # ИНТЕЛЛЕКТУАЛЬНЫЙ ПОИСК ПУНКТА ЧЕРЕЗ LLM
            llm_location = None
            if punkt_number:
                try:
                    llm_location = await self._find_paragraph_location_with_llm(
                        doc=doc,
                        description=description,
                        target_text=target_text,
                        punkt_number=punkt_number
                    )
                    if llm_location:
                        logger.info(f"   ✅ LLM определил местоположение пункта {punkt_number}: {llm_location['location_type']}")
                except Exception as e:
                    logger.warning(f"   ⚠️ Ошибка при LLM поиске пункта, продолжаем с алгоритмическим подходом: {e}")
            
            # Если LLM определил местоположение в таблице, используем эту информацию
            if llm_location and llm_location.get("location_type") == "table":
                table_idx = llm_location.get("table_index")
                row_idx = llm_location.get("row_index")
                
                if table_idx is not None and row_idx is not None and table_idx < len(doc.tables):
                    table = doc.tables[table_idx]
                    if row_idx < len(table.rows):
                        row = table.rows[row_idx]
                        logger.info(f"   📍 LLM указал: Таблица {table_idx}, Строка {row_idx}")
                        
                        # Определяем, в какой ячейке находится номер пункта, чтобы исключить её из поиска
                        punkt_cell_idx_llm = None
                        if punkt_number:
                            punkt_patterns_llm = [
                                f"{punkt_number}.",
                                f"{punkt_number})",
                                f"{punkt_number}:",
                                f"{punkt_number} ",
                                f"пункт {punkt_number}",
                                f"п.{punkt_number}",
                            ]
                            for idx, cell in enumerate(row.cells):
                                cell_text = cell.text.strip()
                                for pattern in punkt_patterns_llm:
                                    if re.match(rf"^{re.escape(pattern)}\b", cell_text) or re.search(rf"\b{re.escape(pattern)}\b", cell_text):
                                        punkt_cell_idx_llm = idx
                                        logger.info(f"   📍 Номер пункта найден в ячейке {idx}, исключаем её из поиска target_text (LLM)")
                                        break
                                if punkt_cell_idx_llm is not None:
                                    break
                        
                        # Ищем target_text в указанной строке
                        found_cell = None
                        cell_idx = llm_location.get("cell_index")
                        
                        if cell_idx is not None and cell_idx < len(row.cells):
                            # LLM указал конкретную ячейку, но проверяем, что это не ячейка с номером пункта
                            if punkt_cell_idx_llm is not None and cell_idx == punkt_cell_idx_llm:
                                logger.warning(f"   ⚠️ LLM указал ячейку {cell_idx} с номером пункта, ищем target_text в других ячейках")
                                found_cell = None
                            else:
                                found_cell = row.cells[cell_idx]
                                if target_text in found_cell.text:
                                    logger.info(f"   ✅ Найден target_text в указанной ячейке {cell_idx}")
                                else:
                                    logger.warning(f"   ⚠️ target_text не найден в указанной ячейке {cell_idx}, ищем во всей строке")
                                    found_cell = None
                        
                        if not found_cell:
                            # Ищем target_text во всех ячейках строки, ИСКЛЮЧАЯ ячейку с номером пункта
                            for idx, cell in enumerate(row.cells):
                                # Пропускаем ячейку с номером пункта
                                if punkt_cell_idx_llm is not None and idx == punkt_cell_idx_llm:
                                    logger.info(f"   ⏭️ Пропускаем ячейку {idx} (содержит номер пункта {punkt_number})")
                                    continue
                                
                                if target_text in cell.text:
                                    found_cell = cell
                                    cell_idx = idx
                                    logger.info(f"   ✅ Найден target_text в ячейке {idx}")
                                    break
                        
                        if found_cell:
                            # Определяем тип замены
                            use_structured = self._should_use_structured_replacement(description)
                            table_location = None
                            
                            if not use_structured:
                                # Простая замена
                                para = found_cell.paragraphs[0] if found_cell.paragraphs else None
                                if self._smart_replace_in_paragraph(para, target_text, new_text, cell=found_cell, punkt_number=punkt_number):
                                    replacements_made += 1
                                    logger.info(f"   ✅ Простая замена выполнена в ячейке Table {table_idx}, Row {row_idx}, Column {cell_idx} (LLM)")
                                    table_paragraph_index = self._find_paragraph_for_table(doc, table_idx)
                                    if table_paragraph_index >= 0:
                                        table_location = {
                                            "table_idx": table_idx,
                                            "row_idx": row_idx,
                                            "cell_idx": cell_idx,
                                            "paragraph_index": table_paragraph_index
                                        }
                            else:
                                # Структурированная замена
                                row_structure = self._analyze_row_structure(row, row_idx)
                                table_context = self._get_table_context(table, row_idx)
                                distribution = await self._map_new_text_to_structure(
                                    new_text=new_text,
                                    target_text=target_text,
                                    row_structure=row_structure,
                                    description=description,
                                    table_context=table_context
                                )
                                if self._apply_structured_replacement(row, target_text, distribution):
                                    replacements_made += 1
                                    logger.info(f"   ✅ Структурированная замена выполнена в строке {row_idx} таблицы {table_idx} (LLM)")
                                    table_paragraph_index = self._find_paragraph_for_table(doc, table_idx)
                                    if table_paragraph_index >= 0:
                                        table_location = {
                                            "table_idx": table_idx,
                                            "row_idx": row_idx,
                                            "cell_idx": cell_idx,
                                            "paragraph_index": table_paragraph_index
                                        }
                            
                            # Сохраняем результат
                            if replacements_made > 0:
                                doc.save(filename)
                                result = {
                                    "success": True,
                                    "message": f"Замена выполнена (LLM: таблица {table_idx}, строка {row_idx})",
                                    "replacements_made": replacements_made,
                                    "method": "llm_guided_table_replace" if use_structured else "llm_guided_cell_replace",
                                    "is_table_change": True,
                                }
                                if table_location:
                                    result["table_location"] = table_location
                                    result["paragraph_index"] = table_location.get("paragraph_index", -1)
                                return result
            
            # Если LLM определил местоположение в параграфе, проверяем сначала таблицы, если matches содержат таблицы
            # Это важно, потому что LLM может ошибиться, и пункт может быть в таблице
            elif llm_location and llm_location.get("location_type") == "paragraph":
                # Если matches содержат таблицы и есть punkt_number, сначала проверяем таблицы
                table_matches_check = [m for m in matches if hasattr(m, 'paragraph_index') and m.paragraph_index == -1] if matches else []
                if table_matches_check and punkt_number:
                    logger.info(f"   ⚠️ LLM определил параграф, но matches содержат таблицы, сначала проверяем таблицы")
                    # Пропускаем проверку параграфа LLM и продолжаем к поиску в таблицах
                else:
                    para_idx = llm_location.get("paragraph_index")
                    if para_idx is not None and para_idx < len(doc.paragraphs):
                        para = doc.paragraphs[para_idx]
                        if target_text in para.text:
                            logger.info(f"   📍 LLM указал: Параграф {para_idx}")
                            if self._smart_replace_in_paragraph(para, target_text, new_text, punkt_number=punkt_number):
                                replacements_made += 1
                                logger.info(f"   ✅ Замена выполнена в параграфе {para_idx} (LLM)")
                                doc.save(filename)
                                return {
                                    "success": True,
                                    "message": f"Замена выполнена (LLM: параграф {para_idx})",
                                    "replacements_made": replacements_made,
                                    "method": "llm_guided_paragraph_replace",
                                    "is_table_change": False,
                                }
            
            # Если LLM не определил или не использовался, продолжаем с алгоритмическим подходом
            # ПРИОРИТЕТ: Если есть punkt_number и matches содержат таблицы, сначала проверяем таблицы
            table_matches_priority = [m for m in matches if hasattr(m, 'paragraph_index') and m.paragraph_index == -1] if matches else []
            should_check_tables_first = punkt_number and table_matches_priority and replacements_made == 0
            
            # Если нужно проверить таблицы в первую очередь, делаем это перед поиском в параграфах
            table_location_priority = None
            if should_check_tables_first:
                logger.info(f"   🔍 ПРИОРИТЕТ: Сначала проверяем таблицы для пункта {punkt_number} (matches содержат таблицы)")
                for table_idx, table in enumerate(doc.tables):
                    for row_idx, row in enumerate(table.rows):
                        # Проверяем, содержит ли строка номер пункта
                        row_contains_punkt = False
                        punkt_patterns = [
                            f"{punkt_number}.",
                            f"{punkt_number})",
                            f"{punkt_number}:",
                            f"{punkt_number} ",
                            f"пункт {punkt_number}",
                            f"п.{punkt_number}",
                        ]
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            for pattern in punkt_patterns:
                                # Используем точное сопоставление в начале ячейки или после пробела
                                if re.match(rf"^{re.escape(pattern)}\b", cell_text) or re.search(rf"\b{re.escape(pattern)}\b", cell_text):
                                    row_contains_punkt = True
                                    logger.info(f"   ✅ Строка {row_idx} в таблице {table_idx} содержит номер пункта {punkt_number} (паттерн: '{pattern}')")
                                    break
                            if row_contains_punkt:
                                break
                        
                        # Определяем, в какой строке находится пункт (может быть в текущей или соседней)
                        punkt_row_idx = row_idx if row_contains_punkt else None
                        
                        # Если не нашли в текущей строке, проверяем соседние строки
                        if not row_contains_punkt:
                            for offset in [-1, 1]:
                                check_row_idx = row_idx + offset
                                if 0 <= check_row_idx < len(table.rows):
                                    check_row = table.rows[check_row_idx]
                                    for cell in check_row.cells:
                                        cell_text = cell.text.strip()
                                        for pattern in punkt_patterns:
                                            if re.match(rf"^{re.escape(pattern)}\b", cell_text) or re.search(rf"\b{re.escape(pattern)}\b", cell_text):
                                                row_contains_punkt = True
                                                punkt_row_idx = check_row_idx
                                                logger.info(f"   ✅ Пункт {punkt_number} найден в соседней строке {check_row_idx} таблицы {table_idx}")
                                                break
                                        if row_contains_punkt:
                                            break
                                    if row_contains_punkt:
                                        break
                        
                        # Если нашли строку с пунктом, ищем target_text ТОЛЬКО в ячейках этой строки
                        if row_contains_punkt and punkt_row_idx is not None:
                            punkt_row = table.rows[punkt_row_idx]
                            logger.info(f"   🔍 Поиск target_text '{target_text}' в ячейках строки {punkt_row_idx} таблицы {table_idx} (пункт {punkt_number})")
                            
                            # Определяем, в какой ячейке находится номер пункта (обычно первая ячейка)
                            punkt_cell_idx = None
                            for cell_idx, cell in enumerate(punkt_row.cells):
                                cell_text = cell.text.strip()
                                for pattern in punkt_patterns:
                                    if re.match(rf"^{re.escape(pattern)}\b", cell_text) or re.search(rf"\b{re.escape(pattern)}\b", cell_text):
                                        punkt_cell_idx = cell_idx
                                        logger.info(f"   📍 Номер пункта найден в ячейке {cell_idx}, исключаем её из поиска target_text")
                                        break
                                if punkt_cell_idx is not None:
                                    break
                            
                            # Ищем target_text только в ячейках БЕЗ номера пункта
                            for cell_idx, cell in enumerate(punkt_row.cells):
                                # Пропускаем ячейку с номером пункта
                                if punkt_cell_idx is not None and cell_idx == punkt_cell_idx:
                                    logger.info(f"   ⏭️ Пропускаем ячейку {cell_idx} (содержит номер пункта {punkt_number})")
                                    continue
                                
                                cell_text = cell.text
                                # Проверяем, что target_text не является частью номера пункта
                                # (например, если target_text = "32", не заменяем в ячейке с "32.")
                                is_target_part_of_punkt_number = False
                                if punkt_number:
                                    punkt_variants = [f"{punkt_number}.", f"{punkt_number})", f"{punkt_number}:", f"{punkt_number} ", punkt_number]
                                    for variant in punkt_variants:
                                        if variant in cell_text:
                                            # Проверяем, что target_text является частью номера пункта, а не отдельным текстом
                                            if target_text == punkt_number or target_text in variant:
                                                is_target_part_of_punkt_number = True
                                                logger.info(f"   ⏭️ target_text '{target_text}' является частью номера пункта в ячейке {cell_idx}, пропускаем")
                                                break
                                
                                if not is_target_part_of_punkt_number and target_text in cell_text:
                                    logger.info(f"   ✅ Найдена ячейка Table {table_idx}, Row {punkt_row_idx}, Column {cell_idx} с текстом '{target_text[:50]}...' (пункт {punkt_number})")
                                    use_structured = self._should_use_structured_replacement(description)
                                    
                                    if not use_structured:
                                        para = cell.paragraphs[0] if cell.paragraphs else None
                                        if self._smart_replace_in_paragraph(para, target_text, new_text, cell=cell, punkt_number=punkt_number):
                                            replacements_made += 1
                                            logger.info(f"   ✅ Простая замена выполнена в ячейке Table {table_idx}, Row {punkt_row_idx}, Column {cell_idx}")
                                            if table_location_priority is None:
                                                table_paragraph_index = self._find_paragraph_for_table(doc, table_idx)
                                                if table_paragraph_index >= 0:
                                                    table_location_priority = {
                                                        "table_idx": table_idx,
                                                        "row_idx": punkt_row_idx,
                                                        "cell_idx": cell_idx,
                                                        "paragraph_index": table_paragraph_index
                                                    }
                                            break
                                        else:
                                            row_structure = self._analyze_row_structure(punkt_row, punkt_row_idx)
                                            table_context = self._get_table_context(table, punkt_row_idx)
                                            distribution = await self._map_new_text_to_structure(
                                                new_text=new_text,
                                                target_text=target_text,
                                                row_structure=row_structure,
                                                description=description,
                                                table_context=table_context
                                            )
                                            if self._apply_structured_replacement(punkt_row, target_text, distribution):
                                                replacements_made += 1
                                                logger.info(f"   ✅ Структурированная замена выполнена в строке {punkt_row_idx} таблицы {table_idx}")
                                                if table_location_priority is None:
                                                    table_paragraph_index = self._find_paragraph_for_table(doc, table_idx)
                                                    if table_paragraph_index >= 0:
                                                        table_location_priority = {
                                                            "table_idx": table_idx,
                                                            "row_idx": punkt_row_idx,
                                                            "cell_idx": cell_idx,
                                                            "paragraph_index": table_paragraph_index
                                                        }
                                                break
                            if replacements_made > 0:
                                break
                    if replacements_made > 0:
                        break
                
                # Если замена выполнена в таблице, возвращаем результат
                if replacements_made > 0:
                    doc.save(filename)
                    result = {
                        "success": True,
                        "message": f"Замена выполнена в таблице (пункт {punkt_number})",
                        "replacements_made": replacements_made,
                        "method": "priority_table_replace",
                        "is_table_change": True,
                    }
                    if table_location_priority:
                        result["table_location"] = table_location_priority
                        result["paragraph_index"] = table_location_priority.get("paragraph_index", -1)
                    return result
            
            # Ищем параграфы, которые начинаются с этого номера пункта
            target_paragraphs = []
            for idx, para in enumerate(doc.paragraphs):
                para_text = para.text.strip()
                if punkt_number and para_text.startswith(f"{punkt_number}."):
                    target_paragraphs.append((idx, para))
                    logger.info(f"   Найден пункт {punkt_number} в параграфе {idx}: '{para_text[:50]}...'")
            
            # Если не нашли по номеру, используем matches
            if not target_paragraphs and matches:
                for match in matches:
                    if hasattr(match, 'paragraph_index') and match.paragraph_index >= 0 and match.paragraph_index < len(doc.paragraphs):
                        para = doc.paragraphs[match.paragraph_index]
                        target_paragraphs.append((match.paragraph_index, para))
                        logger.info(f"   Используем параграф из match: {match.paragraph_index}")
            
            # Если все еще не нашли и есть номер пункта, ищем текст в параграфах с этим номером
            if not target_paragraphs and punkt_number:
                logger.info(f"   🔍 Поиск текста '{target_text}' в параграфах с номером пункта {punkt_number}")
                for idx, para in enumerate(doc.paragraphs):
                    para_text = para.text.strip()
                    # Проверяем, что параграф начинается с номера пункта и содержит target_text
                    if para_text.startswith(f"{punkt_number}.") and target_text in para_text:
                        target_paragraphs.append((idx, para))
                        logger.info(f"   ✅ Найден параграф {idx} с номером {punkt_number} и текстом '{target_text[:50]}...'")
            
            # Если все еще не нашли, ищем target_text во всех параграфах
            if not target_paragraphs:
                logger.info(f"   🔍 Поиск текста '{target_text}' во всех параграфах")
                for idx, para in enumerate(doc.paragraphs):
                    if target_text in para.text:
                        target_paragraphs.append((idx, para))
                        logger.info(f"   ✅ Найден параграф {idx} с текстом '{target_text[:50]}...'")
                        # Ограничиваем поиск первым найденным, если есть номер пункта
                        if punkt_number:
                            break
            
            # Если текст найден в таблице (matches содержат таблицы), ищем в таблицах
            if not target_paragraphs and matches:
                table_matches = [m for m in matches if hasattr(m, 'paragraph_index') and m.paragraph_index == -1]
                if table_matches and punkt_number:
                    logger.info(f"   🔍 Текст найден в таблице, ищем параграф с номером пункта {punkt_number}")
                    # Ищем параграф, который начинается с номера пункта
                    for idx, para in enumerate(doc.paragraphs):
                        para_text = para.text.strip()
                        if para_text.startswith(f"{punkt_number}."):
                            # Проверяем, содержит ли этот параграф или следующий target_text
                            # (текст может быть в следующем параграфе после номера)
                            check_paras = [para]
                            if idx + 1 < len(doc.paragraphs):
                                check_paras.append(doc.paragraphs[idx + 1])
                            
                            for check_para in check_paras:
                                if target_text in check_para.text:
                                    target_paragraphs.append((doc.paragraphs.index(check_para), check_para))
                                    logger.info(f"   ✅ Найден параграф {doc.paragraphs.index(check_para)} с номером {punkt_number} и текстом '{target_text[:50]}...'")
                                    break
                            
                            if target_paragraphs:
                                break
            
            # Заменяем в найденных параграфах
            for para_idx, para in target_paragraphs:
                if target_text in para.text:
                    logger.info(f"   Замена в параграфе {para_idx}: '{target_text}' → '{new_text}'")
                    if self._smart_replace_in_paragraph(para, target_text, new_text, punkt_number=punkt_number):
                        replacements_made += 1
                        logger.info(f"   ✅ Успешная замена в параграфе {para_idx}")
                        # Если есть номер пункта, обрабатываем только первый найденный
                        if punkt_number:
                            break
            
            # Если не нашли в параграфах, но текст был найден в таблице, ищем в таблицах
            table_location = None  # Для сохранения информации о местоположении замены в таблице
            if replacements_made == 0 and matches:
                table_matches = [m for m in matches if hasattr(m, 'paragraph_index') and m.paragraph_index == -1]
                if table_matches and punkt_number:
                    logger.info(f"   🔍 Текст найден в таблице, ищем в таблицах с номером пункта {punkt_number}")
                    # Улучшенная логика: находим ячейки с target_text, затем проверяем, содержит ли их строка номер пункта
                    for table_idx, table in enumerate(doc.tables):
                        for row_idx, row in enumerate(table.rows):
                            # Сначала проверяем, содержит ли строка номер пункта (в любой ячейке)
                            # Проверяем различные форматы: "32.", "32)", "32:", "32", "пункт 32" и т.д.
                            row_contains_punkt = False
                            punkt_patterns = [
                                f"{punkt_number}.",
                                f"{punkt_number})",
                                f"{punkt_number}:",
                                f"{punkt_number} ",
                                f"пункт {punkt_number}",
                                f"п.{punkt_number}",
                            ]
                            for cell in row.cells:
                                cell_text = cell.text
                                for pattern in punkt_patterns:
                                    if pattern in cell_text:
                                        row_contains_punkt = True
                                        logger.info(f"   ✅ Строка {row_idx} в таблице {table_idx} содержит номер пункта {punkt_number} (паттерн: '{pattern}')")
                                        break
                                if row_contains_punkt:
                                    break
                            
                            # Если строка содержит номер пункта, ищем target_text ТОЛЬКО в ячейках этой строки
                            if row_contains_punkt:
                                logger.info(f"   🔍 Поиск target_text '{target_text}' в ячейках строки {row_idx} таблицы {table_idx} (пункт {punkt_number})")
                                
                                # Определяем, в какой ячейке находится номер пункта (обычно первая ячейка)
                                punkt_cell_idx = None
                                for check_cell_idx, check_cell in enumerate(row.cells):
                                    check_cell_text = check_cell.text.strip()
                                    for pattern in punkt_patterns:
                                        if re.match(rf"^{re.escape(pattern)}\b", check_cell_text) or re.search(rf"\b{re.escape(pattern)}\b", check_cell_text):
                                            punkt_cell_idx = check_cell_idx
                                            logger.info(f"   📍 Номер пункта найден в ячейке {check_cell_idx}, исключаем её из поиска target_text")
                                            break
                                    if punkt_cell_idx is not None:
                                        break
                                
                                # Ищем target_text только в ячейках БЕЗ номера пункта
                                for cell_idx, cell in enumerate(row.cells):
                                    # Пропускаем ячейку с номером пункта
                                    if punkt_cell_idx is not None and cell_idx == punkt_cell_idx:
                                        logger.info(f"   ⏭️ Пропускаем ячейку {cell_idx} (содержит номер пункта {punkt_number})")
                                        continue
                                    
                                    cell_text = cell.text
                                    # Проверяем, что target_text не является частью номера пункта
                                    # (например, если target_text = "32", не заменяем в ячейке с "32.")
                                    is_target_part_of_punkt_number = False
                                    if punkt_number:
                                        punkt_variants = [f"{punkt_number}.", f"{punkt_number})", f"{punkt_number}:", f"{punkt_number} ", punkt_number]
                                        for variant in punkt_variants:
                                            if variant in cell_text:
                                                # Проверяем, что target_text является частью номера пункта, а не отдельным текстом
                                                if target_text == punkt_number or target_text in variant:
                                                    is_target_part_of_punkt_number = True
                                                    logger.info(f"   ⏭️ target_text '{target_text}' является частью номера пункта в ячейке {cell_idx}, пропускаем")
                                                    break
                                    
                                    if not is_target_part_of_punkt_number and target_text in cell_text:
                                        logger.info(f"   ✅ Найдена ячейка Table {table_idx}, Row {row_idx}, Column {cell_idx} с текстом '{target_text[:50]}...' (пункт {punkt_number})")
                                        # Определяем тип замены на основе описания инструкции
                                        use_structured = self._should_use_structured_replacement(description)
                                        
                                        if not use_structured:
                                            # Простая замена фразы в найденной ячейке
                                            logger.info(f"   🔄 Простая замена фразы в ячейке (не распределение по столбцам)")
                                            para = cell.paragraphs[0] if cell.paragraphs else None
                                            if self._smart_replace_in_paragraph(para, target_text, new_text, cell=cell, punkt_number=punkt_number):
                                                replacements_made += 1
                                                logger.info(f"   ✅ Простая замена выполнена в ячейке Table {table_idx}, Row {row_idx}, Column {cell_idx}")
                                                # Сохраняем информацию о местоположении для аннотаций
                                                if table_location is None:
                                                    table_paragraph_index = self._find_paragraph_for_table(doc, table_idx)
                                                    if table_paragraph_index >= 0:
                                                        table_location = {
                                                            "table_idx": table_idx,
                                                            "row_idx": row_idx,
                                                            "cell_idx": cell_idx,
                                                            "paragraph_index": table_paragraph_index
                                                        }
                                                        logger.info(f"   📍 Сохранено местоположение для аннотации: Table {table_idx}, Row {row_idx}, Para {table_paragraph_index}")
                                                break
                                        else:
                                            # Структурированная замена (распределение по столбцам)
                                            logger.info(f"   🔄 Структурированная замена (распределение по столбцам)")
                                            # Анализируем структуру строки
                                            row_structure = self._analyze_row_structure(row, row_idx)
                                            # Получаем контекст таблицы для LLM
                                            table_context = self._get_table_context(table, row_idx)
                                            # Распределяем новый текст по структуре (алгоритм + LLM проверка)
                                            distribution = await self._map_new_text_to_structure(
                                                new_text=new_text,
                                                target_text=target_text,
                                                row_structure=row_structure,
                                                description=description,
                                                table_context=table_context
                                            )
                                            # Применяем структурированную замену
                                            if self._apply_structured_replacement(row, target_text, distribution):
                                                replacements_made += 1
                                                logger.info(f"   ✅ Структурированная замена выполнена в строке {row_idx} таблицы {table_idx}")
                                                # Сохраняем информацию о местоположении для аннотаций
                                                if table_location is None:
                                                    table_paragraph_index = self._find_paragraph_for_table(doc, table_idx)
                                                    if table_paragraph_index >= 0:
                                                        table_location = {
                                                            "table_idx": table_idx,
                                                            "row_idx": row_idx,
                                                            "cell_idx": cell_idx,
                                                            "paragraph_index": table_paragraph_index
                                                        }
                                                        logger.info(f"   📍 Сохранено местоположение для аннотации: Table {table_idx}, Row {row_idx}, Para {table_paragraph_index}")
                                                break
                                if replacements_made > 0:
                                    break
                        if replacements_made > 0:
                            break
                    
                    # Если не нашли через проверку строк, пробуем поискать ячейки, содержащие оба текста
                    if replacements_made == 0:
                        logger.info(f"   🔍 Поиск ячеек, содержащих и номер пункта, и target_text")
                        for table_idx, table in enumerate(doc.tables):
                            for row_idx, row in enumerate(table.rows):
                                for cell_idx, cell in enumerate(row.cells):
                                    cell_text = cell.text
                                    # Проверяем, содержит ли ячейка номер пункта и target_text
                                    if f"{punkt_number}." in cell_text and target_text in cell_text:
                                        logger.info(f"   ✅ Найдена ячейка Table {table_idx}, Row {row_idx}, Column {cell_idx} с номером {punkt_number} и текстом '{target_text[:50]}...'")
                                        # Определяем тип замены на основе описания инструкции
                                        use_structured = self._should_use_structured_replacement(description)
                                        
                                        if not use_structured:
                                            # Простая замена фразы в найденной ячейке
                                            logger.info(f"   🔄 Простая замена фразы в ячейке (не распределение по столбцам)")
                                            para = cell.paragraphs[0] if cell.paragraphs else None
                                            if self._smart_replace_in_paragraph(para, target_text, new_text, cell=cell, punkt_number=punkt_number):
                                                replacements_made += 1
                                                logger.info(f"   ✅ Простая замена выполнена в ячейке Table {table_idx}, Row {row_idx}, Column {cell_idx}")
                                                # Сохраняем информацию о местоположении для аннотаций
                                                if table_location is None:
                                                    table_paragraph_index = self._find_paragraph_for_table(doc, table_idx)
                                                    if table_paragraph_index >= 0:
                                                        table_location = {
                                                            "table_idx": table_idx,
                                                            "row_idx": row_idx,
                                                            "cell_idx": cell_idx,
                                                            "paragraph_index": table_paragraph_index
                                                        }
                                                        logger.info(f"   📍 Сохранено местоположение для аннотации: Table {table_idx}, Row {row_idx}, Para {table_paragraph_index}")
                                                break
                                        else:
                                            # Структурированная замена (распределение по столбцам)
                                            logger.info(f"   🔄 Структурированная замена (распределение по столбцам)")
                                            # Анализируем структуру строки
                                            row_structure = self._analyze_row_structure(row, row_idx)
                                            # Получаем контекст таблицы для LLM
                                            table_context = self._get_table_context(table, row_idx)
                                            # Распределяем новый текст по структуре (алгоритм + LLM проверка)
                                            distribution = await self._map_new_text_to_structure(
                                                new_text=new_text,
                                                target_text=target_text,
                                                row_structure=row_structure,
                                                description=description,
                                                table_context=table_context
                                            )
                                            # Применяем структурированную замену
                                            if self._apply_structured_replacement(row, target_text, distribution):
                                                replacements_made += 1
                                                logger.info(f"   ✅ Структурированная замена выполнена в строке {row_idx} таблицы {table_idx}")
                                                # Сохраняем информацию о местоположении для аннотаций
                                                if table_location is None:
                                                    table_paragraph_index = self._find_paragraph_for_table(doc, table_idx)
                                                    if table_paragraph_index >= 0:
                                                        table_location = {
                                                            "table_idx": table_idx,
                                                            "row_idx": row_idx,
                                                            "cell_idx": cell_idx,
                                                            "paragraph_index": table_paragraph_index
                                                        }
                                                        logger.info(f"   📍 Сохранено местоположение для аннотации: Table {table_idx}, Row {row_idx}, Para {table_paragraph_index}")
                                                break
                                if replacements_made > 0:
                                    break
                            if replacements_made > 0:
                                break
            
            if replacements_made > 0:
                doc.save(filename)
                result = {
                    "success": True,
                    "message": f"Интеллектуальная замена в пункте выполнена в {replacements_made} параграфах/ячейках",
                    "replacements_made": replacements_made,
                    "method": "intelligent_paragraph_replace",
                    "is_table_change": False,  # По умолчанию не табличное изменение
                    "details": {}
                }
                
                # Проверяем, было ли изменение в таблице (для аннотаций)
                if table_location:
                    # Если изменение было в таблице, добавляем информацию о местоположении
                    result["is_table_change"] = True
                    result["details"]["is_table_change"] = True
                    result["details"]["table_location"] = table_location
                    result["paragraph_index"] = table_location.get("paragraph_index", -1)
                
                return result
            else:
                return {
                    "success": False,
                    "error": "NO_PARAGRAPH_REPLACEMENTS",
                    "message": f"Текст '{target_text}' не найден в пунктах"
                }
                
        except Exception as e:
            logger.error(f"Ошибка интеллектуальной замены в пункте: {e}")
            return {
                "success": False,
                "error": "INTELLIGENT_PARAGRAPH_ERROR",
                "message": f"Ошибка интеллектуальной замены в пункте: {e}"
            }

    def _is_target_part_of_punkt_number(self, target_text: str, cell_text: str, punkt_number: Optional[str] = None) -> bool:
        """
        Проверяет, является ли target_text частью номера пункта в cell_text.
        
        Args:
            target_text: Текст для поиска
            cell_text: Текст ячейки
            punkt_number: Номер пункта (опционально)
            
        Returns:
            True если target_text является частью номера пункта
        """
        
        # Если указан номер пункта, проверяем напрямую
        if punkt_number:
            punkt_variants = [
                f"{punkt_number}.",
                f"{punkt_number})",
                f"{punkt_number}:",
                f"{punkt_number} ",
                punkt_number
            ]
            target_clean = target_text.strip().replace(".", "").replace(")", "").replace(":", "").replace(" ", "")
            
            # Проверяем, совпадает ли target_text с номером пункта
            if target_clean == punkt_number:
                return True
            
            # Проверяем, является ли target_text частью варианта номера пункта
            for variant in punkt_variants:
                if target_text == variant or target_text in variant:
                    # Дополнительная проверка: убеждаемся, что это действительно номер пункта в тексте
                    if variant in cell_text:
                        # Проверяем, что это в начале строки или после пробела
                        if re.match(rf"^{re.escape(variant)}\b", cell_text.strip()) or \
                           re.search(rf"\b{re.escape(variant)}\b", cell_text.strip()):
                            return True
        
        # Если punkt_number не передан, пытаемся определить его из cell_text
        if not punkt_number:
            punkt_match = re.match(r'^(\d+[\.\):]?\s*)', cell_text.strip())
            if punkt_match:
                punkt_prefix = punkt_match.group(1)
                punkt_num_match = re.match(r'^(\d+)', punkt_prefix)
                if punkt_num_match:
                    punkt_number = punkt_num_match.group(1)
        
        # Если теперь есть punkt_number, проверяем еще раз
        if punkt_number:
            punkt_variants = [
                f"{punkt_number}.",
                f"{punkt_number})",
                f"{punkt_number}:",
                f"{punkt_number} ",
                punkt_number
            ]
            target_clean = target_text.strip().replace(".", "").replace(")", "").replace(":", "").replace(" ", "")
            punkt_clean = punkt_number.strip().replace(".", "").replace(")", "").replace(":", "").replace(" ", "")
            
            # Если target_text равен номеру пункта
            if target_clean == punkt_clean:
                return True
            
            # Проверяем, является ли target_text частью варианта номера пункта
            for variant in punkt_variants:
                if target_text == variant or target_text in variant:
                    # Дополнительная проверка: убеждаемся, что это действительно номер пункта в тексте
                    if variant in cell_text:
                        # Проверяем, что это в начале строки или после пробела
                        if re.match(rf"^{re.escape(variant)}\b", cell_text.strip()) or \
                           re.search(rf"\b{re.escape(variant)}\b", cell_text.strip()):
                            return True
        
        # Проверяем, начинается ли cell_text с номера пункта, содержащего target_text
        punkt_match = re.match(r'^(\d+[\.\):]?\s*)', cell_text.strip())
        if punkt_match:
            punkt_prefix = punkt_match.group(1)
            punkt_num_match = re.match(r'^(\d+)', punkt_prefix)
            if punkt_num_match:
                punkt_num = punkt_num_match.group(1)
                target_clean = target_text.strip().replace(".", "").replace(")", "").replace(":", "").replace(" ", "")
                if target_clean == punkt_num or target_text in punkt_prefix:
                    return True
        
        return False

    def _smart_replace_in_paragraph(self, paragraph, old: str, new: str, cell=None, punkt_number: Optional[str] = None) -> bool:
        """
        Умная замена в параграфе - заменяет только содержимое, не трогая номер пункта.
        
        Args:
            paragraph: Параграф для замены (может быть None, если передан cell)
            old: Старый текст
            new: Новый текст
            cell: Ячейка таблицы (опционально)
            
        Returns:
            True если замена была выполнена
        """
        replaced = False
        
        # Если передан cell, работаем с его параграфами
        if cell is not None:
            # КРИТИЧЕСКАЯ ПРОВЕРКА: Убеждаемся, что target_text не является частью номера пункта
            if punkt_number and self._is_target_part_of_punkt_number(old, cell.text, punkt_number):
                logger.warning(f"   ⚠️ ПРОПУСК: target_text '{old}' является частью номера пункта {punkt_number} в ячейке, не заменяем!")
                return False
            
            # Проверяем все параграфы в ячейке
            if cell.paragraphs:
                # Сначала пробуем найти текст в каждом параграфе и заменить его
                for para in cell.paragraphs:
                    if old in para.text:
                        # Текст найден в этом параграфе
                        original_text = para.text
                        
                        # Проверяем, начинается ли параграф с номера пункта
                        # Более точная проверка номера пункта: "32.", "32)", "32:", "32 "
                        punkt_match = re.match(r'^(\d+[\.\):]?\s*)', original_text)
                        
                        if punkt_match:
                            # Параграф начинается с номера - заменяем только в содержимой части
                            punkt_prefix = punkt_match.group(1)
                            content_part = original_text[len(punkt_prefix):]
                            
                            # Проверяем, что old не является частью номера пункта
                            punkt_number_match = re.match(r'^(\d+)', punkt_prefix)
                            if punkt_number_match:
                                punkt_num = punkt_number_match.group(1)
                                # Если old равен номеру пункта или является его частью, пропускаем замену
                                if old == punkt_num or old == punkt_prefix.strip() or self._is_target_part_of_punkt_number(old, original_text, punkt_number):
                                    logger.warning(f"   ⚠️ target_text '{old}' является номером пункта, пропускаем замену в этой ячейке")
                                    continue
                            
                            if old in content_part:
                                new_content = content_part.replace(old, new)
                                new_full_text = punkt_prefix + new_content
                                
                                logger.info(f"   Умная замена в ячейке (с номером пункта):")
                                logger.info(f"     Префикс пункта: '{punkt_prefix}'")
                                logger.info(f"     Старое содержимое: '{content_part[:50]}...'")
                                logger.info(f"     Новое содержимое: '{new_content[:50]}...'")
                                
                                # Заменяем через runs для сохранения форматирования, но только в content_part
                                found_in_runs = False
                                for run in para.runs:
                                    run_text = run.text
                                    # Проверяем, что run не содержит только номер пункта
                                    if old in run_text and not run_text.strip().startswith(punkt_prefix.strip()):
                                        run.text = run_text.replace(old, new)
                                        found_in_runs = True
                                
                                # Если не удалось через runs, заменяем весь текст
                                if not found_in_runs:
                                    para.text = new_full_text
                                
                                replaced = True
                                break
                        else:
                            # Обычная замена для параграфов без номера пункта
                            # Пробуем через runs сначала
                            found_in_runs = False
                            for run in para.runs:
                                if old in run.text:
                                    run.text = run.text.replace(old, new)
                                    found_in_runs = True
                                    replaced = True
                            
                            # Если не удалось через runs, заменяем весь текст параграфа
                            if not found_in_runs and old in para.text:
                                para.text = para.text.replace(old, new)
                                replaced = True
                            
                            if replaced:
                                logger.info(f"   Умная замена в ячейке (обычный параграф): '{old}' → '{new}'")
                                break
                
                # Если не нашли текст в отдельных параграфах, пробуем заменить во всей ячейке
                if not replaced and old in cell.text:
                    # КРИТИЧЕСКАЯ ПРОВЕРКА: Убеждаемся, что мы не заменяем номер пункта
                    cell_text_full = cell.text
                    
                    # Проверяем, начинается ли текст ячейки с номера пункта
                    punkt_num_match = re.match(r'^(\d+[\.\):]?\s*)', cell_text_full)
                    if punkt_num_match:
                        punkt_prefix = punkt_num_match.group(1)
                        punkt_num_clean = re.match(r'^(\d+)', punkt_prefix)
                        if punkt_num_clean:
                            punkt_num = punkt_num_clean.group(1)
                            # Если old равен номеру пункта, пропускаем замену
                            old_clean = old.strip().replace(".", "").replace(")", "").replace(":", "").replace(" ", "")
                            if old_clean == punkt_num:
                                logger.warning(f"   ⚠️ ПРОПУСК: target_text '{old}' является номером пункта в ячейке, не заменяем!")
                                return False
                            
                            # Если old является частью номера пункта, пропускаем
                            if old in punkt_prefix or old == punkt_num:
                                logger.warning(f"   ⚠️ ПРОПУСК: target_text '{old}' является частью номера пункта '{punkt_prefix}', не заменяем!")
                                return False
                            
                            # Заменяем только в части после номера пункта
                            content_part = cell_text_full[len(punkt_prefix):]
                            if old in content_part:
                                new_content = content_part.replace(old, new)
                                cell.text = punkt_prefix + new_content
                                logger.info(f"   Замена в ячейке (по всему тексту, сохраняя номер пункта): '{old}' → '{new}'")
                                replaced = True
                            else:
                                logger.warning(f"   ⚠️ target_text '{old}' не найден в содержимом ячейки (после номера пункта)")
                        else:
                            # Если не удалось определить номер, заменяем осторожно
                            logger.info(f"   Замена в ячейке (по всему тексту): '{old}' → '{new}'")
                            cell.text = cell.text.replace(old, new)
                            replaced = old not in cell.text
                    else:
                        # Нет номера пункта в начале - обычная замена
                        logger.info(f"   Замена в ячейке (по всему тексту): '{old}' → '{new}'")
                        cell.text = cell.text.replace(old, new)
                        replaced = old not in cell.text
            else:
                # Если нет параграфов, заменяем напрямую в ячейке, но проверяем номер пункта
                if old in cell.text:
                    cell_text_full = cell.text
                    
                    # Проверяем, начинается ли текст ячейки с номера пункта
                    punkt_num_match = re.match(r'^(\d+[\.\):]?\s*)', cell_text_full)
                    if punkt_num_match:
                        punkt_prefix = punkt_num_match.group(1)
                        punkt_num_clean = re.match(r'^(\d+)', punkt_prefix)
                        if punkt_num_clean:
                            punkt_num = punkt_num_clean.group(1)
                            # Если old равен номеру пункта, пропускаем замену
                            old_clean = old.strip().replace(".", "").replace(")", "").replace(":", "").replace(" ", "")
                            if old_clean == punkt_num:
                                logger.warning(f"   ⚠️ ПРОПУСК: target_text '{old}' является номером пункта в ячейке без параграфов, не заменяем!")
                                return False
                            
                            # Если old является частью номера пункта, пропускаем
                            if old in punkt_prefix or old == punkt_num:
                                logger.warning(f"   ⚠️ ПРОПУСК: target_text '{old}' является частью номера пункта '{punkt_prefix}', не заменяем!")
                                return False
                            
                            # Заменяем только в части после номера пункта
                            content_part = cell_text_full[len(punkt_prefix):]
                            if old in content_part:
                                new_content = content_part.replace(old, new)
                                cell.text = punkt_prefix + new_content
                                logger.info(f"   Замена в ячейке (нет параграфов, сохраняя номер пункта): '{old}' → '{new}'")
                                replaced = True
                            else:
                                logger.warning(f"   ⚠️ target_text '{old}' не найден в содержимом ячейки (после номера пункта)")
                        else:
                            logger.info(f"   Замена в ячейке (нет параграфов): '{old}' → '{new}'")
                            cell.text = cell.text.replace(old, new)
                            replaced = old not in cell.text
                    else:
                        logger.info(f"   Замена в ячейке (нет параграфов): '{old}' → '{new}'")
                        cell.text = cell.text.replace(old, new)
                        replaced = old not in cell.text
        
            return replaced
        
        # Если cell не передан, работаем только с paragraph
        if paragraph is None:
            return False
            
        original_text = paragraph.text
        
        # Проверяем, начинается ли параграф с номера пункта
        punkt_match = re.match(r'^(\d+\.?\s*)', original_text)
        
        if punkt_match:
            # Параграф начинается с номера - заменяем только в содержимой части
            punkt_prefix = punkt_match.group(1)
            content_part = original_text[len(punkt_prefix):]
            
            if old in content_part:
                new_content = content_part.replace(old, new)
                new_full_text = punkt_prefix + new_content
                
                logger.info(f"   Умная замена:")
                logger.info(f"     Префикс пункта: '{punkt_prefix}'")
                logger.info(f"     Старое содержимое: '{content_part[:50]}...'")
                logger.info(f"     Новое содержимое: '{new_content[:50]}...'")
                
                # Заменяем текст через runs для сохранения форматирования
                found_in_runs = False
                for run in paragraph.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)
                        found_in_runs = True
                        replaced = True
                
                # Если не удалось через runs, заменяем весь текст
                if not found_in_runs:
                    paragraph.text = new_full_text
                    replaced = True
        else:
            # Обычная замена для параграфов без номера пункта
            found_in_runs = False
            for run in paragraph.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    found_in_runs = True
                    replaced = True
            
            # Если не удалось через runs, заменяем весь текст
            if not found_in_runs and old in paragraph.text:
                paragraph.text = paragraph.text.replace(old, new)
                replaced = True
        
        return replaced

    async def _handle_replace_point_text(self, filename: str, change: Dict[str, Any]) -> Dict[str, Any]:
        """
        Замена всего текста пункта/подпункта новым текстом.
        Находит начало пункта и заменяет весь его текст до следующего пункта.
        """
        target = change.get("target", {})
        payload = change.get("payload", {})
        point_start = target.get("text")  # Например, "36." или "8)"
        new_text = payload.get("new_text")
        
        if not point_start or not new_text:
            return {
                "success": False,
                "error": "INVALID_PAYLOAD",
                "message": "Для REPLACE_POINT_TEXT необходимы target.text и payload.new_text",
            }
        
        # Нормализация текста для поиска
        normalized_start = " ".join(point_start.split())
        logger.debug(f"Поиск пункта для замены: '{normalized_start}'")
        
        matches = await self._safe_find_text(filename, normalized_start, match_case=False)
        
        if not matches:
            # Пробуем варианты
            for variant in [f"{normalized_start.replace('.', '')}.", f"{normalized_start.replace(')', ')')}"]:
                variant_matches = await self._safe_find_text(filename, variant, match_case=False)
                if variant_matches:
                    matches = variant_matches
                    break
        
        if not matches:
            return {
                "success": False,
                "error": "POINT_NOT_FOUND",
                "message": f"Пункт '{point_start}' не найден в документе",
            }
        
        paragraph_index = matches[0].paragraph_index
        doc = Document(filename)
        
        if paragraph_index >= len(doc.paragraphs):
            return {
                "success": False,
                "error": "PARAGRAPH_INDEX_OUT_OF_RANGE",
                "message": f"Неверный индекс параграфа: {paragraph_index}",
            }
        
        # Находим конец пункта (до следующего пункта или раздела)
        start_idx = paragraph_index
        end_idx = self._find_section_end(doc, paragraph_index)
        
        # Заменяем весь текст пункта
        # Удаляем старые параграфы пункта
        removed_texts = []
        for idx in range(start_idx, end_idx):
            if start_idx < len(doc.paragraphs):
                removed_texts.append(doc.paragraphs[start_idx].text)
                self._delete_paragraph(doc.paragraphs[start_idx])
        
        # Вставляем новый текст
        insert_after_idx = max(0, start_idx - 1)
        if insert_after_idx < len(doc.paragraphs):
            insert_after = doc.paragraphs[insert_after_idx]
            # Разбиваем новый текст на параграфы
            new_paragraphs = new_text.split('\n')
            current_para = insert_after
            for para_text in new_paragraphs:
                if para_text.strip():
                    current_para = self._insert_paragraph_after(current_para, para_text.strip())
        else:
            # Если некуда вставлять, добавляем в конец
            for para_text in new_text.split('\n'):
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())
        
        doc.save(filename)
        
        if change.get("annotation", True):
            await self._add_annotation(
                filename,
                insert_after_idx,
                change,
                extra=f"Заменен пункт {point_start}",
            )
        
        return {"success": True, "paragraph_index": start_idx}

    async def _handle_delete_paragraph(self, filename: str, change: Dict[str, Any]) -> Dict[str, Any]:
        target = change.get("target", {})
        text_to_remove = target.get("text")
        match_case = target.get("match_case", False)
        description = change.get("description", "")

        if not text_to_remove:
            return {
                "success": False,
                "error": "INVALID_TARGET",
                "message": "Для DELETE_PARAGRAPH необходим target.text",
            }

        doc = Document(filename)

        # Проверяем, является ли text_to_remove номером пункта
        paragraph_num = None
        if self._is_paragraph_number(text_to_remove):
            # Извлекаем номер пункта (убираем точку, скобку и т.д.)
            num_match = re.match(r'^(\d+)', text_to_remove.replace(".", "").replace(")", "").replace(":", ""))
            if num_match:
                paragraph_num = num_match.group(1)
                logger.info(f"🔍 УДАЛЕНИЕ ПУНКТА {paragraph_num}: ищем в таблицах и параграфах")
        
        # ПРИОРИТЕТ 1: Если это номер пункта, сначала проверяем таблицы
        if paragraph_num:
            logger.info(f"🔍 ПРИОРИТЕТ 1: Удаление пункта {paragraph_num} - проверка таблиц")
            
            # Используем LLM для определения целевой таблицы (если в описании упоминается таблица)
            llm_target_table_indices = None
            if "таблице" in description.lower() or "таблиц" in description.lower():
                try:
                    llm_target_table_indices = await self._identify_target_table_with_llm(
                        doc=doc,
                        description=description,
                        target_text=text_to_remove
                    )
                    if llm_target_table_indices:
                        logger.info(f"✅ LLM определил целевые таблицы для удаления: {llm_target_table_indices}")
                except Exception as e:
                    logger.warning(f"⚠️ Ошибка при LLM определении таблицы, продолжаем поиск: {e}")
            
            # Ищем пункт в таблицах по номеру в первой ячейке строки
            row_deleted = False
            table_found_idx = None
            row_found_idx = None
            
            logger.info(f"   🔍 Поиск пункта {paragraph_num} в {len(doc.tables)} таблицах...")
            for table_idx, table in enumerate(doc.tables):
                # Если LLM определил целевые таблицы, обрабатываем только их
                if llm_target_table_indices is not None:
                    if table_idx not in llm_target_table_indices:
                        logger.info(f"   ⏭️ Пропускаем таблицу {table_idx} (не определена LLM как целевая)")
                        continue
                    else:
                        logger.info(f"   ✅ Таблица {table_idx} определена LLM как целевая")
                else:
                    logger.info(f"   🔍 Проверяем таблицу {table_idx} (строка {len(table.rows)} строк)")
                
                for row_idx, row in enumerate(table.rows):
                    if row.cells:
                        # Проверяем первую ячейку строки на наличие номера пункта
                        first_cell_text = row.cells[0].text.strip()
                        
                        # Проверяем различные форматы номера пункта с точным совпадением
                        patterns = [
                            rf"^{re.escape(paragraph_num)}\.",
                            rf"^{re.escape(paragraph_num)}\)",
                            rf"^{re.escape(paragraph_num)}:",
                        ]
                        
                        # Проверяем каждый паттерн
                        matched = False
                        for pattern in patterns:
                            if re.match(pattern, first_cell_text):
                                matched = True
                                break
                        
                        if matched:
                            logger.info(f"   ✅ Найден пункт {paragraph_num} в таблице {table_idx}, строка {row_idx} (первая ячейка: '{first_cell_text}')")
                            
                            # Удаляем содержимое строки, оставляя только номер пункта
                            # Сохраняем номер пункта в первой ячейке
                            logger.info(f"   🔧 Очищаем содержимое строки, сохраняя номер пункта...")
                            for run in row.cells[0].paragraphs[0].runs:
                                run.text = ""
                            row.cells[0].paragraphs[0].add_run(f"{paragraph_num}.")
                            
                            # Очищаем остальные ячейки строки
                            for cell_idx in range(1, len(row.cells)):
                                for para in row.cells[cell_idx].paragraphs:
                                    for run in para.runs:
                                        run.text = ""
                            
                            row_deleted = True
                            table_found_idx = table_idx
                            row_found_idx = row_idx
                            logger.info(f"   ✅ Содержимое пункта {paragraph_num} удалено из таблицы {table_idx}, строка {row_idx} (номер пункта сохранен)")
                            break
                
                # Если нашли и удалили, останавливаемся
                if row_deleted:
                    logger.info(f"   ✅ Удаление завершено, прекращаем поиск")
                    break
            
            if not row_deleted:
                logger.info(f"   ⚠️ Пункт {paragraph_num} не найден в таблицах, продолжаем поиск в параграфах")
            
            # Если пункт найден и удален из таблицы
            if row_deleted:
                doc.save(filename)
                
                # Добавляем аннотацию перед таблицей
                table_para_idx = self._find_paragraph_for_table(doc, table_found_idx)
                if change.get("annotation", True) and table_para_idx >= 0:
                    await self._add_annotation(
                        filename,
                        table_para_idx,
                        change,
                        extra=f"Удалено содержимое пункта {paragraph_num} из таблицы (номер пункта сохранен)",
                    )
                
                return {
                    "success": True,
                    "paragraph_index": table_para_idx if table_para_idx >= 0 else 0,
                    "table_location": {
                        "table_idx": table_found_idx,
                        "row_idx": row_found_idx,
                        "paragraph_index": table_para_idx if table_para_idx >= 0 else 0
                    },
                    "message": f"Содержимое пункта {paragraph_num} удалено из таблицы (номер сохранен)"
                }
        
        # ПРИОРИТЕТ 2: Если это не номер пункта, но в описании упоминается "строка" и "таблица",
        # то удаляем строку из таблицы по содержимому ячейки
        is_table_row_delete = ("строка" in description.lower() or "строку" in description.lower()) and \
                               ("таблице" in description.lower() or "таблиц" in description.lower())
        
        if is_table_row_delete and not paragraph_num:
            logger.info(f"🔍 УДАЛЕНИЕ СТРОКИ ИЗ ТАБЛИЦЫ: ищем '{text_to_remove}' в таблицах")
            
            # Извлекаем название таблицы из описания
            table_name = None
            table_name_match = re.search(r'таблиц[еи]\s+[«"]([^«"]+)[»"]', description, re.IGNORECASE)
            if table_name_match:
                table_name = table_name_match.group(1)
                logger.info(f"📋 Ограничение поиска таблицей: '{table_name}'")
            
            # Используем LLM для определения целевой таблицы
            llm_target_table_indices = None
            try:
                llm_target_table_indices = await self._identify_target_table_with_llm(
                    doc=doc,
                    description=description,
                    target_text=text_to_remove,
                    table_name=table_name
                )
                if llm_target_table_indices:
                    logger.info(f"✅ LLM определил целевые таблицы для удаления строки: {llm_target_table_indices}")
            except Exception as e:
                logger.warning(f"⚠️ Ошибка при LLM определении таблицы, продолжаем поиск: {e}")
            
            # Если указано название таблицы, ищем его в тексте документа
            table_name_found_in_text = False
            if table_name:
                logger.info(f"🔍 Поиск названия таблицы '{table_name}' в тексте документа...")
                try:
                    matches = await mcp_client.find_text_in_document(filename, table_name, match_case=False)
                    if matches:
                        logger.info(f"   ✅ Найдено {len(matches)} упоминаний названия таблицы в тексте")
                        table_name_found_in_text = True
                        logger.info(f"   📍 Название найдено в тексте, ограничиваем поиск таблицами с target_text")
                except Exception as e:
                    logger.warning(f"   ⚠️ Ошибка при поиске названия таблицы в тексте: {e}")
            
            # Ищем текст в таблицах и удаляем строку
            row_deleted = False
            table_found_idx = None
            row_found_idx = None
            first_table_processed = False
            
            for table_idx, table in enumerate(doc.tables):
                should_process_this_table = True
                
                # ПРИОРИТЕТ 1: Если LLM определил целевые таблицы, обрабатываем только их
                if llm_target_table_indices is not None:
                    if table_idx not in llm_target_table_indices:
                        logger.info(f"   ⏭️ Пропускаем таблицу {table_idx} (не определена LLM как целевая)")
                        continue
                    else:
                        logger.info(f"   ✅ Таблица {table_idx} определена LLM как целевая")
                
                # ПРИОРИТЕТ 2: Если название найдено в тексте, обрабатываем только первую таблицу с target_text
                if table_name and table_name_found_in_text and llm_target_table_indices is None:
                    if first_table_processed:
                        logger.info(f"   ⏭️ Пропускаем таблицу {table_idx} (уже обработана первая таблица с названием)")
                        continue
                    
                    # Проверяем, содержит ли эта таблица target_text
                    table_contains_target = False
                    for row in table.rows:
                        for cell in row.cells:
                            if text_to_remove in cell.text:
                                table_contains_target = True
                                break
                        if table_contains_target:
                            break
                    
                    if not table_contains_target:
                        logger.info(f"   ⏭️ Пропускаем таблицу {table_idx} (не содержит target_text '{text_to_remove}')")
                        should_process_this_table = False
                    else:
                        logger.info(f"   ✅ Таблица {table_idx} содержит target_text и идет после названия в тексте")
                
                if not should_process_this_table:
                    continue
                
                # Ищем строку с text_to_remove в первой ячейке (для таблиц сокращений)
                for row_idx, row in enumerate(table.rows):
                    if not row.cells:
                        continue
                    
                    # Проверяем первую ячейку для точного совпадения (типично для таблиц сокращений)
                    first_cell_text = row.cells[0].text.strip()
                    if first_cell_text == text_to_remove or text_to_remove in first_cell_text:
                        logger.info(f"   ✅ Найден текст '{text_to_remove}' в таблице {table_idx}, строка {row_idx}, ячейка 0")
                        
                        # Удаляем всю строку (XML элемент строки)
                        tbl = table._tbl
                        tr = row._tr
                        tbl.remove(tr)
                        
                        row_deleted = True
                        table_found_idx = table_idx
                        row_found_idx = row_idx
                        logger.info(f"   ✅ Строка {row_idx} удалена из таблицы {table_idx}")
                        first_table_processed = True
                        break
                    
                    # Если не нашли в первой ячейке, проверяем остальные ячейки
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if cell_text == text_to_remove:
                            logger.info(f"   ✅ Найден текст '{text_to_remove}' в таблице {table_idx}, строка {row_idx}, ячейка {cell_idx}")
                            
                            # Удаляем всю строку (XML элемент строки)
                            tbl = table._tbl
                            tr = row._tr
                            tbl.remove(tr)
                            
                            row_deleted = True
                            table_found_idx = table_idx
                            row_found_idx = row_idx
                            logger.info(f"   ✅ Строка {row_idx} удалена из таблицы {table_idx}")
                            first_table_processed = True
                            break
                    
                    if row_deleted:
                        break
                
                if row_deleted:
                    # Если название найдено в тексте, останавливаемся после первой найденной таблицы
                    if table_name_found_in_text:
                        logger.info(f"   ✅ Найдена целевая таблица с названием в тексте, завершаем обработку")
                        break
                    break
            
            # Если строка удалена из таблицы
            if row_deleted:
                doc.save(filename)
                
                # Добавляем аннотацию перед таблицей
                table_para_idx = self._find_paragraph_for_table(doc, table_found_idx)
                if change.get("annotation", True) and table_para_idx >= 0:
                    await self._add_annotation(
                        filename,
                        table_para_idx,
                        change,
                        extra=f"Удалена строка с '{text_to_remove}' из таблицы",
                    )
                
                return {
                    "success": True,
                    "paragraph_index": table_para_idx if table_para_idx >= 0 else 0,
                    "table_location": {
                        "table_idx": table_found_idx,
                        "row_idx": row_found_idx,
                        "paragraph_index": table_para_idx if table_para_idx >= 0 else 0
                    },
                    "message": f"Строка с '{text_to_remove}' удалена из таблицы"
                }
        
        # ПРИОРИТЕТ 3: Если пункт не найден в таблице или это не номер пункта, ищем в параграфах
        # Нормализация текста для поиска
        normalized_text = " ".join(text_to_remove.split())
        logger.debug(f"Поиск текста для удаления: '{normalized_text}' (оригинал: '{text_to_remove}')")
        
        matches = await self._safe_find_text(filename, normalized_text, match_case=match_case)
        
        # Если не найдено, пробуем оригинальный текст
        if not matches and normalized_text != text_to_remove:
            matches = await self._safe_find_text(filename, text_to_remove, match_case=match_case)
        
        # Для пунктов пробуем разные форматы
        if not matches and (text_to_remove.isdigit() or text_to_remove.replace(".", "").replace(")", "").isdigit()):
            for variant in [f"{text_to_remove}.", f"{text_to_remove})", f"{text_to_remove}."]:
                variant_matches = await self._safe_find_text(filename, variant, match_case=False)
                if variant_matches:
                    matches = variant_matches
                    logger.info(f"Найдено совпадение для варианта '{variant}'")
                    break
        
        if not matches:
            logger.warning(f"Текст для удаления '{text_to_remove}' не найден")
            return {
                "success": False,
                "error": "TEXT_NOT_FOUND",
                "message": f"Текст '{text_to_remove}' не найден в документе",
            }

        paragraph_index = matches[0].paragraph_index
        
        if paragraph_index >= len(doc.paragraphs):
            return {
                "success": False,
                "error": "PARAGRAPH_INDEX_OUT_OF_RANGE",
                "message": f"Неверный индекс параграфа: {paragraph_index}",
            }

        # Если это номер пункта, удаляем только содержимое, сохраняя номер
        if paragraph_num:
            para = doc.paragraphs[paragraph_index]
            para_text = para.text.strip()
            
            # Проверяем, что параграф начинается с номера пункта
            patterns = [
                rf"^{re.escape(paragraph_num)}\.",
                rf"^{re.escape(paragraph_num)}\)",
                rf"^{re.escape(paragraph_num)}:",
            ]
            
            if any(re.match(p, para_text) for p in patterns):
                # Находим номер пункта
                num_pattern = None
                for pattern in patterns:
                    match = re.match(pattern, para_text)
                    if match:
                        num_pattern = match.group(0)
                        break
                
                if not num_pattern:
                    num_pattern = f"{paragraph_num}."
                
                # Очищаем параграф и оставляем только номер
                for run in para.runs:
                    run.text = ""
                para.add_run(num_pattern)
                
                # Ищем и удаляем последующие параграфы, которые являются частью этого пункта
                end_idx = self._find_section_end(doc, paragraph_index)
                removed_preview = [para_text]
                
                for idx in range(paragraph_index + 1, end_idx):
                    if idx < len(doc.paragraphs):
                        next_para = doc.paragraphs[paragraph_index + 1]  # Индекс меняется после удаления
                        removed_preview.append(next_para.text)
                        DocumentChangeAgent._delete_paragraph(next_para)
                
                doc.save(filename)
                
                if change.get("annotation", True) and paragraph_index > 0:
                    preview_text = " ".join(removed_preview)[:120]
                    await self._add_annotation(
                        filename,
                        paragraph_index - 1,
                        change,
                        extra=f"Удалено содержимое пункта {paragraph_num} (номер сохранен): {preview_text}",
                    )
                
                return {"success": True, "paragraph_index": paragraph_index}
        
        # Стандартное удаление параграфа (если это не номер пункта)
        start = paragraph_index
        end = self._find_section_end(doc, paragraph_index)
        removed_preview = []

        for idx in range(start, end):
            para = doc.paragraphs[start]  # список пересчитывается после удаления
            removed_preview.append(para.text)
            DocumentChangeAgent._delete_paragraph(para)

        doc.save(filename)

        if change.get("annotation", True) and start > 0:
            preview_text = " ".join(removed_preview)[:120]
            await self._add_annotation(
                filename,
                start - 1,
                change,
                extra=f"Удален раздел: {preview_text}",
            )

        return {"success": True, "paragraph_index": start}

    async def _handle_insert_paragraph(self, filename: str, change: Dict[str, Any], master_doc: Optional[Document] = None) -> Dict[str, Any]:
        """
        Вставка нового параграфа после указанного текста.
        """
        target = change.get("target", {})
        payload = change.get("payload", {})
        after_text = target.get("after_text")
        new_paragraph = payload.get("text")
        style = payload.get("style")

        if not after_text or not new_paragraph:
            error_msg = "Для INSERT_PARAGRAPH необходимы target.after_text и payload.text"
            logger.warning(f"{change.get('change_id', 'UNKNOWN')}: {error_msg}")
            return {
                "success": False,
                "error": "INVALID_PAYLOAD",
                "message": error_msg,
            }

        # Нормализация текста для поиска
        normalized_after = " ".join(after_text.split())
        logger.debug(f"Поиск якоря для вставки: '{normalized_after}' (оригинал: '{after_text}')")

        matches = await self._safe_find_text(filename, normalized_after, match_case=target.get("match_case", False))
        
        # Если не найдено, пробуем оригинальный текст
        if not matches and normalized_after != after_text:
            matches = await self._safe_find_text(filename, after_text, match_case=target.get("match_case", False))
        
        if not matches:
            error_msg = f"Якорь '{after_text}' не найден в документе"
            logger.warning(f"{change.get('change_id', 'UNKNOWN')}: {error_msg}")
            return {
                "success": False,
                "error": "ANCHOR_NOT_FOUND",
                "message": error_msg,
            }

        anchor_index = matches[0].paragraph_index
        doc = Document(filename)
        if anchor_index >= len(doc.paragraphs):
            return {
                "success": False,
                "error": "PARAGRAPH_INDEX_OUT_OF_RANGE",
                "message": f"Неверный индекс параграфа: {anchor_index}",
            }

        insert_after = doc.paragraphs[anchor_index]
        self._insert_paragraph_after(insert_after, new_paragraph, style)
        doc.save(filename)

        doc = Document(filename)
        insert_position = (
            self._find_paragraph_index_by_text(doc, new_paragraph, start=anchor_index)
            or anchor_index + 1
        )

        if change.get("annotation", True):
            await self._add_annotation(
                filename,
                insert_position,
                change,
                extra=f"Добавлен параграф: {new_paragraph[:120]}",
            )

        return {"success": True, "paragraph_index": insert_position}

    async def _handle_insert_section(self, filename: str, change: Dict[str, Any], master_doc: Optional[Document] = None) -> Dict[str, Any]:
        target = change.get("target", {})
        payload = change.get("payload", {})

        after_heading = target.get("after_heading")
        heading_text = payload.get("heading_text")
        heading_level = payload.get("heading_level", 2)
        paragraphs: List[str] = payload.get("paragraphs", [])

        if not after_heading or not heading_text:
            return {
                "success": False,
                "error": "INVALID_PAYLOAD",
                "message": "Для INSERT_SECTION необходимы target.after_heading и payload.heading_text",
            }

        matches = await self._safe_find_text(filename, after_heading, match_case=target.get("match_case", False))
        if not matches:
            return {"success": False, "error": "ANCHOR_NOT_FOUND"}

        anchor_index = matches[0].paragraph_index
        doc = Document(filename)
        if anchor_index >= len(doc.paragraphs):
            return {
                "success": False,
                "error": "PARAGRAPH_INDEX_OUT_OF_RANGE",
                "message": f"Неверный индекс параграфа: {anchor_index}",
            }

        insert_index = self._find_section_end(doc, anchor_index)
        insert_after = doc.paragraphs[max(insert_index - 1, anchor_index)]
        heading_style = f"Heading {heading_level}"
        new_heading = self._insert_paragraph_after(insert_after, heading_text, heading_style)
        current_para = new_heading

        for paragraph in paragraphs:
            current_para = self._insert_paragraph_after(current_para, paragraph)

        doc.save(filename)

        doc = Document(filename)
        start_index = (
            self._find_paragraph_index_by_text(
                doc, heading_text, start=insert_index, style=heading_style
            )
            or insert_index
        )

        if change.get("annotation", True):
            await self._add_annotation(
                filename,
                start_index,
                change,
                extra=f"Добавлен раздел «{heading_text}»",
            )

        return {
            "success": True,
            "start_index": start_index,
            "paragraphs_added": len(paragraphs) + 1,
        }

    async def _handle_insert_table(self, filename: str, change: Dict[str, Any], master_doc: Optional[Document] = None) -> Dict[str, Any]:
        """
        Вставка таблицы после указанного текста.
        """
        target = change.get("target", {})
        payload = change.get("payload", {})
        after_text = target.get("after_text")
        rows = payload.get("rows", [])
        columns = payload.get("columns")
        
        if not after_text:
            return {
                "success": False,
                "error": "INVALID_TARGET",
                "message": "Для INSERT_TABLE необходим target.after_text",
            }
        
        if not rows or not isinstance(rows, list):
            return {
                "success": False,
                "error": "INVALID_PAYLOAD",
                "message": "Для INSERT_TABLE необходим payload.rows (массив строк таблицы)",
            }
        
        # Нормализация текста для поиска
        normalized_after = " ".join(after_text.split())
        logger.debug(f"Поиск якоря для вставки таблицы: '{normalized_after}'")
        
        matches = await self._safe_find_text(filename, normalized_after, match_case=target.get("match_case", False))
        
        if not matches and normalized_after != after_text:
            matches = await self._safe_find_text(filename, after_text, match_case=target.get("match_case", False))
        
        if not matches:
            error_msg = f"Якорь '{after_text}' не найден в документе для вставки таблицы"
            logger.warning(f"{change.get('change_id', 'UNKNOWN')}: {error_msg}")
            return {
                "success": False,
                "error": "ANCHOR_NOT_FOUND",
                "message": error_msg,
            }
        
        anchor_index = matches[0].paragraph_index
        doc = Document(filename)
        
        if anchor_index >= len(doc.paragraphs):
            return {
                "success": False,
                "error": "PARAGRAPH_INDEX_OUT_OF_RANGE",
                "message": f"Неверный индекс параграфа: {anchor_index}",
            }
        
        # Определяем количество колонок
        if not columns:
            columns = max(len(row) for row in rows) if rows else 0
        
        # Вставляем таблицу через MCP
        success = await mcp_client.add_table(filename, rows, position=anchor_index + 1)
        
        if not success:
            return {
                "success": False,
                "error": "TABLE_INSERT_FAILED",
                "message": "Не удалось вставить таблицу в документ",
            }
        
        if change.get("annotation", True):
            await self._add_annotation(
                filename,
                anchor_index,
                change,
                extra=f"Добавлена таблица ({len(rows)} строк, {columns} колонок)",
            )
        
        return {
            "success": True,
            "paragraph_index": anchor_index,
            "rows_count": len(rows),
            "columns_count": columns,
        }

    async def _handle_add_comment(self, filename: str, change: Dict[str, Any], master_doc: Optional[Document] = None) -> Dict[str, Any]:
        target = change.get("target", {})
        payload = change.get("payload", {})

        # Попытка получить paragraph_hint из разных мест
        paragraph_hint = payload.get("paragraph_hint") or target.get("text") or target.get("paragraph_hint")
        comment_text = payload.get("comment_text") or payload.get("text") or change.get("description")
        is_table_change = payload.get("is_table_change", False)  # Флаг, что изменение было в таблице

        if not paragraph_hint or not comment_text:
            logger.warning(
                f"ADD_COMMENT: отсутствуют обязательные параметры. "
                f"target={target}, payload={payload}, change_id={change.get('change_id', 'N/A')}"
            )
            return {
                "success": False,
                "error": "INVALID_PAYLOAD",
                "message": f"Для ADD_COMMENT необходимы payload.paragraph_hint (или target.text) и payload.comment_text. "
                          f"Получено: paragraph_hint={paragraph_hint}, comment_text={bool(comment_text)}",
            }

        # Если paragraph_index уже указан в payload (например, из table_location), используем его
        paragraph_index = payload.get("paragraph_index")
        if paragraph_index is not None and paragraph_index >= 0:
            logger.info(f"ADD_COMMENT: используется указанный paragraph_index: {paragraph_index}")
        else:
            # Иначе ищем текст в документе
            matches = await self._safe_find_text(filename, paragraph_hint, match_case=target.get("match_case", False))
            if not matches:
                logger.warning(f"ADD_COMMENT: текст '{paragraph_hint[:50]}...' не найден в документе")
                return {"success": False, "error": "ANCHOR_NOT_FOUND", "message": f"Текст '{paragraph_hint[:50]}...' не найден"}

            # Ищем первый match, который не в таблице (paragraph_index != -1)
            paragraph_index = -1
            for match in matches:
                if hasattr(match, 'paragraph_index') and match.paragraph_index != -1:
                    paragraph_index = match.paragraph_index
                    logger.info(f"ADD_COMMENT: найден параграф {paragraph_index} с текстом")
                    break
            
            # Если все matches в таблицах, используем первый
            if paragraph_index == -1:
                paragraph_index = matches[0].paragraph_index
            
            # Если текст найден в таблице (paragraph_index = -1) или изменение было в таблице, ищем параграф ПЕРЕД таблицей
            if paragraph_index == -1 or is_table_change:
                logger.info(f"ADD_COMMENT: текст найден в таблице, ищем параграф ПЕРЕД таблицей")
                doc = Document(filename)
                
                # Получаем информацию о таблице из matches
                table_match = None
                for match in matches:
                    if hasattr(match, 'paragraph_index') and match.paragraph_index == -1:
                        table_match = match
                        break
                
                # Если есть информация о таблице, используем её для поиска
                if table_match and hasattr(table_match, 'location'):
                    location = table_match.location
                    logger.info(f"ADD_COMMENT: информация о таблице из match: {location}")
                    
                    # Извлекаем номер таблицы из location (например, "Table 0")
                    table_num_match = re.search(r'Table\s+(\d+)', location)
                    if table_num_match:
                        table_num = int(table_num_match.group(1))
                        logger.info(f"ADD_COMMENT: найден номер таблицы: {table_num}")
                        
                        # Ищем параграф перед таблицей
                        # Для этого нужно найти таблицу в документе и найти параграф перед ней
                        table_found = False
                        target_paragraph_index = -1
                        
                        # Проходим по элементам документа
                        para_count = 0
                        for i, element in enumerate(doc.element.body):
                            if element.tag.endswith('p'):  # Параграф
                                para_count += 1
                            elif element.tag.endswith('tbl'):  # Таблица
                                # Проверяем, это нужная таблица?
                                # Считаем таблицы с начала документа
                                table_idx = sum(1 for j in range(i) if doc.element.body[j].tag.endswith('tbl'))
                                if table_idx == table_num:
                                    table_found = True
                                    # Ищем первый параграф ПОСЛЕ этой таблицы
                                    for j in range(i+1, len(doc.element.body)):
                                        if doc.element.body[j].tag.endswith('p'):
                                            # Подсчитываем индекс параграфа
                                            target_paragraph_index = sum(1 for k in range(j+1) if doc.element.body[k].tag.endswith('p')) - 1
                                            break
                                    # Если не нашли параграф после таблицы, используем последний параграф в документе
                                    if target_paragraph_index == -1:
                                        target_paragraph_index = sum(1 for k in range(len(doc.element.body)) if doc.element.body[k].tag.endswith('p')) - 1
                                    break
                        
                        if table_found and target_paragraph_index >= 0:
                            paragraph_index = target_paragraph_index
                            logger.info(f"ADD_COMMENT: найден параграф {paragraph_index} ПОСЛЕ таблицы {table_num}")
                        else:
                            # Если не нашли параграф перед таблицей, ищем ближайший параграф с текстом
                            for idx, para in enumerate(doc.paragraphs):
                                if paragraph_hint[:30] in para.text:
                                    paragraph_index = idx
                                    logger.info(f"ADD_COMMENT: найден ближайший параграф {idx} с текстом")
                                    break
                            
                            if paragraph_index == -1:
                                # Если все еще не нашли, используем первый параграф
                                paragraph_index = 0
                                logger.warning(f"ADD_COMMENT: не найден параграф перед таблицей, используем первый (0)")
                    else:
                        # Если не удалось извлечь номер таблицы, ищем ближайший параграф
                        for idx, para in enumerate(doc.paragraphs):
                            if paragraph_hint[:30] in para.text:
                                paragraph_index = idx
                                logger.info(f"ADD_COMMENT: найден ближайший параграф {idx}")
                                break
                        
                        if paragraph_index == -1:
                            paragraph_index = 0
                            logger.warning(f"ADD_COMMENT: не найден ближайший параграф, используем первый (0)")
                else:
                    # Если нет информации о таблице, ищем ближайший параграф
                    for idx, para in enumerate(doc.paragraphs):
                        if paragraph_hint[:30] in para.text:
                            paragraph_index = idx
                            logger.info(f"ADD_COMMENT: найден ближайший параграф {idx}")
                            break
                    
                    if paragraph_index == -1:
                        paragraph_index = 0
                        logger.warning(f"ADD_COMMENT: не найден ближайший параграф, используем первый (0)")
            else:
                # Если не таблица и paragraph_index == -1, используем первый match
                if paragraph_index == -1 and matches:
                    paragraph_index = matches[0].paragraph_index
        
        try:
            comment_id = await mcp_client.add_comment(
                filename,
                paragraph_index,
                comment_text,
            )
            if not comment_id:
                logger.warning(f"ADD_COMMENT: не удалось добавить комментарий (comment_id=None)")
                return {"success": False, "error": "COMMENT_FAILED", "message": "Не удалось добавить комментарий"}
            
            logger.info(f"ADD_COMMENT: комментарий добавлен успешно (paragraph_index={paragraph_index}, comment_id={comment_id})")
            return {"success": True, "paragraph_index": paragraph_index, "comment_id": comment_id}
        except Exception as e:
            logger.error(f"ADD_COMMENT: ошибка при добавлении комментария: {e}")
            return {"success": False, "error": "COMMENT_EXCEPTION", "message": str(e)}

    async def _add_annotation(
        self,
        filename: str,
        paragraph_index: int,
        change: Dict[str, Any],
        extra: Optional[str] = None,
    ) -> None:
        # Создаем компактную аннотацию
        change_id = change.get('change_id', 'CHG')
        operation = change.get('operation', '')
        description = change.get("description", "Нет описания")
        
        # Компактный формат аннотации
        annotation = f"[{change_id}] {operation}: {description}"
        if extra:
            annotation += f" | {extra}"

        # Если paragraph_index == -1, это означает, что изменение произошло в таблице.
        # В этом случае мы добавляем аннотацию ПЕРЕД таблицей.
        if paragraph_index == -1:
            doc = Document(filename)
            table_found = False
            target_paragraph_index = 0
            
            # Ищем первую таблицу и добавляем аннотацию перед ней
            for i, element in enumerate(doc.element.body):
                if element.tag.endswith('tbl'): # Если это таблица
                    table_found = True
                    # Ищем параграф перед таблицей
                    for j in range(i-1, -1, -1):
                        if doc.element.body[j].tag.endswith('p'):
                            target_paragraph_index = j
                            break
                    break
            
            if table_found and target_paragraph_index >= 0:
                paragraph_index = target_paragraph_index
            else:
                # Если не удалось найти параграф перед таблицей, добавляем в начало документа
                paragraph_index = 0
                logger.warning(f"Не удалось найти параграф перед таблицей для аннотации. Аннотация добавлена в начало документа.")

        await mcp_client.add_comment(
            filename,
            paragraph_index,
            annotation,
        )

    @staticmethod
    def _replace_in_paragraph(paragraph, old: str, new: str) -> bool:
        replaced = False

        for run in paragraph.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                replaced = True

        if not replaced and old in paragraph.text:
            paragraph.text = paragraph.text.replace(old, new)
            replaced = True

        return replaced

    @staticmethod
    def _robust_replace_in_paragraph(paragraph: Paragraph, old: str, new: str) -> bool:
        """
        НОВЫЙ ФУНКЦИОНАЛ: Надежная замена текста в параграфе, работающая с текстом,
        разбитым на несколько runs.
        
        Этот метод:
        1. Проверяет, есть ли текст целиком в paragraph.text
        2. Если да, объединяет runs и заменяет текст с сохранением форматирования
        3. Работает даже когда текст пересекает границы runs
        
        Args:
            paragraph: Параграф для замены
            old: Текст для замены
            new: Новый текст
            
        Returns:
            True если замена выполнена, False иначе
        """
        # Сначала проверяем, есть ли текст в параграфе
        para_text = paragraph.text
        logger.debug(f"🔍 _robust_replace_in_paragraph: ищем '{old}' в параграфе '{para_text[:50]}...'")
        
        if old not in para_text:
            logger.debug(f"   ❌ Текст '{old}' не найден в параграфе")
            return False
        
        try:
            # Если текст найден, выполняем замену с сохранением форматирования
            # Стратегия: находим позицию текста в объединенном тексте,
            # затем заменяем в соответствующих runs
            
            # Собираем весь текст из всех runs с сохранением позиций
            runs_text = []
            current_pos = 0
            for run in paragraph.runs:
                run_text = run.text
                runs_text.append({
                    'run': run,
                    'text': run_text,
                    'start': current_pos,
                    'end': current_pos + len(run_text)
                })
                current_pos += len(run_text)
            
            # Находим позицию текста для замены
            old_pos = para_text.find(old)
            if old_pos == -1:
                return False
            
            old_end = old_pos + len(old)
            
            # Находим runs, которые нужно изменить
            affected_runs = []
            for run_info in runs_text:
                # Проверяем пересечение
                if not (run_info['end'] <= old_pos or run_info['start'] >= old_end):
                    affected_runs.append(run_info)
            
            if not affected_runs:
                return False
            
            # Если текст полностью в одном run - простая замена
            if len(affected_runs) == 1:
                run_info = affected_runs[0]
                logger.debug(f"   ✅ Текст найден в одном run, выполняем простую замену")
                run_info['run'].text = run_info['run'].text.replace(old, new)
                return True
            
            # Если текст пересекает несколько runs - более сложная замена
            logger.debug(f"   🔧 Текст пересекает {len(affected_runs)} runs, выполняем сложную замену")
            
            # Собираем текст из всех affected runs и заменяем там
            first_run_info = affected_runs[0]
            last_run_info = affected_runs[-1]
            
            # Извлекаем текст из параграфа между началом первого и концом последнего affected run
            # Это более надежный способ, чем простое объединение runs
            segment_start = first_run_info['start']
            segment_end = last_run_info['end']
            segment_text = para_text[segment_start:segment_end]
            logger.debug(f"   📍 Сегмент текста (позиции {segment_start}-{segment_end}): '{segment_text[:80]}...'")
            
            # Проверяем, что old действительно находится в этом сегменте
            segment_old_pos = segment_text.find(old)
            if segment_old_pos == -1:
                # Если не найдено в сегменте, пробуем найти в полном параграфе и заменить напрямую
                logger.warning(f"   ⚠️ Текст '{old}' не найден в сегменте, используем fallback через paragraph.text")
                if old in para_text:
                    try:
                        # Используем простую замену через paragraph.text (сохранит форматирование первого run)
                        new_para_text = para_text.replace(old, new, 1)  # Заменяем только первое вхождение
                        paragraph.text = new_para_text
                        # Проверяем результат
                        if new in paragraph.text:
                            logger.debug(f"   ✅ Замена выполнена через fallback paragraph.text")
                            return True
                    except Exception as e:
                        logger.warning(f"   ⚠️ Ошибка при замене через paragraph.text: {e}, пробуем альтернативный метод")
                        
                        # Альтернативный метод: очистка всех runs и запись нового текста в первый
                        if paragraph.runs:
                            first_run = paragraph.runs[0]
                            new_para_text = para_text.replace(old, new, 1)
                            
                            # Очищаем все runs
                            for run in paragraph.runs:
                                run.text = ""
                            
                            # Записываем в первый run
                            first_run.text = new_para_text
                            if new in paragraph.text:
                                logger.debug(f"   ✅ Замена выполнена через альтернативный метод")
                                return True
                
                logger.debug(f"   ❌ Текст не найден даже в полном параграфе")
                return False
            
            # Выполняем замену в сегменте
            replacement_segment = segment_text[:segment_old_pos] + new + segment_text[segment_old_pos + len(old):]
            logger.debug(f"   🔄 Замененный сегмент: '{replacement_segment[:80]}...'")
            
            # Теперь нужно правильно распределить замененный текст по runs
            # Простая стратегия: весь замененный текст идет в первый affected run
            first_run_info['run'].text = replacement_segment
            # Очищаем остальные affected runs
            for run_info in affected_runs[1:]:
                run_info['run'].text = ""
            
            logger.debug(f"   ✅ Сложная замена выполнена успешно")
            return True
            
        except Exception as e:
            logger.warning(f"Ошибка в _robust_replace_in_paragraph: {e}, пробуем простую замену")
            # Fallback: простая замена через paragraph.text
            if old in paragraph.text:
                try:
                    # Пробуем заменить через paragraph.text (это пересоздаст runs)
                    original_text = paragraph.text
                    new_para_text = original_text.replace(old, new, 1)  # Заменяем только первое вхождение
                    paragraph.text = new_para_text
                    # Проверяем, что замена действительно произошла
                    if new_para_text in paragraph.text or new in paragraph.text:
                        logger.debug(f"   ✅ Замена выполнена через paragraph.text (fallback)")
                        return True
                except Exception as e2:
                    logger.error(f"   ❌ Ошибка при замене через paragraph.text: {e2}")
            
            # Последняя попытка: прямая работа с текстом через очистку и перезапись
            if old in paragraph.text:
                try:
                    # Получаем первый run и используем его для замены
                    if paragraph.runs:
                        # Сохраняем форматирование первого run
                        first_run = paragraph.runs[0]
                        original_text = paragraph.text
                        new_text_final = original_text.replace(old, new, 1)
                        
                        # Очищаем все runs
                        for run in paragraph.runs:
                            run.text = ""
                        
                        # Записываем новый текст в первый run
                        first_run.text = new_text_final
                        logger.debug(f"   ✅ Замена выполнена через очистку и перезапись runs")
                        return True
                except Exception as e3:
                    logger.error(f"   ❌ Ошибка при замене через очистку runs: {e3}")
            
            return False

    @staticmethod
    def _replace_in_cell(cell, old: str, new: str) -> bool:
        """
        Замена текста в ячейке таблицы с сохранением форматирования.
        """
        replaced = False
        
        # Проходим по всем параграфам в ячейке
        for paragraph in cell.paragraphs:
            # Сохраняем форматирование через runs
            for run in paragraph.runs:
                if old in run.text:
                    # Сохраняем исходное форматирование
                    original_font = run.font
                    run.text = run.text.replace(old, new)
                    replaced = True
            
            # Если не нашли в runs, проверяем весь параграф
        if not replaced and old in paragraph.text:
            paragraph.text = paragraph.text.replace(old, new)
            replaced = True

        return replaced

    def _find_text_locally(self, filename: str, text_to_find: str, match_case: bool = True) -> List[MCPTextMatch]:
        """
        Локальный поиск текста в документе через python-docx (fallback для MCP).
        """
        matches = []
        try:
            doc = Document(filename)
            for idx, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text
                if not match_case:
                    para_text = para_text.lower()
                    search_text = text_to_find.lower()
                else:
                    search_text = text_to_find
                
                if search_text in para_text:
                    matches.append(MCPTextMatch(paragraph_index=idx, text=paragraph.text))
                    logger.debug(f"Найден текст '{text_to_find}' в параграфе {idx}: {paragraph.text[:100]}...")
        except Exception as e:
            logger.error(f"Ошибка локального поиска текста: {e}")
        
        return matches

    async def _safe_find_text(self, filename: str, text_to_find: str, match_case: bool = True) -> List[MCPTextMatch]:
        """
        Безопасный поиск текста с fallback на локальный поиск.
        """
        try:
            return await mcp_client.find_text_in_document(filename, text_to_find, match_case=match_case)
        except RuntimeError as e:
            error_details = str(e)
            logger.warning(f"⚠️ MCP сервер недоступен при поиске текста, используем локальный поиск: {error_details}")
            logger.debug(f"   Файл: {filename}, Текст: {text_to_find[:50]}...")
            return self._find_text_locally(filename, text_to_find, match_case)
        except Exception as e:
            # Ловим все типы ошибок (ConnectionError, TimeoutError, etc.)
            error_type = type(e).__name__
            error_details = str(e)
            logger.warning(f"⚠️ Ошибка MCP при поиске текста ({error_type}): {error_details}, используем локальный поиск")
            logger.debug(f"   Файл: {filename}, Текст: {text_to_find[:50]}...")
            return self._find_text_locally(filename, text_to_find, match_case)

    async def _safe_get_document_text(self, filename: str) -> str:
        """
        Безопасное получение текста документа с fallback на локальное чтение.
        """
        try:
            return await mcp_client.get_document_text(filename)
        except RuntimeError as e:
            logger.warning(f"MCP сервер недоступен, используем локальное чтение: {e}")
            return self._get_document_text_locally(filename)
    
    async def _enhanced_text_search(
        self, 
        filename: str, 
        target_text: str, 
        description: str, 
        match_case: bool = False
    ) -> List[MCPTextMatch]:
        """
        НОВЫЙ ФУНКЦИОНАЛ: Расширенный поиск текста с различными вариантами.
        
        Пытается найти текст различными способами:
        1. Поиск части текста (если target_text длинный)
        2. Поиск без учета регистра (если match_case=False)
        3. Поиск с различными вариантами пунктуации
        4. Поиск в таблицах
        5. Поиск по ключевым словам из описания
        
        Args:
            filename: Путь к файлу
            target_text: Искомый текст
            description: Описание изменения
            match_case: Учитывать регистр
            
        Returns:
            Список найденных совпадений или пустой список
        """
        logger.info(f"🔍 Расширенный поиск: '{target_text}'")
        
        # Стратегия 1: Поиск части текста (если текст длинный)
        if len(target_text) > 20:
            # Пробуем найти по ключевым словам
            words = target_text.split()
            if len(words) > 3:
                # Берем первые 3-4 слова
                partial_text = " ".join(words[:4])
                logger.info(f"   Попытка 1: поиск по части текста '{partial_text}'")
                matches = await self._safe_find_text(filename, partial_text, match_case=False)
                if matches:
                    logger.info(f"   ✅ Найдено по части текста: {len(matches)} совпадений")
                    return matches
        
        # Стратегия 2: Поиск без учета регистра (если еще не пробовали)
        if match_case:
            logger.info(f"   Попытка 2: поиск без учета регистра")
            matches = await self._safe_find_text(filename, target_text, match_case=False)
            if matches:
                logger.info(f"   ✅ Найдено без учета регистра: {len(matches)} совпадений")
                return matches
        
        # Стратегия 3: Поиск с различными вариантами пунктуации
        # Удаляем лишние пробелы и пунктуацию для поиска
        cleaned_text = re.sub(r'[^\w\s]', '', target_text)
        cleaned_text = " ".join(cleaned_text.split())
        if cleaned_text != target_text and len(cleaned_text) > 5:
            logger.info(f"   Попытка 3: поиск очищенного текста '{cleaned_text}'")
            matches = await self._safe_find_text(filename, cleaned_text, match_case=False)
            if matches:
                logger.info(f"   ✅ Найдено по очищенному тексту: {len(matches)} совпадений")
                return matches
        
        # Стратегия 4: Поиск в таблицах (если описание указывает на таблицу)
        if "таблиц" in description.lower():
            logger.info(f"   Попытка 4: поиск в таблицах")
            try:
                doc = Document(filename)
                for table_idx, table in enumerate(doc.tables):
                    for row_idx, row in enumerate(table.rows):
                        for cell_idx, cell in enumerate(row.cells):
                            cell_text = cell.text.strip()
                            if target_text.lower() in cell_text.lower() or cell_text.lower() in target_text.lower():
                                logger.info(f"   ✅ Найдено в таблице {table_idx}, строка {row_idx}, ячейка {cell_idx}")
                                # Создаем псевдо-совпадение для таблицы
                                # Возвращаем пустой список, так как для таблиц используется другая логика
                                return []
            except Exception as e:
                logger.debug(f"   Ошибка при поиске в таблицах: {e}")
        
        # Стратегия 5: Поиск по ключевым словам из описания
        if "пункт" in description.lower():
            # Извлекаем номер пункта из описания
            punkt_match = re.search(r'пункт[еа]?\s+(\d+)', description, re.IGNORECASE)
            if punkt_match:
                punkt_num = punkt_match.group(1)
                logger.info(f"   Попытка 5: поиск пункта {punkt_num}")
                # Пробуем различные форматы номера пункта
                for variant in [f"{punkt_num}.", f"{punkt_num})", f"{punkt_num} ", f" {punkt_num}."]:
                    matches = await self._safe_find_text(filename, variant, match_case=False)
                    if matches:
                        logger.info(f"   ✅ Найдено по номеру пункта '{variant}': {len(matches)} совпадений")
                        return matches
        
        # Стратегия 6: Поиск по первым словам target_text
        if len(target_text.split()) > 1:
            first_words = " ".join(target_text.split()[:2])
            if len(first_words) > 5:
                logger.info(f"   Попытка 6: поиск по первым словам '{first_words}'")
                matches = await self._safe_find_text(filename, first_words, match_case=False)
                if matches:
                    logger.info(f"   ✅ Найдено по первым словам: {len(matches)} совпадений")
                    return matches
        
        logger.info(f"   ❌ Расширенный поиск не дал результатов")
        return []
    
    async def _enhanced_replace_attempt(
        self, 
        doc: Document, 
        target_text: str, 
        new_text: str, 
        paragraph_index: int
    ) -> bool:
        """
        НОВЫЙ ФУНКЦИОНАЛ: Расширенная попытка замены текста с различными вариантами.
        
        Пытается заменить текст различными способами:
        1. Замена с нормализацией пробелов
        2. Замена с учетом различных вариантов пунктуации
        3. Замена по части текста
        4. Замена в соседних параграфах
        
        Args:
            doc: Документ
            target_text: Искомый текст
            new_text: Новый текст
            paragraph_index: Индекс параграфа для начала поиска
            
        Returns:
            True если замена выполнена, False иначе
        """
        logger.info(f"🔧 Расширенная попытка замены: '{target_text}' → '{new_text}'")
        
        # Стратегия 1: Нормализация пробелов
        normalized_target = " ".join(target_text.split())
        if normalized_target != target_text:
            logger.info(f"   Попытка 1: замена с нормализованными пробелами")
            # Пробуем сначала в указанном параграфе
            if paragraph_index < len(doc.paragraphs):
                if self._robust_replace_in_paragraph(doc.paragraphs[paragraph_index], normalized_target, new_text):
                    logger.info(f"   ✅ Замена выполнена с нормализованными пробелами (надежный метод)")
                    return True
            # Затем по всему документу
            for para in doc.paragraphs:
                if self._replace_in_paragraph(para, normalized_target, new_text):
                    logger.info(f"   ✅ Замена выполнена с нормализованными пробелами")
                    return True
                # Также пробуем надежную замену
                if self._robust_replace_in_paragraph(para, normalized_target, new_text):
                    logger.info(f"   ✅ Замена выполнена с нормализованными пробелами (надежный метод)")
                    return True
        
        # Стратегия 2: Замена с различными вариантами пунктуации
        # Удаляем пунктуацию для поиска
        cleaned_target = re.sub(r'[^\w\s]', '', target_text)
        cleaned_target = " ".join(cleaned_target.split())
        if cleaned_target != target_text and len(cleaned_target) > 5:
            logger.info(f"   Попытка 2: замена очищенного текста '{cleaned_target}'")
            # Ищем параграфы, содержащие очищенный текст
            for para in doc.paragraphs:
                para_text_cleaned = re.sub(r'[^\w\s]', '', para.text)
                para_text_cleaned = " ".join(para_text_cleaned.split())
                if cleaned_target.lower() in para_text_cleaned.lower():
                    # Пытаемся заменить оригинальный текст в параграфе (сначала надежная замена)
                    if self._robust_replace_in_paragraph(para, target_text, new_text):
                        logger.info(f"   ✅ Замена выполнена по очищенному тексту (надежный метод)")
                        return True
                    if self._replace_in_paragraph(para, target_text, new_text):
                        logger.info(f"   ✅ Замена выполнена по очищенному тексту")
                        return True
        
        # Стратегия 3: Замена по части текста (если текст длинный)
        if len(target_text) > 20:
            words = target_text.split()
            if len(words) > 3:
                # Берем первые 3-4 слова
                partial_text = " ".join(words[:4])
                logger.info(f"   Попытка 3: замена по части текста '{partial_text}'")
                # Ищем параграфы с этой частью текста
                for para in doc.paragraphs:
                    if partial_text.lower() in para.text.lower():
                        # Пытаемся заменить полный текст (сначала надежная замена)
                        if self._robust_replace_in_paragraph(para, target_text, new_text):
                            logger.info(f"   ✅ Замена выполнена по части текста (надежный метод)")
                            return True
                        if self._replace_in_paragraph(para, target_text, new_text):
                            logger.info(f"   ✅ Замена выполнена по части текста")
                            return True
        
        # Стратегия 4: Замена в соседних параграфах (вокруг найденного)
        logger.info(f"   Попытка 4: замена в соседних параграфах (индекс {paragraph_index})")
        start_idx = max(0, paragraph_index - 2)
        end_idx = min(len(doc.paragraphs), paragraph_index + 3)
        for idx in range(start_idx, end_idx):
            if idx != paragraph_index:
                # Пробуем сначала надежную замену
                if self._robust_replace_in_paragraph(doc.paragraphs[idx], target_text, new_text):
                    logger.info(f"   ✅ Замена выполнена в соседнем параграфе {idx} (надежный метод)")
                    return True
                if self._replace_in_paragraph(doc.paragraphs[idx], target_text, new_text):
                    logger.info(f"   ✅ Замена выполнена в соседнем параграфе {idx}")
                    return True
        
        logger.info(f"   ❌ Расширенная попытка замены не дала результатов")
        return False

    async def _sync_heading_with_table_of_contents(
        self,
        filename: str,
        old_heading_text: str,
        new_heading_text: str,
        is_heading_change: bool = True,
        master_doc: Optional[Document] = None  # Единый объект документа для всех изменений
    ) -> None:
        """
        НОВЫЙ ФУНКЦИОНАЛ: Синхронизация изменений между заголовками разделов и содержанием (оглавлением).
        
        Если изменяется заголовок раздела, обновляется соответствующий элемент в содержании.
        Если изменяется элемент в содержании, обновляется соответствующий заголовок раздела.
        
        Args:
            filename: Путь к файлу
            old_heading_text: Старый текст заголовка
            new_heading_text: Новый текст заголовка
            is_heading_change: True если изменяется заголовок, False если изменяется содержание
        """
        try:
            logger.info(f"🔄 Синхронизация заголовка с содержанием: '{old_heading_text}' → '{new_heading_text}'")
            
            # Используем master_doc, если передан, иначе создаем новый
            if master_doc is not None:
                doc = master_doc
                logger.info(f"📄 Используем master_doc для синхронизации")
            else:
                doc = Document(filename)
            synced_count = 0
            
            # Ищем текст в таблицах (содержание обычно хранится в таблицах)
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        
                        # Проверяем, содержит ли ячейка старый текст заголовка
                        # Учитываем, что в содержании может быть только часть текста (без номера или с номером страницы)
                        if old_heading_text in cell_text or cell_text in old_heading_text:
                            # Найдено совпадение в ячейке содержания
                            logger.info(f"   📋 Найдено в содержании (таблица {table_idx}, строка {row_idx}, ячейка {cell_idx}): '{cell_text}'")
                            
                            # Определяем, какая часть ячейки содержит заголовок
                            # В содержании может быть формат: "1. Название раздела ........ 5"
                            # Или просто: "1. Название раздела"
                            
                            # Ищем позицию старого текста в ячейке
                            if old_heading_text in cell_text:
                                # Заменяем старый текст на новый
                                new_cell_text = cell_text.replace(old_heading_text, new_heading_text, 1)
                                
                                # Обновляем ячейку (очищаем и записываем новый текст)
                                cell.text = new_cell_text
                                synced_count += 1
                                logger.info(f"   ✅ Обновлено в содержании: '{cell_text}' → '{new_cell_text}'")
                            elif cell_text in old_heading_text:
                                # Ячейка содержит только часть заголовка, заменяем всю ячейку
                                # Сохраняем форматирование (номер страницы, точки и т.д.)
                                # Пытаемся сохранить номер страницы и точки, если они есть
                                page_match = re.search(r'([. ]+)(\d+)$', cell_text)
                                if page_match:
                                    # Сохраняем разделитель и номер страницы
                                    separator = page_match.group(1)
                                    page_num = page_match.group(2)
                                    # Извлекаем номер раздела из старого текста, если есть
                                    heading_num_match = re.match(r'^(\d+\.?\s*)', old_heading_text)
                                    if heading_num_match:
                                        heading_num = heading_num_match.group(1)
                                        new_cell_text = heading_num + new_heading_text.replace(heading_num, '').strip() + separator + page_num
                                    else:
                                        new_cell_text = new_heading_text + separator + page_num
                                else:
                                    # Просто заменяем текст
                                    heading_num_match = re.match(r'^(\d+\.?\s*)', cell_text)
                                    if heading_num_match:
                                        heading_num = heading_num_match.group(1)
                                        new_cell_text = heading_num + new_heading_text.replace(heading_num, '').strip() if heading_num in new_heading_text else heading_num + new_heading_text
                                    else:
                                        new_cell_text = new_heading_text
                                
                                cell.text = new_cell_text
                                synced_count += 1
                                logger.info(f"   ✅ Обновлено в содержании (с сохранением форматирования): '{cell_text}' → '{new_cell_text}'")
            
            # Если изменяется содержание и не найдено в таблицах, ищем в параграфах (на случай полей TOC)
            if not is_heading_change or synced_count == 0:
                # Ищем текст в параграфах (для полей TOC или обычного текстового содержания)
                for para_idx, para in enumerate(doc.paragraphs):
                    para_text = para.text.strip()
                    if old_heading_text in para_text:
                        # Проверяем, не является ли это самим заголовком раздела
                        if is_heading_change and self._is_heading(para):
                            continue  # Пропускаем сам заголовок
                        
                        # Заменяем в параграфе
                        new_para_text = para_text.replace(old_heading_text, new_heading_text, 1)
                        para.text = new_para_text
                        synced_count += 1
                        logger.info(f"   ✅ Обновлено в содержании (параграф {para_idx}): '{para_text}' → '{new_para_text}'")
            
            # НОВЫЙ ФУНКЦИОНАЛ: Если изменяется содержание, обновляем соответствующий заголовок раздела
            if not is_heading_change and synced_count > 0:
                logger.info(f"   🔄 Ищем соответствующий заголовок раздела для синхронизации...")
                # Ищем заголовок раздела, содержащий старый текст
                for para_idx, para in enumerate(doc.paragraphs):
                    if self._is_heading(para):
                        para_text = para.text.strip()
                        # Извлекаем текст без номера для сравнения
                        # Убираем номер раздела из начала для сравнения
                        heading_text_no_num = re.sub(r'^\d+\.?\s*', '', para_text).strip()
                        old_text_no_num = re.sub(r'^\d+\.?\s*', '', old_heading_text).strip()
                        
                        # Проверяем совпадение текста заголовка
                        if old_text_no_num in heading_text_no_num or heading_text_no_num in old_text_no_num or old_heading_text in para_text:
                            logger.info(f"   📌 Найден заголовок раздела (параграф {para_idx}): '{para_text}'")
                            # Обновляем заголовок
                            if old_heading_text in para_text:
                                new_para_text = para_text.replace(old_heading_text, new_heading_text, 1)
                            else:
                                # Заменяем с сохранением номера раздела
                                heading_num_match = re.match(r'^(\d+\.?\s*)', para_text)
                                if heading_num_match:
                                    heading_num = heading_num_match.group(1)
                                    # Убираем номер из нового текста, если он там есть
                                    new_text_clean = re.sub(r'^\d+\.?\s*', '', new_heading_text).strip()
                                    new_para_text = heading_num + new_text_clean
                                else:
                                    new_para_text = new_heading_text
                            
                            heading_style = para.style
                            para.text = new_para_text
                            if heading_style:
                                para.style = heading_style
                            synced_count += 1
                            logger.info(f"   ✅ Обновлен заголовок раздела: '{para_text}' → '{new_para_text}'")
                            break
            
            if synced_count > 0:
                # КРИТИЧЕСКОЕ: НЕ сохраняем файл здесь, если используется master_doc
                # Файл будет сохранен один раз в конце после всех изменений
                if master_doc is None:
                    doc.save(filename)
                    logger.info(f"💾 Файл сохранен после синхронизации (master_doc не использовался)")
                if is_heading_change:
                    logger.info(f"✅ Синхронизация завершена: обновлено {synced_count} элементов в содержании")
                else:
                    logger.info(f"✅ Синхронизация завершена: обновлено {synced_count} элементов (содержание и заголовок)")
            else:
                logger.info(f"ℹ️ Элементы для синхронизации не найдены")
                
        except Exception as e:
            logger.warning(f"⚠️ Ошибка при синхронизации с содержанием: {e}", exc_info=True)

    def _get_document_text_locally(self, filename: str) -> str:
        """
        Локальное получение текста документа через python-docx.
        """
        try:
            doc = Document(filename)
            text_parts = []
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Ошибка локального чтения документа: {e}")
            return ""

    @staticmethod
    def _is_heading(paragraph: Paragraph) -> bool:
        try:
            style_name = paragraph.style.name if paragraph.style else ""
        except ValueError:
            style_name = ""
        return style_name.startswith("Heading")

    def _find_section_end(self, doc: Document, start_index: int) -> int:
        """
        Находит индекс первого параграфа после текущего раздела.
        """
        start_para = doc.paragraphs[start_index]
        if not self._is_heading(start_para):
            return start_index + 1

        for idx in range(start_index + 1, len(doc.paragraphs)):
            if self._is_heading(doc.paragraphs[idx]):
                return idx
        return len(doc.paragraphs)

    @staticmethod
    def _delete_paragraph(paragraph: Paragraph) -> None:
        p = paragraph._element  # noqa: SLF001
        parent = p.getparent()
        if parent is not None:
            parent.remove(p)

    @staticmethod
    def _insert_paragraph_after(paragraph: Paragraph, text: str = "", style: Optional[str] = None) -> Paragraph:
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)  # noqa: SLF001
        new_para = Paragraph(new_p, paragraph._parent)  # noqa: SLF001
        if text:
            new_para.add_run(text)
        if style:
            # Безопасная установка стиля с проверкой наличия
            try:
                # Получаем документ из paragraph
                doc = paragraph._parent  # noqa: SLF001
                if hasattr(doc, 'styles'):
                    # Пробуем сначала указанный стиль
                    if style in doc.styles:
                        new_para.style = style
                    else:
                        # Пробуем альтернативные стили
                        fallback_styles = []
                        if "Heading" in style:
                            # Для заголовков пробуем разные уровни
                            for level in range(1, 10):
                                fallback_styles.append(f"Heading {level}")
                        fallback_styles.extend(["Normal", "Default Paragraph Font"])
                        
                        style_set = False
                        for fallback_style in fallback_styles:
                            try:
                                if fallback_style in doc.styles:
                                    new_para.style = fallback_style
                                    style_set = True
                                    logger.debug(f"Использован альтернативный стиль '{fallback_style}' вместо '{style}'")
                                    break
                            except (KeyError, ValueError, AttributeError):
                                continue
                        
                        if not style_set:
                            logger.warning(f"Не удалось установить стиль '{style}' и альтернативы, параграф будет без стиля")
            except Exception as e:
                logger.warning(f"Ошибка при установке стиля '{style}': {e}, параграф будет без стиля")
        return new_para

    def _find_paragraph_index_by_text(
        self,
        doc: Document,
        text: str,
        start: int = 0,
        style: Optional[str] = None,
    ) -> Optional[int]:
        for idx in range(start, len(doc.paragraphs)):
            para = doc.paragraphs[idx]
            if para.text != text:
                continue
            if style:
                if self._get_style_name(para) != style:
                    continue
            return idx
        return None

    @staticmethod
    def _get_style_name(paragraph: Paragraph) -> str:
        try:
            return paragraph.style.name if paragraph.style else ""
        except ValueError:
            return ""

    async def close(self) -> None:
        """
        Завершение работы агента.
        """
        if self._openai_http_client:
            await self._openai_http_client.aclose()
            self._openai_http_client = None

    @staticmethod
    def _patch_openai_httpx() -> None:
        """
        Обход несовместимости openai>=1.51.2 (ожидает httpx с параметром proxies)
        и httpx>=0.28 (параметр proxies удалён).
        """
        if "proxies" in inspect.signature(httpx.AsyncClient.__init__).parameters:
            return

        try:
            import openai._base_client as base_client  # type: ignore
        except Exception:
            return

        def strip_proxies(init):
            @wraps(init)
            def wrapper(self, *args, **kwargs):
                kwargs.pop("proxies", None)
                return init(self, *args, **kwargs)

            return wrapper

        target_inits = []

        # Default async client used inside openai sdk
        if hasattr(base_client, "_DefaultAsyncHttpxClient"):
            target_inits.append(
                ("_DefaultAsyncHttpxClient", base_client._DefaultAsyncHttpxClient)
            )
        if hasattr(base_client, "AsyncHttpxClientWrapper"):
            target_inits.append(
                ("AsyncHttpxClientWrapper", base_client.AsyncHttpxClientWrapper)
            )
        if hasattr(base_client, "AsyncAPIClient"):
            target_inits.append(("AsyncAPIClient", base_client.AsyncAPIClient))

        for _, cls in target_inits:
            init = getattr(cls, "__init__", None)
            if not callable(init):
                continue
            patched = strip_proxies(init)
            setattr(cls, "__init__", patched)


document_agent = DocumentChangeAgent()

