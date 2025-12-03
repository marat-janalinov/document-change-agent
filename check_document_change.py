#!/usr/bin/env python3
"""
Простой скрипт для проверки, применилось ли изменение в документе
"""
from docx import Document
from pathlib import Path

def check_document():
    """Проверка содержимого документа"""
    
    base_dir = Path(__file__).parent / "data" / "Пилотный проект"
    source_file = base_dir / "Регламент  03122025.docx"
    
    if not source_file.exists():
        print(f"❌ Файл не найден: {source_file}")
        return
    
    print(f"📄 Проверка файла: {source_file}\n")
    
    try:
        doc = Document(str(source_file))
        
        # Ищем текст "Глава 1. ОПРЕДЕЛЕНИЯ И ТОЛКОВАНИЯ"
        target_text = "Глава 1. ОПРЕДЕЛЕНИЯ И ТОЛКОВАНИЯ"
        expected_new_text = "Глава 1. ОПРЕДЕЛЕНИЯ И ТОЛКОВАНИЯ теста"
        
        print(f"🔍 Ищем: '{target_text}'")
        print(f"🔍 Ожидаем после замены: '{expected_new_text}'\n")
        
        found_original = False
        found_modified = False
        
        print("="*60)
        print("ПРОВЕРКА ПАРАГРАФОВ:")
        print("="*60)
        
        for idx, para in enumerate(doc.paragraphs[:20]):  # Первые 20 параграфов
            para_text = para.text.strip()
            
            if not para_text:
                continue
            
            if target_text in para_text:
                found_original = True
                print(f"\n📍 Параграф {idx}:")
                print(f"   Текст: '{para_text[:100]}...'")
                print(f"   Длина: {len(para_text)}")
                print(f"   Количество runs: {len(para.runs)}")
                
                # Проверяем runs
                if len(para.runs) > 0:
                    print(f"   Детали runs:")
                    for run_idx, run in enumerate(para.runs):
                        run_text = run.text
                        if run_text:
                            print(f"      Run {run_idx}: '{run_text[:50]}...' (длина: {len(run_text)})")
            
            if expected_new_text in para_text:
                found_modified = True
                print(f"\n✅ НАЙДЕНО ИЗМЕНЕНИЕ в параграфе {idx}:")
                print(f"   Текст: '{para_text}'")
        
        print("\n" + "="*60)
        print("РЕЗУЛЬТАТЫ ПРОВЕРКИ:")
        print("="*60)
        
        if found_modified:
            print("✅ ИЗМЕНЕНИЕ ПРИМЕНИЛОСЬ!")
        elif found_original:
            print("⚠️  Исходный текст найден, но изменение не применено")
        else:
            print("❌ Исходный текст не найден в первых 20 параграфах")
        
        # Дополнительная проверка - поиск по всему документу
        print("\n" + "="*60)
        print("ПОЛНАЯ ПРОВЕРКА ДОКУМЕНТА:")
        print("="*60)
        
        all_text = "\n".join([p.text for p in doc.paragraphs])
        
        if target_text in all_text:
            print(f"✅ Исходный текст найден в документе")
            count = all_text.count(target_text)
            print(f"   Вхождений: {count}")
        
        if expected_new_text in all_text:
            print(f"✅ Измененный текст найден в документе")
            count = all_text.count(expected_new_text)
            print(f"   Вхождений: {count}")
        else:
            print(f"❌ Измененный текст НЕ найден в документе")
        
    except Exception as e:
        print(f"❌ Ошибка: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    check_document()

